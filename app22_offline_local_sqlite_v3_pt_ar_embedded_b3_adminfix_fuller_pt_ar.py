import io
import json
import hashlib
import hmac
import secrets
import os
import re
import sqlite3
import zipfile
import uuid
import time
from datetime import datetime, timezone
from typing import Any, Dict, List, Tuple, Optional
from pathlib import Path

import numpy as np
import pandas as pd
import streamlit as st

# =========================
# Configuration Streamlit
# =========================

# IMPORTANT : doit être le tout premier appel Streamlit
st.set_page_config(
    page_title="Consultation STATAFRIC",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# Stabilisation UI (évite les variations de largeur dues à l’apparition/disparition de la barre de défilement)
st.markdown("""
<style>
html { overflow-y: scroll; }
div[data-testid="stVerticalBlock"] { gap: 0.75rem; }
</style>
""", unsafe_allow_html=True)


# CSS léger pour stabiliser le rendu (évite les variations de largeur perçues)
st.markdown(
    """
<style>
/* Largeur stable du conteneur principal */
.block-container { max-width: 1200px; padding-top: 1rem; padding-bottom: 2rem; }
@media (max-width: 1200px) { .block-container { max-width: 100%; } }

/* Réduit les effets de "sauts" quand des messages apparaissent/disparaissent */
div[data-testid="stVerticalBlock"] { gap: 0.75rem; }

/* Boutons : largeur stable */
button[kind="primary"], button[kind="secondary"] { white-space: nowrap; }
</style>
""",
    unsafe_allow_html=True,
)


# =========================
# Configuration
# =========================

APP_TITLE_FR = "Questionnaire de consultation pour l'identification des statistiques prioritaires"
APP_TITLE_EN = "Consultation questionnaire for identifying priority statistics"
APP_TITLE_PT = "Questionário de consulta para a identificação de estatísticas prioritárias"
APP_TITLE_AR = "استبيان تشاوري لتحديد الإحصاءات ذات الأولوية"

APP_BUILD = "2026-02-28 b3 embedded-pt-ar"

EMBEDDED_COUNTRIES_ROWS = [
  {
    "COUNTRY_ISO3": "DZA",
    "COUNTRY_NAME_FR": "Algérie",
    "COUNTRY_VALUE": "DZA | Algérie",
    "COUNTRY_NAME_EN": "Algeria",
    "COUNTRY_VALUE_EN": "DZA | Algeria",
    "COUNTRY_NAME_PT": "Argélia",
    "COUNTRY_VALUE_PT": "DZA | Argélia",
    "COUNTRY_NAME_AR": "الجزائر",
    "COUNTRY_VALUE_AR": "DZA | الجزائر"
  },
  {
    "COUNTRY_ISO3": "AGO",
    "COUNTRY_NAME_FR": "Angola",
    "COUNTRY_VALUE": "AGO | Angola",
    "COUNTRY_NAME_EN": "Angola",
    "COUNTRY_VALUE_EN": "AGO | Angola",
    "COUNTRY_NAME_PT": "Angola",
    "COUNTRY_VALUE_PT": "AGO | Angola",
    "COUNTRY_NAME_AR": "أنغولا",
    "COUNTRY_VALUE_AR": "AGO | أنغولا"
  },
  {
    "COUNTRY_ISO3": "BEN",
    "COUNTRY_NAME_FR": "Bénin",
    "COUNTRY_VALUE": "BEN | Bénin",
    "COUNTRY_NAME_EN": "Benin",
    "COUNTRY_VALUE_EN": "BEN | Benin",
    "COUNTRY_NAME_PT": "Benim",
    "COUNTRY_VALUE_PT": "BEN | Benim",
    "COUNTRY_NAME_AR": "بنين",
    "COUNTRY_VALUE_AR": "BEN | بنين"
  },
  {
    "COUNTRY_ISO3": "BWA",
    "COUNTRY_NAME_FR": "Botswana",
    "COUNTRY_VALUE": "BWA | Botswana",
    "COUNTRY_NAME_EN": "Botswana",
    "COUNTRY_VALUE_EN": "BWA | Botswana",
    "COUNTRY_NAME_PT": "Botsuana",
    "COUNTRY_VALUE_PT": "BWA | Botsuana",
    "COUNTRY_NAME_AR": "بوتسوانا",
    "COUNTRY_VALUE_AR": "BWA | بوتسوانا"
  },
  {
    "COUNTRY_ISO3": "BFA",
    "COUNTRY_NAME_FR": "Burkina Faso",
    "COUNTRY_VALUE": "BFA | Burkina Faso",
    "COUNTRY_NAME_EN": "Burkina Faso",
    "COUNTRY_VALUE_EN": "BFA | Burkina Faso",
    "COUNTRY_NAME_PT": "Burkina Faso",
    "COUNTRY_VALUE_PT": "BFA | Burkina Faso",
    "COUNTRY_NAME_AR": "بوركينا فاسو",
    "COUNTRY_VALUE_AR": "BFA | بوركينا فاسو"
  },
  {
    "COUNTRY_ISO3": "BDI",
    "COUNTRY_NAME_FR": "Burundi",
    "COUNTRY_VALUE": "BDI | Burundi",
    "COUNTRY_NAME_EN": "Burundi",
    "COUNTRY_VALUE_EN": "BDI | Burundi",
    "COUNTRY_NAME_PT": "Burundi",
    "COUNTRY_VALUE_PT": "BDI | Burundi",
    "COUNTRY_NAME_AR": "بوروندي",
    "COUNTRY_VALUE_AR": "BDI | بوروندي"
  },
  {
    "COUNTRY_ISO3": "CMR",
    "COUNTRY_NAME_FR": "Cameroun",
    "COUNTRY_VALUE": "CMR | Cameroun",
    "COUNTRY_NAME_EN": "Cameroon",
    "COUNTRY_VALUE_EN": "CMR | Cameroon",
    "COUNTRY_NAME_PT": "Camarões",
    "COUNTRY_VALUE_PT": "CMR | Camarões",
    "COUNTRY_NAME_AR": "الكاميرون",
    "COUNTRY_VALUE_AR": "CMR | الكاميرون"
  },
  {
    "COUNTRY_ISO3": "CPV",
    "COUNTRY_NAME_FR": "Cabo Verde",
    "COUNTRY_VALUE": "CPV | Cabo Verde",
    "COUNTRY_NAME_EN": "Cabo Verde",
    "COUNTRY_VALUE_EN": "CPV | Cabo Verde",
    "COUNTRY_NAME_PT": "Cabo Verde",
    "COUNTRY_VALUE_PT": "CPV | Cabo Verde",
    "COUNTRY_NAME_AR": "الرأس الأخضر",
    "COUNTRY_VALUE_AR": "CPV | الرأس الأخضر"
  },
  {
    "COUNTRY_ISO3": "CAF",
    "COUNTRY_NAME_FR": "République centrafricaine",
    "COUNTRY_VALUE": "CAF | République centrafricaine",
    "COUNTRY_NAME_EN": "Central African Republic",
    "COUNTRY_VALUE_EN": "CAF | Central African Republic",
    "COUNTRY_NAME_PT": "República Centro-Africana",
    "COUNTRY_VALUE_PT": "CAF | República Centro-Africana",
    "COUNTRY_NAME_AR": "جمهورية أفريقيا الوسطى",
    "COUNTRY_VALUE_AR": "CAF | جمهورية أفريقيا الوسطى"
  },
  {
    "COUNTRY_ISO3": "TCD",
    "COUNTRY_NAME_FR": "Tchad",
    "COUNTRY_VALUE": "TCD | Tchad",
    "COUNTRY_NAME_EN": "Chad",
    "COUNTRY_VALUE_EN": "TCD | Chad",
    "COUNTRY_NAME_PT": "Chade",
    "COUNTRY_VALUE_PT": "TCD | Chade",
    "COUNTRY_NAME_AR": "تشاد",
    "COUNTRY_VALUE_AR": "TCD | تشاد"
  },
  {
    "COUNTRY_ISO3": "COM",
    "COUNTRY_NAME_FR": "Comores",
    "COUNTRY_VALUE": "COM | Comores",
    "COUNTRY_NAME_EN": "Comoros",
    "COUNTRY_VALUE_EN": "COM | Comoros",
    "COUNTRY_NAME_PT": "Comores",
    "COUNTRY_VALUE_PT": "COM | Comores",
    "COUNTRY_NAME_AR": "جزر القمر",
    "COUNTRY_VALUE_AR": "COM | جزر القمر"
  },
  {
    "COUNTRY_ISO3": "COG",
    "COUNTRY_NAME_FR": "Congo",
    "COUNTRY_VALUE": "COG | Congo",
    "COUNTRY_NAME_EN": "Republic of the Congo",
    "COUNTRY_VALUE_EN": "COG | Republic of the Congo",
    "COUNTRY_NAME_PT": "República do Congo",
    "COUNTRY_VALUE_PT": "COG | República do Congo",
    "COUNTRY_NAME_AR": "جمهورية الكونغو",
    "COUNTRY_VALUE_AR": "COG | جمهورية الكونغو"
  },
  {
    "COUNTRY_ISO3": "CIV",
    "COUNTRY_NAME_FR": "Côte d'Ivoire",
    "COUNTRY_VALUE": "CIV | Côte d'Ivoire",
    "COUNTRY_NAME_EN": "Côte d'Ivoire",
    "COUNTRY_VALUE_EN": "CIV | Côte d'Ivoire",
    "COUNTRY_NAME_PT": "Costa do Marfim",
    "COUNTRY_VALUE_PT": "CIV | Costa do Marfim",
    "COUNTRY_NAME_AR": "كوت ديفوار",
    "COUNTRY_VALUE_AR": "CIV | كوت ديفوار"
  },
  {
    "COUNTRY_ISO3": "COD",
    "COUNTRY_NAME_FR": "République démocratique du Congo",
    "COUNTRY_VALUE": "COD | République démocratique du Congo",
    "COUNTRY_NAME_EN": "Democratic Republic of the Congo",
    "COUNTRY_VALUE_EN": "COD | Democratic Republic of the Congo",
    "COUNTRY_NAME_PT": "República Democrática do Congo",
    "COUNTRY_VALUE_PT": "COD | República Democrática do Congo",
    "COUNTRY_NAME_AR": "جمهورية الكونغو الديمقراطية",
    "COUNTRY_VALUE_AR": "COD | جمهورية الكونغو الديمقراطية"
  },
  {
    "COUNTRY_ISO3": "DJI",
    "COUNTRY_NAME_FR": "Djibouti",
    "COUNTRY_VALUE": "DJI | Djibouti",
    "COUNTRY_NAME_EN": "Djibouti",
    "COUNTRY_VALUE_EN": "DJI | Djibouti",
    "COUNTRY_NAME_PT": "Djibuti",
    "COUNTRY_VALUE_PT": "DJI | Djibuti",
    "COUNTRY_NAME_AR": "جيبوتي",
    "COUNTRY_VALUE_AR": "DJI | جيبوتي"
  },
  {
    "COUNTRY_ISO3": "EGY",
    "COUNTRY_NAME_FR": "Egypte",
    "COUNTRY_VALUE": "EGY | Egypte",
    "COUNTRY_NAME_EN": "Egypt",
    "COUNTRY_VALUE_EN": "EGY | Egypt",
    "COUNTRY_NAME_PT": "Egito",
    "COUNTRY_VALUE_PT": "EGY | Egito",
    "COUNTRY_NAME_AR": "مصر",
    "COUNTRY_VALUE_AR": "EGY | مصر"
  },
  {
    "COUNTRY_ISO3": "GNQ",
    "COUNTRY_NAME_FR": "Guinée équatoriale",
    "COUNTRY_VALUE": "GNQ | Guinée équatoriale",
    "COUNTRY_NAME_EN": "Equatorial Guinea",
    "COUNTRY_VALUE_EN": "GNQ | Equatorial Guinea",
    "COUNTRY_NAME_PT": "Guiné Equatorial",
    "COUNTRY_VALUE_PT": "GNQ | Guiné Equatorial",
    "COUNTRY_NAME_AR": "غينيا الاستوائية",
    "COUNTRY_VALUE_AR": "GNQ | غينيا الاستوائية"
  },
  {
    "COUNTRY_ISO3": "ERI",
    "COUNTRY_NAME_FR": "Érythrée",
    "COUNTRY_VALUE": "ERI | Érythrée",
    "COUNTRY_NAME_EN": "Eritrea",
    "COUNTRY_VALUE_EN": "ERI | Eritrea",
    "COUNTRY_NAME_PT": "Eritreia",
    "COUNTRY_VALUE_PT": "ERI | Eritreia",
    "COUNTRY_NAME_AR": "إريتريا",
    "COUNTRY_VALUE_AR": "ERI | إريتريا"
  },
  {
    "COUNTRY_ISO3": "ETH",
    "COUNTRY_NAME_FR": "Ethiopie",
    "COUNTRY_VALUE": "ETH | Ethiopie",
    "COUNTRY_NAME_EN": "Ethiopia",
    "COUNTRY_VALUE_EN": "ETH | Ethiopia",
    "COUNTRY_NAME_PT": "Etiópia",
    "COUNTRY_VALUE_PT": "ETH | Etiópia",
    "COUNTRY_NAME_AR": "إثيوبيا",
    "COUNTRY_VALUE_AR": "ETH | إثيوبيا"
  },
  {
    "COUNTRY_ISO3": "GAB",
    "COUNTRY_NAME_FR": "Gabon",
    "COUNTRY_VALUE": "GAB | Gabon",
    "COUNTRY_NAME_EN": "Gabon",
    "COUNTRY_VALUE_EN": "GAB | Gabon",
    "COUNTRY_NAME_PT": "Gabão",
    "COUNTRY_VALUE_PT": "GAB | Gabão",
    "COUNTRY_NAME_AR": "الغابون",
    "COUNTRY_VALUE_AR": "GAB | الغابون"
  },
  {
    "COUNTRY_ISO3": "GMB",
    "COUNTRY_NAME_FR": "Gambie",
    "COUNTRY_VALUE": "GMB | Gambie",
    "COUNTRY_NAME_EN": "Gambia",
    "COUNTRY_VALUE_EN": "GMB | Gambia",
    "COUNTRY_NAME_PT": "Gâmbia",
    "COUNTRY_VALUE_PT": "GMB | Gâmbia",
    "COUNTRY_NAME_AR": "غامبيا",
    "COUNTRY_VALUE_AR": "GMB | غامبيا"
  },
  {
    "COUNTRY_ISO3": "GHA",
    "COUNTRY_NAME_FR": "Ghana",
    "COUNTRY_VALUE": "GHA | Ghana",
    "COUNTRY_NAME_EN": "Ghana",
    "COUNTRY_VALUE_EN": "GHA | Ghana",
    "COUNTRY_NAME_PT": "Gana",
    "COUNTRY_VALUE_PT": "GHA | Gana",
    "COUNTRY_NAME_AR": "غانا",
    "COUNTRY_VALUE_AR": "GHA | غانا"
  },
  {
    "COUNTRY_ISO3": "GIN",
    "COUNTRY_NAME_FR": "Guinée",
    "COUNTRY_VALUE": "GIN | Guinée",
    "COUNTRY_NAME_EN": "Guinea",
    "COUNTRY_VALUE_EN": "GIN | Guinea",
    "COUNTRY_NAME_PT": "Guiné",
    "COUNTRY_VALUE_PT": "GIN | Guiné",
    "COUNTRY_NAME_AR": "غينيا",
    "COUNTRY_VALUE_AR": "GIN | غينيا"
  },
  {
    "COUNTRY_ISO3": "GNB",
    "COUNTRY_NAME_FR": "Guinée-Bissau",
    "COUNTRY_VALUE": "GNB | Guinée-Bissau",
    "COUNTRY_NAME_EN": "Guinea-Bissau",
    "COUNTRY_VALUE_EN": "GNB | Guinea-Bissau",
    "COUNTRY_NAME_PT": "Guiné-Bissau",
    "COUNTRY_VALUE_PT": "GNB | Guiné-Bissau",
    "COUNTRY_NAME_AR": "غينيا بيساو",
    "COUNTRY_VALUE_AR": "GNB | غينيا بيساو"
  },
  {
    "COUNTRY_ISO3": "KEN",
    "COUNTRY_NAME_FR": "Kenya",
    "COUNTRY_VALUE": "KEN | Kenya",
    "COUNTRY_NAME_EN": "Kenya",
    "COUNTRY_VALUE_EN": "KEN | Kenya",
    "COUNTRY_NAME_PT": "Quénia",
    "COUNTRY_VALUE_PT": "KEN | Quénia",
    "COUNTRY_NAME_AR": "كينيا",
    "COUNTRY_VALUE_AR": "KEN | كينيا"
  },
  {
    "COUNTRY_ISO3": "LSO",
    "COUNTRY_NAME_FR": "Lesotho",
    "COUNTRY_VALUE": "LSO | Lesotho",
    "COUNTRY_NAME_EN": "Lesotho",
    "COUNTRY_VALUE_EN": "LSO | Lesotho",
    "COUNTRY_NAME_PT": "Lesoto",
    "COUNTRY_VALUE_PT": "LSO | Lesoto",
    "COUNTRY_NAME_AR": "ليسوتو",
    "COUNTRY_VALUE_AR": "LSO | ليسوتو"
  },
  {
    "COUNTRY_ISO3": "LBR",
    "COUNTRY_NAME_FR": "Libéria",
    "COUNTRY_VALUE": "LBR | Libéria",
    "COUNTRY_NAME_EN": "Liberia",
    "COUNTRY_VALUE_EN": "LBR | Liberia",
    "COUNTRY_NAME_PT": "Libéria",
    "COUNTRY_VALUE_PT": "LBR | Libéria",
    "COUNTRY_NAME_AR": "ليبيريا",
    "COUNTRY_VALUE_AR": "LBR | ليبيريا"
  },
  {
    "COUNTRY_ISO3": "LBY",
    "COUNTRY_NAME_FR": "Libye",
    "COUNTRY_VALUE": "LBY | Libye",
    "COUNTRY_NAME_EN": "Libya",
    "COUNTRY_VALUE_EN": "LBY | Libya",
    "COUNTRY_NAME_PT": "Líbia",
    "COUNTRY_VALUE_PT": "LBY | Líbia",
    "COUNTRY_NAME_AR": "ليبيا",
    "COUNTRY_VALUE_AR": "LBY | ليبيا"
  },
  {
    "COUNTRY_ISO3": "MDG",
    "COUNTRY_NAME_FR": "Madagascar",
    "COUNTRY_VALUE": "MDG | Madagascar",
    "COUNTRY_NAME_EN": "Madagascar",
    "COUNTRY_VALUE_EN": "MDG | Madagascar",
    "COUNTRY_NAME_PT": "Madagáscar",
    "COUNTRY_VALUE_PT": "MDG | Madagáscar",
    "COUNTRY_NAME_AR": "مدغشقر",
    "COUNTRY_VALUE_AR": "MDG | مدغشقر"
  },
  {
    "COUNTRY_ISO3": "MWI",
    "COUNTRY_NAME_FR": "Malawi",
    "COUNTRY_VALUE": "MWI | Malawi",
    "COUNTRY_NAME_EN": "Malawi",
    "COUNTRY_VALUE_EN": "MWI | Malawi",
    "COUNTRY_NAME_PT": "Maláui",
    "COUNTRY_VALUE_PT": "MWI | Maláui",
    "COUNTRY_NAME_AR": "ملاوي",
    "COUNTRY_VALUE_AR": "MWI | ملاوي"
  },
  {
    "COUNTRY_ISO3": "MLI",
    "COUNTRY_NAME_FR": "Mali",
    "COUNTRY_VALUE": "MLI | Mali",
    "COUNTRY_NAME_EN": "Mali",
    "COUNTRY_VALUE_EN": "MLI | Mali",
    "COUNTRY_NAME_PT": "Mali",
    "COUNTRY_VALUE_PT": "MLI | Mali",
    "COUNTRY_NAME_AR": "مالي",
    "COUNTRY_VALUE_AR": "MLI | مالي"
  },
  {
    "COUNTRY_ISO3": "MRT",
    "COUNTRY_NAME_FR": "Mauritanie",
    "COUNTRY_VALUE": "MRT | Mauritanie",
    "COUNTRY_NAME_EN": "Mauritania",
    "COUNTRY_VALUE_EN": "MRT | Mauritania",
    "COUNTRY_NAME_PT": "Mauritânia",
    "COUNTRY_VALUE_PT": "MRT | Mauritânia",
    "COUNTRY_NAME_AR": "موريتانيا",
    "COUNTRY_VALUE_AR": "MRT | موريتانيا"
  },
  {
    "COUNTRY_ISO3": "MUS",
    "COUNTRY_NAME_FR": "Île Maurice",
    "COUNTRY_VALUE": "MUS | Île Maurice",
    "COUNTRY_NAME_EN": "Mauritius",
    "COUNTRY_VALUE_EN": "MUS | Mauritius",
    "COUNTRY_NAME_PT": "Maurícia",
    "COUNTRY_VALUE_PT": "MUS | Maurícia",
    "COUNTRY_NAME_AR": "موريشيوس",
    "COUNTRY_VALUE_AR": "MUS | موريشيوس"
  },
  {
    "COUNTRY_ISO3": "MAR",
    "COUNTRY_NAME_FR": "Maroc",
    "COUNTRY_VALUE": "MAR | Maroc",
    "COUNTRY_NAME_EN": "Morocco",
    "COUNTRY_VALUE_EN": "MAR | Morocco",
    "COUNTRY_NAME_PT": "Marrocos",
    "COUNTRY_VALUE_PT": "MAR | Marrocos",
    "COUNTRY_NAME_AR": "المغرب",
    "COUNTRY_VALUE_AR": "MAR | المغرب"
  },
  {
    "COUNTRY_ISO3": "MOZ",
    "COUNTRY_NAME_FR": "Mozambique",
    "COUNTRY_VALUE": "MOZ | Mozambique",
    "COUNTRY_NAME_EN": "Mozambique",
    "COUNTRY_VALUE_EN": "MOZ | Mozambique",
    "COUNTRY_NAME_PT": "Moçambique",
    "COUNTRY_VALUE_PT": "MOZ | Moçambique",
    "COUNTRY_NAME_AR": "موزمبيق",
    "COUNTRY_VALUE_AR": "MOZ | موزمبيق"
  },
  {
    "COUNTRY_ISO3": "NAM",
    "COUNTRY_NAME_FR": "Namibie",
    "COUNTRY_VALUE": "NAM | Namibie",
    "COUNTRY_NAME_EN": "Namibia",
    "COUNTRY_VALUE_EN": "NAM | Namibia",
    "COUNTRY_NAME_PT": "Namíbia",
    "COUNTRY_VALUE_PT": "NAM | Namíbia",
    "COUNTRY_NAME_AR": "ناميبيا",
    "COUNTRY_VALUE_AR": "NAM | ناميبيا"
  },
  {
    "COUNTRY_ISO3": "NER",
    "COUNTRY_NAME_FR": "Niger",
    "COUNTRY_VALUE": "NER | Niger",
    "COUNTRY_NAME_EN": "Niger",
    "COUNTRY_VALUE_EN": "NER | Niger",
    "COUNTRY_NAME_PT": "Níger",
    "COUNTRY_VALUE_PT": "NER | Níger",
    "COUNTRY_NAME_AR": "النيجر",
    "COUNTRY_VALUE_AR": "NER | النيجر"
  },
  {
    "COUNTRY_ISO3": "NGA",
    "COUNTRY_NAME_FR": "Nigeria",
    "COUNTRY_VALUE": "NGA | Nigeria",
    "COUNTRY_NAME_EN": "Nigeria",
    "COUNTRY_VALUE_EN": "NGA | Nigeria",
    "COUNTRY_NAME_PT": "Nigéria",
    "COUNTRY_VALUE_PT": "NGA | Nigéria",
    "COUNTRY_NAME_AR": "نيجيريا",
    "COUNTRY_VALUE_AR": "NGA | نيجيريا"
  },
  {
    "COUNTRY_ISO3": "RWA",
    "COUNTRY_NAME_FR": "Rwanda",
    "COUNTRY_VALUE": "RWA | Rwanda",
    "COUNTRY_NAME_EN": "Rwanda",
    "COUNTRY_VALUE_EN": "RWA | Rwanda",
    "COUNTRY_NAME_PT": "Ruanda",
    "COUNTRY_VALUE_PT": "RWA | Ruanda",
    "COUNTRY_NAME_AR": "رواندا",
    "COUNTRY_VALUE_AR": "RWA | رواندا"
  },
  {
    "COUNTRY_ISO3": "STP",
    "COUNTRY_NAME_FR": "Sao Tomé-et-Principe",
    "COUNTRY_VALUE": "STP | Sao Tomé-et-Principe",
    "COUNTRY_NAME_EN": "Sao Tome and Principe",
    "COUNTRY_VALUE_EN": "STP | Sao Tome and Principe",
    "COUNTRY_NAME_PT": "São Tomé e Príncipe",
    "COUNTRY_VALUE_PT": "STP | São Tomé e Príncipe",
    "COUNTRY_NAME_AR": "ساو تومي وبرينسيبي",
    "COUNTRY_VALUE_AR": "STP | ساو تومي وبرينسيبي"
  },
  {
    "COUNTRY_ISO3": "SEN",
    "COUNTRY_NAME_FR": "Sénégal",
    "COUNTRY_VALUE": "SEN | Sénégal",
    "COUNTRY_NAME_EN": "Senegal",
    "COUNTRY_VALUE_EN": "SEN | Senegal",
    "COUNTRY_NAME_PT": "Senegal",
    "COUNTRY_VALUE_PT": "SEN | Senegal",
    "COUNTRY_NAME_AR": "السنغال",
    "COUNTRY_VALUE_AR": "SEN | السنغال"
  },
  {
    "COUNTRY_ISO3": "SYC",
    "COUNTRY_NAME_FR": "Seychelles",
    "COUNTRY_VALUE": "SYC | Seychelles",
    "COUNTRY_NAME_EN": "Seychelles",
    "COUNTRY_VALUE_EN": "SYC | Seychelles",
    "COUNTRY_NAME_PT": "Seicheles",
    "COUNTRY_VALUE_PT": "SYC | Seicheles",
    "COUNTRY_NAME_AR": "سيشل",
    "COUNTRY_VALUE_AR": "SYC | سيشل"
  },
  {
    "COUNTRY_ISO3": "SLE",
    "COUNTRY_NAME_FR": "Sierra Leone",
    "COUNTRY_VALUE": "SLE | Sierra Leone",
    "COUNTRY_NAME_EN": "Sierra Leone",
    "COUNTRY_VALUE_EN": "SLE | Sierra Leone",
    "COUNTRY_NAME_PT": "Serra Leoa",
    "COUNTRY_VALUE_PT": "SLE | Serra Leoa",
    "COUNTRY_NAME_AR": "سيراليون",
    "COUNTRY_VALUE_AR": "SLE | سيراليون"
  },
  {
    "COUNTRY_ISO3": "SOM",
    "COUNTRY_NAME_FR": "Somalie",
    "COUNTRY_VALUE": "SOM | Somalie",
    "COUNTRY_NAME_EN": "Somalia",
    "COUNTRY_VALUE_EN": "SOM | Somalia",
    "COUNTRY_NAME_PT": "Somália",
    "COUNTRY_VALUE_PT": "SOM | Somália",
    "COUNTRY_NAME_AR": "الصومال",
    "COUNTRY_VALUE_AR": "SOM | الصومال"
  },
  {
    "COUNTRY_ISO3": "SSD",
    "COUNTRY_NAME_FR": "Soudan du Sud",
    "COUNTRY_VALUE": "SSD | Soudan du Sud",
    "COUNTRY_NAME_EN": "South Sudan",
    "COUNTRY_VALUE_EN": "SSD | South Sudan",
    "COUNTRY_NAME_PT": "Sudão do Sul",
    "COUNTRY_VALUE_PT": "SSD | Sudão do Sul",
    "COUNTRY_NAME_AR": "جنوب السودان",
    "COUNTRY_VALUE_AR": "SSD | جنوب السودان"
  },
  {
    "COUNTRY_ISO3": "ZAF",
    "COUNTRY_NAME_FR": "Afrique du Sud",
    "COUNTRY_VALUE": "ZAF | Afrique du Sud",
    "COUNTRY_NAME_EN": "South Africa",
    "COUNTRY_VALUE_EN": "ZAF | South Africa",
    "COUNTRY_NAME_PT": "África do Sul",
    "COUNTRY_VALUE_PT": "ZAF | África do Sul",
    "COUNTRY_NAME_AR": "جنوب أفريقيا",
    "COUNTRY_VALUE_AR": "ZAF | جنوب أفريقيا"
  },
  {
    "COUNTRY_ISO3": "SDN",
    "COUNTRY_NAME_FR": "Soudan",
    "COUNTRY_VALUE": "SDN | Soudan",
    "COUNTRY_NAME_EN": "Sudan",
    "COUNTRY_VALUE_EN": "SDN | Sudan",
    "COUNTRY_NAME_PT": "Sudão",
    "COUNTRY_VALUE_PT": "SDN | Sudão",
    "COUNTRY_NAME_AR": "السودان",
    "COUNTRY_VALUE_AR": "SDN | السودان"
  },
  {
    "COUNTRY_ISO3": "SWZ",
    "COUNTRY_NAME_FR": "Eswatini",
    "COUNTRY_VALUE": "SWZ | Eswatini",
    "COUNTRY_NAME_EN": "Eswatini",
    "COUNTRY_VALUE_EN": "SWZ | Eswatini",
    "COUNTRY_NAME_PT": "Essuatíni",
    "COUNTRY_VALUE_PT": "SWZ | Essuatíni",
    "COUNTRY_NAME_AR": "إسواتيني",
    "COUNTRY_VALUE_AR": "SWZ | إسواتيني"
  },
  {
    "COUNTRY_ISO3": "TZA",
    "COUNTRY_NAME_FR": "République-Unie de Tanzanie",
    "COUNTRY_VALUE": "TZA | République-Unie de Tanzanie",
    "COUNTRY_NAME_EN": "Tanzania",
    "COUNTRY_VALUE_EN": "TZA | Tanzania",
    "COUNTRY_NAME_PT": "Tanzânia",
    "COUNTRY_VALUE_PT": "TZA | Tanzânia",
    "COUNTRY_NAME_AR": "تنزانيا",
    "COUNTRY_VALUE_AR": "TZA | تنزانيا"
  },
  {
    "COUNTRY_ISO3": "TGO",
    "COUNTRY_NAME_FR": "Togo",
    "COUNTRY_VALUE": "TGO | Togo",
    "COUNTRY_NAME_EN": "Togo",
    "COUNTRY_VALUE_EN": "TGO | Togo",
    "COUNTRY_NAME_PT": "Togo",
    "COUNTRY_VALUE_PT": "TGO | Togo",
    "COUNTRY_NAME_AR": "توغو",
    "COUNTRY_VALUE_AR": "TGO | توغو"
  },
  {
    "COUNTRY_ISO3": "TUN",
    "COUNTRY_NAME_FR": "Tunisie",
    "COUNTRY_VALUE": "TUN | Tunisie",
    "COUNTRY_NAME_EN": "Tunisia",
    "COUNTRY_VALUE_EN": "TUN | Tunisia",
    "COUNTRY_NAME_PT": "Tunísia",
    "COUNTRY_VALUE_PT": "TUN | Tunísia",
    "COUNTRY_NAME_AR": "تونس",
    "COUNTRY_VALUE_AR": "TUN | تونس"
  },
  {
    "COUNTRY_ISO3": "UGA",
    "COUNTRY_NAME_FR": "Ouganda",
    "COUNTRY_VALUE": "UGA | Ouganda",
    "COUNTRY_NAME_EN": "Uganda",
    "COUNTRY_VALUE_EN": "UGA | Uganda",
    "COUNTRY_NAME_PT": "Uganda",
    "COUNTRY_VALUE_PT": "UGA | Uganda",
    "COUNTRY_NAME_AR": "أوغندا",
    "COUNTRY_VALUE_AR": "UGA | أوغندا"
  },
  {
    "COUNTRY_ISO3": "ZMB",
    "COUNTRY_NAME_FR": "Zambie",
    "COUNTRY_VALUE": "ZMB | Zambie",
    "COUNTRY_NAME_EN": "Zambia",
    "COUNTRY_VALUE_EN": "ZMB | Zambia",
    "COUNTRY_NAME_PT": "Zâmbia",
    "COUNTRY_VALUE_PT": "ZMB | Zâmbia",
    "COUNTRY_NAME_AR": "زامبيا",
    "COUNTRY_VALUE_AR": "ZMB | زامبيا"
  },
  {
    "COUNTRY_ISO3": "ZWE",
    "COUNTRY_NAME_FR": "Zimbabwe",
    "COUNTRY_VALUE": "ZWE | Zimbabwe",
    "COUNTRY_NAME_EN": "Zimbabwe",
    "COUNTRY_VALUE_EN": "ZWE | Zimbabwe",
    "COUNTRY_NAME_PT": "Zimbábue",
    "COUNTRY_VALUE_PT": "ZWE | Zimbábue",
    "COUNTRY_NAME_AR": "زيمبابوي",
    "COUNTRY_VALUE_AR": "ZWE | زيمبابوي"
  }
]

EMBEDDED_LONGLIST_ROWS = [
  {
    "Domain_code": "D01",
    "Domain_label_fr": "D01|Croissance Économique, Transformation Structurelle et Commerce",
    "Stat_label_fr": "D01S01|Taux de croissance du PIB réel",
    "Domain_label_en": "D01|Economic Growth, Structural Transformation and Trade",
    "Stat_label_en": "D01S01|Real GDP growth rate",
    "Domain_label_pt": "D01|Crescimento económico, transformação estrutural e industrialização",
    "Stat_label_pt": "D01S01|Taxa de crescimento do PIB real",
    "Domain_label_ar": "D01|النمو الاقتصادي والتحول الهيكلي والتصنيع",
    "Stat_label_ar": "D01S01|معدل نمو الناتج المحلي الإجمالي الحقيقي"
  },
  {
    "Domain_code": "D01",
    "Domain_label_fr": "D01|Croissance Économique, Transformation Structurelle et Commerce",
    "Stat_label_fr": "D01S02|PIB par habitant (USD constant)",
    "Domain_label_en": "D01|Economic Growth, Structural Transformation and Trade",
    "Stat_label_en": "D01S02|GDP per capita (constant USD)",
    "Domain_label_pt": "D01|Crescimento económico, transformação estrutural e industrialização",
    "Stat_label_pt": "D01S02|PIB per capita (USD constantes)",
    "Domain_label_ar": "D01|النمو الاقتصادي والتحول الهيكلي والتصنيع",
    "Stat_label_ar": "D01S02|نصيب الفرد من الناتج المحلي الإجمالي (دولار أمريكي ثابت)"
  },
  {
    "Domain_code": "D01",
    "Domain_label_fr": "D01|Croissance Économique, Transformation Structurelle et Commerce",
    "Stat_label_fr": "D01S03|Part de la Valeur Ajoutée Manufacturière (VAM) dans le PIB",
    "Domain_label_en": "D01|Economic Growth, Structural Transformation and Trade",
    "Stat_label_en": "D01S03|Manufacturing Value Added (MVA) share in GDP",
    "Domain_label_pt": "D01|Crescimento económico, transformação estrutural e industrialização",
    "Stat_label_pt": "D01S03|Participação do valor acrescentado da indústria transformadora (VAI) no PIB",
    "Domain_label_ar": "D01|النمو الاقتصادي والتحول الهيكلي والتصنيع",
    "Stat_label_ar": "D01S03|حصة القيمة المضافة للصناعة التحويلية (MVA) من الناتج المحلي الإجمالي"
  },
  {
    "Domain_code": "D01",
    "Domain_label_fr": "D01|Croissance Économique, Transformation Structurelle et Commerce",
    "Stat_label_fr": "D01S04|Part des produits de haute technologie dans les exportations manufacturées",
    "Domain_label_en": "D01|Economic Growth, Structural Transformation and Trade",
    "Stat_label_en": "D01S04|Share of high-tech products in manufactured exports",
    "Domain_label_pt": "D01|Crescimento económico, transformação estrutural e industrialização",
    "Stat_label_pt": "D01S04|Participação de produtos de alta tecnologia nas exportações manufaturadas",
    "Domain_label_ar": "D01|النمو الاقتصادي والتحول الهيكلي والتصنيع",
    "Stat_label_ar": "D01S04|حصة المنتجات عالية التكنولوجيا في الصادرات المصنعة"
  },
  {
    "Domain_code": "D01",
    "Domain_label_fr": "D01|Croissance Économique, Transformation Structurelle et Commerce",
    "Stat_label_fr": "D01S05|Part de l'Afrique dans la valeur ajoutée manufacturière mondiale",
    "Domain_label_en": "D01|Economic Growth, Structural Transformation and Trade",
    "Stat_label_en": "D01S05|Africa's share in global manufacturing value added",
    "Domain_label_pt": "D01|Crescimento económico, transformação estrutural e industrialização",
    "Stat_label_pt": "D01S05|Participação de África no valor acrescentado manufatureiro mundial",
    "Domain_label_ar": "D01|النمو الاقتصادي والتحول الهيكلي والتصنيع",
    "Stat_label_ar": "D01S05|حصة أفريقيا من القيمة المضافة للصناعة التحويلية العالمية"
  },
  {
    "Domain_code": "D01",
    "Domain_label_fr": "D01|Croissance Économique, Transformation Structurelle et Commerce",
    "Stat_label_fr": "D01S06|Part du commerce intra-africain dans le commerce total",
    "Domain_label_en": "D01|Economic Growth, Structural Transformation and Trade",
    "Stat_label_en": "D01S06|Intra-African trade share in total trade",
    "Domain_label_pt": "D01|Crescimento económico, transformação estrutural e industrialização",
    "Stat_label_pt": "D01S06|Participação do comércio intra-africano no comércio total",
    "Domain_label_ar": "D01|النمو الاقتصادي والتحول الهيكلي والتصنيع",
    "Stat_label_ar": "D01S06|حصة التجارة البينية الأفريقية من إجمالي التجارة"
  },
  {
    "Domain_code": "D01",
    "Domain_label_fr": "D01|Croissance Économique, Transformation Structurelle et Commerce",
    "Stat_label_fr": "D01S07|Indice de diversification économique",
    "Domain_label_en": "D01|Economic Growth, Structural Transformation and Trade",
    "Stat_label_en": "D01S07|Economic diversification index",
    "Domain_label_pt": "D01|Crescimento económico, transformação estrutural e industrialização",
    "Stat_label_pt": "D01S07|Índice de diversificação económica",
    "Domain_label_ar": "D01|النمو الاقتصادي والتحول الهيكلي والتصنيع",
    "Stat_label_ar": "D01S07|مؤشر التنويع الاقتصادي"
  },
  {
    "Domain_code": "D01",
    "Domain_label_fr": "D01|Croissance Économique, Transformation Structurelle et Commerce",
    "Stat_label_fr": "D01S08|Contribution directe du tourisme au PIB",
    "Domain_label_en": "D01|Economic Growth, Structural Transformation and Trade",
    "Stat_label_en": "D01S08|Direct contribution of tourism to GDP",
    "Domain_label_pt": "D01|Crescimento económico, transformação estrutural e industrialização",
    "Stat_label_pt": "D01S08|Contribuição direta do turismo para o PIB",
    "Domain_label_ar": "D01|النمو الاقتصادي والتحول الهيكلي والتصنيع",
    "Stat_label_ar": "D01S08|المساهمة المباشرة للسياحة في الناتج المحلي الإجمالي"
  },
  {
    "Domain_code": "D01",
    "Domain_label_fr": "D01|Croissance Économique, Transformation Structurelle et Commerce",
    "Stat_label_fr": "D01S09|Formation brute de capital fixe (% PIB)",
    "Domain_label_en": "D01|Economic Growth, Structural Transformation and Trade",
    "Stat_label_en": "D01S09|Gross fixed capital formation (% of GDP)",
    "Domain_label_pt": "D01|Crescimento económico, transformação estrutural e industrialização",
    "Stat_label_pt": "D01S09|Formação bruta de capital fixo (% do PIB)",
    "Domain_label_ar": "D01|النمو الاقتصادي والتحول الهيكلي والتصنيع",
    "Stat_label_ar": "D01S09|تكوين رأس المال الثابت الإجمالي (% من الناتج المحلي الإجمالي)"
  },
  {
    "Domain_code": "D01",
    "Domain_label_fr": "D01|Croissance Économique, Transformation Structurelle et Commerce",
    "Stat_label_fr": "D01S10|Pays commerçant sous le régime ZLECA",
    "Domain_label_en": "D01|Economic Growth, Structural Transformation and Trade",
    "Stat_label_en": "D01S10|Countries trading under the AfCFTA regime",
    "Domain_label_pt": "D01|Crescimento económico, transformação estrutural e industrialização",
    "Stat_label_pt": "D01S10|Países que comercializam no âmbito do regime AfCFTA",
    "Domain_label_ar": "D01|النمو الاقتصادي والتحول الهيكلي والتصنيع",
    "Stat_label_ar": "D01S10|البلدان التي تتاجر في إطار نظام AfCFTA"
  },
  {
    "Domain_code": "D02",
    "Domain_label_fr": "D02|Emploi, Travail Décent et Protection Sociale",
    "Stat_label_fr": "D02S11|Taux de chômage (total, jeunes, femmes)",
    "Domain_label_en": "D02|Employment, Decent Work and Social Protection",
    "Stat_label_en": "D02S11|Unemployment rate (total, youth, women)",
    "Domain_label_pt": "D02|Emprego, trabalho decente e proteção social",
    "Stat_label_pt": "D02S11|Taxa de desemprego (total, jovens, mulheres)",
    "Domain_label_ar": "D02|العمالة والعمل اللائق والحماية الاجتماعية",
    "Stat_label_ar": "D02S11|معدل البطالة (الإجمالي، الشباب، النساء)"
  },
  {
    "Domain_code": "D02",
    "Domain_label_fr": "D02|Emploi, Travail Décent et Protection Sociale",
    "Stat_label_fr": "D02S12|Nombre d'emplois créés (dont jeunes)",
    "Domain_label_en": "D02|Employment, Decent Work and Social Protection",
    "Stat_label_en": "D02S12|Number of jobs created (including youth)",
    "Domain_label_pt": "D02|Emprego, trabalho decente e proteção social",
    "Stat_label_pt": "D02S12|Número de empregos criados (incluindo jovens)",
    "Domain_label_ar": "D02|العمالة والعمل اللائق والحماية الاجتماعية",
    "Stat_label_ar": "D02S12|عدد فرص العمل التي تم خلقها (بما في ذلك للشباب)"
  },
  {
    "Domain_code": "D02",
    "Domain_label_fr": "D02|Emploi, Travail Décent et Protection Sociale",
    "Stat_label_fr": "D02S13|Taux de jeunes NEET (Ni en emploi, éducation ou formation) (15-24 ans)",
    "Domain_label_en": "D02|Employment, Decent Work and Social Protection",
    "Stat_label_en": "D02S13|Share of youth NEET (Not in Employment, Education or Training) (15–24)",
    "Domain_label_pt": "D02|Emprego, trabalho decente e proteção social",
    "Stat_label_pt": "D02S13|Percentagem de jovens NEET (Nem em emprego, educação ou formação) (15–24)",
    "Domain_label_ar": "D02|العمالة والعمل اللائق والحماية الاجتماعية",
    "Stat_label_ar": "D02S13|نسبة الشباب خارج العمل والتعليم أو التدريب (NEET) (15–24)"
  },
  {
    "Domain_code": "D02",
    "Domain_label_fr": "D02|Emploi, Travail Décent et Protection Sociale",
    "Stat_label_fr": "D02S14|Taux de couverture de la protection sociale",
    "Domain_label_en": "D02|Employment, Decent Work and Social Protection",
    "Stat_label_en": "D02S14|Social protection coverage rate",
    "Domain_label_pt": "D02|Emprego, trabalho decente e proteção social",
    "Stat_label_pt": "D02S14|Taxa de cobertura da proteção social",
    "Domain_label_ar": "D02|العمالة والعمل اللائق والحماية الاجتماعية",
    "Stat_label_ar": "D02S14|معدل التغطية بالحماية الاجتماعية"
  },
  {
    "Domain_code": "D02",
    "Domain_label_fr": "D02|Emploi, Travail Décent et Protection Sociale",
    "Stat_label_fr": "D02S15|Pourcentage d'enfants (5-17 ans) au travail",
    "Domain_label_en": "D02|Employment, Decent Work and Social Protection",
    "Stat_label_en": "D02S15|Percentage of children (5–17) in child labour",
    "Domain_label_pt": "D02|Emprego, trabalho decente e proteção social",
    "Stat_label_pt": "D02S15|Percentagem de crianças (5–17) em trabalho infantil",
    "Domain_label_ar": "D02|العمالة والعمل اللائق والحماية الاجتماعية",
    "Stat_label_ar": "D02S15|نسبة الأطفال (5–17) في عمالة الأطفال"
  },
  {
    "Domain_code": "D02",
    "Domain_label_fr": "D02|Emploi, Travail Décent et Protection Sociale",
    "Stat_label_fr": "D02S16|Part de l'emploi informel dans l'emploi total",
    "Domain_label_en": "D02|Employment, Decent Work and Social Protection",
    "Stat_label_en": "D02S16|Share of informal employment in total employment",
    "Domain_label_pt": "D02|Emprego, trabalho decente e proteção social",
    "Stat_label_pt": "D02S16|Percentagem de emprego informal no emprego total",
    "Domain_label_ar": "D02|العمالة والعمل اللائق والحماية الاجتماعية",
    "Stat_label_ar": "D02S16|حصة العمالة غير الرسمية من إجمالي العمالة"
  },
  {
    "Domain_code": "D02",
    "Domain_label_fr": "D02|Emploi, Travail Décent et Protection Sociale",
    "Stat_label_fr": "D02S17|Écart de rémunération horaire hommes-femmes",
    "Domain_label_en": "D02|Employment, Decent Work and Social Protection",
    "Stat_label_en": "D02S17|Gender gap in hourly earnings",
    "Domain_label_pt": "D02|Emprego, trabalho decente e proteção social",
    "Stat_label_pt": "D02S17|Diferença de género nos ganhos horários",
    "Domain_label_ar": "D02|العمالة والعمل اللائق والحماية الاجتماعية",
    "Stat_label_ar": "D02S17|الفجوة بين الجنسين في الأجور بالساعة"
  },
  {
    "Domain_code": "D03",
    "Domain_label_fr": "D03|Agriculture durable, Sécurité alimentaire et Nutrition",
    "Stat_label_fr": "D03S18|Prévalence de la sous-alimentation (Faim)",
    "Domain_label_en": "D03|Sustainable Agriculture, Food Security and Nutrition",
    "Stat_label_en": "D03S18|Prevalence of undernourishment (hunger)",
    "Domain_label_pt": "D03|Agricultura sustentável, segurança alimentar e nutrição",
    "Stat_label_pt": "D03S18|Prevalência de subalimentação (fome)",
    "Domain_label_ar": "D03|الزراعة المستدامة والأمن الغذائي والتغذية",
    "Stat_label_ar": "D03S18|انتشار نقص التغذية (الجوع)"
  },
  {
    "Domain_code": "D03",
    "Domain_label_fr": "D03|Agriculture durable, Sécurité alimentaire et Nutrition",
    "Stat_label_fr": "D03S19|Prévalence du retard de croissance (enfants <5 ans)",
    "Domain_label_en": "D03|Sustainable Agriculture, Food Security and Nutrition",
    "Stat_label_en": "D03S19|Prevalence of stunting (children under 5)",
    "Domain_label_pt": "D03|Agricultura sustentável, segurança alimentar e nutrição",
    "Stat_label_pt": "D03S19|Prevalência de atraso de crescimento (crianças menores de 5 anos)",
    "Domain_label_ar": "D03|الزراعة المستدامة والأمن الغذائي والتغذية",
    "Stat_label_ar": "D03S19|انتشار التقزم (الأطفال دون 5 سنوات)"
  },
  {
    "Domain_code": "D03",
    "Domain_label_fr": "D03|Agriculture durable, Sécurité alimentaire et Nutrition",
    "Stat_label_fr": "D03S20|Taux de croissance des rendements agricoles",
    "Domain_label_en": "D03|Sustainable Agriculture, Food Security and Nutrition",
    "Stat_label_en": "D03S20|Growth rate of agricultural yields",
    "Domain_label_pt": "D03|Agricultura sustentável, segurança alimentar e nutrição",
    "Stat_label_pt": "D03S20|Taxa de crescimento dos rendimentos agrícolas",
    "Domain_label_ar": "D03|الزراعة المستدامة والأمن الغذائي والتغذية",
    "Stat_label_ar": "D03S20|معدل نمو غلة المحاصيل الزراعية"
  },
  {
    "Domain_code": "D03",
    "Domain_label_fr": "D03|Agriculture durable, Sécurité alimentaire et Nutrition",
    "Stat_label_fr": "D03S21|Réduction des pertes post-récolte",
    "Domain_label_en": "D03|Sustainable Agriculture, Food Security and Nutrition",
    "Stat_label_en": "D03S21|Reduction in post-harvest losses",
    "Domain_label_pt": "D03|Agricultura sustentável, segurança alimentar e nutrição",
    "Stat_label_pt": "D03S21|Redução das perdas pós-colheita",
    "Domain_label_ar": "D03|الزراعة المستدامة والأمن الغذائي والتغذية",
    "Stat_label_ar": "D03S21|خفض خسائر ما بعد الحصاد"
  },
  {
    "Domain_code": "D03",
    "Domain_label_fr": "D03|Agriculture durable, Sécurité alimentaire et Nutrition",
    "Stat_label_fr": "D03S22|Part des importations alimentaires / importations totales",
    "Domain_label_en": "D03|Sustainable Agriculture, Food Security and Nutrition",
    "Stat_label_en": "D03S22|Share of food imports in total imports",
    "Domain_label_pt": "D03|Agricultura sustentável, segurança alimentar e nutrição",
    "Stat_label_pt": "D03S22|Percentagem das importações alimentares no total das importações",
    "Domain_label_ar": "D03|الزراعة المستدامة والأمن الغذائي والتغذية",
    "Stat_label_ar": "D03S22|حصة واردات الغذاء من إجمالي الواردات"
  },
  {
    "Domain_code": "D03",
    "Domain_label_fr": "D03|Agriculture durable, Sécurité alimentaire et Nutrition",
    "Stat_label_fr": "D03S23|Terres sous gestion durable",
    "Domain_label_en": "D03|Sustainable Agriculture, Food Security and Nutrition",
    "Stat_label_en": "D03S23|Land under sustainable management",
    "Domain_label_pt": "D03|Agricultura sustentável, segurança alimentar e nutrição",
    "Stat_label_pt": "D03S23|Terras sob gestão sustentável",
    "Domain_label_ar": "D03|الزراعة المستدامة والأمن الغذائي والتغذية",
    "Stat_label_ar": "D03S23|الأراضي المُدارة إدارة مستدامة"
  },
  {
    "Domain_code": "D03",
    "Domain_label_fr": "D03|Agriculture durable, Sécurité alimentaire et Nutrition",
    "Stat_label_fr": "D03S24|Résilience climatique des ménages agricoles",
    "Domain_label_en": "D03|Sustainable Agriculture, Food Security and Nutrition",
    "Stat_label_en": "D03S24|Climate resilience of agricultural households",
    "Domain_label_pt": "D03|Agricultura sustentável, segurança alimentar e nutrição",
    "Stat_label_pt": "D03S24|Resiliência climática dos agregados familiares agrícolas",
    "Domain_label_ar": "D03|الزراعة المستدامة والأمن الغذائي والتغذية",
    "Stat_label_ar": "D03S24|القدرة على الصمود المناخي للأسر الزراعية"
  },
  {
    "Domain_code": "D03",
    "Domain_label_fr": "D03|Agriculture durable, Sécurité alimentaire et Nutrition",
    "Stat_label_fr": "D03S25|Zones de Transformation Agro-industrielle (SAPZ) créées",
    "Domain_label_en": "D03|Sustainable Agriculture, Food Security and Nutrition",
    "Stat_label_en": "D03S25|Special Agro-Industrial Processing Zones (SAPZ) established",
    "Domain_label_pt": "D03|Agricultura sustentável, segurança alimentar e nutrição",
    "Stat_label_pt": "D03S25|Zonas especiais de processamento agroindustrial (SAPZ) estabelecidas",
    "Domain_label_ar": "D03|الزراعة المستدامة والأمن الغذائي والتغذية",
    "Stat_label_ar": "D03S25|إنشاء مناطق معالجة زراعية-صناعية خاصة (SAPZ)"
  },
  {
    "Domain_code": "D03",
    "Domain_label_fr": "D03|Agriculture durable, Sécurité alimentaire et Nutrition",
    "Stat_label_fr": "D03S26|Utilisation d'engrais (kg/ha)",
    "Domain_label_en": "D03|Sustainable Agriculture, Food Security and Nutrition",
    "Stat_label_en": "D03S26|Fertilizer use (kg/ha)",
    "Domain_label_pt": "D03|Agricultura sustentável, segurança alimentar e nutrição",
    "Stat_label_pt": "D03S26|Uso de fertilizantes (kg/ha)",
    "Domain_label_ar": "D03|الزراعة المستدامة والأمن الغذائي والتغذية",
    "Stat_label_ar": "D03S26|استخدام الأسمدة (كغ/هكتار)"
  },
  {
    "Domain_code": "D04",
    "Domain_label_fr": "D04|Infrastructures, Industrialisation et Innovation",
    "Stat_label_fr": "D04S27|Taux d’accès à l'électricité (Ménages)",
    "Domain_label_en": "D04|Infrastructure, Industrialization and Innovation",
    "Stat_label_en": "D04S27|Access to electricity (households)",
    "Domain_label_pt": "D04|Infraestruturas, industrialização e inovação",
    "Stat_label_pt": "D04S27|Acesso à eletricidade (agregados familiares)",
    "Domain_label_ar": "D04|البنية التحتية والتصنيع والابتكار",
    "Stat_label_ar": "D04S27|الحصول على الكهرباء (الأسر)"
  },
  {
    "Domain_code": "D04",
    "Domain_label_fr": "D04|Infrastructures, Industrialisation et Innovation",
    "Stat_label_fr": "D04S28|Accès Internet haut débit (Qualité)",
    "Domain_label_en": "D04|Infrastructure, Industrialization and Innovation",
    "Stat_label_en": "D04S28|Broadband internet access (quality)",
    "Domain_label_pt": "D04|Infraestruturas, industrialização e inovação",
    "Stat_label_pt": "D04S28|Acesso à internet de banda larga (qualidade)",
    "Domain_label_ar": "D04|البنية التحتية والتصنيع والابتكار",
    "Stat_label_ar": "D04S28|الوصول إلى الإنترنت عريض النطاق (الجودة)"
  },
  {
    "Domain_code": "D04",
    "Domain_label_fr": "D04|Infrastructures, Industrialisation et Innovation",
    "Stat_label_fr": "D04S29|Connectivité routière interafricaine",
    "Domain_label_en": "D04|Infrastructure, Industrialization and Innovation",
    "Stat_label_en": "D04S29|Inter-African road connectivity",
    "Domain_label_pt": "D04|Infraestruturas, industrialização e inovação",
    "Stat_label_pt": "D04S29|Conectividade rodoviária interafricana",
    "Domain_label_ar": "D04|البنية التحتية والتصنيع والابتكار",
    "Stat_label_ar": "D04S29|الترابط الطرقي بين البلدان الأفريقية"
  },
  {
    "Domain_code": "D04",
    "Domain_label_fr": "D04|Infrastructures, Industrialisation et Innovation",
    "Stat_label_fr": "D04S30|Connectivité ferroviaire interafricaine",
    "Domain_label_en": "D04|Infrastructure, Industrialization and Innovation",
    "Stat_label_en": "D04S30|Inter-African rail connectivity",
    "Domain_label_pt": "D04|Infraestruturas, industrialização e inovação",
    "Stat_label_pt": "D04S30|Conectividade ferroviária interafricana",
    "Domain_label_ar": "D04|البنية التحتية والتصنيع والابتكار",
    "Stat_label_ar": "D04S30|الترابط السككي بين البلدان الأفريقية"
  },
  {
    "Domain_code": "D04",
    "Domain_label_fr": "D04|Infrastructures, Industrialisation et Innovation",
    "Stat_label_fr": "D04S31|Dépenses de R&D (% PIB)",
    "Domain_label_en": "D04|Infrastructure, Industrialization and Innovation",
    "Stat_label_en": "D04S31|R&D expenditure (% of GDP)",
    "Domain_label_pt": "D04|Infraestruturas, industrialização e inovação",
    "Stat_label_pt": "D04S31|Despesa em I&D (% do PIB)",
    "Domain_label_ar": "D04|البنية التحتية والتصنيع والابتكار",
    "Stat_label_ar": "D04S31|الإنفاق على البحث والتطوير (% من الناتج المحلي الإجمالي)"
  },
  {
    "Domain_code": "D04",
    "Domain_label_fr": "D04|Infrastructures, Industrialisation et Innovation",
    "Stat_label_fr": "D04S32|Capacité de production électrique installée",
    "Domain_label_en": "D04|Infrastructure, Industrialization and Innovation",
    "Stat_label_en": "D04S32|Installed electricity generation capacity",
    "Domain_label_pt": "D04|Infraestruturas, industrialização e inovação",
    "Stat_label_pt": "D04S32|Capacidade instalada de geração de eletricidade",
    "Domain_label_ar": "D04|البنية التحتية والتصنيع والابتكار",
    "Stat_label_ar": "D04S32|القدرة المركبة لتوليد الكهرباء"
  },
  {
    "Domain_code": "D04",
    "Domain_label_fr": "D04|Infrastructures, Industrialisation et Innovation",
    "Stat_label_fr": "D04S33|Contribution des services numériques au PIB",
    "Domain_label_en": "D04|Infrastructure, Industrialization and Innovation",
    "Stat_label_en": "D04S33|Contribution of digital services to GDP",
    "Domain_label_pt": "D04|Infraestruturas, industrialização e inovação",
    "Stat_label_pt": "D04S33|Contribuição dos serviços digitais para o PIB",
    "Domain_label_ar": "D04|البنية التحتية والتصنيع والابتكار",
    "Stat_label_ar": "D04S33|مساهمة الخدمات الرقمية في الناتج المحلي الإجمالي"
  },
  {
    "Domain_code": "D04",
    "Domain_label_fr": "D04|Infrastructures, Industrialisation et Innovation",
    "Stat_label_fr": "D04S34|Part des énergies renouvelables",
    "Domain_label_en": "D04|Infrastructure, Industrialization and Innovation",
    "Stat_label_en": "D04S34|Share of renewable energy",
    "Domain_label_pt": "D04|Infraestruturas, industrialização e inovação",
    "Stat_label_pt": "D04S34|Percentagem de energias renováveis",
    "Domain_label_ar": "D04|البنية التحتية والتصنيع والابتكار",
    "Stat_label_ar": "D04S34|حصة الطاقة المتجددة"
  },
  {
    "Domain_code": "D05",
    "Domain_label_fr": "D05|Inclusion, Pauvreté et Inégalités",
    "Stat_label_fr": "D05S35|Taux de pauvreté (seuil national)",
    "Domain_label_en": "D05|Inclusion, Poverty and Inequalities",
    "Stat_label_en": "D05S35|Poverty rate (national poverty line)",
    "Domain_label_pt": "D05|Inclusão, pobreza e desigualdades",
    "Stat_label_pt": "D05S35|Taxa de pobreza (linha nacional de pobreza)",
    "Domain_label_ar": "D05|الإدماج والفقر وعدم المساواة",
    "Stat_label_ar": "D05S35|معدل الفقر (خط الفقر الوطني)"
  },
  {
    "Domain_code": "D05",
    "Domain_label_fr": "D05|Inclusion, Pauvreté et Inégalités",
    "Stat_label_fr": "D05S36|Indice de Gini (Inégalités de revenus)",
    "Domain_label_en": "D05|Inclusion, Poverty and Inequalities",
    "Stat_label_en": "D05S36|Gini index (income inequality)",
    "Domain_label_pt": "D05|Inclusão, pobreza e desigualdades",
    "Stat_label_pt": "D05S36|Índice de Gini (desigualdade de rendimentos)",
    "Domain_label_ar": "D05|الإدماج والفقر وعدم المساواة",
    "Stat_label_ar": "D05S36|مؤشر جيني (عدم المساواة في الدخل)"
  },
  {
    "Domain_code": "D05",
    "Domain_label_fr": "D05|Inclusion, Pauvreté et Inégalités",
    "Stat_label_fr": "D05S37|Accès à l'eau potable salubre",
    "Domain_label_en": "D05|Inclusion, Poverty and Inequalities",
    "Stat_label_en": "D05S37|Access to safely managed drinking water",
    "Domain_label_pt": "D05|Inclusão, pobreza e desigualdades",
    "Stat_label_pt": "D05S37|Acesso a água potável gerida com segurança",
    "Domain_label_ar": "D05|الإدماج والفقر وعدم المساواة",
    "Stat_label_ar": "D05S37|الحصول على مياه شرب مُدارة بأمان"
  },
  {
    "Domain_code": "D05",
    "Domain_label_fr": "D05|Inclusion, Pauvreté et Inégalités",
    "Stat_label_fr": "D05S38|Accès à l'assainissement amélioré",
    "Domain_label_en": "D05|Inclusion, Poverty and Inequalities",
    "Stat_label_en": "D05S38|Access to improved sanitation",
    "Domain_label_pt": "D05|Inclusão, pobreza e desigualdades",
    "Stat_label_pt": "D05S38|Acesso a saneamento melhorado",
    "Domain_label_ar": "D05|الإدماج والفقر وعدم المساواة",
    "Stat_label_ar": "D05S38|الحصول على خدمات صرف صحي محسّنة"
  },
  {
    "Domain_code": "D05",
    "Domain_label_fr": "D05|Inclusion, Pauvreté et Inégalités",
    "Stat_label_fr": "D05S39|Population urbaine vivant dans des taudis",
    "Domain_label_en": "D05|Inclusion, Poverty and Inequalities",
    "Stat_label_en": "D05S39|Urban population living in slums",
    "Domain_label_pt": "D05|Inclusão, pobreza e desigualdades",
    "Stat_label_pt": "D05S39|População urbana a viver em bairros degradados (slums)",
    "Domain_label_ar": "D05|الإدماج والفقر وعدم المساواة",
    "Stat_label_ar": "D05S39|السكان الحضريون الذين يعيشون في الأحياء العشوائية"
  },
  {
    "Domain_code": "D05",
    "Domain_label_fr": "D05|Inclusion, Pauvreté et Inégalités",
    "Stat_label_fr": "D05S40|Accès à un logement décent",
    "Domain_label_en": "D05|Inclusion, Poverty and Inequalities",
    "Stat_label_en": "D05S40|Access to adequate housing",
    "Domain_label_pt": "D05|Inclusão, pobreza e desigualdades",
    "Stat_label_pt": "D05S40|Acesso a habitação adequada",
    "Domain_label_ar": "D05|الإدماج والفقر وعدم المساواة",
    "Stat_label_ar": "D05S40|الحصول على سكن ملائم"
  },
  {
    "Domain_code": "D05",
    "Domain_label_fr": "D05|Inclusion, Pauvreté et Inégalités",
    "Stat_label_fr": "D05S41|Accès à l'eau (Pays en transition/fragiles)",
    "Domain_label_en": "D05|Inclusion, Poverty and Inequalities",
    "Stat_label_en": "D05S41|Access to water (transition/fragile countries)",
    "Domain_label_pt": "D05|Inclusão, pobreza e desigualdades",
    "Stat_label_pt": "D05S41|Acesso à água (países em transição / frágeis)",
    "Domain_label_ar": "D05|الإدماج والفقر وعدم المساواة",
    "Stat_label_ar": "D05S41|الحصول على المياه (بلدان انتقالية / هشة)"
  },
  {
    "Domain_code": "D06",
    "Domain_label_fr": "D06|Éducation, Compétences et Capital Humain",
    "Stat_label_fr": "D06S42|Taux net de scolarisation (Primaire)",
    "Domain_label_en": "D06|Education, Skills and Human Capital",
    "Stat_label_en": "D06S42|Net enrolment rate (primary)",
    "Domain_label_pt": "D06|Educação, competências e capital humano",
    "Stat_label_pt": "D06S42|Taxa líquida de escolarização (primário)",
    "Domain_label_ar": "D06|التعليم والمهارات ورأس المال البشري",
    "Stat_label_ar": "D06S42|معدل الالتحاق الصافي (الابتدائي)"
  },
  {
    "Domain_code": "D06",
    "Domain_label_fr": "D06|Éducation, Compétences et Capital Humain",
    "Stat_label_fr": "D06S43|Taux net de scolarisation (Secondaire)",
    "Domain_label_en": "D06|Education, Skills and Human Capital",
    "Stat_label_en": "D06S43|Net enrolment rate (secondary)",
    "Domain_label_pt": "D06|Educação, competências e capital humano",
    "Stat_label_pt": "D06S43|Taxa líquida de escolarização (secundário)",
    "Domain_label_ar": "D06|التعليم والمهارات ورأس المال البشري",
    "Stat_label_ar": "D06S43|معدل الالتحاق الصافي (الثانوي)"
  },
  {
    "Domain_code": "D06",
    "Domain_label_fr": "D06|Éducation, Compétences et Capital Humain",
    "Stat_label_fr": "D06S44|Taux de scolarisation pré-primaire",
    "Domain_label_en": "D06|Education, Skills and Human Capital",
    "Stat_label_en": "D06S44|Pre-primary enrolment rate",
    "Domain_label_pt": "D06|Educação, competências e capital humano",
    "Stat_label_pt": "D06S44|Taxa de matrícula no pré-escolar",
    "Domain_label_ar": "D06|التعليم والمهارات ورأس المال البشري",
    "Stat_label_ar": "D06S44|معدل الالتحاق بالتعليم قبل الابتدائي"
  },
  {
    "Domain_code": "D06",
    "Domain_label_fr": "D06|Éducation, Compétences et Capital Humain",
    "Stat_label_fr": "D06S45|Proportion de diplômés en STEM",
    "Domain_label_en": "D06|Education, Skills and Human Capital",
    "Stat_label_en": "D06S45|Share of graduates in STEM",
    "Domain_label_pt": "D06|Educação, competências e capital humano",
    "Stat_label_pt": "D06S45|Percentagem de diplomados em STEM (CTEM)",
    "Domain_label_ar": "D06|التعليم والمهارات ورأس المال البشري",
    "Stat_label_ar": "D06S45|حصة الخريجين في مجالات STEM"
  },
  {
    "Domain_code": "D06",
    "Domain_label_fr": "D06|Éducation, Compétences et Capital Humain",
    "Stat_label_fr": "D06S46|Taux de compétence minimale (Lecture/Maths)",
    "Domain_label_en": "D06|Education, Skills and Human Capital",
    "Stat_label_en": "D06S46|Minimum proficiency level (reading/mathematics)",
    "Domain_label_pt": "D06|Educação, competências e capital humano",
    "Stat_label_pt": "D06S46|Nível mínimo de proficiência (leitura/matemática)",
    "Domain_label_ar": "D06|التعليم والمهارات ورأس المال البشري",
    "Stat_label_ar": "D06S46|الحد الأدنى لمستوى الكفاءة (القراءة/الرياضيات)"
  },
  {
    "Domain_code": "D06",
    "Domain_label_fr": "D06|Éducation, Compétences et Capital Humain",
    "Stat_label_fr": "D06S47|Ratio Élèves / Enseignant qualifié",
    "Domain_label_en": "D06|Education, Skills and Human Capital",
    "Stat_label_en": "D06S47|Pupil-to-qualified-teacher ratio",
    "Domain_label_pt": "D06|Educação, competências e capital humano",
    "Stat_label_pt": "D06S47|Rácio aluno/professor qualificado",
    "Domain_label_ar": "D06|التعليم والمهارات ورأس المال البشري",
    "Stat_label_ar": "D06S47|نسبة التلاميذ إلى المعلمين المؤهلين"
  },
  {
    "Domain_code": "D06",
    "Domain_label_fr": "D06|Éducation, Compétences et Capital Humain",
    "Stat_label_fr": "D06S48|Taux brut de scolarisation (Supérieur)",
    "Domain_label_en": "D06|Education, Skills and Human Capital",
    "Stat_label_en": "D06S48|Gross enrolment rate (tertiary)",
    "Domain_label_pt": "D06|Educação, competências e capital humano",
    "Stat_label_pt": "D06S48|Taxa bruta de escolarização (ensino superior)",
    "Domain_label_ar": "D06|التعليم والمهارات ورأس المال البشري",
    "Stat_label_ar": "D06S48|معدل الالتحاق الإجمالي (التعليم العالي)"
  },
  {
    "Domain_code": "D07",
    "Domain_label_fr": "D07|Santé, Bien-être et Accès Universel",
    "Stat_label_fr": "D07S49|Accès aux soins de santé primaires",
    "Domain_label_en": "D07|Health, Well-being and Universal Access",
    "Stat_label_en": "D07S49|Access to primary health care",
    "Domain_label_pt": "D07|Saúde, bem-estar e acesso universal",
    "Stat_label_pt": "D07S49|Acesso aos cuidados de saúde primários",
    "Domain_label_ar": "D07|الصحة والرفاه وإتاحة الخدمات للجميع",
    "Stat_label_ar": "D07S49|الحصول على الرعاية الصحية الأولية"
  },
  {
    "Domain_code": "D07",
    "Domain_label_fr": "D07|Santé, Bien-être et Accès Universel",
    "Stat_label_fr": "D07S50|Taux de mortalité maternelle",
    "Domain_label_en": "D07|Health, Well-being and Universal Access",
    "Stat_label_en": "D07S50|Maternal mortality ratio",
    "Domain_label_pt": "D07|Saúde, bem-estar e acesso universal",
    "Stat_label_pt": "D07S50|Razão de mortalidade materna",
    "Domain_label_ar": "D07|الصحة والرفاه وإتاحة الخدمات للجميع",
    "Stat_label_ar": "D07S50|معدل وفيات الأمهات"
  },
  {
    "Domain_code": "D07",
    "Domain_label_fr": "D07|Santé, Bien-être et Accès Universel",
    "Stat_label_fr": "D07S51|Taux de mortalité infantile (<5 ans)",
    "Domain_label_en": "D07|Health, Well-being and Universal Access",
    "Stat_label_en": "D07S51|Under-five mortality rate",
    "Domain_label_pt": "D07|Saúde, bem-estar e acesso universal",
    "Stat_label_pt": "D07S51|Taxa de mortalidade de menores de cinco anos",
    "Domain_label_ar": "D07|الصحة والرفاه وإتاحة الخدمات للجميع",
    "Stat_label_ar": "D07S51|معدل وفيات الأطفال دون الخامسة"
  },
  {
    "Domain_code": "D07",
    "Domain_label_fr": "D07|Santé, Bien-être et Accès Universel",
    "Stat_label_fr": "D07S52|Incidence VIH / Couverture ARV",
    "Domain_label_en": "D07|Health, Well-being and Universal Access",
    "Stat_label_en": "D07S52|HIV incidence / ART coverage",
    "Domain_label_pt": "D07|Saúde, bem-estar e acesso universal",
    "Stat_label_pt": "D07S52|Incidência do VIH / cobertura de TAR (ART)",
    "Domain_label_ar": "D07|الصحة والرفاه وإتاحة الخدمات للجميع",
    "Stat_label_ar": "D07S52|معدل الإصابة بفيروس HIV / تغطية العلاج بمضادات الفيروسات (ART)"
  },
  {
    "Domain_code": "D07",
    "Domain_label_fr": "D07|Santé, Bien-être et Accès Universel",
    "Stat_label_fr": "D07S53|Incidence Paludisme",
    "Domain_label_en": "D07|Health, Well-being and Universal Access",
    "Stat_label_en": "D07S53|Malaria incidence",
    "Domain_label_pt": "D07|Saúde, bem-estar e acesso universal",
    "Stat_label_pt": "D07S53|Incidência de malária",
    "Domain_label_ar": "D07|الصحة والرفاه وإتاحة الخدمات للجميع",
    "Stat_label_ar": "D07S53|معدل الإصابة بالملاريا"
  },
  {
    "Domain_code": "D07",
    "Domain_label_fr": "D07|Santé, Bien-être et Accès Universel",
    "Stat_label_fr": "D07S54|Indice de Couverture Sanitaire Universelle",
    "Domain_label_en": "D07|Health, Well-being and Universal Access",
    "Stat_label_en": "D07S54|Universal Health Coverage (UHC) service coverage index",
    "Domain_label_pt": "D07|Saúde, bem-estar e acesso universal",
    "Stat_label_pt": "D07S54|Índice de cobertura de serviços da Cobertura Universal de Saúde (UHC)",
    "Domain_label_ar": "D07|الصحة والرفاه وإتاحة الخدمات للجميع",
    "Stat_label_ar": "D07S54|مؤشر تغطية خدمات التغطية الصحية الشاملة (UHC)"
  },
  {
    "Domain_code": "D08",
    "Domain_label_fr": "D08|Égalité des Genres et Autonomisation",
    "Stat_label_fr": "D08S55|Femmes dans les instances décisionnelles",
    "Domain_label_en": "D08|Gender Equality and Empowerment",
    "Stat_label_en": "D08S55|Women in decision-making bodies",
    "Domain_label_pt": "D08|Igualdade de género e empoderamento",
    "Stat_label_pt": "D08S55|Mulheres em órgãos de decisão",
    "Domain_label_ar": "D08|المساواة بين الجنسين والتمكين",
    "Stat_label_ar": "D08S55|النساء في هيئات صنع القرار"
  },
  {
    "Domain_code": "D08",
    "Domain_label_fr": "D08|Égalité des Genres et Autonomisation",
    "Stat_label_fr": "D08S56|Droits fonciers et propriété des femmes",
    "Domain_label_en": "D08|Gender Equality and Empowerment",
    "Stat_label_en": "D08S56|Women's land and property rights",
    "Domain_label_pt": "D08|Igualdade de género e empoderamento",
    "Stat_label_pt": "D08S56|Direitos das mulheres à terra e à propriedade",
    "Domain_label_ar": "D08|المساواة بين الجنسين والتمكين",
    "Stat_label_ar": "D08S56|حقوق المرأة في الأرض والملكية"
  },
  {
    "Domain_code": "D08",
    "Domain_label_fr": "D08|Égalité des Genres et Autonomisation",
    "Stat_label_fr": "D08S57|Accès au financement (Femmes entrepreneurs)",
    "Domain_label_en": "D08|Gender Equality and Empowerment",
    "Stat_label_en": "D08S57|Access to finance (women entrepreneurs)",
    "Domain_label_pt": "D08|Igualdade de género e empoderamento",
    "Stat_label_pt": "D08S57|Acesso ao financiamento (mulheres empresárias)",
    "Domain_label_ar": "D08|المساواة بين الجنسين والتمكين",
    "Stat_label_ar": "D08S57|الحصول على التمويل (رائدات الأعمال)"
  },
  {
    "Domain_code": "D08",
    "Domain_label_fr": "D08|Égalité des Genres et Autonomisation",
    "Stat_label_fr": "D08S58|Prévalence des Violences Basées sur le Genre",
    "Domain_label_en": "D08|Gender Equality and Empowerment",
    "Stat_label_en": "D08S58|Prevalence of gender-based violence",
    "Domain_label_pt": "D08|Igualdade de género e empoderamento",
    "Stat_label_pt": "D08S58|Prevalência de violência baseada no género",
    "Domain_label_ar": "D08|المساواة بين الجنسين والتمكين",
    "Stat_label_ar": "D08S58|انتشار العنف القائم على النوع الاجتماعي"
  },
  {
    "Domain_code": "D08",
    "Domain_label_fr": "D08|Égalité des Genres et Autonomisation",
    "Stat_label_fr": "D08S59|Mutilations Génitales Féminines (MGF)",
    "Domain_label_en": "D08|Gender Equality and Empowerment",
    "Stat_label_en": "D08S59|Female genital mutilation (FGM)",
    "Domain_label_pt": "D08|Igualdade de género e empoderamento",
    "Stat_label_pt": "D08S59|Mutilação genital feminina (MGF)",
    "Domain_label_ar": "D08|المساواة بين الجنسين والتمكين",
    "Stat_label_ar": "D08S59|تشويه الأعضاء التناسلية الأنثوية (FGM)"
  },
  {
    "Domain_code": "D08",
    "Domain_label_fr": "D08|Égalité des Genres et Autonomisation",
    "Stat_label_fr": "D08S60|Temps de travail domestique non rémunéré",
    "Domain_label_en": "D08|Gender Equality and Empowerment",
    "Stat_label_en": "D08S60|Time spent on unpaid domestic work",
    "Domain_label_pt": "D08|Igualdade de género e empoderamento",
    "Stat_label_pt": "D08S60|Tempo dedicado ao trabalho doméstico não remunerado",
    "Domain_label_ar": "D08|المساواة بين الجنسين والتمكين",
    "Stat_label_ar": "D08S60|الوقت المخصص للعمل المنزلي غير المأجور"
  },
  {
    "Domain_code": "D09",
    "Domain_label_fr": "D09|Environnement, Résilience Climatique et Villes Durables",
    "Stat_label_fr": "D09S61|Pertes économiques dues aux catastrophes",
    "Domain_label_en": "D09|Environment, Climate Resilience and Sustainable Cities",
    "Stat_label_en": "D09S61|Economic losses due to disasters",
    "Domain_label_pt": "D09|Ambiente, resiliência climática e gestão sustentável",
    "Stat_label_pt": "D09S61|Perdas económicas devido a desastres",
    "Domain_label_ar": "D09|البيئة والقدرة على الصمود المناخي والإدارة المستدامة",
    "Stat_label_ar": "D09S61|الخسائر الاقتصادية الناجمة عن الكوارث"
  },
  {
    "Domain_code": "D09",
    "Domain_label_fr": "D09|Environnement, Résilience Climatique et Villes Durables",
    "Stat_label_fr": "D09S62|Couverture forestière",
    "Domain_label_en": "D09|Environment, Climate Resilience and Sustainable Cities",
    "Stat_label_en": "D09S62|Forest cover",
    "Domain_label_pt": "D09|Ambiente, resiliência climática e gestão sustentável",
    "Stat_label_pt": "D09S62|Cobertura florestal",
    "Domain_label_ar": "D09|البيئة والقدرة على الصمود المناخي والإدارة المستدامة",
    "Stat_label_ar": "D09S62|الغطاء الحرجي"
  },
  {
    "Domain_code": "D09",
    "Domain_label_fr": "D09|Environnement, Résilience Climatique et Villes Durables",
    "Stat_label_fr": "D09S63|Ménages avec résilience climatique renforcée",
    "Domain_label_en": "D09|Environment, Climate Resilience and Sustainable Cities",
    "Stat_label_en": "D09S63|Households with strengthened climate resilience",
    "Domain_label_pt": "D09|Ambiente, resiliência climática e gestão sustentável",
    "Stat_label_pt": "D09S63|Agregados familiares com resiliência climática reforçada",
    "Domain_label_ar": "D09|البيئة والقدرة على الصمود المناخي والإدارة المستدامة",
    "Stat_label_ar": "D09S63|أسر لديها قدرة معززة على الصمود المناخي"
  },
  {
    "Domain_code": "D09",
    "Domain_label_fr": "D09|Environnement, Résilience Climatique et Villes Durables",
    "Stat_label_fr": "D09S64|Finance Climatique mobilisée",
    "Domain_label_en": "D09|Environment, Climate Resilience and Sustainable Cities",
    "Stat_label_en": "D09S64|Climate finance mobilized",
    "Domain_label_pt": "D09|Ambiente, resiliência climática e gestão sustentável",
    "Stat_label_pt": "D09S64|Financiamento climático mobilizado",
    "Domain_label_ar": "D09|البيئة والقدرة على الصمود المناخي والإدارة المستدامة",
    "Stat_label_ar": "D09S64|التمويل المناخي المُعبأ"
  },
  {
    "Domain_code": "D09",
    "Domain_label_fr": "D09|Environnement, Résilience Climatique et Villes Durables",
    "Stat_label_fr": "D09S65|Pollution marine (Plastiques/Eutrophisation)",
    "Domain_label_en": "D09|Environment, Climate Resilience and Sustainable Cities",
    "Stat_label_en": "D09S65|Marine pollution (plastics/eutrophication)",
    "Domain_label_pt": "D09|Ambiente, resiliência climática e gestão sustentável",
    "Stat_label_pt": "D09S65|Poluição marinha (plásticos/eutrofização)",
    "Domain_label_ar": "D09|البيئة والقدرة على الصمود المناخي والإدارة المستدامة",
    "Stat_label_ar": "D09S65|التلوث البحري (البلاستيك/الإثراء الغذائي)"
  },
  {
    "Domain_code": "D10",
    "Domain_label_fr": "D10|Gouvernance, Paix et Institutions",
    "Stat_label_fr": "D10S66|Score de Perception de la Corruption (IPC)",
    "Domain_label_en": "D10|Governance, Peace and Institutions",
    "Stat_label_en": "D10S66|Corruption Perceptions Index (CPI) score",
    "Domain_label_pt": "D10|Governação, paz e instituições",
    "Stat_label_pt": "D10S66|Pontuação do Índice de Perceção da Corrupção (IPC)",
    "Domain_label_ar": "D10|الحوكمة والسلام والمؤسسات",
    "Stat_label_ar": "D10S66|درجة مؤشر مدركات الفساد (CPI)"
  },
  {
    "Domain_code": "D10",
    "Domain_label_fr": "D10|Gouvernance, Paix et Institutions",
    "Stat_label_fr": "D10S67|Flux Financiers Illicites (% PIB)",
    "Domain_label_en": "D10|Governance, Peace and Institutions",
    "Stat_label_en": "D10S67|Illicit financial flows (% of GDP)",
    "Domain_label_pt": "D10|Governação, paz e instituições",
    "Stat_label_pt": "D10S67|Fluxos financeiros ilícitos (% do PIB)",
    "Domain_label_ar": "D10|الحوكمة والسلام والمؤسسات",
    "Stat_label_ar": "D10S67|التدفقات المالية غير المشروعة (% من الناتج المحلي الإجمالي)"
  },
  {
    "Domain_code": "D10",
    "Domain_label_fr": "D10|Gouvernance, Paix et Institutions",
    "Stat_label_fr": "D10S68|Perception d'indépendance de la justice",
    "Domain_label_en": "D10|Governance, Peace and Institutions",
    "Stat_label_en": "D10S68|Perceived independence of the judiciary",
    "Domain_label_pt": "D10|Governação, paz e instituições",
    "Stat_label_pt": "D10S68|Independência percebida do poder judicial",
    "Domain_label_ar": "D10|الحوكمة والسلام والمؤسسات",
    "Stat_label_ar": "D10S68|الاستقلال المُدرك للسلطة القضائية"
  },
  {
    "Domain_code": "D10",
    "Domain_label_fr": "D10|Gouvernance, Paix et Institutions",
    "Stat_label_fr": "D10S69|Ratio recettes fiscales / PIB",
    "Domain_label_en": "D10|Governance, Peace and Institutions",
    "Stat_label_en": "D10S69|Tax revenue-to-GDP ratio",
    "Domain_label_pt": "D10|Governação, paz e instituições",
    "Stat_label_pt": "D10S69|Rácio receitas fiscais/PIB",
    "Domain_label_ar": "D10|الحوكمة والسلام والمؤسسات",
    "Stat_label_ar": "D10S69|نسبة الإيرادات الضريبية إلى الناتج المحلي الإجمالي"
  },
  {
    "Domain_code": "D10",
    "Domain_label_fr": "D10|Gouvernance, Paix et Institutions",
    "Stat_label_fr": "D10S70|Taux d'homicide volontaire",
    "Domain_label_en": "D10|Governance, Peace and Institutions",
    "Stat_label_en": "D10S70|Intentional homicide rate",
    "Domain_label_pt": "D10|Governação, paz e instituições",
    "Stat_label_pt": "D10S70|Taxa de homicídios intencionais",
    "Domain_label_ar": "D10|الحوكمة والسلام والمؤسسات",
    "Stat_label_ar": "D10S70|معدل جرائم القتل العمد"
  },
  {
    "Domain_code": "D11",
    "Domain_label_fr": "D11|Économie Bleue et Gestion des Océans",
    "Stat_label_fr": "D11S71|Valeur ajoutée de l'économie bleue (VAB)",
    "Domain_label_en": "D11|Blue Economy and Ocean Management",
    "Stat_label_en": "D11S71|Blue economy value added (GVA)",
    "Domain_label_pt": "D11|Economia azul e gestão dos oceanos",
    "Stat_label_pt": "D11S71|Valor acrescentado da economia azul (VAB)",
    "Domain_label_ar": "D11|الاقتصاد الأزرق وإدارة المحيطات",
    "Stat_label_ar": "D11S71|القيمة المضافة للاقتصاد الأزرق (GVA)"
  },
  {
    "Domain_code": "D11",
    "Domain_label_fr": "D11|Économie Bleue et Gestion des Océans",
    "Stat_label_fr": "D11S72|Emplois liés à l'économie bleue",
    "Domain_label_en": "D11|Blue Economy and Ocean Management",
    "Stat_label_en": "D11S72|Blue economy jobs",
    "Domain_label_pt": "D11|Economia azul e gestão dos oceanos",
    "Stat_label_pt": "D11S72|Empregos na economia azul",
    "Domain_label_ar": "D11|الاقتصاد الأزرق وإدارة المحيطات",
    "Stat_label_ar": "D11S72|فرص العمل في الاقتصاد الأزرق"
  },
  {
    "Domain_code": "D11",
    "Domain_label_fr": "D11|Économie Bleue et Gestion des Océans",
    "Stat_label_fr": "D11S73|Aires marines protégées (% zones territoriales)",
    "Domain_label_en": "D11|Blue Economy and Ocean Management",
    "Stat_label_en": "D11S73|Marine protected areas (% of territorial waters)",
    "Domain_label_pt": "D11|Economia azul e gestão dos oceanos",
    "Stat_label_pt": "D11S73|Áreas marinhas protegidas (% das águas territoriais)",
    "Domain_label_ar": "D11|الاقتصاد الأزرق وإدارة المحيطات",
    "Stat_label_ar": "D11S73|المناطق البحرية المحمية (% من المياه الإقليمية)"
  },
  {
    "Domain_code": "D11",
    "Domain_label_fr": "D11|Économie Bleue et Gestion des Océans",
    "Stat_label_fr": "D11S74|Pénétration de l'énergie bleue",
    "Domain_label_en": "D11|Blue Economy and Ocean Management",
    "Stat_label_en": "D11S74|Blue energy penetration",
    "Domain_label_pt": "D11|Economia azul e gestão dos oceanos",
    "Stat_label_pt": "D11S74|Penetração da energia azul",
    "Domain_label_ar": "D11|الاقتصاد الأزرق وإدارة المحيطات",
    "Stat_label_ar": "D11S74|انتشار الطاقة الزرقاء"
  },
  {
    "Domain_code": "D11",
    "Domain_label_fr": "D11|Économie Bleue et Gestion des Océans",
    "Stat_label_fr": "D11S75|Stocks de poissons durables",
    "Domain_label_en": "D11|Blue Economy and Ocean Management",
    "Stat_label_en": "D11S75|Sustainable fish stocks",
    "Domain_label_pt": "D11|Economia azul e gestão dos oceanos",
    "Stat_label_pt": "D11S75|Stocks de peixe sustentáveis",
    "Domain_label_ar": "D11|الاقتصاد الأزرق وإدارة المحيطات",
    "Stat_label_ar": "D11S75|مخزونات سمكية مستدامة"
  },
  {
    "Domain_code": "D12",
    "Domain_label_fr": "D12|Partenariats et Financement du Développement",
    "Stat_label_fr": "D12S76|Ratio recettes fiscales / PIB",
    "Domain_label_en": "D12|Partnerships and Development Financing",
    "Stat_label_en": "D12S76|Tax revenue-to-GDP ratio",
    "Domain_label_pt": "D12|Parcerias e financiamento do desenvolvimento",
    "Stat_label_pt": "D12S76|Rácio receitas fiscais/PIB",
    "Domain_label_ar": "D12|الشراكات وتمويل التنمية",
    "Stat_label_ar": "D12S76|نسبة الإيرادات الضريبية إلى الناتج المحلي الإجمالي"
  },
  {
    "Domain_code": "D12",
    "Domain_label_fr": "D12|Partenariats et Financement du Développement",
    "Stat_label_fr": "D12S77|Flux financiers illicites (% PIB)",
    "Domain_label_en": "D12|Partnerships and Development Financing",
    "Stat_label_en": "D12S77|Illicit financial flows (% of GDP)",
    "Domain_label_pt": "D12|Parcerias e financiamento do desenvolvimento",
    "Stat_label_pt": "D12S77|Fluxos financeiros ilícitos (% do PIB)",
    "Domain_label_ar": "D12|الشراكات وتمويل التنمية",
    "Stat_label_ar": "D12S77|التدفقات المالية غير المشروعة (% من الناتج المحلي الإجمالي)"
  },
  {
    "Domain_code": "D12",
    "Domain_label_fr": "D12|Partenariats et Financement du Développement",
    "Stat_label_fr": "D12S78|Part du budget financé par impôts nationaux",
    "Domain_label_en": "D12|Partnerships and Development Financing",
    "Stat_label_en": "D12S78|Share of budget financed by domestic taxes",
    "Domain_label_pt": "D12|Parcerias e financiamento do desenvolvimento",
    "Stat_label_pt": "D12S78|Percentagem do orçamento financiado por impostos internos",
    "Domain_label_ar": "D12|الشراكات وتمويل التنمية",
    "Stat_label_ar": "D12S78|حصة الميزانية الممولة بالضرائب المحلية"
  },
  {
    "Domain_code": "D12",
    "Domain_label_fr": "D12|Partenariats et Financement du Développement",
    "Stat_label_fr": "D12S79|Envois de fonds migrants (% PIB)",
    "Domain_label_en": "D12|Partnerships and Development Financing",
    "Stat_label_en": "D12S79|Migrant remittances (% of GDP)",
    "Domain_label_pt": "D12|Parcerias e financiamento do desenvolvimento",
    "Stat_label_pt": "D12S79|Remessas de migrantes (% do PIB)",
    "Domain_label_ar": "D12|الشراكات وتمويل التنمية",
    "Stat_label_ar": "D12S79|تحويلات المهاجرين (% من الناتج المحلي الإجمالي)"
  },
  {
    "Domain_code": "D12",
    "Domain_label_fr": "D12|Partenariats et Financement du Développement",
    "Stat_label_fr": "D12S80|Service de la dette (% exportations)",
    "Domain_label_en": "D12|Partnerships and Development Financing",
    "Stat_label_en": "D12S80|Debt service (% of exports)",
    "Domain_label_pt": "D12|Parcerias e financiamento do desenvolvimento",
    "Stat_label_pt": "D12S80|Serviço da dívida (% das exportações)",
    "Domain_label_ar": "D12|الشراكات وتمويل التنمية",
    "Stat_label_ar": "D12S80|خدمة الدين (% من الصادرات)"
  }
]


APP_DIR = os.path.dirname(os.path.abspath(__file__))
TMP_DIR = os.environ.get("TMPDIR", "/tmp")


def _choose_db_path() -> str:
    """Choisit un emplacement SQLite writable.

    Priorité :
    1. variable d'environnement APP_DB_PATH si fournie ;
    2. dossier de l'app si writable ;
    3. /tmp (plus sûr sur Streamlit Cloud).
    """
    env_path = os.environ.get("APP_DB_PATH", "").strip()
    if env_path:
        return env_path

    app_candidate = os.path.join(APP_DIR, "responses.db")
    try:
        probe = os.path.join(APP_DIR, ".write_test_tmp")
        with open(probe, "w", encoding="utf-8") as f:
            f.write("ok")
        os.remove(probe)
        return app_candidate
    except Exception:
        pass

    return os.path.join(TMP_DIR, "responses.db")


DB_PATH = _choose_db_path()
LONG_LIST_CSV = os.path.join("data", "indicator_longlist.csv")
LONG_LIST_XLSX = os.path.join("data", "longlist.xlsx")


COUNTRY_XLSX = os.path.join("data", "COUNTRY_ISO3_with_EN.xlsx")

REPO_SEARCH_DIRS = [
    Path.cwd(),
    Path.cwd() / "data",
    Path(APP_DIR),
    Path(APP_DIR) / "data",
]


def _dedupe_paths(paths: List[Path]) -> List[Path]:
    out: List[Path] = []
    seen = set()
    for p in paths:
        try:
            key = str(p.resolve())
        except Exception:
            key = str(p)
        if key not in seen:
            seen.add(key)
            out.append(p)
    return out


def candidate_paths(names: List[str]) -> List[str]:
    """
    Retourne des chemins candidats robustes, en cherchant :
    - à la racine du repo / cwd ;
    - dans ./data ;
    - à côté du script ;
    - dans le dossier data à côté du script.
    """
    out: List[Path] = []
    for directory in REPO_SEARCH_DIRS:
        for name in names:
            out.append(directory / name)
    # Compatibilité ascendante avec les chemins relatifs déjà utilisés
    for name in names:
        out.append(Path(name))
        out.append(Path(".") / name)
        out.append(Path(".") / "data" / name)
    return [str(p) for p in _dedupe_paths(out)]


def _first_existing_path(names: List[str]) -> str:
    for p in candidate_paths(names):
        if os.path.exists(p):
            return p
    return ""


def _non_empty_ratio(s: pd.Series) -> float:
    if s is None or len(s) == 0:
        return 0.0
    ss = s.astype(str).str.strip()
    return float((ss != "").mean())




# =========================
# Embedded PT/AR fallbacks
# =========================
# These dictionaries guarantee Portuguese (Portugal) and Arabic labels even when
# the external Excel/CSV files are present but their PT/AR columns are still empty.
EMBEDDED_COUNTRY_NAME_PT = {
    "DZA": "Argélia",
    "AGO": "Angola",
    "BEN": "Benim",
    "BWA": "Botsuana",
    "BFA": "Burkina Faso",
    "BDI": "Burundi",
    "CMR": "Camarões",
    "CPV": "Cabo Verde",
    "CAF": "República Centro-Africana",
    "TCD": "Chade",
    "COM": "Comores",
    "COG": "República do Congo",
    "CIV": "Costa do Marfim",
    "COD": "República Democrática do Congo",
    "DJI": "Djibuti",
    "EGY": "Egito",
    "GNQ": "Guiné Equatorial",
    "ERI": "Eritreia",
    "ETH": "Etiópia",
    "GAB": "Gabão",
    "GMB": "Gâmbia",
    "GHA": "Gana",
    "GIN": "Guiné",
    "GNB": "Guiné-Bissau",
    "KEN": "Quénia",
    "LSO": "Lesoto",
    "LBR": "Libéria",
    "LBY": "Líbia",
    "MDG": "Madagáscar",
    "MWI": "Maláui",
    "MLI": "Mali",
    "MRT": "Mauritânia",
    "MUS": "Maurícia",
    "MAR": "Marrocos",
    "MOZ": "Moçambique",
    "NAM": "Namíbia",
    "NER": "Níger",
    "NGA": "Nigéria",
    "RWA": "Ruanda",
    "STP": "São Tomé e Príncipe",
    "SEN": "Senegal",
    "SYC": "Seicheles",
    "SLE": "Serra Leoa",
    "SOM": "Somália",
    "SSD": "Sudão do Sul",
    "ZAF": "África do Sul",
    "SDN": "Sudão",
    "SWZ": "Essuatíni",
    "TZA": "Tanzânia",
    "TGO": "Togo",
    "TUN": "Tunísia",
    "UGA": "Uganda",
    "ZMB": "Zâmbia",
    "ZWE": "Zimbábue"
}
EMBEDDED_COUNTRY_NAME_AR = {
    "DZA": "الجزائر",
    "AGO": "أنغولا",
    "BEN": "بنين",
    "BWA": "بوتسوانا",
    "BFA": "بوركينا فاسو",
    "BDI": "بوروندي",
    "CMR": "الكاميرون",
    "CPV": "الرأس الأخضر",
    "CAF": "جمهورية أفريقيا الوسطى",
    "TCD": "تشاد",
    "COM": "جزر القمر",
    "COG": "جمهورية الكونغو",
    "CIV": "كوت ديفوار",
    "COD": "جمهورية الكونغو الديمقراطية",
    "DJI": "جيبوتي",
    "EGY": "مصر",
    "GNQ": "غينيا الاستوائية",
    "ERI": "إريتريا",
    "ETH": "إثيوبيا",
    "GAB": "الغابون",
    "GMB": "غامبيا",
    "GHA": "غانا",
    "GIN": "غينيا",
    "GNB": "غينيا بيساو",
    "KEN": "كينيا",
    "LSO": "ليسوتو",
    "LBR": "ليبيريا",
    "LBY": "ليبيا",
    "MDG": "مدغشقر",
    "MWI": "ملاوي",
    "MLI": "مالي",
    "MRT": "موريتانيا",
    "MUS": "موريشيوس",
    "MAR": "المغرب",
    "MOZ": "موزمبيق",
    "NAM": "ناميبيا",
    "NER": "النيجر",
    "NGA": "نيجيريا",
    "RWA": "رواندا",
    "STP": "ساو تومي وبرينسيبي",
    "SEN": "السنغال",
    "SYC": "سيشل",
    "SLE": "سيراليون",
    "SOM": "الصومال",
    "SSD": "جنوب السودان",
    "ZAF": "جنوب أفريقيا",
    "SDN": "السودان",
    "SWZ": "إسواتيني",
    "TZA": "تنزانيا",
    "TGO": "توغو",
    "TUN": "تونس",
    "UGA": "أوغندا",
    "ZMB": "زامبيا",
    "ZWE": "زيمبابوي"
}
EMBEDDED_DOMAIN_LABEL_PT = {
    "D01": "Crescimento económico, transformação estrutural e industrialização",
    "D02": "Emprego, trabalho decente e proteção social",
    "D03": "Agricultura sustentável, segurança alimentar e nutrição",
    "D04": "Infraestruturas, industrialização e inovação",
    "D05": "Inclusão, pobreza e desigualdades",
    "D06": "Educação, competências e capital humano",
    "D07": "Saúde, bem-estar e acesso universal",
    "D08": "Igualdade de género e empoderamento",
    "D09": "Ambiente, resiliência climática e gestão sustentável",
    "D10": "Governação, paz e instituições",
    "D11": "Economia azul e gestão dos oceanos",
    "D12": "Parcerias e financiamento do desenvolvimento"
}
EMBEDDED_DOMAIN_LABEL_AR = {
    "D01": "النمو الاقتصادي والتحول الهيكلي والتصنيع",
    "D02": "العمالة والعمل اللائق والحماية الاجتماعية",
    "D03": "الزراعة المستدامة والأمن الغذائي والتغذية",
    "D04": "البنية التحتية والتصنيع والابتكار",
    "D05": "الإدماج والفقر وعدم المساواة",
    "D06": "التعليم والمهارات ورأس المال البشري",
    "D07": "الصحة والرفاه وإتاحة الخدمات للجميع",
    "D08": "المساواة بين الجنسين والتمكين",
    "D09": "البيئة والقدرة على الصمود المناخي والإدارة المستدامة",
    "D10": "الحوكمة والسلام والمؤسسات",
    "D11": "الاقتصاد الأزرق وإدارة المحيطات",
    "D12": "الشراكات وتمويل التنمية"
}
EMBEDDED_STAT_LABEL_PT = {
    "D01S01": "Taxa de crescimento do PIB real",
    "D01S02": "PIB per capita (USD constantes)",
    "D01S03": "Participação do valor acrescentado da indústria transformadora (VAI) no PIB",
    "D01S04": "Participação de produtos de alta tecnologia nas exportações manufaturadas",
    "D01S05": "Participação de África no valor acrescentado manufatureiro mundial",
    "D01S06": "Participação do comércio intra-africano no comércio total",
    "D01S07": "Índice de diversificação económica",
    "D01S08": "Contribuição direta do turismo para o PIB",
    "D01S09": "Formação bruta de capital fixo (% do PIB)",
    "D01S10": "Países que comercializam no âmbito do regime AfCFTA",
    "D02S11": "Taxa de desemprego (total, jovens, mulheres)",
    "D02S12": "Número de empregos criados (incluindo jovens)",
    "D02S13": "Percentagem de jovens NEET (Nem em emprego, educação ou formação) (15–24)",
    "D02S14": "Taxa de cobertura da proteção social",
    "D02S15": "Percentagem de crianças (5–17) em trabalho infantil",
    "D02S16": "Percentagem de emprego informal no emprego total",
    "D02S17": "Diferença de género nos ganhos horários",
    "D03S18": "Prevalência de subalimentação (fome)",
    "D03S19": "Prevalência de atraso de crescimento (crianças menores de 5 anos)",
    "D03S20": "Taxa de crescimento dos rendimentos agrícolas",
    "D03S21": "Redução das perdas pós-colheita",
    "D03S22": "Percentagem das importações alimentares no total das importações",
    "D03S23": "Terras sob gestão sustentável",
    "D03S24": "Resiliência climática dos agregados familiares agrícolas",
    "D03S25": "Zonas especiais de processamento agroindustrial (SAPZ) estabelecidas",
    "D03S26": "Uso de fertilizantes (kg/ha)",
    "D04S27": "Acesso à eletricidade (agregados familiares)",
    "D04S28": "Acesso à internet de banda larga (qualidade)",
    "D04S29": "Conectividade rodoviária interafricana",
    "D04S30": "Conectividade ferroviária interafricana",
    "D04S31": "Despesa em I&D (% do PIB)",
    "D04S32": "Capacidade instalada de geração de eletricidade",
    "D04S33": "Contribuição dos serviços digitais para o PIB",
    "D04S34": "Percentagem de energias renováveis",
    "D05S35": "Taxa de pobreza (linha nacional de pobreza)",
    "D05S36": "Índice de Gini (desigualdade de rendimentos)",
    "D05S37": "Acesso a água potável gerida com segurança",
    "D05S38": "Acesso a saneamento melhorado",
    "D05S39": "População urbana a viver em bairros degradados (slums)",
    "D05S40": "Acesso a habitação adequada",
    "D05S41": "Acesso à água (países em transição / frágeis)",
    "D06S42": "Taxa líquida de escolarização (primário)",
    "D06S43": "Taxa líquida de escolarização (secundário)",
    "D06S44": "Taxa de matrícula no pré-escolar",
    "D06S45": "Percentagem de diplomados em STEM (CTEM)",
    "D06S46": "Nível mínimo de proficiência (leitura/matemática)",
    "D06S47": "Rácio aluno/professor qualificado",
    "D06S48": "Taxa bruta de escolarização (ensino superior)",
    "D07S49": "Acesso aos cuidados de saúde primários",
    "D07S50": "Razão de mortalidade materna",
    "D07S51": "Taxa de mortalidade de menores de cinco anos",
    "D07S52": "Incidência do VIH / cobertura de TAR (ART)",
    "D07S53": "Incidência de malária",
    "D07S54": "Índice de cobertura de serviços da Cobertura Universal de Saúde (UHC)",
    "D08S55": "Mulheres em órgãos de decisão",
    "D08S56": "Direitos das mulheres à terra e à propriedade",
    "D08S57": "Acesso ao financiamento (mulheres empresárias)",
    "D08S58": "Prevalência de violência baseada no género",
    "D08S59": "Mutilação genital feminina (MGF)",
    "D08S60": "Tempo dedicado ao trabalho doméstico não remunerado",
    "D09S61": "Perdas económicas devido a desastres",
    "D09S62": "Cobertura florestal",
    "D09S63": "Agregados familiares com resiliência climática reforçada",
    "D09S64": "Financiamento climático mobilizado",
    "D09S65": "Poluição marinha (plásticos/eutrofização)",
    "D10S66": "Pontuação do Índice de Perceção da Corrupção (IPC)",
    "D10S67": "Fluxos financeiros ilícitos (% do PIB)",
    "D10S68": "Independência percebida do poder judicial",
    "D10S69": "Rácio receitas fiscais/PIB",
    "D10S70": "Taxa de homicídios intencionais",
    "D11S71": "Valor acrescentado da economia azul (VAB)",
    "D11S72": "Empregos na economia azul",
    "D11S73": "Áreas marinhas protegidas (% das águas territoriais)",
    "D11S74": "Penetração da energia azul",
    "D11S75": "Stocks de peixe sustentáveis",
    "D12S76": "Rácio receitas fiscais/PIB",
    "D12S77": "Fluxos financeiros ilícitos (% do PIB)",
    "D12S78": "Percentagem do orçamento financiado por impostos internos",
    "D12S79": "Remessas de migrantes (% do PIB)",
    "D12S80": "Serviço da dívida (% das exportações)"
}
EMBEDDED_STAT_LABEL_AR = {
    "D01S01": "معدل نمو الناتج المحلي الإجمالي الحقيقي",
    "D01S02": "نصيب الفرد من الناتج المحلي الإجمالي (دولار أمريكي ثابت)",
    "D01S03": "حصة القيمة المضافة للصناعة التحويلية (MVA) من الناتج المحلي الإجمالي",
    "D01S04": "حصة المنتجات عالية التكنولوجيا في الصادرات المصنعة",
    "D01S05": "حصة أفريقيا من القيمة المضافة للصناعة التحويلية العالمية",
    "D01S06": "حصة التجارة البينية الأفريقية من إجمالي التجارة",
    "D01S07": "مؤشر التنويع الاقتصادي",
    "D01S08": "المساهمة المباشرة للسياحة في الناتج المحلي الإجمالي",
    "D01S09": "تكوين رأس المال الثابت الإجمالي (% من الناتج المحلي الإجمالي)",
    "D01S10": "البلدان التي تتاجر في إطار نظام AfCFTA",
    "D02S11": "معدل البطالة (الإجمالي، الشباب، النساء)",
    "D02S12": "عدد فرص العمل التي تم خلقها (بما في ذلك للشباب)",
    "D02S13": "نسبة الشباب خارج العمل والتعليم أو التدريب (NEET) (15–24)",
    "D02S14": "معدل التغطية بالحماية الاجتماعية",
    "D02S15": "نسبة الأطفال (5–17) في عمالة الأطفال",
    "D02S16": "حصة العمالة غير الرسمية من إجمالي العمالة",
    "D02S17": "الفجوة بين الجنسين في الأجور بالساعة",
    "D03S18": "انتشار نقص التغذية (الجوع)",
    "D03S19": "انتشار التقزم (الأطفال دون 5 سنوات)",
    "D03S20": "معدل نمو غلة المحاصيل الزراعية",
    "D03S21": "خفض خسائر ما بعد الحصاد",
    "D03S22": "حصة واردات الغذاء من إجمالي الواردات",
    "D03S23": "الأراضي المُدارة إدارة مستدامة",
    "D03S24": "القدرة على الصمود المناخي للأسر الزراعية",
    "D03S25": "إنشاء مناطق معالجة زراعية-صناعية خاصة (SAPZ)",
    "D03S26": "استخدام الأسمدة (كغ/هكتار)",
    "D04S27": "الحصول على الكهرباء (الأسر)",
    "D04S28": "الوصول إلى الإنترنت عريض النطاق (الجودة)",
    "D04S29": "الترابط الطرقي بين البلدان الأفريقية",
    "D04S30": "الترابط السككي بين البلدان الأفريقية",
    "D04S31": "الإنفاق على البحث والتطوير (% من الناتج المحلي الإجمالي)",
    "D04S32": "القدرة المركبة لتوليد الكهرباء",
    "D04S33": "مساهمة الخدمات الرقمية في الناتج المحلي الإجمالي",
    "D04S34": "حصة الطاقة المتجددة",
    "D05S35": "معدل الفقر (خط الفقر الوطني)",
    "D05S36": "مؤشر جيني (عدم المساواة في الدخل)",
    "D05S37": "الحصول على مياه شرب مُدارة بأمان",
    "D05S38": "الحصول على خدمات صرف صحي محسّنة",
    "D05S39": "السكان الحضريون الذين يعيشون في الأحياء العشوائية",
    "D05S40": "الحصول على سكن ملائم",
    "D05S41": "الحصول على المياه (بلدان انتقالية / هشة)",
    "D06S42": "معدل الالتحاق الصافي (الابتدائي)",
    "D06S43": "معدل الالتحاق الصافي (الثانوي)",
    "D06S44": "معدل الالتحاق بالتعليم قبل الابتدائي",
    "D06S45": "حصة الخريجين في مجالات STEM",
    "D06S46": "الحد الأدنى لمستوى الكفاءة (القراءة/الرياضيات)",
    "D06S47": "نسبة التلاميذ إلى المعلمين المؤهلين",
    "D06S48": "معدل الالتحاق الإجمالي (التعليم العالي)",
    "D07S49": "الحصول على الرعاية الصحية الأولية",
    "D07S50": "معدل وفيات الأمهات",
    "D07S51": "معدل وفيات الأطفال دون الخامسة",
    "D07S52": "معدل الإصابة بفيروس HIV / تغطية العلاج بمضادات الفيروسات (ART)",
    "D07S53": "معدل الإصابة بالملاريا",
    "D07S54": "مؤشر تغطية خدمات التغطية الصحية الشاملة (UHC)",
    "D08S55": "النساء في هيئات صنع القرار",
    "D08S56": "حقوق المرأة في الأرض والملكية",
    "D08S57": "الحصول على التمويل (رائدات الأعمال)",
    "D08S58": "انتشار العنف القائم على النوع الاجتماعي",
    "D08S59": "تشويه الأعضاء التناسلية الأنثوية (FGM)",
    "D08S60": "الوقت المخصص للعمل المنزلي غير المأجور",
    "D09S61": "الخسائر الاقتصادية الناجمة عن الكوارث",
    "D09S62": "الغطاء الحرجي",
    "D09S63": "أسر لديها قدرة معززة على الصمود المناخي",
    "D09S64": "التمويل المناخي المُعبأ",
    "D09S65": "التلوث البحري (البلاستيك/الإثراء الغذائي)",
    "D10S66": "درجة مؤشر مدركات الفساد (CPI)",
    "D10S67": "التدفقات المالية غير المشروعة (% من الناتج المحلي الإجمالي)",
    "D10S68": "الاستقلال المُدرك للسلطة القضائية",
    "D10S69": "نسبة الإيرادات الضريبية إلى الناتج المحلي الإجمالي",
    "D10S70": "معدل جرائم القتل العمد",
    "D11S71": "القيمة المضافة للاقتصاد الأزرق (GVA)",
    "D11S72": "فرص العمل في الاقتصاد الأزرق",
    "D11S73": "المناطق البحرية المحمية (% من المياه الإقليمية)",
    "D11S74": "انتشار الطاقة الزرقاء",
    "D11S75": "مخزونات سمكية مستدامة",
    "D12S76": "نسبة الإيرادات الضريبية إلى الناتج المحلي الإجمالي",
    "D12S77": "التدفقات المالية غير المشروعة (% من الناتج المحلي الإجمالي)",
    "D12S78": "حصة الميزانية الممولة بالضرائب المحلية",
    "D12S79": "تحويلات المهاجرين (% من الناتج المحلي الإجمالي)",
    "D12S80": "خدمة الدين (% من الصادرات)"
}


def _fill_longlist_pt_ar_from_embedded(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return df
    df = df.copy()
    for col in ["domain_label_pt", "domain_label_ar", "stat_label_pt", "stat_label_ar"]:
        if col not in df.columns:
            df[col] = ""
    if "domain_code" in df.columns:
        dc = df["domain_code"].astype(str).str.strip()
        df["domain_label_pt"] = df["domain_label_pt"].where(df["domain_label_pt"].astype(str).str.strip() != "", dc.map(EMBEDDED_DOMAIN_LABEL_PT).fillna(""))
        df["domain_label_ar"] = df["domain_label_ar"].where(df["domain_label_ar"].astype(str).str.strip() != "", dc.map(EMBEDDED_DOMAIN_LABEL_AR).fillna(""))
    if "stat_code" in df.columns:
        sc = df["stat_code"].astype(str).str.strip()
        df["stat_label_pt"] = df["stat_label_pt"].where(df["stat_label_pt"].astype(str).str.strip() != "", sc.map(EMBEDDED_STAT_LABEL_PT).fillna(""))
        df["stat_label_ar"] = df["stat_label_ar"].where(df["stat_label_ar"].astype(str).str.strip() != "", sc.map(EMBEDDED_STAT_LABEL_AR).fillna(""))
    if "domain_label_en" in df.columns and "domain_label_fr" in df.columns:
        df["domain_label_pt"] = df["domain_label_pt"].where(df["domain_label_pt"].astype(str).str.strip() != "", df["domain_label_en"].where(df["domain_label_en"].astype(str).str.strip() != "", df["domain_label_fr"]))
        df["domain_label_ar"] = df["domain_label_ar"].where(df["domain_label_ar"].astype(str).str.strip() != "", df["domain_label_en"].where(df["domain_label_en"].astype(str).str.strip() != "", df["domain_label_fr"]))
    if "stat_label_en" in df.columns and "stat_label_fr" in df.columns:
        df["stat_label_pt"] = df["stat_label_pt"].where(df["stat_label_pt"].astype(str).str.strip() != "", df["stat_label_en"].where(df["stat_label_en"].astype(str).str.strip() != "", df["stat_label_fr"]))
        df["stat_label_ar"] = df["stat_label_ar"].where(df["stat_label_ar"].astype(str).str.strip() != "", df["stat_label_en"].where(df["stat_label_en"].astype(str).str.strip() != "", df["stat_label_fr"]))
    return df.fillna("")


def _fill_countries_pt_ar_from_embedded(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty or "COUNTRY_ISO3" not in df.columns:
        return df
    df = df.copy()
    for col in ["COUNTRY_NAME_PT", "COUNTRY_NAME_AR"]:
        if col not in df.columns:
            df[col] = ""
    iso = df["COUNTRY_ISO3"].astype(str).str.strip().str.upper()
    df["COUNTRY_NAME_PT"] = df["COUNTRY_NAME_PT"].where(df["COUNTRY_NAME_PT"].astype(str).str.strip() != "", iso.map(EMBEDDED_COUNTRY_NAME_PT).fillna(""))
    df["COUNTRY_NAME_AR"] = df["COUNTRY_NAME_AR"].where(df["COUNTRY_NAME_AR"].astype(str).str.strip() != "", iso.map(EMBEDDED_COUNTRY_NAME_AR).fillna(""))
    if "COUNTRY_NAME_EN" in df.columns and "COUNTRY_NAME_FR" in df.columns:
        df["COUNTRY_NAME_PT"] = df["COUNTRY_NAME_PT"].where(df["COUNTRY_NAME_PT"].astype(str).str.strip() != "", df["COUNTRY_NAME_EN"].where(df["COUNTRY_NAME_EN"].astype(str).str.strip() != "", df["COUNTRY_NAME_FR"]))
        df["COUNTRY_NAME_AR"] = df["COUNTRY_NAME_AR"].where(df["COUNTRY_NAME_AR"].astype(str).str.strip() != "", df["COUNTRY_NAME_EN"].where(df["COUNTRY_NAME_EN"].astype(str).str.strip() != "", df["COUNTRY_NAME_FR"]))
    return df.fillna("")


def _read_excel_flexible(path: str) -> pd.DataFrame:
    """Read xlsx/xls robustly with multiple fallbacks and first non-empty sheet."""
    last_exc = None
    for kwargs in [
        {"dtype": str, "sheet_name": 0},
        {"dtype": str, "engine": "openpyxl", "sheet_name": 0},
        {"dtype": str, "sheet_name": None},
        {"dtype": str, "engine": "openpyxl", "sheet_name": None},
    ]:
        try:
            data = pd.read_excel(path, **kwargs)
            if isinstance(data, dict):
                for _, df in data.items():
                    if df is not None and not df.empty:
                        return df.fillna("")
                return pd.DataFrame()
            return data.fillna("")
        except Exception as exc:
            last_exc = exc
            continue
    raise last_exc if last_exc else RuntimeError("Unable to read Excel file: " + str(path))

def _safe_split_label(series: pd.Series) -> pd.Series:
    return series.astype(str).str.split("|", n=1).str[-1].str.strip()


def _normalize_longlist_csv(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy().fillna("")
    # Harmonisation des noms de colonnes
    rename_map = {
        "Domain_code": "domain_code",
        "Stat_code": "stat_code",
        "Domain_label_fr": "domain_label_fr",
        "Domain_label_en": "domain_label_en",
        "Domain_label_pt": "domain_label_pt",
        "Domain_label_ar": "domain_label_ar",
        "Stat_label_fr": "stat_label_fr",
        "Stat_label_en": "stat_label_en",
        "Stat_label_pt": "stat_label_pt",
        "Stat_label_ar": "stat_label_ar",
    }
    for old, new in rename_map.items():
        if old in df.columns and new not in df.columns:
            df[new] = df[old]
    # Retirer d'éventuels préfixes "D01|" / "D01S01|"
    for col in ["domain_label_fr", "domain_label_en", "domain_label_pt", "domain_label_ar",
                "stat_label_fr", "stat_label_en", "stat_label_pt", "stat_label_ar"]:
        if col in df.columns:
            df[col] = _safe_split_label(df[col])
    if "domain_code" not in df.columns and "stat_code" in df.columns:
        df["domain_code"] = df["stat_code"].astype(str).str.extract(r"^(D\d{2})", expand=False).fillna("")
    for col in ["domain_code", "stat_code"]:
        if col not in df.columns:
            df[col] = ""
        df[col] = df[col].astype(str).str.strip()
    if "domain_label_fr" not in df.columns:
        df["domain_label_fr"] = ""
    if "stat_label_fr" not in df.columns:
        df["stat_label_fr"] = ""
    if "domain_label_en" not in df.columns:
        df["domain_label_en"] = df["domain_label_fr"]
    if "stat_label_en" not in df.columns:
        df["stat_label_en"] = df["stat_label_fr"]
    if "domain_label_pt" not in df.columns:
        df["domain_label_pt"] = df["domain_label_en"].where(df["domain_label_en"].astype(str).str.strip() != "", df["domain_label_fr"])
    if "stat_label_pt" not in df.columns:
        df["stat_label_pt"] = df["stat_label_en"].where(df["stat_label_en"].astype(str).str.strip() != "", df["stat_label_fr"])
    if "domain_label_ar" not in df.columns:
        df["domain_label_ar"] = df["domain_label_en"].where(df["domain_label_en"].astype(str).str.strip() != "", df["domain_label_fr"])
    if "stat_label_ar" not in df.columns:
        df["stat_label_ar"] = df["stat_label_en"].where(df["stat_label_en"].astype(str).str.strip() != "", df["stat_label_fr"])
    df = _fill_longlist_pt_ar_from_embedded(df)
    return df[[
        "domain_code", "domain_label_fr", "domain_label_en", "domain_label_pt", "domain_label_ar",
        "stat_code", "stat_label_fr", "stat_label_en", "stat_label_pt", "stat_label_ar"
    ]].fillna("")


def _normalize_longlist_xlsx(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy().fillna("")
    if not set(["Domain_code", "Domain_label_fr", "Stat_label_fr"]).issubset(df.columns):
        return pd.DataFrame()
    out = pd.DataFrame()
    out["domain_code"] = df["Domain_code"].astype(str).str.strip()
    out["domain_label_fr"] = _safe_split_label(df["Domain_label_fr"])
    out["stat_code"] = df["Stat_label_fr"].astype(str).str.split("|", n=1).str[0].str.strip()
    out["stat_label_fr"] = _safe_split_label(df["Stat_label_fr"])
    out["domain_label_en"] = _safe_split_label(df["Domain_label_en"]) if "Domain_label_en" in df.columns else out["domain_label_fr"]
    out["stat_label_en"] = _safe_split_label(df["Stat_label_en"]) if "Stat_label_en" in df.columns else out["stat_label_fr"]
    out["domain_label_pt"] = _safe_split_label(df["Domain_label_pt"]) if "Domain_label_pt" in df.columns else out["domain_label_en"]
    out["stat_label_pt"] = _safe_split_label(df["Stat_label_pt"]) if "Stat_label_pt" in df.columns else out["stat_label_en"]
    out["domain_label_ar"] = _safe_split_label(df["Domain_label_ar"]) if "Domain_label_ar" in df.columns else out["domain_label_en"]
    out["stat_label_ar"] = _safe_split_label(df["Stat_label_ar"]) if "Stat_label_ar" in df.columns else out["stat_label_en"]
    out = _fill_longlist_pt_ar_from_embedded(out)
    return out[[
        "domain_code", "domain_label_fr", "domain_label_en", "domain_label_pt", "domain_label_ar",
        "stat_code", "stat_label_fr", "stat_label_en", "stat_label_pt", "stat_label_ar"
    ]].fillna("")


def _merge_longlists(base_df: pd.DataFrame, rich_df: pd.DataFrame) -> pd.DataFrame:
    if base_df is None or base_df.empty:
        return rich_df.copy()
    if rich_df is None or rich_df.empty:
        return base_df.copy()
    keys = ["domain_code", "stat_code"]
    base = base_df.copy()
    rich = rich_df.copy()
    base = base.drop_duplicates(subset=keys, keep="first")
    rich = rich.drop_duplicates(subset=keys, keep="first")
    merged = base.merge(
        rich[keys + ["domain_label_fr", "domain_label_en", "domain_label_pt", "domain_label_ar",
                     "stat_label_fr", "stat_label_en", "stat_label_pt", "stat_label_ar"]],
        on=keys, how="outer", suffixes=("", "_rich")
    )
    for col in ["domain_label_fr", "domain_label_en", "domain_label_pt", "domain_label_ar",
                "stat_label_fr", "stat_label_en", "stat_label_pt", "stat_label_ar"]:
        rich_col = f"{col}_rich"
        if rich_col in merged.columns:
            merged[col] = merged[rich_col].where(merged[rich_col].astype(str).str.strip() != "", merged.get(col, ""))
    keep_cols = keys + ["domain_label_fr", "domain_label_en", "domain_label_pt", "domain_label_ar",
                        "stat_label_fr", "stat_label_en", "stat_label_pt", "stat_label_ar"]
    merged = merged[keep_cols].copy()
    for col in ["domain_label_en", "domain_label_pt", "domain_label_ar"]:
        merged[col] = merged[col].where(merged[col].astype(str).str.strip() != "", merged["domain_label_fr"])
    for col in ["stat_label_en", "stat_label_pt", "stat_label_ar"]:
        merged[col] = merged[col].where(merged[col].astype(str).str.strip() != "", merged["stat_label_fr"])
    return merged.fillna("")


def _normalize_countries_df(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy().fillna("")
    df.columns = [str(c).strip() for c in df.columns]
    if "COUNTRY_ISO3" not in df.columns:
        return pd.DataFrame()
    if "COUNTRY_NAME_FR" not in df.columns:
        df["COUNTRY_NAME_FR"] = ""
    if "COUNTRY_NAME_EN" not in df.columns:
        df["COUNTRY_NAME_EN"] = ""
    if "COUNTRY_NAME_PT" not in df.columns:
        df["COUNTRY_NAME_PT"] = ""
    if "COUNTRY_NAME_AR" not in df.columns:
        df["COUNTRY_NAME_AR"] = ""
    # Si seuls COUNTRY_VALUE_* existent, on en extrait le nom après "ISO3 | "
    for name_col, value_col in [
        ("COUNTRY_NAME_FR", "COUNTRY_VALUE"),
        ("COUNTRY_NAME_EN", "COUNTRY_VALUE_EN"),
        ("COUNTRY_NAME_PT", "COUNTRY_VALUE_PT"),
        ("COUNTRY_NAME_AR", "COUNTRY_VALUE_AR"),
    ]:
        if value_col in df.columns:
            extracted = df[value_col].astype(str).str.split("|", n=1).str[-1].str.strip()
            df[name_col] = df[name_col].where(df[name_col].astype(str).str.strip() != "", extracted)
    df["COUNTRY_ISO3"] = df["COUNTRY_ISO3"].astype(str).str.strip().str.upper()
    for col in ["COUNTRY_NAME_FR", "COUNTRY_NAME_EN", "COUNTRY_NAME_PT", "COUNTRY_NAME_AR"]:
        df[col] = df[col].astype(str).str.strip()
    # Fallbacks : PT Portugal => EN puis FR ; AR => EN puis FR
    df["COUNTRY_NAME_EN"] = df["COUNTRY_NAME_EN"].where(df["COUNTRY_NAME_EN"] != "", df["COUNTRY_NAME_FR"])
    df["COUNTRY_NAME_PT"] = df["COUNTRY_NAME_PT"].where(df["COUNTRY_NAME_PT"] != "", df["COUNTRY_NAME_EN"])
    df["COUNTRY_NAME_AR"] = df["COUNTRY_NAME_AR"].where(df["COUNTRY_NAME_AR"] != "", df["COUNTRY_NAME_EN"])
    df = _fill_countries_pt_ar_from_embedded(df)
    return df[df["COUNTRY_ISO3"] != ""][["COUNTRY_ISO3", "COUNTRY_NAME_FR", "COUNTRY_NAME_EN", "COUNTRY_NAME_PT", "COUNTRY_NAME_AR"]].copy()


def _embedded_countries_df() -> pd.DataFrame:
    try:
        df = pd.DataFrame(EMBEDDED_COUNTRIES_ROWS).fillna("")
        df = _normalize_countries_df(df)
        df.attrs["source_path"] = "embedded::COUNTRY_ISO3_4_filled_pt_ar_final.xlsx"
        return df
    except Exception:
        empty = pd.DataFrame(columns=["COUNTRY_ISO3", "COUNTRY_NAME_FR", "COUNTRY_NAME_EN", "COUNTRY_NAME_PT", "COUNTRY_NAME_AR"])
        empty.attrs["source_path"] = ""
        return empty


def _merge_countries(base_df: pd.DataFrame, rich_df: pd.DataFrame) -> pd.DataFrame:
    if base_df is None or base_df.empty:
        return rich_df.copy() if rich_df is not None else pd.DataFrame()
    if rich_df is None or rich_df.empty:
        return base_df.copy()
    keys = ["COUNTRY_ISO3"]
    base = base_df.drop_duplicates(subset=keys, keep="first").copy()
    rich = rich_df.drop_duplicates(subset=keys, keep="first").copy()
    merged = base.merge(rich, on=keys, how="outer", suffixes=("", "_rich"))
    for col in ["COUNTRY_NAME_FR", "COUNTRY_NAME_EN", "COUNTRY_NAME_PT", "COUNTRY_NAME_AR"]:
        rich_col = f"{col}_rich"
        if rich_col in merged.columns:
            merged[col] = merged[rich_col].where(merged[rich_col].astype(str).str.strip() != "", merged.get(col, ""))
    keep = ["COUNTRY_ISO3", "COUNTRY_NAME_FR", "COUNTRY_NAME_EN", "COUNTRY_NAME_PT", "COUNTRY_NAME_AR"]
    merged = merged[[c for c in keep if c in merged.columns]].fillna("")
    return _normalize_countries_df(merged)


def _embedded_longlist_df() -> pd.DataFrame:
    try:
        df = pd.DataFrame(EMBEDDED_LONGLIST_ROWS).fillna("")
        df = _normalize_longlist_xlsx(df)
        df.attrs["source_path"] = "embedded::longlist_filled_pt_ar_final.xlsx"
        return df
    except Exception:
        empty = pd.DataFrame(columns=[
            "domain_code", "domain_label_fr", "domain_label_en", "domain_label_pt", "domain_label_ar",
            "stat_code", "stat_label_fr", "stat_label_en", "stat_label_pt", "stat_label_ar",
        ])
        empty.attrs["source_path"] = ""
        return empty

UK_FR = "NSP (Ne sais pas)"
UK_EN = "DNK (Do not know)"
UK_PT = "NSP (Não sabe)"
UK_AR = "لا أعلم"

# Version du scoring (pour compatibilité ascendante)
# v1 : ancien critère "gap" (écart) ; v2 : disponibilité inversée (Bonne=1) ; v3 : disponibilité directe (Bonne=3)
SCORING_VERSION = 3


# Scores affichés (notation multicritères)
# Barèmes de notation (0–3) par critère (scoring rationalisé)
# Remarque : 0 = NSP / DNK (ne sait pas) ; 1–3 = intensité croissante selon le libellé.
SCORE_SCALES = {
    "demand": {
        "fr": {0: "NSP", 1: "Faible", 2: "Moyenne", 3: "Élevée"},
        "en": {0: "DNK", 1: "Low", 2: "Medium", 3: "High"},
        "pt": {0: "NSP", 1: "Baixa", 2: "Média", 3: "Alta"},
        "ar": {0: "لا أعلم", 1: "منخفض", 2: "متوسط", 3: "مرتفع"},
    },
    "availability": {
        "fr": {0: "NSP", 1: "Faible ou inexistante", 2: "Partielle", 3: "Bonne"},
        "en": {0: "DNK", 1: "Low or none", 2: "Partial", 3: "Good"},
        "pt": {0: "NSP", 1: "Baixa ou inexistente", 2: "Parcial", 3: "Boa"},
        "ar": {0: "لا أعلم", 1: "ضعيفة أو غير متاحة", 2: "جزئية", 3: "جيدة"},
    },
    "feasibility": {
        "fr": {0: "NSP", 1: "Difficile", 2: "Modérée", 3: "Facile"},
        "en": {0: "DNK", 1: "Difficult", 2: "Moderate", 3: "Easy"},
        "pt": {0: "NSP", 1: "Difícil", 2: "Moderada", 3: "Fácil"},
        "ar": {0: "لا أعلم", 1: "صعب", 2: "متوسط", 3: "سهل"},
    },
}

def score_format(lang: str, criterion: str):
    """Formatter for score selectboxes (criterion-aware).

    We include a None option (placeholder) so we don't prefill answers.
    """
    placeholder_fr = "— Sélectionner —"
    placeholder_en = "— Select —"
    scale = SCORE_SCALES.get(criterion, SCORE_SCALES["demand"])
    mapping = scale.get(lang) or scale[lang_base(lang)]

    def _fmt(v):
        if v is None or v == "":
            return t(lang, placeholder_fr, placeholder_en, "— Selecionar —", "— اختر —")
        try:
            iv = int(v)
        except Exception:
            return str(v)
        return mapping.get(iv, str(v))

    return _fmt


ROLE_OPTIONS_FR = [
    "DG/DGA/SG",
    "Directeur",
    "Conseiller",
    "Chef de division",
    "Chef de bureau",
    "Autre",
]
ROLE_OPTIONS_EN = [
    "DG/DGA/SG",
    "Director",
    "Advisor",
    "Head of division",
    "Head of office",
    "Other",
]
ROLE_OPTIONS_PT = [
    "DG/DGA/SG",
    "Diretor",
    "Conselheiro",
    "Chefe de divisão",
    "Chefe de gabinete",
    "Outro",
]
ROLE_OPTIONS_AR = [
    "DG/DGA/SG",
    "مدير",
    "مستشار",
    "رئيس قسم",
    "رئيس مصلحة",
    "أخرى",
]


# =========================
# Helpers : i18n and state
# =========================

LANG_OPTIONS = ["fr", "en", "pt", "ar"]
LANG_LABELS = {
    "fr": "Français",
    "en": "English",
    "pt": "Português (Portugal)",
    "ar": "العربية",
}

PT_TRANSLATIONS = {
    "Langue / Language": "Idioma",
    "— Select —": "— Selecionar —",
    "— Sélectionner —": "— Selecionar —",
    "Section 1: Instructions": "Secção 1: Instruções",
    "Section 2: Respondent identification": "Secção 2: Identificação do respondente",
    "Section 3: Scope of response": "Secção 3: Âmbito da resposta",
    "Section 4: Priority domains": "Secção 4: Domínios prioritários",
    "Section 5: Priority indicators and scoring": "Secção 5: Estatísticas prioritárias e pontuação",
    "Section 6: Gender dimension": "Secção 6: Dimensão de género",
    "Section 7: Gender priorities": "Secção 7: Prioridades de género",
    "Section 8: Capacity and feasibility (12–24 months)": "Secção 8: Capacidade e viabilidade (12–24 meses)",
    "Section 9: Harmonization and quality": "Secção 9: Harmonização e qualidade",
    "Section 10: Dissemination": "Secção 10: Disseminação",
    "Section 11: Relevant data sources": "Secção 11: Fontes de dados pertinentes",
    "Section 12: Open questions": "Secção 12: Questões abertas",
    "SUBMIT": "ENVIAR",
    "SUBMIT questionnaire": "ENVIAR questionário",
    "Organization Name": "Nome da organização",
    "Country of residence": "País de residência",
    "Other country (please specify)": "Outro país (especifique)",
    "Email": "Email",
    "Stakeholder type": "Tipo de ator",
    "Role/Function": "Função",
    "Specify (role)": "Especificar (função)",
    "Please fix the following:": "Por favor, corrija os pontos seguintes :",
    "National": "Nacional",
    "Regional (REC)": "Regional (CER)",
    "Continental (AU)": "Continental (UA)",
    "International": "Internacional",
    "Other": "Outro",
    "Yes": "Sim",
    "No": "Não",
    "In preparation": "Em preparação",
    "Under implementation AND new one in preparation": "Em implementação E nova estratégia em preparação",
    "Status of the current NSDS / national statistical plan": "Estado atual da ENDE / plano estatístico nacional",
    "Pre-selection (5–10 domains)": "Pré-seleção (5–10 domínios)",
    "Rank 1": "Classificação 1",
    "Rank 2": "Classificação 2",
    "Rank 3": "Classificação 3",
    "Rank 4": "Classificação 4",
    "Rank 5": "Classificação 5",
    "Step 1: Pre-selection": "Etapa 1: Pré-seleção",
    "Step 2: Rank TOP 5": "Etapa 2: Classificar o TOP 5",
    "Step A: Select indicators": "Etapa A: Selecionar indicadores",
    "Step B: Multi-criteria scoring": "Etapa B: Pontuação multicritério",
    "Select": "Selecionar",
    "Specify": "Especificar",
    "Specify (Other)": "Especificar (Outro)",
    "Relevant data sources": "Fontes de dados pertinentes",
    "2 to 4 most relevant data sources": "2 a 4 fontes de dados mais pertinentes",
    "Other: please specify": "Outro : especifique",
    "Question 1 / 3": "Pergunta 1 / 3",
    "Question 2 / 3": "Pergunta 2 / 3",
    "Question 3 / 3": "Pergunta 3 / 3",
    "Confirmation": "Confirmação",
    "⬅ Previous question": "⬅ Pergunta anterior",
    "Next question ➡": "Pergunta seguinte ➡",
    "Go to confirmation ➡": "Ir para a confirmação ➡",
    "OK (section completed)": "OK (secção concluída)",
    "⬅ Previous": "⬅ Anterior",
    "Next ➡": "Seguinte ➡",
    "Navigation": "Navegação",
    "Go to": "Ir para",
    "Draft": "Rascunho",
    "Save now": "Guardar agora",
    "ENVOYER": "ENVIAR",
    "✅ SUBMIT and save": "✅ ENVIAR e guardar",
    "Thank you! Your submission has been saved.": "Obrigado ! O seu questionário foi registado.",
    "Everything is ready. Click **SUBMIT** to send your questionnaire.": "Tudo está pronto. Clique em **ENVIAR** para submeter o questionário.",
    "Unified app (FR/EN) – hidden codes – built-in quality controls.": "Aplicação multilingue (FR/EN/PT/AR) – códigos ocultos – controlos de qualidade integrados.",

    "Please provide these details. They are used for analysis and will not be published in a personally identifiable way.": "Por favor, forneça estas informações. Serão utilizadas apenas para análise e não serão publicadas de forma nominativa.",
    "Please provide the full organization name (avoid acronym only).": "Indique a designação completa da organização (evite apenas a sigla).",
    "Select your country of residence (ISO3 list) or ‘Other country (please specify)’.": "Selecione o seu país de residência (lista ISO3) ou ‘Outro país (especifique)’.",
    "First select 5 to 10 domains (pre-selection). Then choose exactly 5 domains within that subset (TOP 5).\n\nTip: choose domains where policy demand is strong.": "Selecione primeiro 5 a 10 domínios (pré-seleção). Em seguida, escolha exatamente 5 domínios dentro desse subconjunto (TOP 5).\n\nConselho : escolha os domínios em que a procura política é forte.",
    "Domain list is not available (longlist not found or empty).": "A lista de domínios não está disponível (longlist não encontrada ou vazia).",
    "Check that the repository contains: data/indicator_longlist.csv (preferred) or data/longlist.xlsx (or these files at repo root).": "Verifique se o repositório contém : data/indicator_longlist.csv (preferencial) ou data/longlist.xlsx (ou estes ficheiros na raiz do repositório).",
    "Select **5 to 10 domains** (no duplicates).": "Selecione **5 a 10 domínios** (sem duplicados).",
    "Choose between 5 and 10 priority domains. The same domain cannot be selected twice.": "Escolha entre 5 e 10 domínios prioritários. O mesmo domínio não pode ser selecionado duas vezes.",
    "### Step 1: Pre-selection\nSelect **5 to 10 domains** (no duplicates).": "### Etapa 1 : Pré-seleção\nSelecione **5 a 10 domínios** (sem duplicados).",
    "### Step 2: Rank TOP 5\nRank exactly **5 domains** from your pre-selection.": "### Etapa 2 : Classificar o TOP 5\nClassifique exatamente **5 domínios** da sua pré-seleção.",
    "Please pre-select at least 5 domains first.": "Primeiro, faça a pré-seleção de pelo menos 5 domínios.",
    "This section is completed in two steps: (A) indicator selection, then (B) multi-criteria scoring.": "Esta secção é preenchida em duas etapas : (A) seleção dos indicadores, depois (B) pontuação multicritério.",
    "For each TOP 5 domain, select 1 to 3 indicators (expected total: 5 to 15). The same indicator must not appear under two domains.": "Para cada domínio do TOP 5, selecione 1 a 3 indicadores (total esperado : 5 a 15). O mesmo indicador não deve aparecer em dois domínios.",
    "Select 1 to 3 indicators": "Selecionar 1 a 3 indicadores",
    "Select at least 1 and at most 3 indicators for this domain.": "Selecione pelo menos 1 e no máximo 3 indicadores para este domínio.",
    "One or more indicators are selected under multiple domains. Please correct.": "Um ou mais indicadores estão selecionados em vários domínios. Corrija, por favor.",
    "Please complete TOP 5 domains first (Section 4).": "Finalize primeiro o TOP 5 dos domínios (Secção 4).",
    "Language": "Idioma",
    "Note: quality checks may prevent moving forward when constraints are not met.": "Nota : os controlos de qualidade podem impedir a progressão se alguma restrição não for respeitada.",
    "UK: Unknown (score 0). Use UK only when information is unavailable.": "NSP : Não sabe (pontuação 0). Utilize NSP apenas quando a informação não estiver disponível.",
    "Draft saved.": "Rascunho guardado.",
    "Draft not saved.": "Rascunho não guardado.",
    "Resume: keep this page URL (rid=... parameter).": "Retoma : conserve o URL desta página (parâmetro rid=...).",
    "Your entry is saved. If you pause for less than 48 hours, resume where you left off by reopening the link containing rid (bookmark / save / find it in your browser history).": "A sua resposta foi guardada. Se interromper por menos de 48 horas, retome onde ficou ao reabrir a ligação contendo rid (guardar nos favoritos / conservar / encontrar no histórico do navegador).",
    'Please provide these details. They are used for analysis and will not be published in a personally identifiable way.': 'Queira fornecer estas informações. São utilizadas apenas para análise e não serão publicadas de forma nominativa.',
    'Please provide the full organization name (avoid acronym only).': 'Indique o nome completo da organização (evite apenas a sigla).',
    'Select your country of residence (ISO3 list) or ‘Other country (please specify)’.': 'Selecione o seu país de residência (lista ISO3) ou ‘Outro país (especificar)’.',
    'Stakeholder type': 'Tipo de ator',
    'Choose the category that best matches your organization.': 'Escolha a categoria que melhor corresponde à sua organização.',
    'Role/Function': 'Função',
    'Indicate your main role in the organization.': 'Indique a sua função principal na organização.',
    'Please fix the following:': 'Corrija os elementos abaixo :',
    'National Statistical Office': 'Instituto Nacional de Estatística',
    'Ministry / Sector statistical unit': 'Ministério / Serviço estatístico setorial',
    'Regional Economic Community': 'Comunidade Económica Regional',
    'African Union (AU)': 'União Africana (UA)',
    'Civil society': 'Sociedade civil',
    'Development partner': 'Parceiro técnico e financeiro',
    'Academia / Research': 'Universidade / Investigação',
    'Other': 'Outro',
    'Director': 'Diretor',
    'Advisor': 'Conselheiro',
    'Head of division': 'Chefe de divisão',
    'Head of office': 'Chefe de gabinete',
    '### Purpose\nThis questionnaire collects your views on **priority socio-economic statistics** to be produced and disseminated at continental level.\n\n### How to answer\n1. **Identify** your organization (Section 2).\n2. **Pre-select 5–10 domains** and rank a **TOP 5** (Section 4).\n3. For each TOP 5 domain: select **1–3 indicators** and provide **scores** (Section 5).\n4. Complete cross-cutting sections: **gender**, **capacity/feasibility**,  **etc.**.\n5. **Feel free to consult the ⍰ tooltips for more details.**\n\n            ': '### Objetivo\nEste questionário recolhe a sua opinião sobre **as estatísticas socioeconómicas prioritárias** a produzir e divulgar ao nível continental.\n\n### Como responder\n1. **Identifique** a sua organização (Secção 2).\n2. **Pré-selecione 5 a 10 domínios** e classifique um **TOP 5** (Secção 4).\n3. Para cada domínio do TOP 5 : selecione **1 a 3 indicadores** e atribua **pontuações** (Secção 5).\n4. Complete as secções transversais : **género**, **capacidade/viabilidade**, **etc.**.\n5. **Consulte livremente as infobolhas ⍰ para obter mais pormenores.**\n\n            ',
    'Scope': 'Âmbito',
    'Indicate the main scope of your response. This helps interpret your priorities.': 'Indique o âmbito principal da sua resposta. Isto ajuda a interpretar as suas prioridades.',
    'National': 'Nacional',
    'Regional (REC)': 'Regional (CER)',
    'Continental (AU)': 'Continental (UA)',
    'International': 'Internacional',
    'Indicate the main scope of your response: national, regional (REC), continental (AU), or international.': 'Indique o âmbito principal da sua resposta : nacional, regional (CER), continental (UA) ou internacional.',
    'Status of the current NSDS / national statistical plan': 'Estado da ENDE / do plano estatístico nacional em vigor',
    'In preparation': 'Em preparação',
    'Under implementation AND new one in preparation': 'Em execução E uma nova em preparação',
    'Indicate whether an NSDS / national statistical plan is current, not in place, under preparation, or DK.': 'Indique se existe uma ENDE / plano estatístico nacional em vigor, inexistente, em preparação ou se não sabe.',
    'Warning: you selected “National Statistical Office” or “Ministry”, but the scope is not “National”. Please check consistency.': 'Aviso : selecionou ‘Instituto Nacional de Estatística’ ou ‘Ministério’, mas o âmbito não é ‘Nacional’. Verifique a coerência.',
    'Specify': 'Especificar',
    'First select 5 to 10 domains (pre-selection). Then choose exactly 5 domains within that subset (TOP 5).\n\nTip: choose domains where policy demand is strong.': 'Primeiro selecione 5 a 10 domínios (pré-seleção). Em seguida, escolha exatamente 5 domínios nesse subconjunto (TOP 5).\n\nConselho : escolha os domínios em que a procura política é mais forte.',
    'Domain list is not available (longlist not found or empty).': 'A lista de domínios não está disponível (longlist não encontrada ou vazia).',
    'Check that the repository contains: data/indicator_longlist.csv (preferred) or data/longlist.xlsx (or these files at repo root).': 'Verifique se o repositório contém : data/indicator_longlist.csv (preferencial) ou data/longlist.xlsx (ou esses ficheiros na raiz do repositório).',
    '### Step 1: Pre-selection\nSelect **5 to 10 domains** (no duplicates).\n            ': '### Etapa 1 : Pré-seleção\nSelecione **5 a 10 domínios** (sem duplicados).\n            ',
    'Pre-selection (5–10 domains)': 'Pré-seleção (5–10 domínios)',
    'Select up to 10 domains. Once 10 domains are selected, additional clicks are ignored.': 'Selecione no máximo 10 domínios. Depois de selecionar 10 domínios, os cliques adicionais serão ignorados.',
    '### Step 2: Rank TOP 5\nRank exactly **5 domains** from your pre-selection.\n            ': '### Etapa 2 : Classificar o TOP 5\nClassifique exatamente **5 domínios** da sua pré-seleção.\n            ',
    'Please pre-select at least 5 domains first.': 'Pré-selecione primeiro pelo menos 5 domínios.',
    'Rank {i+1}': 'Posição {i+1}',
    'Choose a unique domain for each rank. Already selected domains are removed from the next ranks.': 'Escolha um domínio único para cada posição. Os domínios já escolhidos deixam de ser propostos nas posições seguintes.',
    'Quality checks:': 'Controlos de qualidade :',
    'Please complete TOP 5 domains first (Section 4).': 'Complete primeiro o TOP 5 dos domínios (Secção 4).',
    'Step A: Select indicators': 'Etapa A : Seleção dos indicadores',
    'For each TOP 5 domain, select 1 to 3 indicators (expected total: 5 to 15). The same indicator must not appear under two domains.': 'Para cada domínio do TOP 5, selecione 1 a 3 indicadores (total esperado : 5 a 15). O mesmo indicador não deve aparecer em dois domínios.',
    'Select 1 to 3 indicators': 'Selecionar 1 a 3 indicadores',
    'Select at least 1 and at most 3 indicators for this domain.': 'Selecione no mínimo 1 e no máximo 3 indicadores para este domínio.',
    'One or more indicators are selected under multiple domains. Please correct.': 'Um ou mais indicadores foram selecionados em vários domínios. Corrija, por favor.',
    'Step B: Multi-criteria scoring': 'Etapa B : Pontuação multicritério',
    'For each selected indicator, assign a score (0–3) for: political demand, current availability (good = higher score), and feasibility in 12–24 months.': 'Para cada indicador selecionado, atribua uma pontuação (0–3) para : procura política, disponibilidade atual (boa = pontuação mais elevada) e viabilidade em 12–24 meses.',
    'Political demand': 'Procura política',
    'Definition: importance for steering public policies, accountability and priorities.': 'Definição : importância do indicador para a condução das políticas públicas, a prestação de contas e as prioridades.',
    'Current availability': 'Disponibilidade atual',
    'Definition: is the indicator already produced regularly with sufficient coverage and quality, in a usable form? (Good = higher score).': 'Definição : o indicador já é produzido regularmente, com cobertura e qualidade suficientes, e numa forma utilizável? (Boa = pontuação mais elevada).',
    'Feasibility (12–24 months)': 'Viabilidade (12–24 meses)',
    'Definition: realistic ability to improve or produce the indicator within 12–24 months, considering sources, capacities and prerequisites.': 'Definição : capacidade realista para melhorar ou produzir o indicador nos próximos 12–24 meses, tendo em conta as fontes, as capacidades e os pré-requisitos.',
    'demand': 'procura',
    'availability': 'disponibilidade',
    'feasibility': 'viabilidade',
    'Indicate whether priority indicators should integrate these dimensions (Yes/No/Indicator-specific/UK).': 'Indique se os indicadores prioritários devem integrar estas dimensões (Sim/Não/Consoante o indicador/NSP).',
    'Indicator-specific': 'Consoante o indicador',
    'Disaggregation by sex': 'Desagregação por sexo',
    'Disaggregation by age': 'Desagregação por idade',
    'Urban / rural': 'Urbano / rural',
    'Disability': 'Deficiência',
    'Wealth quintile': 'Quintil de riqueza',
    'Gender-based violence (GBV)': 'Violência baseada no género (VBG)',
    'Unpaid domestic work': 'Trabalho doméstico não remunerado',
    'Section 6: please complete the table.': 'Secção 6 : preencha a tabela.',
    'Section 6: missing answer for: {k}.': 'Secção 6 : falta resposta para : {k}.',
    'Select 1 to 3 gender priorities, starting with the most important.': 'Selecione de 1 a 3 prioridades de género, começando pela mais importante.',
    'Economic empowerment': 'Empoderamento económico',
    'Access to services': 'Acesso aos serviços',
    'Participation in decision-making bodies': 'Participação em instâncias de decisão',
    'Unpaid domestic and care work': 'Trabalho doméstico e de cuidados não remunerado',
    'Your three (3) gender priorities – Priority 1 (required)': 'As suas três (3) prioridades de género – Prioridade 1 (obrigatória)',
    'Priority 2 (optional)': 'Prioridade 2 (opcional)',
    'Priority 3 (optional)': 'Prioridade 3 (opcional)',
    'Section 7: please select at least one gender priority (Priority 1).': 'Secção 7 : selecione pelo menos uma prioridade de género (Prioridade 1).',
    'Section 7: please fill Priority 2 before Priority 3.': 'Secção 7 : preencha a Prioridade 2 antes da Prioridade 3.',
    'Section 7: gender priorities must be distinct (no duplicates).': 'Secção 7 : as prioridades de género devem ser distintas (sem duplicados).',
    "Section 7: please specify the 'Other' option.": 'Secção 7 : especifique a opção ‘Outro’.',
    'Assess the availability and adequacy of resources to produce priority statistics in the coming 12–24 months.': 'Avalie a disponibilidade e a adequação dos meios para produzir as estatísticas prioritárias nos próximos 12–24 meses.',
    'High': 'Elevado',
    'Medium': 'Médio',
    'Low': 'Baixo',
    'Scale: High = sufficient and operational; Medium = partially available; Low = insufficient; DK = does not know.': 'Escala : Elevado = capacidade suficiente e operacional ; Médio = parcialmente disponível ; Baixo = insuficiente ; NSP = não sabe.',
    'Available statistical skills': 'Competências estatísticas disponíveis',
    'Access to administrative data': 'Acesso aos dados administrativos',
    'Available funding': 'Financiamento disponível',
    'Digital tools (collection, processing, dissemination)': 'Ferramentas digitais (recolha, tratamento, disseminação)',
    'Legal framework for data sharing': 'Quadro jurídico para a partilha de dados',
    'Inter-institutional coordination': 'Coordenação interinstitucional',
    'Human resources: availability of qualified statisticians/analysts and relevant experience.': 'Recursos humanos : disponibilidade de estatísticos/analistas qualificados e experiência pertinente.',
    'Access to administrative data: availability, quality, timeliness and conditions of access for statistical use.': 'Acesso aos dados administrativos : disponibilidade, qualidade, atualidade e condições de acesso para utilização estatística.',
    'Funding: available and sustainable budget for production, including collection/processing operations.': 'Financiamento : orçamento disponível e sustentável para a produção, incluindo operações de recolha/tratamento.',
    'Digital tools: availability and adequacy of tools for collection, processing, storage, dissemination, interoperability (software, hardware, connectivity, security).': 'Ferramentas digitais : disponibilidade e adequação dos meios de recolha, tratamento, armazenamento, disseminação e interoperabilidade (software, hardware, conectividade, segurança).',
    'Legal framework: existence and enforceability of texts/agreements enabling data sharing for statistical purposes (laws, decrees, protocols, MoUs, confidentiality clauses).': 'Quadro jurídico : existência e aplicabilidade de textos/acordos que permitam a partilha de dados para fins estatísticos (leis, decretos, protocolos, memorandos, cláusulas de confidencialidade).',
    'Coordination: inter-institutional coordination mechanisms (committees, agreements, regular exchanges, shared standards).': 'Coordenação : mecanismos de coordenação interinstitucional (comités, acordos, intercâmbios regulares, normas comuns).',
    'Section 8: please complete the table.': 'Secção 8 : preencha a tabela.',
    'Section 8: missing answer for: {k}.': 'Secção 8 : falta resposta para : {k}.',
    'Indicate 1 to 3 expectations regarding harmonization and quality assurance.': 'Indique de 1 a 3 expectativas em matéria de harmonização e garantia da qualidade.',
    'Manuals on common standards and methods (by domain) available': 'Manuais de normas e métodos comuns (por domínio) disponíveis',
    'Functional quality assurance framework (quality toolkit) ': 'Quadro funcional de garantia da qualidade (ferramentas de qualidade)',
    'Data validation and certification procedures (certified quality) ': 'Procedimentos de validação e certificação de dados (qualidade certificada)',
    'Toolkit / mechanisms for cross-sector consistency of national data': 'Ferramentas / mecanismos para a coerência intersetorial dos dados nacionais',
    'Strengthening NSS technical capacity': 'Reforço da capacidade técnica do SEN',
    'Strengthening NSO leadership within the NSS': 'Reforço da liderança do INE no seio do SEN',
    'Specialized Technical Groups (STGs/AU) operational': 'Grupos Técnicos Especializados (GTE/UA) operacionais',
    'Other (specify) ': 'Outro (especificar)',
    'Section 9: please select at least one option.': 'Secção 9 : selecione pelo menos uma opção.',
    'Section 9: please specify the "Other" option.': 'Secção 9 : especifique a opção ‘Outro’.',
    'Indicate 1 to 3 dissemination channels you find most useful for priority statistics.': 'Indique de 1 a 3 canais de disseminação que considere mais úteis para as estatísticas prioritárias.',
    'Web portal / dashboards': 'Portal web / painéis de bordo',
    'Press releases / bulletins': 'Comunicados / boletins',
    'Anonymized microdata (secure access)': 'Microdados anonimizados (acesso seguro)',
    'Workshops and webinars': 'Ateliês e webinários',
    'Select the most useful dissemination channels.': 'Escolha os canais de disseminação mais úteis.',
    'Section 10: please select at least one option.': 'Secção 10 : selecione pelo menos uma opção.',
    'Section 10: please specify the "Other" option.': 'Secção 10 : especifique a opção ‘Outro’.',
    'Select **2 to 4** of the most important data sources to produce the priority statistics.': 'Selecione **2 a 4** das fontes de dados mais importantes para produzir as estatísticas prioritárias.',
    'Household surveys': 'Inquéritos aos agregados familiares',
    'Enterprise surveys': 'Inquéritos às empresas',
    'Censuses': 'Recenseamentos',
    'Administrative data': 'Dados administrativos',
    'Civil registration and vital statistics (CRVS)': 'Registo civil e estatísticas vitais (CRVS)',
    'Geospatial data': 'Dados geoespaciais',
    'Private data': 'Dados privados',
    'Choose between 2 and 4 options. If you select Other, please specify.': 'Escolha entre 2 e 4 opções. Se selecionar Outra, especifique.',
    'Section 11: please select at least 2 sources.': 'Secção 11 : selecione pelo menos 2 fontes.',
    'Section 11: please select at most 4 sources.': 'Secção 11 : selecione no máximo 4 fontes.',
    "Section 11: please specify the 'Other' option.": 'Secção 11 : especifique a opção ‘Outras’.',
    'These questions are **optional**. You may leave them blank. They are presented **one by one** to facilitate completion.': 'Estas perguntas são **opcionais**. Pode deixá-las em branco. São apresentadas **uma a uma** para facilitar o preenchimento.',
    'Warning: question 1 is empty (you can still proceed).': 'Aviso : a pergunta 1 está vazia (pode ainda assim continuar).',
    'Warning: question 2 is empty (you can still proceed).': 'Aviso : a pergunta 2 está vazia (pode ainda assim continuar).',
    'Warning: question 3 is empty (you can still proceed).': 'Aviso : a pergunta 3 está vazia (pode ainda assim continuar).',
    'Final step: please indicate whether you consulted other colleagues. This question is mandatory.': 'Etapa final : indique se consultou outros colegas para preencher este questionário. Esta pergunta é obrigatória.',
    'Did you consult other colleagues to complete this questionnaire?': 'Consultou outros colegas para preencher este questionário ?',
    'Progress: 1/3 → 2/3 → 3/3 → Confirmation.': 'Progressão : 1/3 → 2/3 → 3/3 → Confirmação.',
    'Section 12: please go through the open questions one by one (use the “Next question” button) until Confirmation.': 'Secção 12 : percorra as questões abertas uma a uma (botão ‘Pergunta seguinte’) até à Confirmação.',
    'Section 12: please indicate whether you consulted other colleagues (Yes/No).': 'Secção 12 : indique se consultou outros colegas (Sim/Não).',
    'Question 1 / 3': 'Pergunta 1 / 3',
    'Question 2 / 3': 'Pergunta 2 / 3',
    'Question 3 / 3': 'Pergunta 3 / 3',
    'Confirmation': 'Confirmação',

}

AR_TRANSLATIONS = {
    "Langue / Language": "اللغة",
    "— Select —": "— اختر —",
    "— Sélectionner —": "— اختر —",
    "Section 1: Instructions": "القسم 1: التعليمات",
    "Section 2: Respondent identification": "القسم 2: تعريف المجيب",
    "Section 3: Scope of response": "القسم 3: نطاق الإجابة",
    "Section 4: Priority domains": "القسم 4: المجالات ذات الأولوية",
    "Section 5: Priority indicators and scoring": "القسم 5: الإحصاءات ذات الأولوية والتقييم",
    "Section 6: Gender dimension": "القسم 6: البعد الجنساني",
    "Section 7: Gender priorities": "القسم 7: أولويات النوع الاجتماعي",
    "Section 8: Capacity and feasibility (12–24 months)": "القسم 8: القدرات وإمكانية التنفيذ (12–24 شهراً)",
    "Section 9: Harmonization and quality": "القسم 9: المواءمة والجودة",
    "Section 10: Dissemination": "القسم 10: النشر",
    "Section 11: Relevant data sources": "القسم 11: مصادر البيانات ذات الصلة",
    "Section 12: Open questions": "القسم 12: الأسئلة المفتوحة",
    "SUBMIT": "إرسال",
    "SUBMIT questionnaire": "إرسال الاستبيان",
    "Organization Name": "اسم المؤسسة",
    "Country of residence": "بلد الإقامة",
    "Other country (please specify)": "بلد آخر (يرجى التحديد)",
    "Email": "البريد الإلكتروني",
    "Stakeholder type": "نوع الجهة الفاعلة",
    "Role/Function": "الصفة / الوظيفة",
    "Specify (role)": "حدد (الوظيفة)",
    "Please fix the following:": "يرجى تصحيح ما يلي :",
    "National": "وطني",
    "Regional (REC)": "إقليمي (التجمع الاقتصادي الإقليمي)",
    "Continental (AU)": "قاري (الاتحاد الأفريقي)",
    "International": "دولي",
    "Other": "أخرى",
    "Yes": "نعم",
    "No": "لا",
    "In preparation": "قيد الإعداد",
    "Under implementation AND new one in preparation": "قيد التنفيذ مع إعداد استراتيجية جديدة",
    "Status of the current NSDS / national statistical plan": "وضع الاستراتيجية الوطنية الحالية لتطوير الإحصاء / الخطة الإحصائية الوطنية",
    "Pre-selection (5–10 domains)": "الاختيار الأولي (5 إلى 10 مجالات)",
    "Rank 1": "الترتيب 1",
    "Rank 2": "الترتيب 2",
    "Rank 3": "الترتيب 3",
    "Rank 4": "الترتيب 4",
    "Rank 5": "الترتيب 5",
    "Step 1: Pre-selection": "المرحلة 1: الاختيار الأولي",
    "Step 2: Rank TOP 5": "المرحلة 2: ترتيب أفضل 5 مجالات",
    "Step A: Select indicators": "المرحلة أ: اختيار المؤشرات",
    "Step B: Multi-criteria scoring": "المرحلة ب: التقييم متعدد المعايير",
    "Select": "اختر",
    "Specify": "حدد",
    "Specify (Other)": "حدد (أخرى)",
    "Relevant data sources": "مصادر البيانات ذات الصلة",
    "2 to 4 most relevant data sources": "من 2 إلى 4 من أكثر مصادر البيانات صلة",
    "Other: please specify": "أخرى : يرجى التحديد",
    "Question 1 / 3": "السؤال 1 / 3",
    "Question 2 / 3": "السؤال 2 / 3",
    "Question 3 / 3": "السؤال 3 / 3",
    "Confirmation": "التأكيد",
    "⬅ Previous question": "⬅ السؤال السابق",
    "Next question ➡": "السؤال التالي ➡",
    "Go to confirmation ➡": "الانتقال إلى التأكيد ➡",
    "OK (section completed)": "حسنًا (اكتمل القسم)",
    "⬅ Previous": "⬅ السابق",
    "Next ➡": "التالي ➡",
    "Navigation": "التنقل",
    "Go to": "اذهب إلى",
    "Draft": "مسودة",
    "Save now": "احفظ الآن",
    "✅ SUBMIT and save": "✅ إرسال وحفظ",
    "Thank you! Your submission has been saved.": "شكرًا ! تم حفظ استجابتكم.",
    "Everything is ready. Click **SUBMIT** to send your questionnaire.": "كل شيء جاهز. انقر على **إرسال** لإرسال الاستبيان.",
    "Unified app (FR/EN) – hidden codes – built-in quality controls.": "تطبيق متعدد اللغات (FR/EN/PT/AR) – رموز مخفية – ضوابط جودة مدمجة.",
    "Household surveys": "المسوحات الأسرية",
    "Enterprise surveys": "مسوحات المنشآت",
    "Censuses": "التعدادات",
    "Administrative data": "البيانات الإدارية",
    "Civil registration and vital statistics (CRVS)": "التسجيل المدني والإحصاءات الحيوية (CRVS)",
    "Geospatial data": "البيانات الجغرافية المكانية",
    "Private data": "البيانات الخاصة",
    "Web portal / dashboards": "بوابة إلكترونية / لوحات معلومات",
    "Press releases / bulletins": "بيانات صحفية / نشرات",
    "Anonymized microdata (secure access)": "بيانات جزئية مجهولة الهوية (ولوج آمن)",
    "Workshops and webinars": "ورشات عمل وندوات عبر الإنترنت",
    "Manuals on common standards and methods (by domain) available": "توفر أدلة للمعايير والمناهج المشتركة (حسب المجال)",
    "Functional quality assurance framework (quality toolkit) ": "إطار وظيفي لضمان الجودة",
    "Data validation and certification procedures (certified quality) ": "إجراءات التحقق من البيانات واعتمادها",
    "Toolkit / mechanisms for cross-sector consistency of national data": "آليات لضمان الاتساق بين القطاعات في البيانات الوطنية",
    "Strengthening NSS technical capacity": "تعزيز القدرات التقنية للنظام الإحصائي الوطني",
    "Strengthening NSO leadership within the NSS": "تعزيز قيادة المكتب الوطني للإحصاء داخل النظام الإحصائي الوطني",
    "Specialized Technical Groups (STGs/AU) operational": "تشغيل المجموعات التقنية المتخصصة (الاتحاد الأفريقي)",
    "Available statistical skills": "المهارات الإحصائية المتاحة",
    "Access to administrative data": "إمكانية الوصول إلى البيانات الإدارية",
    "Available funding": "التمويل المتاح",
    "Digital tools (collection, processing, dissemination)": "الأدوات الرقمية (الجمع والمعالجة والنشر)",
    "Legal framework for data sharing": "الإطار القانوني لتبادل البيانات",
    "Inter-institutional coordination": "التنسيق بين المؤسسات",
    "Disaggregation by sex": "التصنيف حسب الجنس",
    "Disaggregation by age": "التصنيف حسب العمر",
    "Urban / rural": "حضري / ريفي",
    "Disability": "الإعاقة",
    "Wealth quintile": "خمس الثروة",
    "Gender-based violence (GBV)": "العنف القائم على النوع الاجتماعي",
    "Unpaid domestic work": "العمل المنزلي غير المأجور",

    "Please provide these details. They are used for analysis and will not be published in a personally identifiable way.": "يرجى تقديم هذه المعلومات. ستُستخدم فقط لأغراض التحليل ولن تُنشر بشكل يحدد هوية المجيب.",
    "Please provide the full organization name (avoid acronym only).": "يرجى إدخال الاسم الكامل للمؤسسة (وتجنب الاكتفاء بالاختصار).",
    "Select your country of residence (ISO3 list) or ‘Other country (please specify)’.": "اختر بلد إقامتك (قائمة ISO3) أو ‘بلد آخر (يرجى التحديد)’.",
    "First select 5 to 10 domains (pre-selection). Then choose exactly 5 domains within that subset (TOP 5).\n\nTip: choose domains where policy demand is strong.": "اختر أولاً من 5 إلى 10 مجالات (اختيار أولي). ثم اختر بالضبط 5 مجالات داخل هذه المجموعة الفرعية (أفضل 5).\n\nنصيحة : اختر المجالات التي يكون فيها الطلب السياسي مرتفعاً.",
    "Domain list is not available (longlist not found or empty).": "قائمة المجالات غير متاحة (ملف longlist غير موجود أو فارغ).",
    "Check that the repository contains: data/indicator_longlist.csv (preferred) or data/longlist.xlsx (or these files at repo root).": "تحقق من أن المستودع يحتوي على : data/indicator_longlist.csv (مفضل) أو data/longlist.xlsx (أو هذه الملفات في جذر المستودع).",
    "Select **5 to 10 domains** (no duplicates).": "اختر **من 5 إلى 10 مجالات** (من دون تكرار).",
    "Choose between 5 and 10 priority domains. The same domain cannot be selected twice.": "اختر بين 5 و10 مجالات ذات أولوية. لا يمكن اختيار المجال نفسه مرتين.",
    "### Step 1: Pre-selection\nSelect **5 to 10 domains** (no duplicates).": "### المرحلة 1 : الاختيار الأولي\nاختر **من 5 إلى 10 مجالات** (من دون تكرار).",
    "### Step 2: Rank TOP 5\nRank exactly **5 domains** from your pre-selection.": "### المرحلة 2 : ترتيب أفضل 5\nرتب بالضبط **5 مجالات** من اختيارك الأولي.",
    "Please pre-select at least 5 domains first.": "يرجى أولاً إجراء اختيار أولي لا يقل عن 5 مجالات.",
    "This section is completed in two steps: (A) indicator selection, then (B) multi-criteria scoring.": "يُستكمل هذا القسم على مرحلتين : (أ) اختيار المؤشرات، ثم (ب) التقييم متعدد المعايير.",
    "For each TOP 5 domain, select 1 to 3 indicators (expected total: 5 to 15). The same indicator must not appear under two domains.": "لكل مجال من مجالات أفضل 5، اختر من 1 إلى 3 مؤشرات (المجموع المتوقع : من 5 إلى 15). ولا يجب أن يظهر المؤشر نفسه في مجالين.",
    "Select 1 to 3 indicators": "اختر من 1 إلى 3 مؤشرات",
    "Select at least 1 and at most 3 indicators for this domain.": "اختر مؤشراً واحداً على الأقل وثلاثة مؤشرات كحد أقصى لهذا المجال.",
    "One or more indicators are selected under multiple domains. Please correct.": "تم اختيار مؤشر واحد أو أكثر ضمن عدة مجالات. يرجى التصحيح.",
    "Please complete TOP 5 domains first (Section 4).": "يرجى أولاً استكمال أفضل 5 مجالات (القسم 4).",
    "Language": "اللغة",
    "Note: quality checks may prevent moving forward when constraints are not met.": "ملاحظة : قد تمنع ضوابط الجودة التقدم إذا لم يتم احترام بعض القيود.",
    "UK: Unknown (score 0). Use UK only when information is unavailable.": "لا أعلم : استخدم هذه الإجابة فقط عندما تكون المعلومة غير متاحة (الدرجة 0).",
    "Draft saved.": "تم حفظ المسودة.",
    "Draft not saved.": "لم يتم حفظ المسودة.",
    "Resume: keep this page URL (rid=... parameter).": "للاستئناف : احتفظ برابط هذه الصفحة (المعامل rid=...).",
    "Your entry is saved. If you pause for less than 48 hours, resume where you left off by reopening the link containing rid (bookmark / save / find it in your browser history).": "تم حفظ إدخالك. إذا توقفت لمدة تقل عن 48 ساعة، يمكنك المتابعة من حيث توقفت عبر إعادة فتح الرابط الذي يحتوي على rid (احفظه في المفضلة أو ابحث عنه في سجل المتصفح).",
    'Please provide these details. They are used for analysis and will not be published in a personally identifiable way.': 'يُرجى تقديم هذه المعلومات. فهي تُستخدم لأغراض التحليل فقط ولن تُنشر على نحو يتيح التعرف على الأشخاص.',
    'Please provide the full organization name (avoid acronym only).': 'يرجى تقديم الاسم الكامل للمؤسسة (وتجنب الاكتفاء بالاختصار).',
    'Select your country of residence (ISO3 list) or ‘Other country (please specify)’.': 'اختر بلد إقامتك (قائمة ISO3) أو ‘بلد آخر (يرجى التحديد)’.',
    'Stakeholder type': 'نوع الجهة الفاعلة',
    'Choose the category that best matches your organization.': 'اختر الفئة التي تتوافق أكثر مع مؤسستك.',
    'Role/Function': 'المنصب / الوظيفة',
    'Indicate your main role in the organization.': 'حدد وظيفتك الرئيسية داخل المؤسسة.',
    'Please fix the following:': 'يرجى تصحيح العناصر التالية :',
    'National Statistical Office': 'المعهد الوطني للإحصاء',
    'Ministry / Sector statistical unit': 'وزارة / وحدة إحصائية قطاعية',
    'Regional Economic Community': 'جماعة اقتصادية إقليمية',
    'African Union (AU)': 'الاتحاد الأفريقي',
    'Civil society': 'المجتمع المدني',
    'Development partner': 'شريك تقني ومالي',
    'Academia / Research': 'الجامعة / البحث',
    'Other': 'أخرى',
    'Director': 'مدير',
    'Advisor': 'مستشار',
    'Head of division': 'رئيس قسم',
    'Head of office': 'رئيس مصلحة',
    '### Purpose\nThis questionnaire collects your views on **priority socio-economic statistics** to be produced and disseminated at continental level.\n\n### How to answer\n1. **Identify** your organization (Section 2).\n2. **Pre-select 5–10 domains** and rank a **TOP 5** (Section 4).\n3. For each TOP 5 domain: select **1–3 indicators** and provide **scores** (Section 5).\n4. Complete cross-cutting sections: **gender**, **capacity/feasibility**,  **etc.**.\n5. **Feel free to consult the ⍰ tooltips for more details.**\n\n            ': '### الهدف\nيجمع هذا الاستبيان آراءكم بشأن **الإحصاءات الاجتماعية والاقتصادية ذات الأولوية** التي ينبغي إنتاجها ونشرها على المستوى القاري.\n\n### كيفية الإجابة\n1. **حدّدوا** مؤسستكم (القسم 2).\n2. **اختاروا مسبقاً من 5 إلى 10 مجالات** ثم رتّبوا **أفضل 5** (القسم 4).\n3. لكل مجال من المجالات الخمسة الأولى : اختاروا **1 إلى 3 مؤشرات** وقدّموا **درجات** التقييم (القسم 5).\n4. استكملوا الأقسام الأفقية : **النوع الاجتماعي** و**القدرة/إمكانية التنفيذ** و**غير ذلك**.\n5. **يمكنكم الرجوع إلى تلميحات ⍰ للحصول على مزيد من التفاصيل.**\n\n            ',
    'Scope': 'النطاق',
    'Indicate the main scope of your response. This helps interpret your priorities.': 'حدّدوا النطاق الرئيسي لإجابتكم. فهذا يساعد على تفسير أولوياتكم.',
    'National': 'وطني',
    'Regional (REC)': 'إقليمي (جماعة اقتصادية إقليمية)',
    'Continental (AU)': 'قاري (الاتحاد الأفريقي)',
    'International': 'دولي',
    'Indicate the main scope of your response: national, regional (REC), continental (AU), or international.': 'حدّدوا النطاق الرئيسي لإجابتكم : وطني أو إقليمي أو قاري أو دولي.',
    'Status of the current NSDS / national statistical plan': 'وضع الاستراتيجية الوطنية لتنمية الإحصاء / الخطة الإحصائية الوطنية الحالية',
    'In preparation': 'قيد الإعداد',
    'Under implementation AND new one in preparation': 'قيد التنفيذ وتوجد أخرى جديدة قيد الإعداد',
    'Indicate whether an NSDS / national statistical plan is current, not in place, under preparation, or DK.': 'بيّنوا ما إذا كانت هناك استراتيجية / خطة إحصائية وطنية سارية أو غير موجودة أو قيد الإعداد أو إن كنتم لا تعرفون.',
    'Warning: you selected “National Statistical Office” or “Ministry”, but the scope is not “National”. Please check consistency.': 'تنبيه : لقد اخترتم ‘المعهد الوطني للإحصاء’ أو ‘وزارة’، لكن النطاق ليس ‘وطنياً’. يُرجى التحقق من الاتساق.',
    'Specify': 'يرجى التحديد',
    'First select 5 to 10 domains (pre-selection). Then choose exactly 5 domains within that subset (TOP 5).\n\nTip: choose domains where policy demand is strong.': 'اختاروا أولاً من 5 إلى 10 مجالات (اختيار مسبق). ثم اختاروا بالضبط 5 مجالات ضمن هذه المجموعة الفرعية (أفضل 5).\n\nنصيحة : اختاروا المجالات التي يكون فيها الطلب السياسي قوياً.',
    'Domain list is not available (longlist not found or empty).': 'قائمة المجالات غير متاحة (القائمة الطويلة غير موجودة أو فارغة).',
    'Check that the repository contains: data/indicator_longlist.csv (preferred) or data/longlist.xlsx (or these files at repo root).': 'تحققوا من أن المستودع يحتوي على : data/indicator_longlist.csv (المفضل) أو data/longlist.xlsx (أو هذه الملفات في جذر المستودع).',
    '### Step 1: Pre-selection\nSelect **5 to 10 domains** (no duplicates).\n            ': '### الخطوة 1 : الاختيار المسبق\nاختاروا **من 5 إلى 10 مجالات** (من دون تكرار).\n            ',
    'Pre-selection (5–10 domains)': 'الاختيار المسبق (من 5 إلى 10 مجالات)',
    'Select up to 10 domains. Once 10 domains are selected, additional clicks are ignored.': 'اختاروا بحد أقصى 10 مجالات. وبعد اختيار 10 مجالات، سيتم تجاهل النقرات الإضافية.',
    '### Step 2: Rank TOP 5\nRank exactly **5 domains** from your pre-selection.\n            ': '### الخطوة 2 : ترتيب أفضل 5\nرتّبوا بالضبط **5 مجالات** من ضمن اختياركم المسبق.\n            ',
    'Please pre-select at least 5 domains first.': 'يرجى أولاً اختيار ما لا يقل عن 5 مجالات.',
    'Choose a unique domain for each rank. Already selected domains are removed from the next ranks.': 'اختاروا مجالاً فريداً لكل رتبة. وتُزال المجالات المختارة سابقاً من الرتب اللاحقة.',
    'Quality checks:': 'فحوصات الجودة :',
    'Please complete TOP 5 domains first (Section 4).': 'يرجى أولاً استكمال أفضل 5 مجالات (القسم 4).',
    'Step A: Select indicators': 'المرحلة أ : اختيار المؤشرات',
    'For each TOP 5 domain, select 1 to 3 indicators (expected total: 5 to 15). The same indicator must not appear under two domains.': 'لكل مجال من المجالات الخمسة الأولى، اختاروا من 1 إلى 3 مؤشرات (المجموع المتوقع : من 5 إلى 15). ولا ينبغي أن يظهر المؤشر نفسه تحت مجالين.',
    'Select 1 to 3 indicators': 'اختيار 1 إلى 3 مؤشرات',
    'Select at least 1 and at most 3 indicators for this domain.': 'اختاروا مؤشراً واحداً على الأقل وثلاثة مؤشرات كحد أقصى لهذا المجال.',
    'One or more indicators are selected under multiple domains. Please correct.': 'تم اختيار مؤشر واحد أو أكثر ضمن عدة مجالات. يُرجى التصحيح.',
    'Step B: Multi-criteria scoring': 'المرحلة ب : التقييم متعدد المعايير',
    'For each selected indicator, assign a score (0–3) for: political demand, current availability (good = higher score), and feasibility in 12–24 months.': 'لكل مؤشر مختار، امنحوا درجة (0–3) لكل من : الطلب السياسي، والتوافر الحالي (الجيد = درجة أعلى)، وإمكانية التنفيذ خلال 12–24 شهراً.',
    'Political demand': 'الطلب السياسي',
    'Definition: importance for steering public policies, accountability and priorities.': 'التعريف : مدى أهمية المؤشر في توجيه السياسات العامة والمساءلة والأولويات.',
    'Current availability': 'التوافر الحالي',
    'Definition: is the indicator already produced regularly with sufficient coverage and quality, in a usable form? (Good = higher score).': 'التعريف : هل يُنتج المؤشر بالفعل بانتظام، مع تغطية وجودة كافيتين، وفي شكل قابل للاستخدام؟ (الجيد = درجة أعلى).',
    'Feasibility (12–24 months)': 'إمكانية التنفيذ (12–24 شهراً)',
    'Definition: realistic ability to improve or produce the indicator within 12–24 months, considering sources, capacities and prerequisites.': 'التعريف : القدرة الواقعية على تحسين المؤشر أو إنتاجه خلال 12–24 شهراً، مع مراعاة المصادر والقدرات والمتطلبات المسبقة.',
    'demand': 'الطلب',
    'availability': 'التوافر',
    'feasibility': 'إمكانية التنفيذ',
    'Indicate whether priority indicators should integrate these dimensions (Yes/No/Indicator-specific/UK).': 'بيّنوا ما إذا كان ينبغي أن تدمج المؤشرات ذات الأولوية هذه الأبعاد (نعم/لا/حسب المؤشر/لا أعرف).',
    'Indicator-specific': 'حسب المؤشر',
    'Disaggregation by sex': 'التفصيل حسب الجنس',
    'Disaggregation by age': 'التفصيل حسب العمر',
    'Urban / rural': 'حضري / ريفي',
    'Disability': 'الإعاقة',
    'Wealth quintile': 'شريحة الثروة الخمسية',
    'Gender-based violence (GBV)': 'العنف القائم على النوع الاجتماعي',
    'Unpaid domestic work': 'العمل المنزلي غير المأجور',
    'Section 6: please complete the table.': 'القسم 6 : يُرجى استكمال الجدول.',
    'Section 6: missing answer for: {k}.': 'القسم 6 : توجد إجابة ناقصة لـ : {k}.',
    'Select 1 to 3 gender priorities, starting with the most important.': 'اختاروا من 1 إلى 3 أولويات متعلقة بالنوع الاجتماعي، بدءاً بالأهم.',
    'Economic empowerment': 'التمكين الاقتصادي',
    'Access to services': 'الوصول إلى الخدمات',
    'Participation in decision-making bodies': 'المشاركة في هيئات صنع القرار',
    'Unpaid domestic and care work': 'أعمال الرعاية والعمل المنزلي غير المأجور',
    'Your three (3) gender priorities – Priority 1 (required)': 'أولوياتكم الثلاث (3) في مجال النوع الاجتماعي – الأولوية 1 (إلزامية)',
    'Priority 2 (optional)': 'الأولوية 2 (اختيارية)',
    'Priority 3 (optional)': 'الأولوية 3 (اختيارية)',
    'Section 7: please select at least one gender priority (Priority 1).': 'القسم 7 : يُرجى اختيار أولوية واحدة على الأقل في مجال النوع الاجتماعي (الأولوية 1).',
    'Section 7: please fill Priority 2 before Priority 3.': 'القسم 7 : يُرجى ملء الأولوية 2 قبل الأولوية 3.',
    'Section 7: gender priorities must be distinct (no duplicates).': 'القسم 7 : يجب أن تكون أولويات النوع الاجتماعي مختلفة (من دون تكرار).',
    "Section 7: please specify the 'Other' option.": 'القسم 7 : يُرجى تحديد خيار ‘أخرى’.',
    'Assess the availability and adequacy of resources to produce priority statistics in the coming 12–24 months.': 'قيّموا مدى توافر وكفاية الموارد لإنتاج الإحصاءات ذات الأولوية خلال الأشهر 12–24 المقبلة.',
    'High': 'مرتفع',
    'Medium': 'متوسط',
    'Low': 'منخفض',
    'Scale: High = sufficient and operational; Medium = partially available; Low = insufficient; DK = does not know.': 'المقياس : مرتفع = قدرة كافية وقابلة للتشغيل ؛ متوسط = متوافر جزئياً ؛ منخفض = غير كافٍ ؛ لا أعرف = لا يعلم.',
    'Available statistical skills': 'المهارات الإحصائية المتاحة',
    'Access to administrative data': 'الوصول إلى البيانات الإدارية',
    'Available funding': 'التمويل المتاح',
    'Digital tools (collection, processing, dissemination)': 'الأدوات الرقمية (الجمع، المعالجة، النشر)',
    'Legal framework for data sharing': 'الإطار القانوني لتبادل البيانات',
    'Inter-institutional coordination': 'التنسيق بين المؤسسات',
    'Human resources: availability of qualified statisticians/analysts and relevant experience.': 'الموارد البشرية : توافر إحصائيين/محللين مؤهلين وخبرة ذات صلة.',
    'Access to administrative data: availability, quality, timeliness and conditions of access for statistical use.': 'الوصول إلى البيانات الإدارية : التوافر والجودة والآنية وشروط الوصول للاستخدام الإحصائي.',
    'Funding: available and sustainable budget for production, including collection/processing operations.': 'التمويل : ميزانية متاحة ومستدامة للإنتاج، بما في ذلك عمليات الجمع والمعالجة.',
    'Digital tools: availability and adequacy of tools for collection, processing, storage, dissemination, interoperability (software, hardware, connectivity, security).': 'الأدوات الرقمية : توافر وملاءمة الأدوات الخاصة بالجمع والمعالجة والتخزين والنشر وقابلية التشغيل البيني (البرمجيات، العتاد، الاتصال، الأمن).',
    'Legal framework: existence and enforceability of texts/agreements enabling data sharing for statistical purposes (laws, decrees, protocols, MoUs, confidentiality clauses).': 'الإطار القانوني : وجود النصوص/الاتفاقات وقابليتها للتنفيذ بما يتيح تبادل البيانات للأغراض الإحصائية (قوانين، مراسيم، بروتوكولات، مذكرات تفاهم، بنود السرية).',
    'Coordination: inter-institutional coordination mechanisms (committees, agreements, regular exchanges, shared standards).': 'التنسيق : آليات التنسيق بين المؤسسات (لجان، اتفاقات، تبادلات منتظمة، معايير مشتركة).',
    'Section 8: please complete the table.': 'القسم 8 : يُرجى استكمال الجدول.',
    'Section 8: missing answer for: {k}.': 'القسم 8 : توجد إجابة ناقصة لـ : {k}.',
    'Indicate 1 to 3 expectations regarding harmonization and quality assurance.': 'حدّدوا من 1 إلى 3 توقعات تتعلق بالمواءمة وضمان الجودة.',
    'Manuals on common standards and methods (by domain) available': 'توافر أدلة للمعايير والمناهج المشتركة (حسب المجال)',
    'Functional quality assurance framework (quality toolkit) ': 'إطار وظيفي لضمان الجودة (أدوات الجودة)',
    'Data validation and certification procedures (certified quality) ': 'إجراءات التحقق من صحة البيانات واعتمادها (جودة معتمدة)',
    'Toolkit / mechanisms for cross-sector consistency of national data': 'أدوات / آليات لتحقيق اتساق البيانات الوطنية بين القطاعات',
    'Strengthening NSS technical capacity': 'تعزيز القدرات التقنية للنظام الإحصائي الوطني',
    'Strengthening NSO leadership within the NSS': 'تعزيز قيادة المعهد الوطني للإحصاء داخل النظام الإحصائي الوطني',
    'Specialized Technical Groups (STGs/AU) operational': 'مجموعات فنية متخصصة (الاتحاد الأفريقي) عاملة',
    'Other (specify) ': 'أخرى (يرجى التحديد)',
    'Section 9: please select at least one option.': 'القسم 9 : يُرجى اختيار خيار واحد على الأقل.',
    'Section 9: please specify the "Other" option.': 'القسم 9 : يُرجى تحديد خيار ‘أخرى’.',
    'Indicate 1 to 3 dissemination channels you find most useful for priority statistics.': 'حدّدوا من 1 إلى 3 قنوات للنشر ترونها الأكثر فائدة بالنسبة للإحصاءات ذات الأولوية.',
    'Web portal / dashboards': 'بوابة ويب / لوحات متابعة',
    'Press releases / bulletins': 'بلاغات / نشرات',
    'Anonymized microdata (secure access)': 'بيانات جزئية مجهولة الهوية (ولوج آمن)',
    'Workshops and webinars': 'ورشات عمل وندوات عبر الإنترنت',
    'Select the most useful dissemination channels.': 'اختاروا قنوات النشر الأكثر فائدة.',
    'Section 10: please select at least one option.': 'القسم 10 : يُرجى اختيار خيار واحد على الأقل.',
    'Section 10: please specify the "Other" option.': 'القسم 10 : يُرجى تحديد خيار ‘أخرى’.',
    'Select **2 to 4** of the most important data sources to produce the priority statistics.': 'اختاروا **من 2 إلى 4** من أهم مصادر البيانات لإنتاج الإحصاءات ذات الأولوية.',
    'Household surveys': 'المسوحات الأسرية',
    'Enterprise surveys': 'مسوحات المؤسسات',
    'Censuses': 'التعدادات',
    'Administrative data': 'البيانات الإدارية',
    'Civil registration and vital statistics (CRVS)': 'التسجيل المدني وإحصاءات الأحوال المدنية والوقائع الحيوية',
    'Geospatial data': 'البيانات الجغرافية المكانية',
    'Private data': 'بيانات خاصة',
    'Choose between 2 and 4 options. If you select Other, please specify.': 'اختاروا بين 2 و4 خيارات. وإذا اخترتم ‘أخرى’، فيُرجى التحديد.',
    'Section 11: please select at least 2 sources.': 'القسم 11 : يُرجى اختيار مصدرين على الأقل.',
    'Section 11: please select at most 4 sources.': 'القسم 11 : يُرجى اختيار أربعة مصادر كحد أقصى.',
    "Section 11: please specify the 'Other' option.": 'القسم 11 : يُرجى تحديد خيار ‘أخرى’.',
    'These questions are **optional**. You may leave them blank. They are presented **one by one** to facilitate completion.': 'هذه الأسئلة **اختيارية**. ويمكنكم تركها فارغة. وهي تُعرض **واحدة تلو الأخرى** لتسهيل الإجابة.',
    'Warning: question 1 is empty (you can still proceed).': 'تنبيه : السؤال 1 فارغ (ويمكنكم مع ذلك المتابعة).',
    'Warning: question 2 is empty (you can still proceed).': 'تنبيه : السؤال 2 فارغ (ويمكنكم مع ذلك المتابعة).',
    'Warning: question 3 is empty (you can still proceed).': 'تنبيه : السؤال 3 فارغ (ويمكنكم مع ذلك المتابعة).',
    'Final step: please indicate whether you consulted other colleagues. This question is mandatory.': 'الخطوة الأخيرة : يُرجى بيان ما إذا كنتم قد استشرتم زملاء آخرين. هذا السؤال إلزامي.',
    'Did you consult other colleagues to complete this questionnaire?': 'هل استشرتم زملاء آخرين لإكمال هذا الاستبيان ؟',
    'Progress: 1/3 → 2/3 → 3/3 → Confirmation.': 'التقدم : 1/3 ← 2/3 ← 3/3 ← تأكيد.',
    'Section 12: please go through the open questions one by one (use the “Next question” button) until Confirmation.': 'القسم 12 : يُرجى المرور عبر الأسئلة المفتوحة واحداً تلو الآخر (باستخدام زر ‘السؤال التالي’) إلى غاية التأكيد.',
    'Section 12: please indicate whether you consulted other colleagues (Yes/No).': 'القسم 12 : يُرجى بيان ما إذا كنتم قد استشرتم زملاء آخرين (نعم/لا).',
    'Question 1 / 3': 'السؤال 1 / 3',
    'Question 2 / 3': 'السؤال 2 / 3',
    'Question 3 / 3': 'السؤال 3 / 3',
    'Confirmation': 'التأكيد',

}



def _normalize_i18n_key(text: str) -> str:
    return "\n".join([line.strip() for line in str(text).strip().splitlines() if line.strip()])

def lang_base(lang: str) -> str:
    return lang if lang in {"fr", "en", "pt", "ar"} else "en"

def tr(lang: str, text: str) -> str:
    key = str(text)
    norm = _normalize_i18n_key(key)
    if lang == "pt":
        return PT_TRANSLATIONS.get(key) or PT_TRANSLATIONS.get(norm) or key
    if lang == "ar":
        return AR_TRANSLATIONS.get(key) or AR_TRANSLATIONS.get(norm) or key
    return key

def t(lang: str, fr: str, en: str, pt: Optional[str] = None, ar: Optional[str] = None) -> str:
    if lang == "fr":
        return fr
    if lang == "en":
        return en
    fr_key = str(fr)
    en_key = str(en)
    fr_norm = _normalize_i18n_key(fr_key)
    en_norm = _normalize_i18n_key(en_key)
    if lang == "pt":
        return pt if pt is not None else (PT_TRANSLATIONS.get(en_key) or PT_TRANSLATIONS.get(en_norm) or PT_TRANSLATIONS.get(fr_key) or PT_TRANSLATIONS.get(fr_norm) or en)
    if lang == "ar":
        return ar if ar is not None else (AR_TRANSLATIONS.get(en_key) or AR_TRANSLATIONS.get(en_norm) or AR_TRANSLATIONS.get(fr_key) or AR_TRANSLATIONS.get(fr_norm) or en)
    return en

def get_role_options(lang: str) -> List[str]:
    if lang == "fr":
        return ROLE_OPTIONS_FR
    if lang == "pt":
        return ROLE_OPTIONS_PT
    if lang == "ar":
        return ROLE_OPTIONS_AR
    return ROLE_OPTIONS_EN

def is_other_value(v: Any) -> bool:
    return str(v or "").strip() in {"Autre", "Other", "Outro", "أخرى", "Autres", "Outros"}

def has_other_option(values: Any) -> bool:
    if not isinstance(values, list):
        return False
    return any(is_other_value(v) for v in values)


def now_utc_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def get_query_params() -> Dict[str, List[str]]:
    """Compatibility across Streamlit versions."""
    try:
        # Streamlit >= 1.30
        qp = st.query_params  # type: ignore
        return {k: list(v) if isinstance(v, (list, tuple)) else [str(v)] for k, v in qp.items()}
    except Exception:
        try:
            return st.experimental_get_query_params()
        except Exception:
            return {}


def set_query_params(params: Dict[str, Any]) -> None:
    try:
        st.query_params.update(params)  # type: ignore
    except Exception:
        try:
            st.experimental_set_query_params(**params)
        except Exception:
            pass


def init_session() -> None:
    if "lang" not in st.session_state:
        st.session_state.lang = "fr"
    if "nav_idx" not in st.session_state:
        st.session_state.nav_idx = 0
    if "responses" not in st.session_state:
        st.session_state["responses"] = {}
    elif not isinstance(st.session_state.get("responses"), dict):
        st.session_state["responses"] = {}
    if "submission_id" not in st.session_state:
        st.session_state.submission_id = None
    if "admin_authed" not in st.session_state:
        st.session_state.admin_authed = False
    if "admin_role" not in st.session_state:
        st.session_state.admin_role = None  # "admin" | "superadmin"
    if "draft_id" not in st.session_state:
        st.session_state.draft_id = None

    if "draft_exists" not in st.session_state:
        st.session_state.draft_exists = False
    if "draft_resume_notice_shown" not in st.session_state:
        st.session_state.draft_resume_notice_shown = False
    if "draft_restored" not in st.session_state:
        st.session_state.draft_restored = False
    if "last_draft_save_ts" not in st.session_state:
        st.session_state.last_draft_save_ts = 0.0
    if "r12_substep" not in st.session_state:
        st.session_state["r12_substep"] = 0  # 0..2=open questions, 3=confirmation


def ensure_responses() -> None:
    """Garantit l’existence de st.session_state['responses'] (dict)."""
    if "responses" not in st.session_state or not isinstance(st.session_state.get("responses"), dict):
        st.session_state["responses"] = {}


def resp_get(key: str, default=None):
    ensure_responses()
    return st.session_state["responses"].get(key, default)


def resp_set(key: str, value) -> None:
    ensure_responses()
    st.session_state["responses"][key] = value



def normalize_availability(v_raw: Any, scoring_version: Any) -> int:
    """Normalise la disponibilité sur l'échelle 'Bonne=3' (SCORING_VERSION=3).

    - v3+ : on conserve la valeur telle quelle (0–3).
    - v1/v2 ou absence de version : on inverse (1<->3) car l'ancien codage correspondait à un "écart" / ou à une disponibilité inversée.
    """
    try:
        iv = int(v_raw)
    except Exception:
        return 0
    if iv == 0:
        return 0

    try:
        ver = int(scoring_version)
    except Exception:
        ver = 0

    if ver >= SCORING_VERSION:
        return iv

    # Inversion pour les versions antérieures : 1<->3, 2 inchangé
    if iv in (1, 2, 3):
        return 4 - iv
    return iv

def ensure_draft_id() -> Optional[str]:
    """Ensure a stable draft id exists (used for mobile resume)."""
    if st.session_state.get("draft_id"):
        return st.session_state.draft_id
    email = (resp_get("email", "") or "").strip()
    if not email:
        return None
    draft_id = str(uuid.uuid4())
    st.session_state.draft_id = draft_id
    # Keep any existing query params (admin, etc.)
    try:
        qp = get_query_params()
        qp["rid"] = [draft_id]
        set_query_params({k: v[0] if len(v) == 1 else v for k, v in qp.items()})
    except Exception:
        pass
    return draft_id


def autosave_draft(force: bool = False) -> Tuple[bool, str]:
    """Persist current responses to DB to mitigate mobile refresh/resets."""
    draft_id = st.session_state.get("draft_id")
    email = (resp_get("email", "") or "").strip()
    if not draft_id or not email:
        return False, "no_draft_id_or_email"
    now_ts = time.time()
    last_ts = float(st.session_state.get("last_draft_save_ts", 0.0) or 0.0)
    if (not force) and (now_ts - last_ts < 2.0):
        return True, "skipped_rate_limit"
    payload = {
        "responses": st.session_state.responses,
        "nav_idx": int(st.session_state.get("nav_idx", 0)),
        "lang": st.session_state.get("lang", "fr"),
    }
    try:
        db_save_draft(draft_id, email, payload)
        st.session_state.last_draft_save_ts = now_ts
        return True, "saved"
    except Exception as e:
        return False, str(e)


def maybe_restore_draft() -> None:
    """Restore responses from DB if URL contains rid and session is empty."""
    if st.session_state.get("draft_restored"):
        return
    st.session_state.draft_restored = True

    qp = get_query_params()
    rid = None
    if "rid" in qp and qp["rid"]:
        rid = qp["rid"][0]
    if not rid:
        return

    # Do not restore while in admin mode
    if "admin" in qp and qp["admin"] and qp["admin"][0] in ["1", "true", "yes"]:
        return

    payload = db_load_draft(rid)
    st.session_state.draft_exists = bool(payload)

    # Restore only if session is empty (avoid overriding ongoing input)
    if st.session_state.get("responses") and len(st.session_state.responses) > 0:
        st.session_state.draft_id = rid
        return


    if not payload:
        st.session_state.draft_id = rid
        return

    responses = payload.get("responses", {})
    if isinstance(responses, dict):
        st.session_state.responses = responses

    try:
        st.session_state.nav_idx = int(payload.get("nav_idx", 0))
    except Exception:
        st.session_state.nav_idx = 0

    lang = payload.get("lang", None)
    if lang in LANG_OPTIONS:
        st.session_state.lang = lang

    st.session_state.draft_id = rid




# =========================
# Data : longlist loader
# =========================

def load_longlist() -> pd.DataFrame:
    """
    Charge la longlist depuis les emplacements racine et /data.
    Si la lecture des fichiers échoue (par ex. dépendance Excel absente),
    un jeu de données embarqué complet PT/AR est utilisé.
    """
    embedded = _embedded_longlist_df()

    csv_names = ["indicator_longlist.csv"]
    xlsx_names = ["longlist_filled_pt_ar_final.xlsx", "longlist.xlsx"]

    df_csv = None
    df_csv_path = ""
    for p in candidate_paths(csv_names):
        if os.path.exists(p):
            try:
                df_csv = _normalize_longlist_csv(pd.read_csv(p, dtype=str).fillna(""))
                if df_csv is not None and not df_csv.empty:
                    df_csv_path = p
                    break
            except Exception:
                continue

    df_xlsx = None
    df_xlsx_path = ""
    for p in candidate_paths(xlsx_names):
        if os.path.exists(p):
            try:
                df_xlsx = _normalize_longlist_xlsx(_read_excel_flexible(p))
                if df_xlsx is not None and not df_xlsx.empty:
                    df_xlsx_path = p
                    break
            except Exception:
                continue

    final_df = embedded.copy()
    source_bits = []
    if not embedded.empty:
        source_bits.append(getattr(embedded, "attrs", {}).get("source_path", "embedded"))
    if df_xlsx is not None and not df_xlsx.empty:
        final_df = _merge_longlists(final_df, df_xlsx)
        source_bits.append(df_xlsx_path)
    if df_csv is not None and not df_csv.empty:
        final_df = _merge_longlists(final_df, df_csv)
        source_bits.append(df_csv_path)

    if final_df is not None and not final_df.empty:
        final_df.attrs["source_path"] = "merged::" + " + ".join([s for s in source_bits if s])
        return final_df

    empty = pd.DataFrame(columns=[
        "domain_code",
        "domain_label_fr",
        "domain_label_en",
        "domain_label_pt",
        "domain_label_ar",
        "stat_code",
        "stat_label_fr",
        "stat_label_en",
        "stat_label_pt",
        "stat_label_ar",
    ])
    empty.attrs["source_path"] = ""
    return empty


# =========================
# Data : countries loader
# =========================

def load_countries() -> pd.DataFrame:
    """
    Charge la liste des pays depuis plusieurs noms possibles, à la racine et dans /data.
    Si la lecture Excel échoue, le référentiel embarqué complet PT/AR est utilisé.
    """
    embedded = _embedded_countries_df()
    candidate_names = [
        "COUNTRY_ISO3_4_filled_pt_ar_final.xlsx",
        "COUNTRY_ISO3_4.xlsx",
        "COUNTRY_ISO3_with_EN.xlsx",
        "COUNTRY_ISO3.xlsx",
    ]

    df_ext = None
    ext_path = ""
    for p in candidate_paths(candidate_names):
        if os.path.exists(p):
            try:
                df_ext = _normalize_countries_df(_read_excel_flexible(p))
                if df_ext is not None and not df_ext.empty:
                    ext_path = p
                    break
            except Exception:
                continue

    final_df = _merge_countries(embedded, df_ext)
    if final_df is not None and not final_df.empty:
        source_bits = []
        if not embedded.empty:
            source_bits.append(getattr(embedded, "attrs", {}).get("source_path", "embedded"))
        if ext_path:
            source_bits.append(ext_path)
        final_df.attrs["source_path"] = "merged::" + " + ".join([s for s in source_bits if s])
        return final_df

    empty = pd.DataFrame(columns=["COUNTRY_ISO3", "COUNTRY_NAME_FR", "COUNTRY_NAME_EN", "COUNTRY_NAME_PT", "COUNTRY_NAME_AR"])
    empty.attrs["source_path"] = ""
    return empty

def _pick_country_name_col(lang: str) -> str:
    if lang == "fr":
        return "COUNTRY_NAME_FR"
    if lang == "pt":
        return "COUNTRY_NAME_PT"
    if lang == "ar":
        return "COUNTRY_NAME_AR"
    return "COUNTRY_NAME_EN"


def _pick_longlist_domain_col(lang: str) -> str:
    if lang == "fr":
        return "domain_label_fr"
    if lang == "pt":
        return "domain_label_pt"
    if lang == "ar":
        return "domain_label_ar"
    return "domain_label_en"


def _pick_longlist_stat_col(lang: str) -> str:
    if lang == "fr":
        return "stat_label_fr"
    if lang == "pt":
        return "stat_label_pt"
    if lang == "ar":
        return "stat_label_ar"
    return "stat_label_en"


def country_maps(df_c: pd.DataFrame) -> Tuple[List[str], Dict[str, str], Dict[str, str], Dict[str, str], Dict[str, str]]:
    """Retourne (iso3_list, iso3_to_fr, iso3_to_en, iso3_to_pt, iso3_to_ar)."""
    if df_c is None or df_c.empty:
        return [], {}, {}, {}, {}
    iso3_to_fr: Dict[str, str] = {}
    iso3_to_en: Dict[str, str] = {}
    iso3_to_pt: Dict[str, str] = {}
    iso3_to_ar: Dict[str, str] = {}
    for _, r in df_c.iterrows():
        iso3 = str(r.get("COUNTRY_ISO3", "")).strip().upper()
        if not iso3:
            continue
        iso3_to_fr[iso3] = str(r.get("COUNTRY_NAME_FR", "")).strip()
        iso3_to_en[iso3] = str(r.get("COUNTRY_NAME_EN", "")).strip()
        iso3_to_pt[iso3] = str(r.get("COUNTRY_NAME_PT", "")).strip()
        iso3_to_ar[iso3] = str(r.get("COUNTRY_NAME_AR", "")).strip()
    iso3_list = sorted(set(list(iso3_to_fr.keys()) + list(iso3_to_en.keys()) + list(iso3_to_pt.keys()) + list(iso3_to_ar.keys())))
    return iso3_list, iso3_to_fr, iso3_to_en, iso3_to_pt, iso3_to_ar


def country_label(
    iso3: str,
    lang: str,
    iso3_to_fr: Dict[str, str],
    iso3_to_en: Dict[str, str],
    iso3_to_pt: Optional[Dict[str, str]] = None,
    iso3_to_ar: Optional[Dict[str, str]] = None,
) -> str:
    if not iso3:
        return ""
    iso3_to_pt = iso3_to_pt or {}
    iso3_to_ar = iso3_to_ar or {}
    if lang == "fr":
        return (iso3_to_fr.get(iso3) or iso3_to_en.get(iso3) or iso3_to_pt.get(iso3) or iso3_to_ar.get(iso3) or iso3).strip()
    if lang == "pt":
        return (iso3_to_pt.get(iso3) or iso3_to_en.get(iso3) or iso3_to_fr.get(iso3) or iso3_to_ar.get(iso3) or iso3).strip()
    if lang == "ar":
        return (iso3_to_ar.get(iso3) or iso3_to_en.get(iso3) or iso3_to_fr.get(iso3) or iso3_to_pt.get(iso3) or iso3).strip()
    return (iso3_to_en.get(iso3) or iso3_to_fr.get(iso3) or iso3_to_pt.get(iso3) or iso3_to_ar.get(iso3) or iso3).strip()


def domains_from_longlist(df_long: pd.DataFrame, lang: str) -> List[Tuple[str, str]]:
    if df_long.empty:
        return []
    col = _pick_longlist_domain_col(lang)
    if col not in df_long.columns:
        col = "domain_label_en" if "domain_label_en" in df_long.columns else "domain_label_fr"
    tmp = df_long[["domain_code", col]].drop_duplicates().sort_values(["domain_code", col])
    return [(r["domain_code"], r[col]) for _, r in tmp.iterrows()]


def stats_for_domain(df_long: pd.DataFrame, domain_code: str, lang: str) -> List[Tuple[str, str]]:
    if df_long.empty or not domain_code:
        return []
    col = _pick_longlist_stat_col(lang)
    if col not in df_long.columns:
        col = "stat_label_en" if "stat_label_en" in df_long.columns else "stat_label_fr"
    tmp = df_long[df_long["domain_code"] == domain_code][["stat_code", col]].drop_duplicates()
    tmp = tmp.sort_values(["stat_code", col])
    return [(r["stat_code"], (str(r[col]).strip() if str(r[col]).strip() else r["stat_code"])) for _, r in tmp.iterrows()]


# =========================
# Stockage : SQLite (local)
# =========================


def domain_label_map(df_long: pd.DataFrame, lang: str) -> Dict[str, str]:
    """Map domain_code -> label in selected language."""
    if df_long is None or df_long.empty:
        return {}
    col = _pick_longlist_domain_col(lang)
    if col not in df_long.columns:
        col = "domain_label_en" if "domain_label_en" in df_long.columns else "domain_label_fr"
    m = {}
    for _, r in df_long.drop_duplicates("domain_code").iterrows():
        code = str(r["domain_code"])
        lbl = str(r.get(col, "")).strip()
        if not lbl:
            lbl = str(r.get("domain_label_fr", "")).strip() or code
        m[code] = lbl
    return m

def stat_label_map(df_long: pd.DataFrame, lang: str) -> Dict[str, str]:
    """Map stat_code -> label in selected language."""
    if df_long is None or df_long.empty:
        return {}
    col = _pick_longlist_stat_col(lang)
    if col not in df_long.columns:
        col = "stat_label_en" if "stat_label_en" in df_long.columns else "stat_label_fr"
    m = {}
    for _, r in df_long.drop_duplicates("stat_code").iterrows():
        code = str(r["stat_code"])
        lbl = str(r.get(col, "")).strip()
        if not lbl:
            lbl = str(r.get("stat_label_fr", "")).strip() or code
        m[code] = lbl
    return m

def build_publication_report_docx(lang: str, filtered_payloads: pd.DataFrame, by_domain: pd.DataFrame, by_stat: pd.DataFrame, scored_rows: pd.DataFrame) -> bytes:
    """
    Génère un rapport Word 'publication' enrichi :
    - profil des répondants
    - domaines TOP 5 (fréquences)
    - tableau agrégé des statistiques et scores moyens
    - graphiques (bar charts)
    - annexes
    """
    from docx import Document
    from docx.shared import Inches
    import matplotlib.pyplot as plt

    doc = Document()
    title = t(lang, "Rapport de synthèse – Consultation sur les statistiques prioritaires", "Summary report – Consultation on priority statistics")
    doc.add_heading(title, level=0)
    doc.add_paragraph(t(lang, f"Date : {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", f"Date: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"))
    doc.add_paragraph("")

    # Sample profile
    doc.add_heading(t(lang, "Profil des répondants", "Respondent profile"), level=1)
    n = len(filtered_payloads)
    doc.add_paragraph(t(lang, f"Nombre de réponses analysées : {n}", f"Number of responses analyzed: {n}"))

    # Countries
    if "pays" in filtered_payloads.columns:
        vc = filtered_payloads["pays"].fillna("").astype(str)
        vc = vc[vc.str.strip() != ""].value_counts().head(10)
        if len(vc) > 0:
            doc.add_paragraph(t(lang, "Top pays (10 premiers) :", "Top countries (top 10):"))
            for k, v in vc.items():
                doc.add_paragraph(f"- {k} : {v}", style=None)

    # Actor types
    if "type_acteur" in filtered_payloads.columns:
        vc = filtered_payloads["type_acteur"].fillna("").astype(str)
        vc = vc[vc.str.strip() != ""].value_counts()
        if len(vc) > 0:
            doc.add_paragraph(t(lang, "Répartition par type d’acteur :", "Distribution by stakeholder type:"))
            for k, v in vc.items():
                doc.add_paragraph(f"- {k} : {v}", style=None)

    # Domain aggregation
    doc.add_heading(t(lang, "Domaines prioritaires (scores moyens)", "Priority domains (mean scores)"), level=1)
    top_dom = by_domain.head(15).copy()
    # Table
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = t(lang, "Domaine", "Domain")
    hdr[1].text = t(lang, "Nb. soumissions", "Submissions")
    hdr[2].text = t(lang, "Nb. stats notées", "Scored indicators")
    hdr[3].text = t(lang, "Score moyen", "Mean score")
    for _, r in top_dom.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["domain_label"])
        row[1].text = str(int(r["n_submissions"]))
        row[2].text = str(int(r["n_stats"]))
        row[3].text = f"{float(r['mean_overall']):.2f}"

    # Chart domain
    try:
        fig = plt.figure()
        plt.bar(top_dom["domain_label"], top_dom["mean_overall"])
        plt.xticks(rotation=75, ha="right")
        plt.ylabel(t(lang, "Score moyen", "Mean score"))
        plt.tight_layout()
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format="png", dpi=150)
        plt.close(fig)
        img_stream.seek(0)
        doc.add_paragraph(t(lang, "Graphique : score moyen par domaine (top 15).", "Chart: mean score by domain (top 15)."))
        doc.add_picture(img_stream, width=Inches(6.5))
    except Exception:
        pass

    # Statistic aggregation
    doc.add_heading(t(lang, "Statistiques prioritaires (scores moyens)", "Priority indicators (mean scores)"), level=1)
    top_stat = by_stat.sort_values(["mean_overall", "n"], ascending=[False, False]).head(30).copy()
    table2 = doc.add_table(rows=1, cols=6)
    h = table2.rows[0].cells
    h[0].text = t(lang, "Domaine", "Domain")
    h[1].text = t(lang, "Statistique", "Indicator")
    h[2].text = t(lang, "N", "N")
    h[3].text = t(lang, "Demande", "Demand")
    h[4].text = t(lang, "Disponibilité", "Availability")
    h[5].text = t(lang, "Faisabilité", "Feasibility")
    for _, r in top_stat.iterrows():
        row = table2.add_row().cells
        row[0].text = str(r["domain_label"])
        row[1].text = str(r["stat_label"])
        row[2].text = str(int(r["n"]))
        row[3].text = f"{float(r['mean_demand']):.2f}"
        row[4].text = f"{float(r['mean_availability']):.2f}"
        row[5].text = f"{float(r['mean_feasibility']):.2f}"


    # Chart top overall indicators
    try:
        fig = plt.figure()
        plt.barh(top_stat["stat_label"].iloc[::-1], top_stat["mean_overall"].iloc[::-1])
        plt.xlabel(t(lang, "Score moyen", "Mean score"))
        plt.tight_layout()
        img2 = io.BytesIO()
        plt.savefig(img2, format="png", dpi=150)
        plt.close(fig)
        img2.seek(0)
        doc.add_paragraph(t(lang, "Graphique : score moyen par statistique (top 30).", "Chart: mean score by indicator (top 30)."))
        doc.add_picture(img2, width=Inches(6.5))
    except Exception:
        pass

    # Interpretation auto
    doc.add_heading(t(lang, "Interprétations automatiques", "Automatic interpretations"), level=1)
    # Simple rules
    best_dom = top_dom.iloc[0]
    doc.add_paragraph(
        t(
            lang,
            f"Le domaine le mieux noté est « {best_dom['domain_label']} » avec un score moyen de {best_dom['mean_overall']:.2f} (sur 3).",
            f"The highest-rated domain is “{best_dom['domain_label']}” with a mean score of {best_dom['mean_overall']:.2f} (out of 3)."
        )
    )
    best_stat = top_stat.iloc[0]
    doc.add_paragraph(
        t(
            lang,
            f"La statistique la mieux notée est « {best_stat['stat_label']} » (domaine : {best_stat['domain_label']}) avec un score moyen de {best_stat['mean_overall']:.2f}.",
            f"The highest-rated indicator is “{best_stat['stat_label']}” (domain: {best_stat['domain_label']}) with a mean score of {best_stat['mean_overall']:.2f}."
        )
    )

    # Annexes
    doc.add_heading(t(lang, "Annexes", "Annexes"), level=1)
    doc.add_paragraph(t(lang, "A1. Tableau détaillé (statistiques agrégées) – extrait", "A1. Detailed table (aggregated indicators) – excerpt"))
    annex = by_stat.head(50).copy()
    tab3 = doc.add_table(rows=1, cols=5)
    hh = tab3.rows[0].cells
    hh[0].text = t(lang, "Domaine", "Domain")
    hh[1].text = t(lang, "Statistique", "Indicator")
    hh[2].text = t(lang, "N", "N")
    hh[3].text = t(lang, "Score moyen", "Mean score")
    hh[4].text = t(lang, "Détail", "Detail")
    for _, r in annex.iterrows():
        rr = tab3.add_row().cells
        rr[0].text = str(r["domain_label"])
        rr[1].text = str(r["stat_label"])
        rr[2].text = str(int(r["n"]))
        rr[3].text = f"{float(r['mean_overall']):.2f}"
        if lang == "fr":
            rr[4].text = f"Demande={float(r['mean_demand']):.2f}, Disponibilité={float(r['mean_availability']):.2f}, Faisabilité={float(r['mean_feasibility']):.2f}"
        elif lang == "pt":
            rr[4].text = f"Procura={float(r['mean_demand']):.2f}, Disponibilidade={float(r['mean_availability']):.2f}, Viabilidade={float(r['mean_feasibility']):.2f}"
        elif lang == "ar":
            rr[4].text = f"الطلب={float(r['mean_demand']):.2f}, التوفر={float(r['mean_availability']):.2f}, الجدوى={float(r['mean_feasibility']):.2f}"
        else:
            rr[4].text = f"Demand={float(r['mean_demand']):.2f}, Availability={float(r['mean_availability']):.2f}, Feasibility={float(r['mean_feasibility']):.2f}"


    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()




def db_init() -> None:
    os.makedirs(os.path.dirname(DB_PATH) or ".", exist_ok=True)
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS submissions(
            submission_id TEXT PRIMARY KEY,
            submitted_at_utc TEXT,
            lang TEXT,
            email TEXT,
            payload_json TEXT
        )
    """)

    # Backward compatibility : add email column if existing DB was created with older schema
    try:
        cur.execute("PRAGMA table_info(submissions)")
        cols = [r[1] for r in cur.fetchall()]
        if "email" not in cols:
            cur.execute("ALTER TABLE submissions ADD COLUMN email TEXT")
    except Exception:
        pass

    # Helpful index (non-unique) for email lookups
    try:
        cur.execute("CREATE INDEX IF NOT EXISTS idx_submissions_email ON submissions(email)")
    except Exception:
        pass


    # Optionnel : empêcher les doublons email (best-effort)
    # Note : si des doublons existent déjà dans la base, la création de l’index UNIQUE échouera (sans bloquer l’app).
    try:
        cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS uq_submissions_email ON submissions(email)")
    except Exception:
        pass
    # Drafts (for mobile stability / resume)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS drafts(
            draft_id TEXT PRIMARY KEY,
            updated_at_utc TEXT,
            email TEXT,
            payload_json TEXT
        )
    """)
    try:
        cur.execute("CREATE INDEX IF NOT EXISTS idx_drafts_email ON drafts(email)")
    except Exception:
        pass

    # App config (e.g. hashed admin password override)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS app_config(
            k TEXT PRIMARY KEY,
            v TEXT,
            updated_at_utc TEXT
        )
    """)

    con.commit()
    con.close()



# =========================
# Admin auth helpers
# =========================

PBKDF2_ITERS = 200_000

def db_get_config(k: str) -> Optional[str]:
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("SELECT v FROM app_config WHERE k=? LIMIT 1", (k,))
    row = cur.fetchone()
    con.close()
    return row[0] if row and row[0] is not None else None


def db_set_config(k: str, v: str) -> None:
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute(
        "INSERT OR REPLACE INTO app_config(k, v, updated_at_utc) VALUES(?, ?, ?)",
        (k, v, now_utc_iso()),
    )
    con.commit()
    con.close()


def db_delete_config(k: str) -> None:
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("DELETE FROM app_config WHERE k=?", (k,))
    con.commit()
    con.close()


def _pbkdf2_sha256_hex(password: str, salt: bytes, iterations: int = PBKDF2_ITERS) -> str:
    dk = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, iterations)
    return dk.hex()


def _safe_eq(a: str, b: str) -> bool:
    try:
        return hmac.compare_digest(a or "", b or "")
    except Exception:
        return (a or "") == (b or "")


def _get_secret_or_env(name: str) -> Tuple[Optional[str], Optional[str]]:
    """Return (value, source) where source is 'secrets' | 'env' | None.

    Robust to different Streamlit versions and secrets layouts:
    - Direct key : ADMIN_PASSWORD = "..."
    - Nested table : [general] ADMIN_PASSWORD = "..."
    """
    val: Optional[str] = None
    src: Optional[str] = None

    # 1) Streamlit secrets
    try:
        secrets_obj = getattr(st, "secrets", None)
        if secrets_obj is not None:
            # Direct access (preferred)
            try:
                if hasattr(secrets_obj, "__contains__") and name in secrets_obj:
                    v = secrets_obj[name]
                    if v not in (None, ""):
                        val = str(v)
                        src = "secrets"
            except Exception:
                pass

            # Fallback : convert to dict then search (supports nested sections)
            if not val:
                try:
                    d = secrets_obj.to_dict() if hasattr(secrets_obj, "to_dict") else dict(secrets_obj)
                except Exception:
                    d = {}

                if isinstance(d, dict):
                    if name in d and d.get(name) not in (None, ""):
                        val = str(d.get(name))
                        src = "secrets"
                    else:
                        # Search nested dicts
                        for _k, _v in d.items():
                            if isinstance(_v, dict) and name in _v and _v.get(name) not in (None, ""):
                                val = str(_v.get(name))
                                src = "secrets"
                                break
    except Exception:
        # Leave val/src as None
        pass

    # 2) Environment variable (only if not found in secrets)
    if not val:
        env_val = os.environ.get(name, None)
        if env_val not in (None, ""):
            val = str(env_val)
            src = "env"

    return (val, src)


def get_admin_auth_source() -> Tuple[str, str]:
    """Human-readable indicator of current admin password source."""
    h = db_get_config("ADMIN_PASSWORD_HASH")
    s = db_get_config("ADMIN_PASSWORD_SALT")
    if h and s:
        return ("db", t(st.session_state.get("lang", "fr"), "base (haché)", "database (hashed)"))
    v, src = _get_secret_or_env("ADMIN_PASSWORD")
    if v and src:
        return (src, t(st.session_state.get("lang", "fr"), src, src))
    return ("none", t(st.session_state.get("lang", "fr"), "non configuré", "not configured"))


def verify_admin_password(pw: str) -> bool:
    pw = pw or ""
    h = db_get_config("ADMIN_PASSWORD_HASH")
    s = db_get_config("ADMIN_PASSWORD_SALT")
    it = db_get_config("ADMIN_PASSWORD_ITERS")
    if h and s:
        try:
            salt = bytes.fromhex(s)
            iterations = int(it) if it else PBKDF2_ITERS
            calc = _pbkdf2_sha256_hex(pw, salt, iterations)
            return _safe_eq(calc, h)
        except Exception:
            return False

    expected, _src = _get_secret_or_env("ADMIN_PASSWORD")
    if expected:
        return _safe_eq(pw, str(expected))
    return False


def verify_superadmin_password(pw: str) -> bool:
    pw = pw or ""
    expected, _src = _get_secret_or_env("SUPERADMIN_PASSWORD")
    if expected:
        return _safe_eq(pw, str(expected))
    return False


def set_admin_password(new_pw: str) -> None:
    """Set (hashed) admin password override in DB."""
    new_pw = (new_pw or "").strip()
    if not new_pw:
        raise ValueError("empty password")
    salt = secrets.token_bytes(16)
    h = _pbkdf2_sha256_hex(new_pw, salt, PBKDF2_ITERS)
    db_set_config("ADMIN_PASSWORD_HASH", h)
    db_set_config("ADMIN_PASSWORD_SALT", salt.hex())
    db_set_config("ADMIN_PASSWORD_ITERS", str(PBKDF2_ITERS))


def reset_admin_password_to_secrets_env() -> None:
    """Remove DB override so app falls back to secrets/env."""
    db_delete_config("ADMIN_PASSWORD_HASH")
    db_delete_config("ADMIN_PASSWORD_SALT")
    db_delete_config("ADMIN_PASSWORD_ITERS")

def db_email_exists(email: str) -> bool:
    email = (email or "").strip().lower()
    if not email or not os.path.exists(DB_PATH):
        return False
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("SELECT 1 FROM submissions WHERE lower(email)=? LIMIT 1", (email,))
    row = cur.fetchone()
    con.close()
    return row is not None


def db_save_submission(submission_id: str, lang: str, email: str, payload: Dict[str, Any]) -> None:
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
        INSERT OR REPLACE INTO submissions(submission_id, submitted_at_utc, lang, email, payload_json)
        VALUES(?, ?, ?, ?, ?)
    """, (submission_id, now_utc_iso(), lang, (email or "").strip().lower(), json.dumps(payload, ensure_ascii=False)))
    con.commit()
    con.close()



def db_save_draft(draft_id: str, email: str, payload: Dict[str, Any]) -> None:
    draft_id = (draft_id or "").strip()
    email = (email or "").strip().lower()
    if not draft_id or not email:
        return
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute(
        """
        INSERT OR REPLACE INTO drafts(draft_id, updated_at_utc, email, payload_json)
        VALUES(?, ?, ?, ?)
        """,
        (draft_id, now_utc_iso(), email, json.dumps(payload, ensure_ascii=False)),
    )
    con.commit()
    con.close()


def db_load_draft(draft_id: str) -> Optional[Dict[str, Any]]:
    draft_id = (draft_id or "").strip()
    if not draft_id or not os.path.exists(DB_PATH):
        return None
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("SELECT payload_json FROM drafts WHERE draft_id=? LIMIT 1", (draft_id,))
    row = cur.fetchone()
    con.close()
    if not row or not row[0]:
        return None
    try:
        return json.loads(row[0])
    except Exception:
        return None


def db_delete_draft(draft_id: str) -> None:
    draft_id = (draft_id or "").strip()
    if not draft_id or not os.path.exists(DB_PATH):
        return
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("DELETE FROM drafts WHERE draft_id=?", (draft_id,))
    con.commit()
    con.close()

def db_read_submissions(limit: int = 2000) -> pd.DataFrame:
    if not os.path.exists(DB_PATH):
        return pd.DataFrame(columns=["submission_id", "submitted_at_utc", "lang", "email", "payload_json"])
    con = sqlite3.connect(DB_PATH)
    df = pd.read_sql_query(
        "SELECT submission_id, submitted_at_utc, lang, email, payload_json FROM submissions ORDER BY submitted_at_utc DESC LIMIT ?",
        con,
        params=(limit,),
    )
    con.close()
    return df


def db_dump_csv_bytes(limit: int = 2000000) -> bytes:
    """Export the SQLite submissions table to CSV bytes."""
    df = db_read_submissions(limit=limit)
    return df.to_csv(index=False).encode("utf-8-sig")



def flatten_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Create a 'flat' row for exports (comprehensive).
    - Keeps keys stable across FR/EN by mapping table items to canonical ids.
    - Serializes list/dict fields into '; ' / JSON strings as needed.
    """
    def _join_list(v: Any) -> str:
        if isinstance(v, list):
            return "; ".join([str(x) for x in v if x is not None and str(x).strip() != ""])
        return ""

    def _json(v: Any) -> str:
        try:
            return json.dumps(v, ensure_ascii=False)
        except Exception:
            return ""

    # Canonical mappings for table questions (FR/EN)
    GENDER_ITEM_MAP = {
        "Désagrégation par sexe": "sex",
        "Disaggregation by sex": "sex",
        "Sexe": "sex",
        "Sex": "sex",
        "Désagrégation par âge": "age",
        "Disaggregation by age": "age",
        "Âge": "age",
        "Age": "age",
        "Milieu urbain / rural": "urban_rural",
        "Urban / rural": "urban_rural",
        "Milieu urbain/rural": "urban_rural",
        "Urban/rural residence": "urban_rural",
        "Handicap": "disability",
        "Disability": "disability",
        "Quintile de richesse": "wealth_quintile",
        "Wealth quintile": "wealth_quintile",
        "Violences basées sur le genre (VBG)": "gbv",
        "Gender-based violence (GBV)": "gbv",
        "Temps domestique non rémunéré": "unpaid_domestic",
        "Unpaid domestic work": "unpaid_domestic",
    }
    CAPACITY_ITEM_MAP = {
        "Compétences statistiques disponibles": "skills_hr",
        "Available statistical skills": "skills_hr",
        "Accès aux données administratives": "access_admin_data",
        "Access to administrative data": "access_admin_data",
        "Financement disponible": "funding",
        "Available funding": "funding",
        "Financement": "funding",
        "Funding": "funding",
        "Outils numériques (collecte, traitement, diffusion)": "digital_tools",
        "Digital tools (collection, processing, dissemination)": "digital_tools",
        "Outils numériques": "digital_tools",
        "Digital tools": "digital_tools",
        "Cadre juridique pour le partage de données": "legal_framework",
        "Legal framework for data sharing": "legal_framework",
        "Cadre juridique": "legal_framework",
        "Legal framework": "legal_framework",
        "Coordination interinstitutionnelle": "institutional_coordination",
        "Inter-institutional coordination": "institutional_coordination",
        "Coordination institutionnelle": "institutional_coordination",
        "Institutional coordination": "institutional_coordination",
    }

    def _extract_table(table_obj: Any, mapping: Dict[str, str], prefix: str) -> Dict[str, Any]:
        out_tbl: Dict[str, Any] = {}
        # Ensure stable columns even when a respondent skips the section
        canons = sorted(set(mapping.values()))
        for canon in canons:
            out_tbl[f"{prefix}_{canon}"] = ""
            out_tbl[f"{prefix}_{canon}_spec"] = ""
        if not isinstance(table_obj, dict):
            return out_tbl
        for label, canon in mapping.items():
            cell = table_obj.get(label, None)
            if isinstance(cell, dict):
                out_tbl[f"{prefix}_{canon}"] = cell.get("code", "")
                out_tbl[f"{prefix}_{canon}_spec"] = cell.get("spec", "")
            elif isinstance(cell, str):
                out_tbl[f"{prefix}_{canon}"] = cell
        return out_tbl

    out: Dict[str, Any] = {}

    # Identification (Rubrique 2)
    out["organisation"] = payload.get("organisation", "")
    out["pays"] = payload.get("pays", "")
    out["pays_autre"] = payload.get("pays_autre", "")
    out["pays_name_fr"] = payload.get("pays_name_fr", "")
    out["pays_name_en"] = payload.get("pays_name_en", "")
    out["type_acteur"] = payload.get("type_acteur", "")
    out["fonction"] = payload.get("fonction", "")
    out["email"] = payload.get("email", "")
    out["lang"] = payload.get("lang", "")

    # Rubrique 3 : portée
    out["scope"] = payload.get("scope", "")
    out["scope_other"] = payload.get("scope_other", "")

    # Rubrique 4 : domaines
    pre = payload.get("preselected_domains", payload.get("preselection_domains", []))
    out["preselection_domains"] = _join_list(pre)
    out["nb_preselection_domains"] = len(pre) if isinstance(pre, list) else 0

    top5 = payload.get("top5_domains", [])
    for i in range(5):
        out[f"top_domain_{i+1}"] = top5[i] if i < len(top5) else ""

    # Rubrique 5 : stats et notation
    selected_stats = payload.get("selected_stats", [])
    out["nb_stats"] = len(selected_stats) if isinstance(selected_stats, list) else 0
    out["stats_list"] = _join_list(selected_stats)
    out["selected_by_domain_json"] = _json(payload.get("selected_by_domain", {}))
    out["scoring_json"] = _json(payload.get("scoring", {}))
    out["scoring_version"] = payload.get("scoring_version", "")

    # Rubrique 6 : perspective de genre (table)
    out.update(_extract_table(payload.get("gender_table", {}), GENDER_ITEM_MAP, "gender"))

    # Rubrique 8 : capacité & faisabilité (table)
    out.update(_extract_table(payload.get("capacity_table", {}), CAPACITY_ITEM_MAP, "capacity"))

    # Rubrique 9 : harmonisation & qualité
    out["quality_expectations"] = _join_list(payload.get("quality_expectations", []))
    out["quality_other"] = payload.get("quality_other", "")

    # Rubrique 10 : diffusion
    out["dissemination_channels"] = _join_list(payload.get("dissemination_channels", []))
    out["dissemination_other"] = payload.get("dissemination_other", "")

    # Rubrique 12 : questions ouvertes
    out["comment_1"] = payload.get("open_q1", "")
    out["missing_indicators"] = payload.get("open_q2", "")
    out["support_needs"] = payload.get("open_q3", "")

    # Dernière question : consultation d’autres personnes pour remplir le questionnaire
    out["consulted_colleagues"] = payload.get("consulted_colleagues", "")

    return out


# =========================
# Validation logic (quality controls)
# =========================

EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")


def validate_r2(lang: str) -> List[str]:
    errs: List[str] = []
    organisation = resp_get("organisation", "").strip()
    pays = resp_get("pays", "").strip()
    pays_autre = resp_get("pays_autre", "").strip()
    type_acteur = resp_get("type_acteur", "").strip()
    fonction = resp_get("fonction", "").strip()
    fonction_autre = resp_get("fonction_autre", "").strip()
    email = resp_get("email", "").strip()

    if not organisation:
        errs.append(t(lang, "Organisation : champ obligatoire.", "Organization: required field."))
    # Contrôle qualité : éviter les sigles seuls
    elif len(organisation) < 12:
        errs.append(t(lang, "Organisation : indiquez le libellé complet (au moins 12 caractères) et non le sigle.", "Organization: please provide the full name (at least 12 characters), not only an acronym."))
    if not pays:
        errs.append(t(lang, "Pays de résidence : champ obligatoire.", "Country of Residence: required field."))
    elif pays in {"OTHER", "__OTHER__"} and not pays_autre:
        errs.append(t(lang, "Autre pays (à préciser) : champ obligatoire.", "Other country (please specify): required field."))
    if not type_acteur:
        errs.append(t(lang, "Type d’acteur : champ obligatoire.", "Stakeholder type: required field."))
    if not fonction:
        errs.append(t(lang, "Fonction : champ obligatoire.", "Role/Function: required field."))
    if is_other_value(fonction):
        if not fonction_autre:
            errs.append(t(lang, "Fonction (Autre) : précisez.", "Role (Other): please specify."))
    if not email:
        errs.append(t(lang, "Email : champ obligatoire.", "Email: required field."))
    elif not EMAIL_RE.match(email):
        errs.append(t(lang, "Email : format invalide.", "Email: invalid format."))
    return errs



def validate_r3(lang: str) -> List[str]:
    errs: List[str] = []
    scope = (resp_get("scope", "") or "").strip()
    if not scope:
        errs.append(t(lang, "Rubrique 3 : veuillez sélectionner une portée.", "Section 3: please select a scope."))
        return errs
    snds = (resp_get("snds_status", "") or "").strip()
    if not snds:
        errs.append(
            t(
                lang,
                "Rubrique 3 : veuillez indiquer le statut de la SNDS / plan national statistique.",
                "Section 3: please indicate the status of the NSDS / national statistical plan.",
            )
        )

    if scope == "Other":
        other = (resp_get("scope_other", "") or "").strip()
        if not other:
            errs.append(t(lang, "Rubrique 3 : précisez l’option « Autre ».", "Section 3: please specify the \"Other\" option."))
    return errs

def validate_r4(lang: str) -> List[str]:
    errs: List[str] = []
    pre = resp_get("preselected_domains", [])
    top5 = resp_get("top5_domains", [])
    if not isinstance(pre, list):
        pre = []
    if not isinstance(top5, list):
        top5 = []
    if len(pre) < 5 or len(pre) > 10:
        errs.append(t(lang, "Rubrique 4 : pré-sélectionnez entre 5 et 10 domaines.", "Section 4: pre-select 5 to 10 domains."))
    if len(set(pre)) != len(pre):
        errs.append(t(lang, "Rubrique 4 : la pré-sélection contient des doublons.", "Section 4: duplicates found in pre-selection."))
    if len(top5) != 5:
        errs.append(t(lang, "Rubrique 4 : le TOP 5 doit contenir exactement 5 domaines.", "Section 4: TOP 5 must contain exactly 5 domains."))
    else:
        if len(set(top5)) != 5:
            errs.append(t(lang, "Rubrique 4 : le TOP 5 contient des doublons.", "Section 4: TOP 5 contains duplicates."))
        missing = [d for d in top5 if d not in pre]
        if missing:
            errs.append(t(lang, "Rubrique 4 : chaque domaine du TOP 5 doit provenir de la pré-sélection.", "Section 4: TOP 5 must be selected from pre-selection."))
    return errs


def validate_r5(lang: str) -> List[str]:
    errs: List[str] = []
    top5 = resp_get("top5_domains", [])
    selected_by_domain: Dict[str, List[str]] = resp_get("selected_by_domain", {})
    if not isinstance(selected_by_domain, dict):
        selected_by_domain = {}

    all_stats: List[str] = []
    for d in top5:
        stats = selected_by_domain.get(d, [])
        if not isinstance(stats, list):
            stats = []
        if len(stats) < 1:
            errs.append(t(lang, f"Rubrique 5 : choisissez au moins 1 statistique pour {d}.",
                          f"Section 5: select at least 1 indicator for {d}."))
        if len(stats) > 3:
            errs.append(t(lang, f"Rubrique 5 : maximum 3 statistiques pour {d}.",
                          f"Section 5: maximum 3 indicators for {d}."))
        all_stats.extend(stats)

    if len(all_stats) < 5 or len(all_stats) > 15:
        errs.append(t(lang, "Rubrique 5 : le total des statistiques doit être entre 5 et 15.",
                      "Section 5: total number of indicators must be between 5 and 15."))

    if len(set(all_stats)) != len(all_stats):
        errs.append(t(lang, "Rubrique 5 : une même statistique ne doit pas être sélectionnée plusieurs fois.",
                      "Section 5: the same indicator must not be selected more than once."))

    # scoring
    scoring: Dict[str, Dict[str, Any]] = resp_get("scoring", {})
    if not isinstance(scoring, dict):
        scoring = {}

    for s in all_stats:
        if s not in scoring:
            errs.append(t(lang, f"Rubrique 5 : vous devez noter la statistique {s}.",
                          f"Section 5: you must score indicator {s}."))
            continue
        for k in ["demand", "availability", "feasibility"]:
            sc_row = scoring.get(s, {}) or {}
            k_lbl = {
                "demand": t(lang, "demande", "demand"),
                "availability": t(lang, "disponibilité", "availability"),
                "feasibility": t(lang, "faisabilité", "feasibility"),
            }.get(k, k)

            # Backward compatibility: legacy key "gap" -> "availability"
            if k == "availability":
                v_raw = sc_row.get("availability", sc_row.get("gap", None))
            else:
                v_raw = sc_row.get(k, None)

            if v_raw is None or str(v_raw).strip() == "":
                errs.append(t(lang, f"Rubrique 5 : la note '{k_lbl}' manque pour {s}.",
                              f"Section 5: missing score '{k_lbl}' for {s}."))
            else:
                try:
                    v = int(v_raw)
                    if v < 0 or v > 3:
                        errs.append(t(lang, f"Rubrique 5 : note invalide pour {s} ({k_lbl}).",
                                      f"Section 5: invalid score for {s} ({k_lbl})."))
                except Exception:
                    errs.append(t(lang, f"Rubrique 5 : note invalide pour {s} ({k_lbl}).",
                                  f"Section 5: invalid score for {s} ({k_lbl})."))
    return errs


def validate_r6(lang: str) -> List[str]:
    errs: List[str] = []
    tbl = resp_get("gender_table", {})
    if not isinstance(tbl, dict) or not tbl:
        errs.append(t(lang, "Rubrique 6 : veuillez renseigner le tableau.", "Section 6: please complete the table."))
        return errs
    for k, v in tbl.items():
        if not v:
            errs.append(t(lang, f"Rubrique 6 : ligne non renseignée : {k}.", f"Section 6: missing answer for: {k}."))
    return errs


def validate_r8(lang: str) -> List[str]:
    errs: List[str] = []

    # Tableau capacité / faisabilité
    tbl = resp_get("capacity_table", {})
    if not isinstance(tbl, dict) or not tbl:
        errs.append(t(lang, "Rubrique 8 : veuillez renseigner le tableau.", "Section 8: please complete the table."))
        return errs
    for k, v in tbl.items():
        if not v:
            errs.append(t(lang, f"Rubrique 8 : ligne non renseignée : {k}.", f"Section 8: missing answer for: {k}."))

    return errs


def validate_r9(lang: str) -> List[str]:
    errs: List[str] = []
    sel = resp_get("quality_expectations", [])
    if not isinstance(sel, list) or len([x for x in sel if str(x).strip() != ""]) == 0:
        errs.append(t(lang, "Rubrique 9 : veuillez sélectionner au moins une option.", "Section 9: please select at least one option."))
        return errs
    if has_other_option(sel):
        other = (resp_get("quality_other", "") or "").strip()
        if not other:
            errs.append(t(lang, "Rubrique 9 : précisez l’option « Autre ».", "Section 9: please specify the \"Other\" option."))
    return errs


def validate_r10(lang: str) -> List[str]:
    errs: List[str] = []
    sel = resp_get("dissemination_channels", [])
    if not isinstance(sel, list) or len([x for x in sel if str(x).strip() != ""]) == 0:
        errs.append(t(lang, "Rubrique 10 : veuillez sélectionner au moins une option.", "Section 10: please select at least one option."))
        return errs
    if has_other_option(sel):
        other = (resp_get("dissemination_other", "") or "").strip()
        if not other:
            errs.append(t(lang, "Rubrique 10 : précisez l’option « Autre ».", "Section 10: please specify the \"Other\" option."))
    return errs


def validate_r11(lang: str) -> List[str]:
    errs: List[str] = []
    sel = resp_get("data_sources", [])
    if not isinstance(sel, list):
        sel = []

    sel_clean = [str(x).strip() for x in sel if str(x).strip()]
    if len(sel_clean) < 2:
        errs.append(t(lang, "Rubrique 11 : sélectionnez au moins 2 sources.", "Section 11: please select at least 2 sources."))
        return errs
    if len(sel_clean) > 4:
        errs.append(t(lang, "Rubrique 11 : sélectionnez au maximum 4 sources.", "Section 11: please select at most 4 sources."))

    if has_other_option(sel_clean):
        other = (resp_get("data_sources_other", "") or "").strip()
        if not other:
            errs.append(t(lang, "Rubrique 11 : précisez l’option « Autres ».", "Section 11: please specify the 'Other' option."))

    return errs


def validate_r12(lang: str) -> List[str]:
    errs: List[str] = []
    sub = int(st.session_state.get("r12_substep", 0) or 0)
    if sub < 3:
        errs.append(
            t(
                lang,
                "Rubrique 12 : veuillez traiter les questions ouvertes une à une (bouton « Question suivante ») jusqu’à la Confirmation.",
                "Section 12: please go through the open questions one by one (use the “Next question” button) until Confirmation.",
            )
        )

    cc = (resp_get("consulted_colleagues", "") or "").strip()
    if cc not in ("YES", "NO"):
        errs.append(
            t(
                lang,
                "Rubrique 12 : veuillez indiquer si vous avez consulté d’autres collègues (Oui/Non).",
                "Section 12: please indicate whether you consulted other colleagues (Yes/No).",
            )
        )
    return errs

def validate_all(lang: str) -> List[str]:
    errs = []
    errs.extend(validate_r2(lang))
    errs.extend(validate_r3(lang))
    errs.extend(validate_r4(lang))
    errs.extend(validate_r5(lang))
    errs.extend(validate_r6(lang))
    errs.extend(validate_r8(lang))
    errs.extend(validate_r9(lang))
    errs.extend(validate_r10(lang))
    errs.extend(validate_r11(lang))
    errs.extend(validate_r12(lang))
    # Open questions (text fields) remain optional; warnings are shown in Section 12 / Submit.
    return errs


# =========================
# Navigation
# =========================

def get_steps(lang: str) -> List[Tuple[str, str]]:
    # Rubric 7 added, plus final SEND tab
    return [
        ("R1", t(lang, "Rubrique 1 : Instructions", "Section 1: Instructions")),
        ("R2", t(lang, "Rubrique 2 : Identification du répondant", "Section 2: Respondent identification")),
        ("R3", t(lang, "Rubrique 3 : Portée de la réponse", "Section 3: Scope of response")),
        ("R4", t(lang, "Rubrique 4 : Domaines prioritaires", "Section 4: Priority domains")),
        ("R5", t(lang, "Rubrique 5 : Statistiques prioritaires et notation", "Section 5: Priority indicators and scoring")),
        ("R6", t(lang, "Rubrique 6 : Dimension genre", "Section 6: Gender dimension")),
        ("R7", t(lang, "Rubrique 7 : Priorités genre", "Section 7: Gender priorities")),
        ("R8", t(lang, "Rubrique 8 : Capacité et faisabilité (12–24 mois)", "Section 8: Capacity and feasibility (12–24 months)")),
        ("R9", t(lang, "Rubrique 9 : Harmonisation et qualité", "Section 9: Harmonization and quality")),
        ("R10", t(lang, "Rubrique 10 : Diffusion", "Section 10: Dissemination")),
        ("R11", t(lang, "Rubrique 11 : Sources de données pertinentes", "Section 11: Relevant data sources")),
        ("R12", t(lang, "Rubrique 12 : Questions ouvertes", "Section 12: Open questions")),
        ("SEND", t(lang, "ENVOYER", "SUBMIT")),
    ]


def render_sidebar(lang: str, steps: List[Tuple[str, str]]) -> None:
    st.sidebar.header(t(lang, "Navigation", "Navigation"))
    labels = [s[1] for s in steps]

    # Keep sidebar selection in sync with nav_idx
    st.session_state.nav_radio = int(st.session_state.nav_idx)
    chosen = st.sidebar.radio(
        t(lang, "Aller à", "Go to"),
        options=list(range(len(labels))),
        index=int(st.session_state.nav_idx),
        format_func=lambda i: labels[i],
        key="nav_radio"
    )

    # User clicked in sidebar
    if int(chosen) != int(st.session_state.nav_idx):
        st.session_state.nav_idx = int(chosen)

    st.sidebar.divider()
    st.sidebar.caption(
        t(
            lang,
            "Note : les contrôles qualité peuvent bloquer la progression si une contrainte n’est pas respectée.",
            "Note: quality checks may prevent moving forward when constraints are not met."
        )
    )

    st.sidebar.markdown("---")
    st.sidebar.caption(
        t(
            lang,
            "NSP : Ne sait pas (score 0). Utilisez NSP uniquement si l’information est indisponible.",
            "UK: Unknown (score 0). Use UK only when information is unavailable."
        )
    )
    st.sidebar.markdown("---")
    st.sidebar.subheader(t(lang, "Brouillon", "Draft"))
    if st.sidebar.button(t(lang, "Sauvegarder maintenant", "Save now")):
        ok, msg = autosave_draft(force=True)
        if ok:
            st.sidebar.success(t(lang, "Brouillon sauvegardé.", "Draft saved."))
        else:
            st.sidebar.error(t(lang, "Brouillon non sauvegardé.", "Draft not saved."))
    if st.session_state.get("draft_id"):
        st.sidebar.caption(
            t(
                lang,
                "Reprise : conservez l’URL de cette page (paramètre rid=...).",
                "Resume: keep this page URL (rid=... parameter)."
            )
        )



def nav_buttons(lang: str, steps: List[Tuple[str, str]], df_long: pd.DataFrame) -> None:
    """Bottom nav buttons, with blocking based on current step validations."""
    step_key = steps[st.session_state.nav_idx][0]
    errors: List[str] = []

    # Blocking rules per step
    if step_key == "R2":
        errors = validate_r2(lang)
    elif step_key == "R3":
        errors = validate_r3(lang)
    elif step_key == "R4":
        errors = validate_r4(lang)
    elif step_key == "R5":
        errors = validate_r5(lang)
    elif step_key == "R6":
        errors = validate_r6(lang)
    elif step_key == "R7":
        errors = validate_r7(lang)
    elif step_key == "R8":
        errors = validate_r8(lang)
    elif step_key == "R9":
        errors = validate_r9(lang)
    elif step_key == "R10":
        errors = validate_r10(lang)
    elif step_key == "R11":
        errors = validate_r11(lang)
    elif step_key == "R12":
        errors = validate_r12(lang)

    col1, col2, col3 = st.columns([1, 1, 3])
    with col1:
        prev_disabled = st.session_state.nav_idx <= 0
        if st.button(t(lang, "⬅ Précédent", "⬅ Previous"), disabled=prev_disabled):
            autosave_draft(force=True)
            st.session_state.nav_idx = max(0, st.session_state.nav_idx - 1)
            st.rerun()
    with col2:
        next_disabled = (st.session_state.nav_idx >= len(steps) - 1) or bool(errors)
        if st.button(t(lang, "Suivant ➡", "Next ➡"), disabled=next_disabled):
            autosave_draft(force=True)
            st.session_state.nav_idx = min(len(steps) - 1, st.session_state.nav_idx + 1)
            st.rerun()
    with col3:
        if errors:
            st.error("\n".join(errors))


# =========================
# UI : Rubrics
# =========================

def rubric_1(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 1 : Instructions", "Section 1: Instructions", "Secção 1: Instruções", "القسم 1: التعليمات"))
    st.markdown(
        t(
            lang,
            """
### Objectif
Ce questionnaire vise à recueillir votre avis sur **les statistiques socio-économiques prioritaires** à produire et diffuser au niveau continental.

### Comment répondre
1. **Identifiez** votre organisation (Rubrique 2).
2. **Pré-sélectionnez 5 à 10 domaines** et classez un **TOP 5** (Rubrique 4).
3. Pour chaque domaine du TOP 5 : choisissez **1 à 3 statistiques** et attribuez des **notes** (Rubrique 5).
4. Complétez les rubriques transversales : **genre**, **capacité/faisabilité**, **etc.**.
5. **N'hésitez pas à consulter les infobulles ⍰ pour plus de précisions.**

            """,
            """
### Purpose
This questionnaire collects your views on **priority socio-economic statistics** to be produced and disseminated at continental level.

### How to answer
1. **Identify** your organization (Section 2).
2. **Pre-select 5–10 domains** and rank a **TOP 5** (Section 4).
3. For each TOP 5 domain: select **1–3 indicators** and provide **scores** (Section 5).
4. Complete cross-cutting sections: **gender**, **capacity/feasibility**,  **etc.**.
5. **Feel free to consult the ⍰ tooltips for more details.**

            """,
            """
### Objetivo
Este questionário recolhe a sua opinião sobre **as estatísticas socioeconómicas prioritárias** a produzir e divulgar ao nível continental.

### Como responder
1. **Identifique** a sua organização (Secção 2).
2. **Pré-selecione 5 a 10 domínios** e classifique um **TOP 5** (Secção 4).
3. Para cada domínio do TOP 5 : selecione **1 a 3 indicadores** e atribua **pontuações** (Secção 5).
4. Complete as secções transversais : **género**, **capacidade/viabilidade**, **etc.**.
5. **Consulte livremente as infobolhas ⍰ para obter mais pormenores.**

            """,
            """
### الهدف
يجمع هذا الاستبيان آراءكم بشأن **الإحصاءات الاجتماعية والاقتصادية ذات الأولوية** التي ينبغي إنتاجها ونشرها على المستوى القاري.

### كيفية الإجابة
1. **حدّدوا** مؤسستكم (القسم 2).
2. **اختاروا مسبقاً من 5 إلى 10 مجالات** ثم رتّبوا **أفضل 5** (القسم 4).
3. لكل مجال من المجالات الخمسة الأولى : اختاروا **1 إلى 3 مؤشرات** وقدّموا **درجات** التقييم (القسم 5).
4. استكملوا الأقسام الأفقية : **النوع الاجتماعي** و**القدرة/إمكانية التنفيذ** و**غير ذلك**.
5. **يمكنكم الرجوع إلى تلميحات ⍰ للحصول على مزيد من التفاصيل.**

            """
        )
    )


def rubric_2(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 2 : Identification du répondant", "Section 2: Respondent identification", "Secção 2: Identificação do respondente", "القسم 2: تعريف المجيب"))
    st.info(
        t(
            lang,
            "Merci de renseigner ces informations. Elles servent uniquement à l’analyse et ne seront pas publiées nominativement.",
            "Please provide these details. They are used for analysis and will not be published in a personally identifiable way."
        )
    )

    resp_set("lang", lang)

    st.text_input(t(lang, "Nom de l'organisation", "Organization Name"), key="org_input", value=resp_get("organisation", ""))
    resp_set("organisation", st.session_state.get("org_input", "").strip())
    st.caption(t(lang, "Merci d’indiquer le libellé complet (évitez le sigle seul).", "Please provide the full organization name (avoid acronym only)."))
    col1, col2 = st.columns(2)
    with col1:
        # Pays de résidence : liste déroulante (ISO3 + libellés FR/EN) + option Autre
        df_countries = load_countries()
        iso3_list, iso3_to_fr, iso3_to_en, iso3_to_pt, iso3_to_ar = country_maps(df_countries)

        prev_country = (resp_get("pays", "") or "").strip()
        prev_country_other = (resp_get("pays_autre", "") or "").strip()
        prev_iso3 = (prev_country.split("|", 1)[0].strip().upper() if "|" in prev_country else prev_country.strip().upper())

        # Compatibilité ascendante : si ancienne valeur libre non-ISO3, basculer sur "OTHER"
        if prev_iso3 and prev_iso3 not in iso3_list and prev_iso3 not in {"OTHER", "__OTHER__"}:
            if not prev_country_other:
                prev_country_other = prev_country
            prev_iso3 = "OTHER"
        elif prev_iso3 == "__OTHER__":
            prev_iso3 = "OTHER"

        if not iso3_list:
            # Fallback si le fichier pays est introuvable : champ libre
            st.text_input(t(lang, "Pays de résidence", "Country of residence"), key="country_input", value=resp_get("pays", ""))
            resp_set("pays", st.session_state.get("country_input", "").strip())
            resp_set("pays_autre", "")
            resp_set("pays_name_fr", "")
            resp_set("pays_name_en", "")
        else:
            options = [""] + sorted(iso3_list, key=lambda x: country_label(x, lang, iso3_to_fr, iso3_to_en, iso3_to_pt, iso3_to_ar).lower()) + ["OTHER"]

            chosen_iso3 = st.selectbox(
                t(lang, "Pays de résidence", "Country of residence"),
                options=options,
                index=options.index(prev_iso3) if prev_iso3 in options else 0,
                format_func=lambda x: (
                    t(lang, "— Sélectionner —", "— Select —") if x == ""
                    else t(lang, "Autre pays (à préciser)", "Other country (please specify)") if x == "OTHER"
                    else f"{country_label(x, lang, iso3_to_fr, iso3_to_en, iso3_to_pt, iso3_to_ar)} ({x})"
                ),
                help=t(lang, "Choisissez votre pays de résidence (liste ISO3) ou ‘Autre pays (à préciser)’.",
                       "Select your country of residence (ISO3 list) or ‘Other country (please specify)’."),
                key="country_iso3_select",
            )
            resp_set("pays", chosen_iso3)

            if chosen_iso3 == "OTHER":
                other_country = st.text_input(
                    t(lang, "Autre pays (à préciser)", "Other country (please specify)"),
                    key="country_other_input",
                    value=prev_country_other,
                )
                other_country = (other_country or "").strip()
                resp_set("pays_autre", other_country)
                # On remplit aussi les libellés pour l'export, sans forcer une ISO3
                resp_set("pays_name_fr", other_country)
                resp_set("pays_name_en", other_country)
                resp_set("pays_name_pt", other_country)
                resp_set("pays_name_ar", other_country)
            elif chosen_iso3:
                resp_set("pays_autre", "")
                # Libellés normalisés (utile pour les exports / analyses)
                resp_set("pays_name_fr", iso3_to_fr.get(chosen_iso3, "").strip())
                resp_set("pays_name_en", (iso3_to_en.get(chosen_iso3, "") or iso3_to_fr.get(chosen_iso3, "")).strip())
                resp_set("pays_name_pt", (iso3_to_pt.get(chosen_iso3, "") or iso3_to_en.get(chosen_iso3, "") or iso3_to_fr.get(chosen_iso3, "")).strip())
                resp_set("pays_name_ar", (iso3_to_ar.get(chosen_iso3, "") or iso3_to_en.get(chosen_iso3, "") or iso3_to_fr.get(chosen_iso3, "")).strip())
            else:
                resp_set("pays_autre", "")
                resp_set("pays_name_fr", "")
                resp_set("pays_name_en", "")
                resp_set("pays_name_pt", "")
                resp_set("pays_name_ar", "")

    with col2:
        st.text_input(
            t(lang, "Email", "Email"),
            key="email_input",
            value=resp_get("email", ""),
            help=t(
                lang,
                "Saisissez une adresse email valide (ex. nom@domaine.tld).",
                "Enter a valid email address (e.g., name@domain.tld).",
            ),
        )
        resp_set("email", st.session_state.get("email_input", "").strip())

    # Brouillon : crée un identifiant de reprise dès que l’email est renseigné
    ensure_draft_id()
    autosave_draft(force=False)

    # Afficher le message de reprise dès la première session après saisie de l’email
    _email_now = resp_get("email", "")
    if _email_now and ("@" in _email_now) and not st.session_state.get("draft_resume_notice_shown"):
        st.warning(
            t(
                lang,
                "La saisie est sauvegardée. En cas de suspension de moins de 48 heures, reprenez là où vous vous étiez arrêté en ré-ouvrant le lien contenant rid (à conserver / mettre en favori / retrouver dans l'historique).",
                "Your input is saved. If you pause for less than 48 hours, resume where you left off by reopening the link containing rid (keep it / bookmark it / find it in your browser history).",
            )
        )
        st.session_state["draft_resume_notice_shown"] = True
    type_options = [
        ("NSO", {"fr": "Institut national de statistique", "en": "National Statistical Office"}),
        ("Ministry", {"fr": "Ministère / Service statistique sectoriel", "en": "Ministry / Sector statistical unit"}),
        ("REC", {"fr": "Communauté économique régionale", "en": "Regional Economic Community"}),
        ("AU", {"fr": "Union Africaine (UA)", "en": "African Union (AU)"}),
        ("CivilSoc", {"fr": "Société civile", "en": "Civil society"}),
        ("DevPartner", {"fr": "Partenaire technique et financier", "en": "Development partner"}),
        ("Academia", {"fr": "Université / Recherche", "en": "Academia / Research"}),
        ("Other", {"fr": "Autre", "en": "Other"}),
    ]
    type_labels = [t(lang, x[1]["fr"], x[1]["en"]) for x in type_options]
    type_keys = [x[0] for x in type_options]

    # Type d’acteur : pas de pré-remplissage (placeholder)
    type_opts = [""] + type_keys
    prev_type = resp_get("type_acteur", "")
    idx = type_opts.index(prev_type) if prev_type in type_opts else 0

    chosen_type = st.selectbox(
        t(lang, "Type d’acteur", "Stakeholder type"),
        options=type_opts,
        index=idx,
        format_func=lambda k: (t(lang, "— Sélectionner —", "— Select —") if k == "" else type_labels[type_keys.index(k)]),
        help=t(lang, "Choisissez la catégorie correspondant le mieux à votre organisation.", 
               "Choose the category that best matches your organization.")
    )
    resp_set("type_acteur", chosen_type)
# Fonction dropdown : pas de pré-remplissage (placeholder)
    role_opts = get_role_options(lang)
    role_options = [""] + role_opts
    prev_role = resp_get("fonction", "")
    role_idx = role_options.index(prev_role) if prev_role in role_options else 0

    chosen_role = st.selectbox(
        t(lang, "Fonction", "Role/Function"),
        options=role_options,
        index=role_idx,
        format_func=lambda x: (t(lang, "— Sélectionner —", "— Select —") if x == "" else x),
        help=t(lang, "Indiquez votre fonction principale dans l’organisation.", "Indicate your main role in the organization."),
    )
    resp_set("fonction", chosen_role)
    if is_other_value(chosen_role):
        st.text_input(t(lang, "Préciser (fonction)", "Specify (role)"),
                      key="fonction_autre_input", value=resp_get("fonction_autre", ""))
        resp_set("fonction_autre", st.session_state.get("fonction_autre_input", "").strip())
    else:
        resp_set("fonction_autre", "")

    # Live errors
    errs = validate_r2(lang)
    if errs:
        st.warning(t(lang, "Veuillez corriger les éléments ci-dessous :", "Please fix the following:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def rubric_3(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 3 : Portée de la réponse", "Section 3: Scope of response", "Secção 3: Âmbito da resposta", "القسم 3: نطاق الإجابة"))
    st.markdown(
        t(
            lang,
            "Indiquez le périmètre principal de votre réponse. Cela aide à interpréter vos priorités.",
            "Indicate the main scope of your response. This helps interpret your priorities.",
            "Indique o âmbito principal da sua resposta. Isto ajuda a interpretar as suas prioridades.",
            "حددوا النطاق الرئيسي لإجابتكم. فهذا يساعد على تفسير أولوياتكم."
        )
    )

    scope_opts_raw = [
        ("National", {"fr": "National", "en": "National"}),
        ("Regional", {"fr": "Régional (CER)", "en": "Regional (REC)"}),
        ("Continental", {"fr": "Continental (UA)", "en": "Continental (AU)"}),
        ("Global", {"fr": "International", "en": "International"}),
        ("Other", {"fr": "Autre", "en": "Other"}),
    ]
    scope_labels = {k: t(lang, v["fr"], v["en"]) for k, v in scope_opts_raw}
    scope_keys = [k for k, _ in scope_opts_raw]
    scope_options = [""] + scope_keys

    prev_scope = resp_get("scope", "")
    scope_idx = scope_options.index(prev_scope) if prev_scope in scope_options else 0

    chosen_scope = st.selectbox(
        t(lang, "Portée", "Scope"),
        options=scope_options,
        index=scope_idx,
        format_func=lambda k: (t(lang, "— Sélectionner —", "— Select —") if k == "" else scope_labels.get(k, k)),
        help=t(
            lang,
            "Indiquez le périmètre principal de votre réponse : national, régional (CER), continental (UA) ou international.",
            "Indicate the main scope of your response: national, regional (REC), continental (AU), or international."
        )
    )
    resp_set("scope", chosen_scope)

    # SNDS / Plan statistique national (obligatoire)
    snds_opts = ["", "YES", "NO", "PREP", "IMPL_PREP", "NSP"]
    snds_labels = {
        "YES": t(lang, "Oui", "Yes"),
        "NO": t(lang, "Non", "No"),
        "PREP": t(lang, "En préparation", "In preparation"),
        "IMPL_PREP": t(lang, "En cours de mise en œuvre ET nouvelle en préparation", "Under implementation AND new one in preparation"),
        "NSP": t(lang, "NSP", "DK"),
    }
    prev_snds = (resp_get("snds_status", "") or "").strip()
    idx_snds = snds_opts.index(prev_snds) if prev_snds in snds_opts else 0
    chosen_snds = st.selectbox(
        t(
            lang,
            "Statut de la SNDS / plan national statistique en cours",
            "Status of the current NSDS / national statistical plan",
        ),
        options=snds_opts,
        index=idx_snds,
        format_func=lambda k: (
            t(lang, "— Sélectionner —", "— Select —") if k == "" else snds_labels.get(k, k)
        ),
        help=t(
            lang,
            "Indiquez si une stratégie / plan statistique national est en cours, non, en préparation, ou NSP.",
            "Indicate whether an NSDS / national statistical plan is current, not in place, under preparation, or DK.",
        ),
        key="snds_status_select",
    )
    resp_set("snds_status", chosen_snds)


    # Contrôle qualité (alerte) : cohérence acteur × portée
    _actor = (resp_get("type_acteur", "") or "").strip()
    _scope = (resp_get("scope", "") or "").strip()
    if _actor in ["NSO", "Ministry"] and _scope and _scope != "National":
        st.warning(
            t(
                lang,
                "Alerte : vous avez indiqué « Institut national de statistique » ou « Ministère », mais la portée n’est pas « National ». Merci de vérifier la cohérence.",
                "Warning: you selected “National Statistical Office” or “Ministry”, but the scope is not “National”. Please check consistency."
            )
        )

    if resp_get("scope") == "Other":
        st.text_input(t(lang, "Préciser", "Specify"), key="scope_other_input", value=resp_get("scope_other", ""))
        resp_set("scope_other", st.session_state.get("scope_other_input", "").strip())
    else:
        resp_set("scope_other", "")



def rubric_4(lang: str, df_long: pd.DataFrame) -> None:
    st.subheader(t(lang, "Rubrique 4 : Domaines prioritaires", "Section 4: Priority domains", "Secção 4: Domínios prioritários", "القسم 4: المجالات ذات الأولوية"))

    st.info(
        t(
            lang,
            "Veuillez d’abord choisir 5 à 10 domaines (pré-sélection). Ensuite, choisissez exactement 5 domaines dans ce sous-ensemble (TOP 5).\n\nConseil : choisissez les domaines où la demande politique est forte.",
            "First select 5 to 10 domains (pre-selection). Then choose exactly 5 domains within that subset (TOP 5).\n\nTip: choose domains where policy demand is strong."
        )
    )

    domains = domains_from_longlist(df_long, lang)
    if not domains:
        st.error(
            t(
                lang,
                "La liste des domaines n’est pas disponible (longlist introuvable ou vide).",
                "Domain list is not available (longlist not found or empty).",
            )
        )
        st.caption(
            t(
                lang,
                "Vérifiez que le dépôt contient : data/indicator_longlist.csv (prioritaire) ou data/longlist.xlsx (ou ces fichiers à la racine).",
                "Check that the repository contains: data/indicator_longlist.csv (preferred) or data/longlist.xlsx (or these files at repo root).",
            )
        )
        return

    code_to_label = {c: lbl for c, lbl in domains}

    # Build display labels without showing codes (codes are stored internally)
    labels = [code_to_label[c] for c, _ in domains]
    # Disambiguate duplicates if any (rare)
    seen = {}
    for i, (c, _) in enumerate(domains):
        lbl = code_to_label[c]
        seen[lbl] = seen.get(lbl, 0) + 1
    display_labels = []
    label_to_code = {}
    for c, _ in domains:
        lbl = code_to_label[c]
        disp = lbl if seen[lbl] == 1 else f"{lbl} ({c})"
        display_labels.append(disp)
        label_to_code[disp] = c

    st.markdown(
        t(
            lang,
            """
### Étape 1 : Pré-sélection
Sélectionnez **entre 5 et 10 domaines** (sans doublons).
            """,
            """
### Step 1: Pre-selection
Select **5 to 10 domains** (no duplicates).
            """
        )
    )

    pre_default_codes = resp_get("preselected_domains", [])
    pre_default_disp = []
    for c in pre_default_codes:
        lbl = code_to_label.get(c, "")
        if not lbl:
            continue
        disp = lbl if seen.get(lbl, 1) == 1 else f"{lbl} ({c})"
        if disp in label_to_code:
            pre_default_disp.append(disp)

    # Avoid "first click not kept" by initializing widget state once (no default on every rerun)
    if "r4_preselection_ms" not in st.session_state:
        st.session_state["r4_preselection_ms"] = pre_default_disp

    pre_disp = st.multiselect(
        t(lang, "Pré-sélection (5–10 domaines)", "Pre-selection (5–10 domains)"),
        options=display_labels,
        max_selections=10,
        key="r4_preselection_ms",
        help=t(lang, "Choisissez au maximum 10 domaines. Une fois 10 domaines sélectionnés, les nouveaux clics seront ignorés.", 
               "Select up to 10 domains. Once 10 domains are selected, additional clicks are ignored.")
    )

    pre_codes = [label_to_code[x] for x in pre_disp]
    resp_set("preselected_domains", pre_codes)

    st.divider()
    st.markdown(
        t(
            lang,
            """
### Étape 2 : Classement TOP 5
Classez exactement **5 domaines** parmi votre pré-sélection.
            """,
            """
### Step 2: Rank TOP 5
Rank exactly **5 domains** from your pre-selection.
            """
        )
    )

    if len(pre_codes) < 5:
        st.warning(t(lang, "Sélectionnez d’abord au moins 5 domaines dans la pré-sélection.",
                     "Please pre-select at least 5 domains first."))
        resp_set("top5_domains", [])
        return

    top5: List[str] = []

    # Ranking with 5 selectboxes (no prefill + no duplicates)
    chosen_prev: List[str] = []
    for i in range(5):
        key = f"top5_rank_{i+1}"

        # Options for this rank = preselection minus already chosen
        remaining = [c for c in pre_codes if c not in chosen_prev]
        options = [""] + remaining  # "" placeholder (no prefill)

        prev = resp_get(key, "")
        if prev and prev in remaining:
            idx = options.index(prev)
        else:
            idx = 0

        choice = st.selectbox(
            t(lang, f"Rang {i+1}", f"Rank {i+1}"),
            options=options,
            index=idx,
            format_func=lambda c: (t(lang, "— Sélectionner —", "— Select —") if c == "" else code_to_label.get(c, c)),
            help=t(
                lang,
                "Choisissez un domaine unique pour chaque rang. Les domaines déjà choisis ne sont plus proposés aux rangs suivants.",
                "Choose a unique domain for each rank. Already selected domains are removed from the next ranks.",
            ),
            key=key,
        )

        if choice != "":
            top5.append(choice)
            chosen_prev.append(choice)

    resp_set("top5_domains", top5)


    errs = validate_r4(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))



def rubric_5(lang: str, df_long: pd.DataFrame) -> None:
    st.subheader(t(lang, "Rubrique 5 : Statistiques prioritaires et notation", "Section 5: Priority indicators and scoring", "Secção 5: Estatísticas prioritárias e pontuação", "القسم 5: الإحصاءات ذات الأولوية والتقييم"))

    top5 = resp_get("top5_domains", [])
    if not top5 or len(top5) != 5:
        st.warning(t(lang, "Veuillez d’abord finaliser le TOP 5 des domaines (Rubrique 4).",
                     "Please complete TOP 5 domains first (Section 4)."))
        return

    # mapping for domain display
    dom_map = {c: lbl for c, lbl in domains_from_longlist(df_long, lang)}

    st.markdown(
        t(
            lang,
            "La rubrique se remplit en deux temps : (A) sélection des statistiques, puis (B) notation multicritères.",
            "This section is completed in two steps: (A) indicator selection, then (B) multi-criteria scoring.",
            "Esta secção é preenchida em duas etapas : (A) seleção dos indicadores, depois (B) pontuação multicritério.",
            "يُستكمل هذا القسم على مرحلتين : (أ) اختيار المؤشرات، ثم (ب) التقييم متعدد المعايير."
        )
    )

    selected_by_domain: Dict[str, List[str]] = resp_get("selected_by_domain", {})
    if not isinstance(selected_by_domain, dict):
        selected_by_domain = {}

    scoring: Dict[str, Dict[str, Any]] = resp_get("scoring", {})
    if not isinstance(scoring, dict):
        scoring = {}

    # Ensure dict keys exist
    for d in top5:
        if d not in selected_by_domain:
            selected_by_domain[d] = []

    st.markdown("### " + t(lang, "Étape A : Sélection des statistiques", "Step A: Select indicators"))
    st.caption(
        t(
            lang,
            "Pour chaque domaine du TOP 5, choisissez 1 à 3 statistiques (total attendu : 5 à 15). Une statistique ne doit pas apparaître dans deux domaines.",
            "For each TOP 5 domain, select 1 to 3 indicators (expected total: 5 to 15). The same indicator must not appear under two domains.",
        )
    )

    # UI selection per domain (codes hidden)
    for d in top5:
        st.markdown(f"#### {dom_map.get(d, d)}")

        stats_opts = stats_for_domain(df_long, d, lang)
        stat_code_to_label = {c: lbl for c, lbl in stats_opts}

        # build display labels without showing stat codes
        labels = [stat_code_to_label[c] for c, _ in stats_opts]
        seen = {}
        for c, _ in stats_opts:
            lbl = stat_code_to_label[c]
            seen[lbl] = seen.get(lbl, 0) + 1
        display_labels = []
        label_to_code = {}
        for c, _ in stats_opts:
            lbl = stat_code_to_label[c]
            disp = lbl if seen[lbl] == 1 else f"{lbl} ({c})"
            display_labels.append(disp)
            label_to_code[disp] = c

        default_codes = selected_by_domain.get(d, [])
        default_disp = []
        for c in default_codes:
            lbl = stat_code_to_label.get(c, "")
            if not lbl:
                continue
            disp = lbl if seen.get(lbl, 1) == 1 else f"{lbl} ({c})"
            if disp in label_to_code:
                default_disp.append(disp)

        key_ms = f"stats_ms_{d}"

        # Init widget state once (avoid "first click" issues)
        if key_ms not in st.session_state:
            st.session_state[key_ms] = default_disp

        picked_disp = st.multiselect(
            t(lang, "Choisir 1 à 3 statistiques", "Select 1 to 3 indicators"),
            options=display_labels,
            max_selections=3,
            key=key_ms,
            help=t(lang, "Sélectionnez au minimum 1 et au maximum 3 statistiques pour ce domaine.",
                   "Select at least 1 and at most 3 indicators for this domain.")
        )

        picked_codes = [label_to_code[x] for x in picked_disp]
        selected_by_domain[d] = picked_codes


    # Uniqueness check
    flattened = []
    for d in top5:
        flattened.extend(selected_by_domain.get(d, []))
    duplicates = [x for x in set(flattened) if flattened.count(x) > 1]
    if duplicates:
        st.error(
            t(
                lang,
                "Une ou plusieurs statistiques sont sélectionnées dans plusieurs domaines. Veuillez corriger.",
                "One or more indicators are selected under multiple domains. Please correct."
            )
        )

    resp_set("selected_by_domain", selected_by_domain)
    resp_set("selected_stats", flattened)

    # Map codes to labels for display in scoring
    global_map = {}
    for d in top5:
        for c, lbl in stats_for_domain(df_long, d, lang):
            global_map[c] = lbl

    st.divider()
    st.markdown("### " + t(lang, "Étape B : Notation multicritères", "Step B: Multi-criteria scoring"))
    st.caption(
        t(
            lang,
            "Pour chaque statistique sélectionnée, attribuez une note (0–3) sur : demande politique, disponibilité actuelle (bonne = score plus élevé) et faisabilité à 12–24 mois.",
            "For each selected indicator, assign a score (0–3) for: political demand, current availability (good = higher score), and feasibility in 12–24 months.",
        )
    )

    for s in flattened:
        if s not in scoring or not isinstance(scoring.get(s), dict):
            scoring[s] = {}

        # Backward compatibility: legacy key "gap" (écart) -> "availability"
        # On normalise sur l'échelle v3 (Bonne=3)
        if "availability" not in scoring[s] and "gap" in scoring[s]:
            scoring[s]["availability"] = normalize_availability(scoring[s].get("gap", 0), 0)

        # Ensure keys exist
        for k in ["demand", "availability", "feasibility"]:
            if k not in scoring[s]:
                scoring[s][k] = None

        st.markdown(f"**{global_map.get(s, s)}**")

        c1, c2, c3 = st.columns(3)
        opts = [None, 1, 2, 3, 0]  # None = placeholder (no prefill). 0 = NSP / DK

        with c1:
            prev = scoring[s].get("demand", None)
            idx = 0
            try:
                if prev is not None and prev != "":
                    idx = opts.index(int(prev))
            except Exception:
                idx = 0
            scoring[s]["demand"] = st.selectbox(
                t(lang, "Demande politique", "Political demand"),
                options=opts,
                index=idx,
                format_func=score_format(lang, "demand"),
                help=t(
                    lang,
                    "Définition : importance de l’indicateur pour le pilotage des politiques publiques, la redevabilité et les priorités.",
                    "Definition: importance for steering public policies, accountability and priorities.",
                ),
                key=f"sc_dem_{s}",
            )

        with c2:
            prev = scoring[s].get("availability", None)
            idx = 0
            try:
                if prev is not None and prev != "":
                    idx = opts.index(int(prev))
            except Exception:
                idx = 0
            scoring[s]["availability"] = st.selectbox(
                t(lang, "Disponibilité actuelle", "Current availability"),
                options=opts,
                index=idx,
                format_func=score_format(lang, "availability"),
                help=t(
                    lang,
                    "Définition : l’indicateur est-il déjà produit régulièrement, avec une couverture et une qualité suffisantes, et sous une forme utilisable ? (Bonne = score plus élevé).",
                    "Definition: is the indicator already produced regularly with sufficient coverage and quality, in a usable form? (Good = higher score).",
                ),
                key=f"sc_avail_{s}",
            )

        with c3:
            prev = scoring[s].get("feasibility", None)
            idx = 0
            try:
                if prev is not None and prev != "":
                    idx = opts.index(int(prev))
            except Exception:
                idx = 0
            scoring[s]["feasibility"] = st.selectbox(
                t(lang, "Faisabilité (12–24 mois)", "Feasibility (12–24 months)"),
                options=opts,
                index=idx,
                format_func=score_format(lang, "feasibility"),
                help=t(
                    lang,
                    "Définition : capacité réaliste à améliorer ou produire l’indicateur d’ici 12–24 mois, compte tenu des sources, capacités et prérequis.",
                    "Definition: realistic ability to improve or produce the indicator within 12–24 months, considering sources, capacities and prerequisites.",
                ),
                key=f"sc_fea_{s}",
            )

    resp_set("scoring", scoring)

    errs = validate_r5(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def rubric_6(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 6 : Dimension genre", "Section 6: Gender dimension", "Secção 6: Dimensão de género", "القسم 6: البعد الجنساني"))
    st.markdown(
        t(
            lang,
            "Indiquez si les statistiques prioritaires doivent intégrer ces dimensions (Oui/Non/Selon indicateur/NSP).",
            "Indicate whether priority indicators should integrate these dimensions (Yes/No/Indicator-specific/UK)."
        )
    )

    options = [
        (t(lang, "Oui", "Yes"), "YES"),
        (t(lang, "Non", "No"), "NO"),
        (t(lang, "Selon indicateur", "Indicator-specific"), "SPEC"),
        (t(lang, UK_FR, UK_EN, UK_PT, UK_AR), "UK"),
    ]
    labels = [x[0] for x in options]
    code_map = {x[0]: x[1] for x in options}

    items_fr = [
        "Désagrégation par sexe",
        "Désagrégation par âge",
        "Milieu urbain / rural",
        "Handicap",
        "Quintile de richesse",
        "Violences basées sur le genre (VBG)",
        "Temps domestique non rémunéré",
    ]
    items_en = [
        "Disaggregation by sex",
        "Disaggregation by age",
        "Urban / rural",
        "Disability",
        "Wealth quintile",
        "Gender-based violence (GBV)",
        "Unpaid domestic work",
    ]
    items = items_fr if lang == "fr" else items_en

    tbl = resp_get("gender_table", {})
    if not isinstance(tbl, dict):
        tbl = {}

    for it in items:
        rev_map = {v: k for k, v in code_map.items()}
        prev_code = tbl.get(it, None)
        idx = labels.index(rev_map[prev_code]) if prev_code in rev_map else None
        chosen = st.radio(tr(lang, it), options=labels, index=idx, horizontal=True, key=f"gender_{it}")
        tbl[it] = code_map.get(chosen, None)

    resp_set("gender_table", tbl)

    errs = validate_r6(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def validate_r7(lang: str) -> List[str]:
    errs: List[str] = []
    p1 = (resp_get("gender_priority_1", "") or "").strip()
    p2 = (resp_get("gender_priority_2", "") or "").strip()
    p3 = (resp_get("gender_priority_3", "") or "").strip()
    other = (resp_get("gender_priority_other", "") or "").strip()

    if not p1:
        errs.append(
            t(
                lang,
                "Rubrique 7 : veuillez sélectionner au moins une priorité genre (Priorité 1).",
                "Section 7: please select at least one gender priority (Priority 1).",
            )
        )

    # No rank 3 without rank 2
    if p3 and not p2:
        errs.append(
            t(
                lang,
                "Rubrique 7 : veuillez renseigner la Priorité 2 avant la Priorité 3.",
                "Section 7: please fill Priority 2 before Priority 3.",
            )
        )

    # Uniqueness
    chosen = [x for x in [p1, p2, p3] if x]
    if len(set(chosen)) != len(chosen):
        errs.append(
            t(
                lang,
                "Rubrique 7 : les priorités genre doivent être différentes (pas de doublons).",
                "Section 7: gender priorities must be distinct (no duplicates).",
            )
        )

    # Other text required if OTHER selected
    if "OTHER" in chosen and not other:
        errs.append(
            t(
                lang,
                "Rubrique 7 : précisez l’option « Autre ».",
                "Section 7: please specify the 'Other' option.",
            )
        )

    return errs


def rubric_7(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 7 : Priorités genre", "Section 7: Gender priorities", "Secção 7: Prioridades de género", "القسم 7: أولويات النوع الاجتماعي"))
    st.markdown(
        t(
            lang,
            "Sélectionnez de 1 à 3 priorités genre en commençant par la plus importante.",
            "Select 1 to 3 gender priorities, starting with the most important.",
        )
    )

    gp_opts = ["", "ECO", "SERV", "GBV", "PART_DEC", "CARE", "OTHER"]
    gp_labels = {
        "ECO": t(lang, "Autonomisation économique", "Economic empowerment"),
        "SERV": t(lang, "Accès aux services", "Access to services"),
        "GBV": t(lang, "Violences basées sur le genre (VBG)", "Gender based violence (GBV)"),
        "PART_DEC": t(lang, "Participation aux instances décisionnelles", "Participation in decision-making bodies"),
        "CARE": t(lang, "Temps domestique non rémunéré", "Unpaid domestic and care work"),
        "OTHER": t(lang, "Autre", "Other"),
        "": t(lang, "— Sélectionner —", "— Select —"),
    }

    # Rank 1 (required)
    prev1 = (resp_get("gender_priority_1", "") or "").strip()
    idx1 = gp_opts.index(prev1) if prev1 in gp_opts else 0
    p1 = st.selectbox(
        t(lang, "Vos trois (3) priorités genre – Priorité 1 (obligatoire)", "Your three (3) gender priorities – Priority 1 (required)"),
        options=gp_opts,
        index=idx1,
        format_func=lambda k: gp_labels.get(k, k),
        key="gender_priority_1_select",
    )
    resp_set("gender_priority_1", p1)
    # Backward compatibility (previous single-priority field)
    resp_set("gender_priority_main", p1)

    # Rank 2 (optional), exclude already chosen (except empty)
    opts2 = [""] + [k for k in gp_opts if k not in ["", p1]]
    prev2 = (resp_get("gender_priority_2", "") or "").strip()
    idx2 = opts2.index(prev2) if prev2 in opts2 else 0
    p2 = st.selectbox(
        t(lang, "Priorité 2 (optionnelle)", "Priority 2 (optional)"),
        options=opts2,
        index=idx2,
        format_func=lambda k: gp_labels.get(k, k),
        key="gender_priority_2_select",
    )
    resp_set("gender_priority_2", p2)

    # Rank 3 (optional), exclude already chosen (except empty)
    opts3 = [""] + [k for k in gp_opts if k not in ["", p1, p2]]
    prev3 = (resp_get("gender_priority_3", "") or "").strip()
    idx3 = opts3.index(prev3) if prev3 in opts3 else 0
    p3 = st.selectbox(
        t(lang, "Priorité 3 (optionnelle)", "Priority 3 (optional)"),
        options=opts3,
        index=idx3,
        format_func=lambda k: gp_labels.get(k, k),
        key="gender_priority_3_select",
    )
    resp_set("gender_priority_3", p3)

    chosen_any = [x for x in [p1, p2, p3] if x]
    if "OTHER" in chosen_any:
        other = st.text_input(
            t(lang, "Autre : préciser", "Other: please specify"),
            key="gender_priority_other_input",
            value=resp_get("gender_priority_other", ""),
        )
        resp_set("gender_priority_other", (other or "").strip())
    else:
        resp_set("gender_priority_other", "")

    errs = validate_r7(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))

def rubric_8(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 8 : Capacité et faisabilité (12–24 mois)", "Section 8: Capacity and feasibility (12–24 months)", "Secção 8: Capacidade e viabilidade (12–24 meses)", "القسم 8: القدرات وإمكانية التنفيذ (12–24 شهراً)"))
    st.markdown(
        t(
            lang,
            "Évaluez le niveau de disponibilité et d’adéquation des moyens pour produire les statistiques prioritaires dans les 12–24 mois à venir.",
            "Assess the availability and adequacy of resources to produce priority statistics in the coming 12–24 months."
        )
    )

    scale = [
        (t(lang, "Élevé", "High"), "HIGH"),
        (t(lang, "Moyen", "Medium"), "MED"),
        (t(lang, "Faible", "Low"), "LOW"),
        (t(lang, UK_FR, UK_EN, UK_PT, UK_AR), "UK"),
    ]
    labels = [x[0] for x in scale]
    code_map = {x[0]: x[1] for x in scale}

    st.caption(t(lang, "Échelle : Élevé = capacité suffisante et opérationnelle ; Moyen = partiellement disponible ; Faible = insuffisant ; NSP = ne sait pas.",
                   "Scale: High = sufficient and operational; Medium = partially available; Low = insufficient; DK = does not know."))

    items_fr = [
        "Compétences statistiques disponibles",
        "Accès aux données administratives",
        "Financement disponible",
        "Outils numériques (collecte, traitement, diffusion)",
        "Cadre juridique pour le partage de données",
        "Coordination interinstitutionnelle",
    ]
    items_en = [
        "Available statistical skills",
        "Access to administrative data",
        "Available funding",
        "Digital tools (collection, processing, dissemination)",
        "Legal framework for data sharing",
        "Inter-institutional coordination",
    ]
    items = items_fr if lang == "fr" else items_en

    helps_fr = [
        "Ressources humaines : disponibilité de statisticiens/analystes qualifiés et expérience pertinente.",
        "Accès aux données administratives : disponibilité, qualité, régularité et conditions d’accès pour usage statistique.",
        "Financement : budget disponible et soutenable pour la production, y compris opérations de collecte/traitement.",
        "Outils numériques : disponibilité et adéquation des outils pour collecte, traitement, stockage, diffusion, interopérabilité (logiciels, matériel, connectivité, sécurité).",
        "Cadre juridique : existence et applicabilité des textes/accords permettant le partage de données à des fins statistiques (lois, décrets, protocoles, MoU, clauses de confidentialité).",
        "Coordination : mécanismes de coordination interinstitutionnelle (comités, conventions, échanges réguliers, standards communs).",
    ]
    helps_en = [
        "Human resources: availability of qualified statisticians/analysts and relevant experience.",
        "Access to administrative data: availability, quality, timeliness and conditions of access for statistical use.",
        "Funding: available and sustainable budget for production, including collection/processing operations.",
        "Digital tools: availability and adequacy of tools for collection, processing, storage, dissemination, interoperability (software, hardware, connectivity, security).",
        "Legal framework: existence and enforceability of texts/agreements enabling data sharing for statistical purposes (laws, decrees, protocols, MoUs, confidentiality clauses).",
        "Coordination: inter-institutional coordination mechanisms (committees, agreements, regular exchanges, shared standards).",
    ]
    helps = helps_fr if lang == "fr" else helps_en

    tbl = resp_get("capacity_table", {})
    if not isinstance(tbl, dict):
        tbl = {}

    rev_map = {v: k for k, v in code_map.items()}
    for it, hp in zip(items, helps):
        prev_code = tbl.get(it, None)
        idx = labels.index(rev_map[prev_code]) if prev_code in rev_map else None
        chosen = st.radio(tr(lang, it), options=labels, index=idx, horizontal=True, key=f"cap_{it}", help=tr(lang, hp))
        tbl[it] = code_map.get(chosen, None)

    resp_set("capacity_table", tbl)

    errs = validate_r8(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def rubric_9(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 9 : Harmonisation et qualité", "Section 9: Harmonization and quality", "Secção 9: Harmonização e qualidade", "القسم 9: المواءمة والجودة"))
    st.markdown(
        t(
            lang,
            "Indiquez 1 à 3 exigences attendues en matière d’harmonisation et d’assurance qualité.",
            "Indicate 1 to 3 expectations regarding harmonization and quality assurance."
        )
    )

    opts_fr = [
        "Manuels de normes et méthodes communes (par domaine) disponibles",
        "Cadre d’assurance qualité fonctionnel",
        "Procédures de validation et certification des données",
        "Mécanismes de cohérence des données nationales entre secteurs",
        "Renforcement des capacités techniques du SSN",
        "Renforcement du leadership de l’INS au sein du SSN",
        "Groupes techniques spécialisés (GTS/UA) opérationnels",
        "Autre (préciser) ",
     ]
    opts_en = [
        "Manuals on common standards and methods (by domain) available",
        "Functional quality assurance framework (quality toolkit) ",
        "Data validation and certification procedures (certified quality) ",
        "Toolkit / mechanisms for cross-sector consistency of national data",
        "Strengthening NSS technical capacity",
        "Strengthening NSO leadership within the NSS",
        "Specialized Technical Groups (STGs/AU) operational",
        "Other (specify) ",
    ]
    opts = opts_fr if lang == "fr" else opts_en
    # Stabilité mobile : initialiser le widget une seule fois
    if "r9_multiselect" not in st.session_state:
        st.session_state["r9_multiselect"] = resp_get("quality_expectations", [])
    sel = st.multiselect(t(lang, "Sélectionnez", "Select"), options=opts, key="r9_multiselect", max_selections=3, format_func=lambda x: tr(lang, x))
    resp_set("quality_expectations", sel)
    if has_other_option(sel):
        st.text_input(t(lang, "Préciser (Autre)", "Specify (Other)"),
                      key="q9_other_input", value=resp_get("quality_other", ""))
        resp_set("quality_other", st.session_state.get("q9_other_input", "").strip())
    else:
        resp_set("quality_other", "")

    errs = validate_r9(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))

    # Auto-save draft (mobile)
    autosave_draft(force=False)



def rubric_10(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 10 : Diffusion", "Section 10: Dissemination", "Secção 10: Disseminação", "القسم 10: النشر"))
    st.markdown(
        t(
            lang,
            "Indiquez 1 à 3 canaux de diffusion jugés les plus utiles pour les statistiques prioritaires.",
            "Indicate 1 to 3 dissemination channels you find most useful for priority statistics."
        )
    )
    opts_fr = [
        "Portail web / tableaux de bord",
        "Communiqués / notes de conjoncture",
        "Microdonnées anonymisées (accès sécurisé)",
        "API / Open data",
        "Ateliers et webinaires",
        "Autre",
    ]
    opts_en = [
        "Web portal / dashboards",
        "Press releases / bulletins",
        "Anonymized microdata (secure access)",
        "API / Open data",
        "Workshops and webinars",
        "Other",
    ]
    opts = opts_fr if lang == "fr" else opts_en
    # Éviter les problèmes de clic (init du widget une seule fois)
    if "r10_multiselect" not in st.session_state:
        st.session_state["r10_multiselect"] = resp_get("dissemination_channels", [])
    sel = st.multiselect(
        t(lang, "Sélectionnez", "Select"),
        options=opts,
        max_selections=3,
        key="r10_multiselect",
        format_func=lambda x: tr(lang, x),
        help=t(lang, "Choisissez les canaux de diffusion les plus utiles.", "Select the most useful dissemination channels.")
    )
    resp_set("dissemination_channels", sel)
    if has_other_option(sel):
        st.text_input(t(lang, "Préciser (Autre)", "Specify (Other)"),
                      key="q10_other_input", value=resp_get("dissemination_other", ""))
        resp_set("dissemination_other", st.session_state.get("q10_other_input", "").strip())
    else:
        resp_set("dissemination_other", "")

    errs = validate_r10(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))

    # Auto-save draft (mobile)
    autosave_draft(force=False)



def rubric_11(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 11 : Sources de données pertinentes", "Section 11: Relevant data sources", "Secção 11: Fontes de dados pertinentes", "القسم 11: مصادر البيانات ذات الصلة"))
    st.markdown(
        t(
            lang,
            "Sélectionnez **2 à 4** sources de données les plus importantes pour produire les statistiques prioritaires.",
            "Select **2 to 4** of the most important data sources to produce the priority statistics.",
        )
    )

    opts_fr = [
        "Enquêtes ménages",
        "Enquêtes entreprises",
        "Recensements",
        "Données administratives",
        "Registres état-civil",
        "Données géospatiales",
        "Données privées",
        "Autres",
    ]
    opts_en = [
        "Household surveys",
        "Enterprise surveys",
        "Censuses",
        "Administrative data",
        "Civil registration and vital statistics (CRVS)",
        "Geospatial data",
        "Private data",
        "Other",
    ]

    options = opts_fr if lang == "fr" else opts_en

    prev = resp_get("data_sources", [])
    if not isinstance(prev, list):
        prev = []

    sel = st.multiselect(
        t(
            lang,
            "2 à 4 sources de données les plus pertinentes",
            "2 to 4 most relevant data sources",
        ),
        options=options,
        default=[x for x in prev if x in options],
        max_selections=4,
        format_func=lambda x: tr(lang, x),
        help=t(
            lang,
            "Choisissez entre 2 et 4 options. Si vous choisissez Autres, précisez.",
            "Choose between 2 and 4 options. If you select Other, please specify.",
        ),
        key="data_sources_multiselect",
    )
    resp_set("data_sources", sel)

    other_label = t(lang, "Autres", "Other", "Outras", "أخرى")
    if has_other_option(sel):
        other = st.text_input(
            t(lang, "Autres : préciser", "Other: please specify"),
            key="data_sources_other_input",
            value=resp_get("data_sources_other", ""),
        )
        resp_set("data_sources_other", (other or "").strip())
    else:
        resp_set("data_sources_other", "")

    errs = validate_r11(lang)
    if errs:
        for e in errs:
            st.error(e)

def rubric_12(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 12 : Questions ouvertes", "Section 12: Open questions", "Secção 12: Questões abertas", "القسم 12: الأسئلة المفتوحة"))
    st.markdown(
        t(
            lang,
            "Ces questions sont **optionnelles**. Vous pouvez les laisser vides. Toutefois, elles sont présentées **une à une** pour faciliter la saisie de vos opinions ou sentiments sur le système statistique.",
            "These questions are **optional**. You may leave them blank. They are presented **one by one** to facilitate completion.",
            "Estas perguntas são **opcionais**. Pode deixá-las em branco. Contudo, são apresentadas **uma a uma** para facilitar a introdução das suas opiniões ou perceções sobre o sistema estatístico.",
            "هذه الأسئلة **اختيارية**. يمكنكم تركها فارغة. ومع ذلك، تُعرض **واحدة تلو الأخرى** لتسهيل إدخال آرائكم أو انطباعاتكم حول النظام الإحصائي."
        )
    )

    if "r12_substep" not in st.session_state:
        st.session_state["r12_substep"] = 0  # 0..2=open questions, 3=confirmation

    sub = int(st.session_state.get("r12_substep", 0) or 0)

    if sub == 0:
        st.markdown("#### " + t(lang, "Question 1 / 3", "Question 1 / 3"))
        q1 = st.text_area(
            t(lang, "1) Commentaires / recommandations clés", "1) Key comments / recommendations", "1) Comentários / recomendações principais", "1) التعليقات / التوصيات الأساسية"),
            value=resp_get("open_q1", ""),
            height=160,
            key="open_q1_input",
        )
        resp_set("open_q1", (q1 or "").strip())
        if not resp_get("open_q1", ""):
            st.warning(t(lang, "Alerte : la question 1 est vide (vous pouvez tout de même continuer).",
                         "Warning: question 1 is empty (you can still proceed)."))

    elif sub == 1:
        st.markdown("#### " + t(lang, "Question 2 / 3", "Question 2 / 3"))
        q2 = st.text_area(
            t(
                lang,
                "2) Un ou des indicateur(s) statistique(s) socio-économique(s) essentiel(s) manquant(s) et justification(s)",
                "2) One or more missing essential socio-economic statistical indicator(s) and justification(s)",
                "2) Um ou mais indicadores estatísticos socioeconómicos essenciais em falta e respetiva justificação",
                "2) مؤشر أو أكثر من المؤشرات الإحصائية الاجتماعية والاقتصادية الأساسية غير المدرجة مع المبررات",
            ),
            value=resp_get("open_q2", ""),
            height=160,
            key="open_q2_input",
        )
        resp_set("open_q2", (q2 or "").strip())
        if not resp_get("open_q2", ""):
            st.warning(t(lang, "Alerte : la question 2 est vide (vous pouvez tout de même continuer).",
                         "Warning: question 2 is empty (you can still proceed)."))

    elif sub == 2:
        st.markdown("#### " + t(lang, "Question 3 / 3", "Question 3 / 3"))
        q3 = st.text_area(
            t(lang, "3) Besoins de soutien (technique, financier, etc.)", "3) Support needs (technical, financial, etc.)", "3) Necessidades de apoio (técnico, financeiro, etc.)", "3) احتياجات الدعم (الفني، المالي، إلخ)"),
            value=resp_get("open_q3", ""),
            height=160,
            key="open_q3_input",
        )
        resp_set("open_q3", (q3 or "").strip())
        if not resp_get("open_q3", ""):
            st.warning(t(lang, "Alerte : la question 3 est vide (vous pouvez tout de même continuer).",
                         "Warning: question 3 is empty (you can still proceed)."))

    else:
        st.markdown("#### " + t(lang, "Confirmation", "Confirmation"))
        st.info(
            t(
                lang,
                "Dernière étape : merci d’indiquer si vous avez consulté d’autres collègues. Cette question est obligatoire.",
                "Final step: please indicate whether you consulted other colleagues. This question is mandatory.",
            )
        )

        cc_opts = ["", "YES", "NO"]
        cc_labels = {"YES": t(lang, "Oui", "Yes"), "NO": t(lang, "Non", "No")}
        prev_cc = (resp_get("consulted_colleagues", "") or "").strip()
        idx_cc = cc_opts.index(prev_cc) if prev_cc in cc_opts else 0
        chosen_cc = st.selectbox(
            t(
                lang,
                "Avez-vous consulté d’autres collègues pour remplir ce questionnaire ?",
                "Did you consult other colleagues to complete this questionnaire?",
            ),
            options=cc_opts,
            index=idx_cc,
            format_func=lambda k: (t(lang, "— Sélectionner —", "— Select —") if k == "" else cc_labels.get(k, k)),
            key="consulted_colleagues_select",
        )
        resp_set("consulted_colleagues", chosen_cc)

    st.markdown("---")
    col1, col2, col3 = st.columns([1, 1, 3])
    with col1:
        prev_disabled = (sub <= 0)
        if st.button(t(lang, "⬅ Question précédente", "⬅ Previous question"), disabled=prev_disabled, key="r12_prev_btn"):
            st.session_state["r12_substep"] = max(0, sub - 1)
            st.rerun()
    with col2:
        if sub < 2:
            next_label = t(lang, "Question suivante ➡", "Next question ➡")
        elif sub == 2:
            next_label = t(lang, "Aller à la confirmation ➡", "Go to confirmation ➡")
        else:
            next_label = t(lang, "OK (rubrique terminée)", "OK (section completed)")
        next_disabled = (sub >= 3)
        if st.button(next_label, disabled=next_disabled, key="r12_next_btn"):
            st.session_state["r12_substep"] = min(3, sub + 1)
            st.rerun()
    with col3:
        st.caption(t(lang, "Progression : 1/3 → 2/3 → 3/3 → Confirmation.",
                     "Progress: 1/3 → 2/3 → 3/3 → Confirmation."))

    errs = validate_r12(lang)
    if errs:
        for e in errs:
            st.error(e)

def rubric_send(lang: str, df_long: pd.DataFrame) -> None:
    st.subheader(t(lang, "ENVOYER le questionnaire", "SUBMIT questionnaire", "ENVIAR questionário", "إرسال الاستبيان"))

    errors = validate_all(lang)
    if errors:
        st.error(t(lang, "Le questionnaire contient des erreurs bloquantes :",
                   "There are blocking errors:"))
        st.write("\n".join([f"- {e}" for e in errors]))
        st.info(t(lang, "Retournez aux rubriques concernées via la navigation.",
                  "Go back to the relevant sections using navigation."))
        return

    # Optional warnings
    if not resp_get("open_q1", "") or not resp_get("open_q2", "") or not resp_get("open_q3", ""):
        st.warning(t(lang,
                     "Certaines questions ouvertes sont vides (optionnel).",
                     "Some open questions are empty (optional)."))

    st.info(t(
        lang,
        "Tout est prêt. Cliquez sur **ENVOYER** pour soumettre votre questionnaire.",
        "Everything is ready. Click **SUBMIT** to send your questionnaire."
    ))

    # Empêcher les envois multiples (par email + par session)
    email = (resp_get("email", "") or "").strip()
    already_in_session = bool(st.session_state.get("submitted_once", False))

    # Pré-contrôle email (local SQLite)
    already_in_db = False
    if email:
        try:
            already_in_db = db_email_exists(email)
        except Exception:
            # Ne pas bloquer si la base locale est momentanément inaccessible (rare)
            st.warning(t(
                lang,
                "Impossible de vérifier pour l’instant si cet email a déjà répondu (base locale). "
                "Vous pouvez continuer ; en cas de doublon, l’envoi pourra être refusé au moment de l’enregistrement.",
                "Cannot check right now whether this email already submitted (local database). "
                "You can continue; duplicates may be rejected during saving."
            ))
            already_in_db = False

    if already_in_db:
        st.error(t(
            lang,
            "Ce questionnaire a déjà été envoyé avec cet email. Un seul envoi est autorisé.",
            "This questionnaire has already been submitted with this email. Only one submission is allowed."
        ))

    if already_in_session and not already_in_db:
        st.info(t(
            lang,
            "Ce navigateur a déjà effectué un envoi. Pour un nouvel envoi, utilisez un autre email / session.",
            "This browser session already submitted once. For a new submission, use another email / session."
        ))

    disable_submit = already_in_db or already_in_session

    if st.button(t(lang, "✅ ENVOYER et enregistrer", "✅ SUBMIT and save"), disabled=disable_submit):
        submission_id = str(uuid.uuid4())
        payload = st.session_state.responses.copy()
        payload["submission_id"] = submission_id
        payload["submitted_at_utc"] = now_utc_iso()
        payload["scoring_version"] = SCORING_VERSION

        try:
            db_save_submission(submission_id, lang, email, payload)
        except Exception as e:
            st.error(t(lang, "Échec d'enregistrement (base locale) : ",
                       "Save failed (local database): ") + str(e))

            # Anti-blocage : fournir une sauvegarde JSON immédiate
            json_text = json.dumps(payload, ensure_ascii=False, indent=2)
            st.download_button(
                label=t(lang, "⬇️ Télécharger la sauvegarde JSON", "⬇️ Download JSON backup"),
                data=json_text.encode("utf-8"),
                file_name=f"submission_{submission_id}.json",
                mime="application/json",
                key=f"dl_json_{submission_id}",
            )
            st.text_area(
                t(lang, "Copie du JSON (à conserver)", "JSON copy (keep it)"),
                value=json_text,
                height=260,
                key=f"json_copy_{submission_id}",
            )
            st.info(t(lang,
                      "Conservez ce fichier. Vous pourrez le transmettre pour intégration ultérieure.",
                      "Keep this file. You can share it later for integration."))
            return

        st.success(t(
            lang,
            "Merci ! Votre questionnaire a été enregistré.",
            "Thank you! Your submission has been saved."
        ))

        # Fournir aussi une copie JSON (utile si l’hébergement redémarre)
        json_text = json.dumps(payload, ensure_ascii=False, indent=2)
        st.download_button(
            label=t(lang, "⬇️ Télécharger une copie JSON", "⬇️ Download a JSON copy"),
            data=json_text.encode("utf-8"),
            file_name=f"submission_{submission_id}.json",
            mime="application/json",
            key=f"dl_json_ok_{submission_id}",
        )

        # Bloquer les envois multiples pour cette session
        st.session_state.submitted_once = True
        st.caption(f"ID : {submission_id}")
        st.info(t(lang, "Envoi terminé. Vous pouvez fermer cette page.",
                  "Submission complete. You can close this page."))
        st.session_state.submission_id = submission_id


# =========================
# Admin dashboard
# =========================

def admin_login(lang: str) -> None:
    st.subheader(t(lang, "Administration", "Administration"))
    pw = st.text_input(t(lang, "Mot de passe admin", "Admin password"), type="password")

    # Si le répondant saisit le nom du secret au lieu du mot de passe
    if (pw or "").strip() in ("ADMIN_PASSWORD", "SUPERADMIN_PASSWORD"):
        st.info(
            t(
                lang,
                "Astuce : saisissez le mot de passe réel, pas le nom du secret (ex. ADMIN_PASSWORD).",
                "Tip: enter the actual password value, not the secret name (e.g., ADMIN_PASSWORD).",
            )
        )

    # Diagnostic discret : présence du mot de passe (sans afficher la valeur)
    src_code, src_lbl = get_admin_auth_source()
    if src_code == "none":
        st.caption(
            t(
                lang,
                "ADMIN_PASSWORD : non configuré (secrets, variable d’environnement ou mot de passe haché en base).",
                "ADMIN_PASSWORD : not configured (secrets, environment variable, or hashed password in DB).",
            )
        )
    else:
        st.caption(
            t(
                lang,
                f"ADMIN_PASSWORD : configuré via {src_lbl}.",
                f"ADMIN_PASSWORD : configured via {src_lbl}.",
            )
        )


    # Diagnostic avancé (invisible par défaut) : ajouter ?diag=1 à l’URL
    # Utile pour vérifier rapidement que SUPERADMIN_PASSWORD est bien injecté (sans afficher la valeur).
    qp = get_query_params()
    if qp.get("diag", ["0"])[0] == "1":
        v_sa, src_sa = _get_secret_or_env("SUPERADMIN_PASSWORD")
        if v_sa and src_sa:
            st.caption(
                t(
                    lang,
                    f"SUPERADMIN_PASSWORD : configuré via {src_sa}.",
                    f"SUPERADMIN_PASSWORD : configured via {src_sa}.",
                )
            )
        else:
            st.caption(
                t(
                    lang,
                    "SUPERADMIN_PASSWORD : non configuré (secrets ou variable d’environnement).",
                    "SUPERADMIN_PASSWORD : not configured (secrets or environment variable).",
                )
            )

    if st.button(t(lang, "Se connecter", "Login")):
        # Superadmin : même écran, mais rôle différent (invisible tant que non authentifié)
        if verify_superadmin_password(pw):
            st.session_state.admin_authed = True
            st.session_state.admin_role = "superadmin"
            st.success(t(lang, "Connexion réussie.", "Logged in."))
            st.rerun()
        elif verify_admin_password(pw):
            st.session_state.admin_authed = True
            st.session_state.admin_role = "admin"
            st.success(t(lang, "Connexion réussie.", "Logged in."))
            st.rerun()
        else:
            st.error(
                t(
                    lang,
                    "Mot de passe incorrect ou secret ADMIN_PASSWORD manquant.",
                    "Incorrect password or missing ADMIN_PASSWORD secret.",
                )
            )


def admin_dashboard(lang: str) -> None:
    st.subheader(t(lang, "Tableau de bord admin", "Admin dashboard"))

    # Load data from SQLite
    df = db_read_submissions(limit=20000)
    st.metric(t(lang, "Nombre de réponses", "Number of responses"), len(df))

    if df.empty:
        st.info(t(lang, "Aucune réponse pour le moment.", "No responses yet."))
        return

    # Parse payloads
    payloads = []
    for _, r in df.iterrows():
        try:
            payloads.append(json.loads(r["payload_json"]))
        except Exception:
            payloads.append({})

    # Flat view for quick export
    flat = pd.DataFrame([flatten_payload(p) for p in payloads])
    flat.insert(0, "submission_id", df["submission_id"].values)
    flat.insert(1, "submitted_at_utc", df["submitted_at_utc"].values)

    is_super = st.session_state.get("admin_role") == "superadmin"

    if is_super:
        tab_quick, tab_super, tab_sec = st.tabs([
            t(lang, "Vue rapide", "Quick view"),
            t(lang, "Analyse avancée", "Advanced analysis"),
            t(lang, "Sécurité", "Security"),
        ])
    else:
        tab_quick, = st.tabs([t(lang, "Vue rapide", "Quick view")])

    with tab_quick:
        st.dataframe(flat, use_container_width=True)

        # Export Excel (flat + raw) avec repli si openpyxl/xlsxwriter n'est pas installé
        excel_engine = None
        try:
            import openpyxl  # noqa: F401
            excel_engine = "openpyxl"
        except Exception:
            try:
                import xlsxwriter  # noqa: F401
                excel_engine = "xlsxwriter"
            except Exception:
                excel_engine = None

        if excel_engine is not None:
            try:
                out = io.BytesIO()
                with pd.ExcelWriter(out, engine=excel_engine) as writer:
                    flat.to_excel(writer, sheet_name="submissions", index=False)
                    df.to_excel(writer, sheet_name="raw_json", index=False)
                out.seek(0)

                st.download_button(
                    t(lang, "Exporter en Excel", "Export to Excel"),
                    data=out.getvalue(),
                    file_name="consultation_stat_niang_export.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            except Exception as e:
                st.warning(
                    t(
                        lang,
                        f"L'export Excel n'est pas disponible pour le moment ({e}). Utilisez les exports CSV ci-dessous ou ajoutez openpyxl au fichier requirements.txt.",
                        f"Excel export is temporarily unavailable ({e}). Please use the CSV exports below or add openpyxl to requirements.txt.",
                    )
                )
        else:
            st.info(
                t(
                    lang,
                    "L'export Excel n'est pas disponible sur ce déploiement car ni openpyxl ni xlsxwriter n'est installé. Les exports CSV restent disponibles ci-dessous.",
                    "Excel export is not available on this deployment because neither openpyxl nor xlsxwriter is installed. CSV exports remain available below.",
                )
            )

        # Export CSV (UTF-8 avec BOM pour Excel)
        csv_flat = flat.to_csv(index=False).encode("utf-8-sig")
        csv_raw = df.to_csv(index=False).encode("utf-8-sig")

        st.download_button(
            t(lang, "Exporter en CSV (vue à plat)", "Export CSV (flat view)"),
            data=csv_flat,
            file_name="consultation_submissions_flat.csv",
            mime="text/csv",
        )

        st.download_button(
            t(lang, "Exporter en CSV (brut : JSON)", "Export CSV (raw : JSON)"),
            data=csv_raw,
            file_name="consultation_submissions_raw.csv",
            mime="text/csv",
        )

        # Export JSON Lines (JSONL) : 1 ligne = 1 soumission
        jsonl_bytes = ("\n".join([json.dumps(p, ensure_ascii=False) for p in payloads if isinstance(p, dict)]) + "\n").encode("utf-8")
        st.download_button(
            t(lang, "Exporter en JSONL (1 ligne = 1 soumission)", "Export JSONL (1 line per submission)"),
            data=jsonl_bytes,
            file_name="consultation_submissions.jsonl",
            mime="application/json",
        )

        # Télécharger la base SQLite
        st.caption(t(lang, f"Base locale : {DB_PATH}", f"Local database : {DB_PATH}"))
        try:
            with open(DB_PATH, "rb") as f:
                db_bytes = f.read()
            st.download_button(
                t(lang, "Télécharger la base locale (responses.db)", "Download local database (responses.db)"),
                data=db_bytes,
                file_name="responses.db",
                mime="application/x-sqlite3",
            )
        except Exception:
            st.warning(t(lang, "Base SQLite introuvable pour le moment.", "SQLite database not found yet."))

        # Export ZIP complet (DB + CSV + JSON par soumission)
        zip_buf = io.BytesIO()
        try:
            with zipfile.ZipFile(zip_buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
                # DB
                try:
                    with open(DB_PATH, "rb") as f:
                        z.writestr("responses.db", f.read())
                except Exception:
                    pass

                # CSV
                z.writestr("consultation_submissions_flat.csv", csv_flat)
                z.writestr("consultation_submissions_raw.csv", csv_raw)
                z.writestr("consultation_submissions.jsonl", jsonl_bytes)

                # JSON individuels
                for p in payloads:
                    sid = str(p.get("submission_id") or "")
                    if not sid:
                        continue
                    z.writestr(
                        f"json/submission_{sid}.json",
                        json.dumps(p, ensure_ascii=False, indent=2).encode("utf-8"),
                    )

            zip_buf.seek(0)
            st.download_button(
                t(lang, "Exporter en ZIP (DB + CSV + JSON)", "Export ZIP (DB + CSV + JSON)"),
                data=zip_buf.getvalue(),
                file_name="consultation_export.zip",
                mime="application/zip",
            )
        except Exception:
            st.warning(t(lang, "Export ZIP indisponible (erreur technique).", "ZIP export unavailable (technical error)."))

        st.info(
            t(
                lang,
                "Astuce : utilisez ?admin=1 dans l’URL pour afficher l’espace admin.",
                "Tip: use ?admin=1 in the URL to access the admin space."
            )
        )

    if not is_super:
        return

    with tab_super:
        # --- Build a richer dataset for analysis
        df_super = pd.DataFrame(payloads)
        df_super["submission_id"] = df["submission_id"].values
        df_super["submitted_at_utc"] = pd.to_datetime(df["submitted_at_utc"], errors="coerce", utc=True)

        # Filters
        st.markdown("### " + t(lang, "Filtres", "Filters"))

        colf1, colf2, colf3 = st.columns(3)
        with colf1:
            countries = sorted([c for c in df_super.get("pays", pd.Series(dtype=str)).dropna().unique().tolist() if str(c).strip() != ""])
            sel_countries = st.multiselect(t(lang, "Pays de résidence", "Country of residence"), options=countries, default=[], key="f_country")
        with colf2:
            actors = sorted([a for a in df_super.get("type_acteur", pd.Series(dtype=str)).dropna().unique().tolist() if str(a).strip() != ""])
            sel_actors = st.multiselect(t(lang, "Type d’acteur", "Stakeholder type"), options=actors, default=[], key="f_actor")
        with colf3:
            # Period filter
            min_dt = df_super["submitted_at_utc"].min()
            max_dt = df_super["submitted_at_utc"].max()
            if pd.isna(min_dt) or pd.isna(max_dt):
                min_dt = pd.Timestamp.utcnow() - pd.Timedelta(days=30)
                max_dt = pd.Timestamp.utcnow()
            date_range = st.date_input(
                t(lang, "Période", "Period"),
                value=(min_dt.date(), max_dt.date()),
                key="f_period"
            )

        filtered = df_super.copy()
        if sel_countries:
            filtered = filtered[filtered["pays"].isin(sel_countries)]
        if sel_actors:
            filtered = filtered[filtered["type_acteur"].isin(sel_actors)]
        if isinstance(date_range, (list, tuple)) and len(date_range) == 2:
            # Streamlit date_input returns datetime.date; our column may be tz-aware (UTC).
            col = pd.to_datetime(filtered["submitted_at_utc"], utc=True, errors="coerce")

            start_d = pd.Timestamp(date_range[0])
            end_d = pd.Timestamp(date_range[1]) + pd.Timedelta(days=1) - pd.Timedelta(seconds=1)

            if start_d.tz is None:
                start_d = start_d.tz_localize("UTC")
            else:
                start_d = start_d.tz_convert("UTC")

            if end_d.tz is None:
                end_d = end_d.tz_localize("UTC")
            else:
                end_d = end_d.tz_convert("UTC")

            filtered = filtered[(col >= start_d) & (col <= end_d)]

        st.caption(t(lang, f"Réponses filtrées : {len(filtered)}", f"Filtered responses: {len(filtered)}"))

        if filtered.empty:
            st.warning(t(lang, "Aucune réponse dans ce filtre.", "No responses match these filters."))
            return

        # Longlist for labels (domain/stat)
        df_long = load_longlist()
        dom_lbl = domain_label_map(df_long, lang)
        stat_lbl = stat_label_map(df_long, lang)

        # --- Build aggregated prioritization table
        rows = []
        for _, p in filtered.iterrows():
            top5 = p.get("top5_domains", []) or []
            sel_by_dom = p.get("selected_by_domain", {}) or {}
            scoring = p.get("scoring", {}) or {}
            sid = p.get("submission_id", "")
            for d, stats in (sel_by_dom.items() if isinstance(sel_by_dom, dict) else []):
                if not isinstance(stats, list):
                    continue
                for s in stats:
                    sc = scoring.get(s, {})
                    avail_raw = (sc.get("availability", sc.get("gap", 0)))
                    avail = normalize_availability(avail_raw, p.get("scoring_version", 0))
                    dem = int(sc.get("demand", 0) or 0)
                    fea = int(sc.get("feasibility", 0) or 0)
                    overall = (avail + dem + fea) / 3.0
                    rows.append({
                        "submission_id": sid,
                        "pays": p.get("pays", ""),
                        "type_acteur": p.get("type_acteur", ""),
                        "domain_code": d,
                        "domain_label": dom_lbl.get(d, d),
                        "stat_code": s,
                        "stat_label": stat_lbl.get(s, s),
                        "availability": avail, "demand": dem, "feasibility": fea, "overall": overall
                    })

        df_rows = pd.DataFrame(rows)
        if df_rows.empty:
            st.warning(t(lang, "Aucune statistique notée dans ces réponses.", "No scored indicators in these responses."))
            return

        # Aggregation
        by_stat = df_rows.groupby(["domain_code", "domain_label", "stat_code", "stat_label"], as_index=False).agg(
            n=("submission_id", "nunique"),
            mean_availability=("availability", "mean"),
            mean_demand=("demand", "mean"),
            mean_feasibility=("feasibility", "mean"),
            mean_overall=("overall", "mean"),
        ).sort_values(["domain_code", "mean_overall", "n"], ascending=[True, False, False])

        by_domain = df_rows.groupby(["domain_code", "domain_label"], as_index=False).agg(
            n_stats=("stat_code", "count"),
            n_submissions=("submission_id", "nunique"),
            mean_overall=("overall", "mean"),
        ).sort_values(["mean_overall", "n_submissions"], ascending=[False, False])

        st.markdown("### " + t(lang, "Tableau de priorisation agrégé", "Aggregated prioritization table"))
        st.dataframe(by_stat, use_container_width=True, height=420)

        # Export aggregated Excel
        out2 = io.BytesIO()
        with pd.ExcelWriter(out2, engine="openpyxl") as writer:
            by_domain.to_excel(writer, sheet_name="by_domain", index=False)
            by_stat.to_excel(writer, sheet_name="by_statistic", index=False)
            df_rows.to_excel(writer, sheet_name="scored_rows", index=False)
        out2.seek(0)

        st.download_button(
            t(lang, "Télécharger l’agrégé (Excel)", "Download aggregated (Excel)"),
            data=out2.getvalue(),
            file_name="prioritization_aggregated.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        # Rich Word report
        st.markdown("### " + t(lang, "Rapport Word (publication)", "Word report (publication)"))
        st.caption(t(lang, "Génère un rapport enrichi avec graphiques et annexes.", "Generates an enriched report with charts and annexes."))

        if st.button(t(lang, "Générer le rapport Word", "Generate Word report"), key="btn_word_publication"):
            try:
                doc_bytes = build_publication_report_docx(lang, filtered, by_domain, by_stat, df_rows)
                st.download_button(
                    t(lang, "Télécharger le rapport (.docx)", "Download report (.docx)"),
                    data=doc_bytes,
                    file_name="rapport_publication_priorisation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                st.success(t(lang, "Rapport généré.", "Report generated."))


            except Exception as e:


                st.error(f"Word : {e}")


    with tab_sec:
        st.subheader(t(lang, "Sécurité", "Security"))

        st.markdown("### " + t(lang, "Mot de passe admin", "Admin password"))
        src_code, src_lbl = get_admin_auth_source()
        if src_code == "db":
            st.info(t(lang, "Mot de passe admin : défini en base (haché).", "Admin password: set in database (hashed)."))
        elif src_code in ("secrets", "env"):
            st.info(t(lang, f"Mot de passe admin : défini via {src_lbl}.", f"Admin password: set via {src_lbl}."))
        else:
            st.warning(t(lang, "Mot de passe admin : non configuré.", "Admin password: not configured."))

        st.markdown(t(
            lang,
            "#### Changer le mot de passe admin (réservé au superadmin)",
            "#### Change admin password (superadmin only)"
        ))

        new1 = st.text_input(t(lang, "Nouveau mot de passe admin", "New admin password"), type="password", key="new_admin_pw1")
        new2 = st.text_input(t(lang, "Confirmer le nouveau mot de passe", "Confirm new password"), type="password", key="new_admin_pw2")

        if st.button(t(lang, "Mettre à jour le mot de passe admin", "Update admin password")):
            p1 = (new1 or "").strip()
            p2 = (new2 or "").strip()
            if len(p1) < 10:
                st.error(t(lang, "Mot de passe trop court (≥ 10 caractères).", "Password too short (≥ 10 characters)."))
            elif p1 != p2:
                st.error(t(lang, "Les deux champs ne correspondent pas.", "The two fields do not match."))
            else:
                try:
                    set_admin_password(p1)
                    st.success(t(lang, "Mot de passe admin mis à jour.", "Admin password updated."))
                    st.info(t(lang, "Le nouveau mot de passe est actif immédiatement (sans redeploy).",
                              "The new password is active immediately (no redeploy needed)."))
                except Exception as e:
                    st.error(f"{e}")

        st.divider()
        st.markdown(t(
            lang,
            "#### Réinitialiser vers Secrets / variable d’environnement",
            "#### Reset to Secrets / environment variable"
        ))
        st.caption(t(
            lang,
            "Cette action supprime le mot de passe haché stocké en base. L’app utilisera alors ADMIN_PASSWORD (secrets/env).",
            "This removes the hashed password stored in DB. The app will then use ADMIN_PASSWORD (secrets/env)."
        ))

        if st.button(t(lang, "Réinitialiser le mot de passe admin", "Reset admin password")):
            reset_admin_password_to_secrets_env()
            v, src = _get_secret_or_env("ADMIN_PASSWORD")
            if v and src:
                st.success(t(lang, f"Réinitialisation effectuée : retour à {src}.", f"Reset done: back to {src}."))
            else:
                st.warning(t(lang, "Réinitialisation effectuée, mais ADMIN_PASSWORD n’est pas défini dans secrets/env.",
                            "Reset done, but ADMIN_PASSWORD is not defined in secrets/env."))





def apply_language_direction(lang: str) -> None:
    if lang == "ar":
        st.markdown(
            """
            <style>
            .block-container, .stMarkdown, .stText, label, p, h1, h2, h3, h4, h5, h6 { direction: rtl; text-align: right; }
            div[data-testid="stHorizontalBlock"] { direction: rtl; }
            </style>
            """,
            unsafe_allow_html=True,
        )


def render_startup_diagnostics(lang: str) -> None:
    """Affiche des informations utiles en mode discret pour diagnostiquer le déploiement."""
    try:
        with st.expander(t(lang, "Diagnostic technique", "Technical diagnostics", "Diagnóstico técnico", "تشخيص تقني"), expanded=False):
            df_long_diag = load_longlist()
            df_c_diag = load_countries()
            st.write({
                "db_path": DB_PATH,
                "db_exists": os.path.exists(DB_PATH),
                "app_dir": APP_DIR,
                "cwd": str(Path.cwd()),
                "longlist_csv_candidates": candidate_paths(["indicator_longlist.csv"]),
                "longlist_xlsx_candidates": candidate_paths(["longlist_filled_pt_ar_final.xlsx", "longlist.xlsx"]),
                "country_xlsx_candidates": candidate_paths(["COUNTRY_ISO3_4_filled_pt_ar_final.xlsx", "COUNTRY_ISO3_4.xlsx", "COUNTRY_ISO3_with_EN.xlsx", "COUNTRY_ISO3.xlsx"]),
                "longlist_source_used": getattr(df_long_diag, "attrs", {}).get("source_path", ""),
                "countries_source_used": getattr(df_c_diag, "attrs", {}).get("source_path", ""),
                "longlist_rows": int(len(df_long_diag)) if df_long_diag is not None else 0,
                "country_rows": int(len(df_c_diag)) if df_c_diag is not None else 0,
                "embedded_longlist_rows": len(EMBEDDED_LONGLIST_ROWS),
                "embedded_country_rows": len(EMBEDDED_COUNTRIES_ROWS),
                "openpyxl_available": bool(__import__("importlib").util.find_spec("openpyxl")),
            })
    except Exception:
        pass


def render_language_switch(lang: str) -> str:
    st.markdown(
        """
        <style>
        .lang-banner {
            border: 1px solid rgba(49, 51, 63, 0.2);
            padding: 0.9rem 1rem 0.35rem 1rem;
            border-radius: 0.9rem;
            margin-bottom: 1rem;
            background: rgba(240, 242, 246, 0.65);
        }
        .lang-banner-title {
            font-size: 1rem;
            font-weight: 700;
            margin-bottom: 0.35rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    st.markdown(f"<div class='lang-banner'><div class='lang-banner-title'>{t(lang, '🌐 Choisissez votre langue', '🌐 Choose your language', '🌐 Escolha o seu idioma', '🌐 اختر اللغة')}</div></div>", unsafe_allow_html=True)
    options = LANG_OPTIONS
    idx = options.index(lang) if lang in options else 0
    chosen = st.radio(
        label=t(lang, "Langue / Language", "Language", "Idioma", "اللغة"),
        options=options,
        index=idx,
        horizontal=True,
        format_func=lambda x: LANG_LABELS.get(x, x),
        key="visible_lang_switch",
        label_visibility="collapsed",
    )
    st.session_state.lang = chosen
    return chosen


# =========================
# Main
# =========================

def main() -> None:
    # set_page_config déjà défini en haut du fichier
    init_session()
    # Initialise la base locale (créée dans le même dossier que ce script)
    db_init()
    maybe_restore_draft()

    # Language toggle (prominent on main screen + compact reminder in sidebar)
    lang = render_language_switch(st.session_state.lang)
    st.sidebar.markdown("### 🌐 " + t(lang, "Langue", "Language", "Idioma", "اللغة"))
    st.sidebar.caption(" / ".join([LANG_LABELS[k] for k in LANG_OPTIONS]))
    apply_language_direction(lang)
    render_startup_diagnostics(lang)

    # Admin access
    qp = get_query_params()
    is_admin = "admin" in qp and qp["admin"] and qp["admin"][0] in ["1", "true", "yes"]

    df_long = load_longlist()

    # Header
    st.title(t(lang, APP_TITLE_FR, APP_TITLE_EN, APP_TITLE_PT, APP_TITLE_AR))
    st.caption(t(lang, "Application unifiée (FR/EN/PT/AR) – codes masqués – contrôles qualité intégrés.",
                 "Unified app (FR/EN/PT/AR) – hidden codes – built-in quality controls.",
                 "Aplicação multilingue (FR/EN/PT/AR) – códigos ocultos – controlos de qualidade integrados.",
                 "تطبيق متعدد اللغات (FR/EN/PT/AR) – رموز مخفية – ضوابط جودة مدمجة."))
    st.caption(t(lang, f"Version déployée : {APP_BUILD}", f"Deployed version: {APP_BUILD}", f"Versão instalada : {APP_BUILD}", f"الإصدار المنشور : {APP_BUILD}"))
    if st.session_state.get("draft_exists") and not st.session_state.get("draft_resume_notice_shown"):
        st.warning(
            t(
                lang,
                "La saisie est sauvegardée. En cas de suspension de moins de 48 heures, reprenez-la là où vous vous étiez arrêté en ré-ouvrant le lien contenant rid (à conserver / mettre en favori / retrouver dans l’historique).",
                "Your entry is saved. If you pause for less than 48 hours, resume where you left off by reopening the link containing rid (bookmark / save / find it in your browser history).",
            ),
            icon="💾",
        )
        st.session_state.draft_resume_notice_shown = True

    # Sidebar navigation
    steps = get_steps(lang)
    render_sidebar(lang, steps)

    # Admin view
    if is_admin:
        if not st.session_state.admin_authed:
            admin_login(lang)
            return
        admin_dashboard(lang)
        return

    # Normal flow
    step_key = steps[st.session_state.nav_idx][0]

    if step_key == "R1":
        rubric_1(lang)
    elif step_key == "R2":
        rubric_2(lang)
    elif step_key == "R3":
        rubric_3(lang)
    elif step_key == "R4":
        rubric_4(lang, df_long)
    elif step_key == "R5":
        rubric_5(lang, df_long)
    elif step_key == "R6":
        rubric_6(lang)
    elif step_key == "R7":
        rubric_7(lang)
    elif step_key == "R8":
        rubric_8(lang)
    elif step_key == "R9":
        rubric_9(lang)
    elif step_key == "R10":
        rubric_10(lang)
    elif step_key == "R11":
        rubric_11(lang)
    elif step_key == "R12":
        rubric_12(lang)
    elif step_key == "SEND":
        rubric_send(lang, df_long)

    st.divider()
    nav_buttons(lang, steps, df_long)


if __name__ == "__main__":
    main()
