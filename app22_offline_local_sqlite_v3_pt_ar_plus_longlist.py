import io
import json
import hashlib
import hmac
import secrets
import os
import re
import sqlite3
import zipfile
import uuid
import time
from datetime import datetime, timezone
from typing import Any, Dict, List, Tuple, Optional

import numpy as np
import pandas as pd
import streamlit as st

# =========================
# Configuration Streamlit
# =========================

# IMPORTANT : doit être le tout premier appel Streamlit
st.set_page_config(
    page_title="Consultation STATAFRIC",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# Stabilisation UI (évite les variations de largeur dues à l’apparition/disparition de la barre de défilement)
st.markdown("""
<style>
html { overflow-y: scroll; }
div[data-testid="stVerticalBlock"] { gap: 0.75rem; }
</style>
""", unsafe_allow_html=True)


# CSS léger pour stabiliser le rendu (évite les variations de largeur perçues)
st.markdown(
    """
<style>
/* Largeur stable du conteneur principal */
.block-container { max-width: 1200px; padding-top: 1rem; padding-bottom: 2rem; }
@media (max-width: 1200px) { .block-container { max-width: 100%; } }

/* Réduit les effets de "sauts" quand des messages apparaissent/disparaissent */
div[data-testid="stVerticalBlock"] { gap: 0.75rem; }

/* Boutons : largeur stable */
button[kind="primary"], button[kind="secondary"] { white-space: nowrap; }
</style>
""",
    unsafe_allow_html=True,
)


# =========================
# Configuration
# =========================

APP_TITLE_FR = "Questionnaire de consultation pour l'identification des statistiques prioritaires"
APP_TITLE_EN = "Consultation questionnaire for identifying priority statistics"
APP_TITLE_PT = "Questionário de consulta para a identificação de estatísticas prioritárias"
APP_TITLE_AR = "استبيان تشاوري لتحديد الإحصاءات ذات الأولوية"

APP_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(APP_DIR, "responses.db")
LONG_LIST_CSV = os.path.join("data", "indicator_longlist.csv")
LONG_LIST_XLSX = os.path.join("data", "longlist.xlsx")


COUNTRY_XLSX = os.path.join("data", "COUNTRY_ISO3_with_EN.xlsx")
UK_FR = "NSP (Ne sais pas)"
UK_EN = "DNK (Do not know)"
UK_PT = "NSP (Não sabe)"
UK_AR = "لا أعلم"

# Version du scoring (pour compatibilité ascendante)
# v1 : ancien critère "gap" (écart) ; v2 : disponibilité inversée (Bonne=1) ; v3 : disponibilité directe (Bonne=3)
SCORING_VERSION = 3


# Scores affichés (notation multicritères)
# Barèmes de notation (0–3) par critère (scoring rationalisé)
# Remarque : 0 = NSP / DNK (ne sait pas) ; 1–3 = intensité croissante selon le libellé.
SCORE_SCALES = {
    "demand": {
        "fr": {0: "NSP", 1: "Faible", 2: "Moyenne", 3: "Élevée"},
        "en": {0: "DNK", 1: "Low", 2: "Medium", 3: "High"},
        "pt": {0: "NSP", 1: "Baixa", 2: "Média", 3: "Alta"},
        "ar": {0: "لا أعلم", 1: "منخفض", 2: "متوسط", 3: "مرتفع"},
    },
    "availability": {
        "fr": {0: "NSP", 1: "Faible ou inexistante", 2: "Partielle", 3: "Bonne"},
        "en": {0: "DNK", 1: "Low or none", 2: "Partial", 3: "Good"},
        "pt": {0: "NSP", 1: "Baixa ou inexistente", 2: "Parcial", 3: "Boa"},
        "ar": {0: "لا أعلم", 1: "ضعيفة أو غير متاحة", 2: "جزئية", 3: "جيدة"},
    },
    "feasibility": {
        "fr": {0: "NSP", 1: "Difficile", 2: "Modérée", 3: "Facile"},
        "en": {0: "DNK", 1: "Difficult", 2: "Moderate", 3: "Easy"},
        "pt": {0: "NSP", 1: "Difícil", 2: "Moderada", 3: "Fácil"},
        "ar": {0: "لا أعلم", 1: "صعب", 2: "متوسط", 3: "سهل"},
    },
}

def score_format(lang: str, criterion: str):
    """Formatter for score selectboxes (criterion-aware).

    We include a None option (placeholder) so we don't prefill answers.
    """
    placeholder_fr = "— Sélectionner —"
    placeholder_en = "— Select —"
    scale = SCORE_SCALES.get(criterion, SCORE_SCALES["demand"])
    mapping = scale.get(lang) or scale[lang_base(lang)]

    def _fmt(v):
        if v is None or v == "":
            return t(lang, placeholder_fr, placeholder_en, "— Selecionar —", "— اختر —")
        try:
            iv = int(v)
        except Exception:
            return str(v)
        return mapping.get(iv, str(v))

    return _fmt


ROLE_OPTIONS_FR = [
    "DG/DGA/SG",
    "Directeur",
    "Conseiller",
    "Chef de division",
    "Chef de bureau",
    "Autre",
]
ROLE_OPTIONS_EN = [
    "DG/DGA/SG",
    "Director",
    "Advisor",
    "Head of division",
    "Head of office",
    "Other",
]
ROLE_OPTIONS_PT = [
    "DG/DGA/SG",
    "Diretor",
    "Conselheiro",
    "Chefe de divisão",
    "Chefe de gabinete",
    "Outro",
]
ROLE_OPTIONS_AR = [
    "DG/DGA/SG",
    "مدير",
    "مستشار",
    "رئيس قسم",
    "رئيس مصلحة",
    "أخرى",
]


# =========================
# Helpers : i18n and state
# =========================

LANG_OPTIONS = ["fr", "en", "pt", "ar"]
LANG_LABELS = {
    "fr": "Français",
    "en": "English",
    "pt": "Português",
    "ar": "العربية",
}

PT_TRANSLATIONS = {
    "Langue / Language": "Idioma",
    "— Select —": "— Selecionar —",
    "— Sélectionner —": "— Selecionar —",
    "Section 1: Instructions": "Secção 1: Instruções",
    "Section 2: Respondent identification": "Secção 2: Identificação do respondente",
    "Section 3: Scope of response": "Secção 3: Âmbito da resposta",
    "Section 4: Priority domains": "Secção 4: Domínios prioritários",
    "Section 5: Priority indicators and scoring": "Secção 5: Estatísticas prioritárias e pontuação",
    "Section 6: Gender dimension": "Secção 6: Dimensão de género",
    "Section 7: Gender priorities": "Secção 7: Prioridades de género",
    "Section 8: Capacity and feasibility (12–24 months)": "Secção 8: Capacidade e viabilidade (12–24 meses)",
    "Section 9: Harmonization and quality": "Secção 9: Harmonização e qualidade",
    "Section 10: Dissemination": "Secção 10: Disseminação",
    "Section 11: Relevant data sources": "Secção 11: Fontes de dados pertinentes",
    "Section 12: Open questions": "Secção 12: Questões abertas",
    "SUBMIT": "ENVIAR",
    "SUBMIT questionnaire": "ENVIAR questionário",
    "Organization Name": "Nome da organização",
    "Country of residence": "País de residência",
    "Other country (please specify)": "Outro país (especifique)",
    "Email": "Email",
    "Stakeholder type": "Tipo de ator",
    "Role/Function": "Função",
    "Specify (role)": "Especificar (função)",
    "Please fix the following:": "Por favor, corrija os pontos seguintes :",
    "National": "Nacional",
    "Regional (REC)": "Regional (CER)",
    "Continental (AU)": "Continental (UA)",
    "International": "Internacional",
    "Other": "Outro",
    "Yes": "Sim",
    "No": "Não",
    "In preparation": "Em preparação",
    "Under implementation AND new one in preparation": "Em implementação E nova estratégia em preparação",
    "Status of the current NSDS / national statistical plan": "Estado atual da ENDE / plano estatístico nacional",
    "Pre-selection (5–10 domains)": "Pré-seleção (5–10 domínios)",
    "Rank 1": "Classificação 1",
    "Rank 2": "Classificação 2",
    "Rank 3": "Classificação 3",
    "Rank 4": "Classificação 4",
    "Rank 5": "Classificação 5",
    "Step 1: Pre-selection": "Etapa 1: Pré-seleção",
    "Step 2: Rank TOP 5": "Etapa 2: Classificar o TOP 5",
    "Step A: Select indicators": "Etapa A: Selecionar indicadores",
    "Step B: Multi-criteria scoring": "Etapa B: Pontuação multicritério",
    "Select": "Selecionar",
    "Specify": "Especificar",
    "Specify (Other)": "Especificar (Outro)",
    "Relevant data sources": "Fontes de dados pertinentes",
    "2 to 4 most relevant data sources": "2 a 4 fontes de dados mais pertinentes",
    "Other: please specify": "Outro : especifique",
    "Question 1 / 3": "Pergunta 1 / 3",
    "Question 2 / 3": "Pergunta 2 / 3",
    "Question 3 / 3": "Pergunta 3 / 3",
    "Confirmation": "Confirmação",
    "⬅ Previous question": "⬅ Pergunta anterior",
    "Next question ➡": "Pergunta seguinte ➡",
    "Go to confirmation ➡": "Ir para a confirmação ➡",
    "OK (section completed)": "OK (secção concluída)",
    "⬅ Previous": "⬅ Anterior",
    "Next ➡": "Seguinte ➡",
    "Navigation": "Navegação",
    "Go to": "Ir para",
    "Draft": "Rascunho",
    "Save now": "Guardar agora",
    "ENVOYER": "ENVIAR",
    "✅ SUBMIT and save": "✅ ENVIAR e guardar",
    "Thank you! Your submission has been saved.": "Obrigado ! O seu questionário foi registado.",
    "Everything is ready. Click **SUBMIT** to send your questionnaire.": "Tudo está pronto. Clique em **ENVIAR** para submeter o questionário.",
    "Unified app (FR/EN) – hidden codes – built-in quality controls.": "Aplicação multilingue (FR/EN/PT/AR) – códigos ocultos – controlos de qualidade integrados.",
}

AR_TRANSLATIONS = {
    "Langue / Language": "اللغة",
    "— Select —": "— اختر —",
    "— Sélectionner —": "— اختر —",
    "Section 1: Instructions": "القسم 1: التعليمات",
    "Section 2: Respondent identification": "القسم 2: تعريف المجيب",
    "Section 3: Scope of response": "القسم 3: نطاق الإجابة",
    "Section 4: Priority domains": "القسم 4: المجالات ذات الأولوية",
    "Section 5: Priority indicators and scoring": "القسم 5: الإحصاءات ذات الأولوية والتقييم",
    "Section 6: Gender dimension": "القسم 6: البعد الجنساني",
    "Section 7: Gender priorities": "القسم 7: أولويات النوع الاجتماعي",
    "Section 8: Capacity and feasibility (12–24 months)": "القسم 8: القدرات وإمكانية التنفيذ (12–24 شهراً)",
    "Section 9: Harmonization and quality": "القسم 9: المواءمة والجودة",
    "Section 10: Dissemination": "القسم 10: النشر",
    "Section 11: Relevant data sources": "القسم 11: مصادر البيانات ذات الصلة",
    "Section 12: Open questions": "القسم 12: الأسئلة المفتوحة",
    "SUBMIT": "إرسال",
    "SUBMIT questionnaire": "إرسال الاستبيان",
    "Organization Name": "اسم المؤسسة",
    "Country of residence": "بلد الإقامة",
    "Other country (please specify)": "بلد آخر (يرجى التحديد)",
    "Email": "البريد الإلكتروني",
    "Stakeholder type": "نوع الجهة الفاعلة",
    "Role/Function": "الصفة / الوظيفة",
    "Specify (role)": "حدد (الوظيفة)",
    "Please fix the following:": "يرجى تصحيح ما يلي :",
    "National": "وطني",
    "Regional (REC)": "إقليمي (التجمع الاقتصادي الإقليمي)",
    "Continental (AU)": "قاري (الاتحاد الأفريقي)",
    "International": "دولي",
    "Other": "أخرى",
    "Yes": "نعم",
    "No": "لا",
    "In preparation": "قيد الإعداد",
    "Under implementation AND new one in preparation": "قيد التنفيذ مع إعداد استراتيجية جديدة",
    "Status of the current NSDS / national statistical plan": "وضع الاستراتيجية الوطنية الحالية لتطوير الإحصاء / الخطة الإحصائية الوطنية",
    "Pre-selection (5–10 domains)": "الاختيار الأولي (5 إلى 10 مجالات)",
    "Rank 1": "الترتيب 1",
    "Rank 2": "الترتيب 2",
    "Rank 3": "الترتيب 3",
    "Rank 4": "الترتيب 4",
    "Rank 5": "الترتيب 5",
    "Step 1: Pre-selection": "المرحلة 1: الاختيار الأولي",
    "Step 2: Rank TOP 5": "المرحلة 2: ترتيب أفضل 5 مجالات",
    "Step A: Select indicators": "المرحلة أ: اختيار المؤشرات",
    "Step B: Multi-criteria scoring": "المرحلة ب: التقييم متعدد المعايير",
    "Select": "اختر",
    "Specify": "حدد",
    "Specify (Other)": "حدد (أخرى)",
    "Relevant data sources": "مصادر البيانات ذات الصلة",
    "2 to 4 most relevant data sources": "من 2 إلى 4 من أكثر مصادر البيانات صلة",
    "Other: please specify": "أخرى : يرجى التحديد",
    "Question 1 / 3": "السؤال 1 / 3",
    "Question 2 / 3": "السؤال 2 / 3",
    "Question 3 / 3": "السؤال 3 / 3",
    "Confirmation": "التأكيد",
    "⬅ Previous question": "⬅ السؤال السابق",
    "Next question ➡": "السؤال التالي ➡",
    "Go to confirmation ➡": "الانتقال إلى التأكيد ➡",
    "OK (section completed)": "حسنًا (اكتمل القسم)",
    "⬅ Previous": "⬅ السابق",
    "Next ➡": "التالي ➡",
    "Navigation": "التنقل",
    "Go to": "اذهب إلى",
    "Draft": "مسودة",
    "Save now": "احفظ الآن",
    "✅ SUBMIT and save": "✅ إرسال وحفظ",
    "Thank you! Your submission has been saved.": "شكرًا ! تم حفظ استجابتكم.",
    "Everything is ready. Click **SUBMIT** to send your questionnaire.": "كل شيء جاهز. انقر على **إرسال** لإرسال الاستبيان.",
    "Unified app (FR/EN) – hidden codes – built-in quality controls.": "تطبيق متعدد اللغات (FR/EN/PT/AR) – رموز مخفية – ضوابط جودة مدمجة.",
    "Household surveys": "المسوحات الأسرية",
    "Enterprise surveys": "مسوحات المنشآت",
    "Censuses": "التعدادات",
    "Administrative data": "البيانات الإدارية",
    "Civil registration and vital statistics (CRVS)": "التسجيل المدني والإحصاءات الحيوية (CRVS)",
    "Geospatial data": "البيانات الجغرافية المكانية",
    "Private data": "البيانات الخاصة",
    "Web portal / dashboards": "بوابة إلكترونية / لوحات معلومات",
    "Press releases / bulletins": "بيانات صحفية / نشرات",
    "Anonymized microdata (secure access)": "بيانات جزئية مجهولة الهوية (ولوج آمن)",
    "Workshops and webinars": "ورشات عمل وندوات عبر الإنترنت",
    "Manuals on common standards and methods (by domain) available": "توفر أدلة للمعايير والمناهج المشتركة (حسب المجال)",
    "Functional quality assurance framework (quality toolkit) ": "إطار وظيفي لضمان الجودة",
    "Data validation and certification procedures (certified quality) ": "إجراءات التحقق من البيانات واعتمادها",
    "Toolkit / mechanisms for cross-sector consistency of national data": "آليات لضمان الاتساق بين القطاعات في البيانات الوطنية",
    "Strengthening NSS technical capacity": "تعزيز القدرات التقنية للنظام الإحصائي الوطني",
    "Strengthening NSO leadership within the NSS": "تعزيز قيادة المكتب الوطني للإحصاء داخل النظام الإحصائي الوطني",
    "Specialized Technical Groups (STGs/AU) operational": "تشغيل المجموعات التقنية المتخصصة (الاتحاد الأفريقي)",
    "Available statistical skills": "المهارات الإحصائية المتاحة",
    "Access to administrative data": "إمكانية الوصول إلى البيانات الإدارية",
    "Available funding": "التمويل المتاح",
    "Digital tools (collection, processing, dissemination)": "الأدوات الرقمية (الجمع والمعالجة والنشر)",
    "Legal framework for data sharing": "الإطار القانوني لتبادل البيانات",
    "Inter-institutional coordination": "التنسيق بين المؤسسات",
    "Disaggregation by sex": "التصنيف حسب الجنس",
    "Disaggregation by age": "التصنيف حسب العمر",
    "Urban / rural": "حضري / ريفي",
    "Disability": "الإعاقة",
    "Wealth quintile": "خمس الثروة",
    "Gender-based violence (GBV)": "العنف القائم على النوع الاجتماعي",
    "Unpaid domestic work": "العمل المنزلي غير المأجور",
}

def lang_base(lang: str) -> str:
    return "fr" if lang == "fr" else "en"

def tr(lang: str, text: str) -> str:
    if lang == "pt":
        return PT_TRANSLATIONS.get(text, text)
    if lang == "ar":
        return AR_TRANSLATIONS.get(text, text)
    return text

def t(lang: str, fr: str, en: str, pt: Optional[str] = None, ar: Optional[str] = None) -> str:
    if lang == "fr":
        return fr
    if lang == "en":
        return en
    if lang == "pt":
        return pt if pt is not None else (PT_TRANSLATIONS.get(en) or PT_TRANSLATIONS.get(fr) or en)
    if lang == "ar":
        return ar if ar is not None else (AR_TRANSLATIONS.get(en) or AR_TRANSLATIONS.get(fr) or en)
    return en

def get_role_options(lang: str) -> List[str]:
    if lang == "fr":
        return ROLE_OPTIONS_FR
    if lang == "pt":
        return ROLE_OPTIONS_PT
    if lang == "ar":
        return ROLE_OPTIONS_AR
    return ROLE_OPTIONS_EN

def is_other_value(v: Any) -> bool:
    return str(v or "").strip() in {"Autre", "Other", "Outro", "أخرى", "Autres", "Outros"}

def has_other_option(values: Any) -> bool:
    if not isinstance(values, list):
        return False
    return any(is_other_value(v) for v in values)


def now_utc_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def get_query_params() -> Dict[str, List[str]]:
    """Compatibility across Streamlit versions."""
    try:
        # Streamlit >= 1.30
        qp = st.query_params  # type: ignore
        return {k: list(v) if isinstance(v, (list, tuple)) else [str(v)] for k, v in qp.items()}
    except Exception:
        try:
            return st.experimental_get_query_params()
        except Exception:
            return {}


def set_query_params(params: Dict[str, Any]) -> None:
    try:
        st.query_params.update(params)  # type: ignore
    except Exception:
        try:
            st.experimental_set_query_params(**params)
        except Exception:
            pass


def init_session() -> None:
    if "lang" not in st.session_state:
        st.session_state.lang = "fr"
    if "nav_idx" not in st.session_state:
        st.session_state.nav_idx = 0
    if "responses" not in st.session_state:
        st.session_state["responses"] = {}
    elif not isinstance(st.session_state.get("responses"), dict):
        st.session_state["responses"] = {}
    if "submission_id" not in st.session_state:
        st.session_state.submission_id = None
    if "admin_authed" not in st.session_state:
        st.session_state.admin_authed = False
    if "admin_role" not in st.session_state:
        st.session_state.admin_role = None  # "admin" | "superadmin"
    if "draft_id" not in st.session_state:
        st.session_state.draft_id = None

    if "draft_exists" not in st.session_state:
        st.session_state.draft_exists = False
    if "draft_resume_notice_shown" not in st.session_state:
        st.session_state.draft_resume_notice_shown = False
    if "draft_restored" not in st.session_state:
        st.session_state.draft_restored = False
    if "last_draft_save_ts" not in st.session_state:
        st.session_state.last_draft_save_ts = 0.0
    if "r12_substep" not in st.session_state:
        st.session_state["r12_substep"] = 0  # 0..2=open questions, 3=confirmation


def ensure_responses() -> None:
    """Garantit l’existence de st.session_state['responses'] (dict)."""
    if "responses" not in st.session_state or not isinstance(st.session_state.get("responses"), dict):
        st.session_state["responses"] = {}


def resp_get(key: str, default=None):
    ensure_responses()
    return st.session_state["responses"].get(key, default)


def resp_set(key: str, value) -> None:
    ensure_responses()
    st.session_state["responses"][key] = value



def normalize_availability(v_raw: Any, scoring_version: Any) -> int:
    """Normalise la disponibilité sur l'échelle 'Bonne=3' (SCORING_VERSION=3).

    - v3+ : on conserve la valeur telle quelle (0–3).
    - v1/v2 ou absence de version : on inverse (1<->3) car l'ancien codage correspondait à un "écart" / ou à une disponibilité inversée.
    """
    try:
        iv = int(v_raw)
    except Exception:
        return 0
    if iv == 0:
        return 0

    try:
        ver = int(scoring_version)
    except Exception:
        ver = 0

    if ver >= SCORING_VERSION:
        return iv

    # Inversion pour les versions antérieures : 1<->3, 2 inchangé
    if iv in (1, 2, 3):
        return 4 - iv
    return iv

def ensure_draft_id() -> Optional[str]:
    """Ensure a stable draft id exists (used for mobile resume)."""
    if st.session_state.get("draft_id"):
        return st.session_state.draft_id
    email = (resp_get("email", "") or "").strip()
    if not email:
        return None
    draft_id = str(uuid.uuid4())
    st.session_state.draft_id = draft_id
    # Keep any existing query params (admin, etc.)
    try:
        qp = get_query_params()
        qp["rid"] = [draft_id]
        set_query_params({k: v[0] if len(v) == 1 else v for k, v in qp.items()})
    except Exception:
        pass
    return draft_id


def autosave_draft(force: bool = False) -> Tuple[bool, str]:
    """Persist current responses to DB to mitigate mobile refresh/resets."""
    draft_id = st.session_state.get("draft_id")
    email = (resp_get("email", "") or "").strip()
    if not draft_id or not email:
        return False, "no_draft_id_or_email"
    now_ts = time.time()
    last_ts = float(st.session_state.get("last_draft_save_ts", 0.0) or 0.0)
    if (not force) and (now_ts - last_ts < 2.0):
        return True, "skipped_rate_limit"
    payload = {
        "responses": st.session_state.responses,
        "nav_idx": int(st.session_state.get("nav_idx", 0)),
        "lang": st.session_state.get("lang", "fr"),
    }
    try:
        db_save_draft(draft_id, email, payload)
        st.session_state.last_draft_save_ts = now_ts
        return True, "saved"
    except Exception as e:
        return False, str(e)


def maybe_restore_draft() -> None:
    """Restore responses from DB if URL contains rid and session is empty."""
    if st.session_state.get("draft_restored"):
        return
    st.session_state.draft_restored = True

    qp = get_query_params()
    rid = None
    if "rid" in qp and qp["rid"]:
        rid = qp["rid"][0]
    if not rid:
        return

    # Do not restore while in admin mode
    if "admin" in qp and qp["admin"] and qp["admin"][0] in ["1", "true", "yes"]:
        return

    payload = db_load_draft(rid)
    st.session_state.draft_exists = bool(payload)

    # Restore only if session is empty (avoid overriding ongoing input)
    if st.session_state.get("responses") and len(st.session_state.responses) > 0:
        st.session_state.draft_id = rid
        return


    if not payload:
        st.session_state.draft_id = rid
        return

    responses = payload.get("responses", {})
    if isinstance(responses, dict):
        st.session_state.responses = responses

    try:
        st.session_state.nav_idx = int(payload.get("nav_idx", 0))
    except Exception:
        st.session_state.nav_idx = 0

    lang = payload.get("lang", None)
    if lang in LANG_OPTIONS:
        st.session_state.lang = lang

    st.session_state.draft_id = rid




# =========================
# Data : longlist loader
# =========================

@st.cache_data(show_spinner=False)
def load_longlist() -> pd.DataFrame:
    """
    Charge la longlist (statistiques par domaine) depuis :
    - CSV : data/indicator_longlist.csv (prioritaire)
    - XLSX : data/longlist.xlsx (fallback)
    Tolère aussi les fichiers placés à la racine du dépôt.

    Si aucun fichier n'est trouvé, l'application démarre quand même,
    mais les listes déroulantes de la Rubrique 4/5 seront vides.
    """
    csv_candidates = [
        LONG_LIST_CSV,
        "indicator_longlist.csv",
        os.path.join(".", "indicator_longlist.csv"),
        os.path.join(".", "data", "indicator_longlist.csv"),
    ]
    xlsx_candidates = [
        LONG_LIST_XLSX,
        "longlist.xlsx",
        os.path.join(".", "longlist.xlsx"),
        os.path.join(".", "data", "longlist.xlsx"),
    ]

    # 1) CSV (prioritaire si la traduction EN est suffisamment complète)
    df_csv = None
    df_csv_path = None
    for p in csv_candidates:
        if os.path.exists(p):
            df_csv = pd.read_csv(p, dtype=str).fillna("")
            df_csv_path = p
            break

    if df_csv is not None:
        # Sanity check : si beaucoup de libellés EN sont vides, on préfère l'XLSX (souvent plus à jour)
        if "stat_label_en" in df_csv.columns:
            miss_ratio = (df_csv["stat_label_en"].astype(str).str.strip() == "").mean()
        else:
            miss_ratio = 1.0

        if "domain_label_pt" not in df_csv.columns:
            df_csv["domain_label_pt"] = df_csv.get("domain_label_en", df_csv.get("domain_label_fr", ""))
        if "domain_label_ar" not in df_csv.columns:
            df_csv["domain_label_ar"] = df_csv.get("domain_label_en", df_csv.get("domain_label_fr", ""))
        if "stat_label_pt" not in df_csv.columns:
            df_csv["stat_label_pt"] = df_csv.get("stat_label_en", df_csv.get("stat_label_fr", ""))
        if "stat_label_ar" not in df_csv.columns:
            df_csv["stat_label_ar"] = df_csv.get("stat_label_en", df_csv.get("stat_label_fr", ""))

        if miss_ratio <= 0.20:
            df_csv.attrs["source_path"] = df_csv_path
            return df_csv

    # 2) XLSX (format utilisateur) (format utilisateur)
    for p in xlsx_candidates:
        if os.path.exists(p):
            df = pd.read_excel(p, dtype=str).fillna("")
            df.attrs["source_path"] = p

            # Colonnes attendues (minimum) : Domain_code, Domain_label_fr, Stat_label_fr
            if set(["Domain_code", "Domain_label_fr", "Stat_label_fr"]).issubset(df.columns):
                out = pd.DataFrame()
                out["domain_code"] = df["Domain_code"].astype(str).str.strip()

                # Labels FR (on retire le préfixe code "D01|...")
                out["domain_label_fr"] = df["Domain_label_fr"].astype(str).str.split("|", n=1).str[-1].str.strip()
                out["stat_code"] = df["Stat_label_fr"].astype(str).str.split("|", n=1).str[0].str.strip()
                out["stat_label_fr"] = df["Stat_label_fr"].astype(str).str.split("|", n=1).str[-1].str.strip()

                # Labels EN/PT/AR si disponibles, sinon fallback EN puis FR
                if "Domain_label_en" in df.columns:
                    out["domain_label_en"] = df["Domain_label_en"].astype(str).str.split("|", n=1).str[-1].str.strip()
                else:
                    out["domain_label_en"] = out["domain_label_fr"]

                if "Stat_label_en" in df.columns:
                    out["stat_label_en"] = df["Stat_label_en"].astype(str).str.split("|", n=1).str[-1].str.strip()
                else:
                    out["stat_label_en"] = out["stat_label_fr"]

                if "Domain_label_pt" in df.columns:
                    out["domain_label_pt"] = df["Domain_label_pt"].astype(str).str.split("|", n=1).str[-1].str.strip()
                else:
                    out["domain_label_pt"] = out["domain_label_en"]

                if "Stat_label_pt" in df.columns:
                    out["stat_label_pt"] = df["Stat_label_pt"].astype(str).str.split("|", n=1).str[-1].str.strip()
                else:
                    out["stat_label_pt"] = out["stat_label_en"]

                if "Domain_label_ar" in df.columns:
                    out["domain_label_ar"] = df["Domain_label_ar"].astype(str).str.split("|", n=1).str[-1].str.strip()
                else:
                    out["domain_label_ar"] = out["domain_label_en"]

                if "Stat_label_ar" in df.columns:
                    out["stat_label_ar"] = df["Stat_label_ar"].astype(str).str.split("|", n=1).str[-1].str.strip()
                else:
                    out["stat_label_ar"] = out["stat_label_en"]

                out.attrs["source_path"] = p
                return out[[
                    "domain_code",
                    "domain_label_fr",
                    "domain_label_en",
                    "domain_label_pt",
                    "domain_label_ar",
                    "stat_code",
                    "stat_label_fr",
                    "stat_label_en",
                    "stat_label_pt",
                    "stat_label_ar",
                ]]


    # Fallback final : si un CSV a été trouvé (même avec traduction EN incomplète), on le renvoie
    if df_csv is not None:
        for c_new, c_fallback in [
            ("domain_label_pt", "domain_label_en"),
            ("domain_label_ar", "domain_label_en"),
            ("stat_label_pt", "stat_label_en"),
            ("stat_label_ar", "stat_label_en"),
        ]:
            if c_new not in df_csv.columns:
                df_csv[c_new] = df_csv.get(c_fallback, df_csv.get(c_fallback.replace("_en", "_fr"), ""))
        df_csv.attrs["source_path"] = df_csv_path or ""
        return df_csv
# Aucun fichier trouvé : dataframe vide
    empty = pd.DataFrame(columns=[
        "domain_code",
        "domain_label_fr",
        "domain_label_en",
        "domain_label_pt",
        "domain_label_ar",
        "stat_code",
        "stat_label_fr",
        "stat_label_en",
        "stat_label_pt",
        "stat_label_ar",
    ])
    empty.attrs["source_path"] = ""
    return empty



# =========================
# Data : countries loader
# =========================

@st.cache_data(show_spinner=False)
def load_countries() -> pd.DataFrame:
    """
    Charge la liste des pays (ISO3 + noms FR/EN) depuis :
    - data/COUNTRY_ISO3_with_EN.xlsx (prioritaire)
    Tolère aussi le fichier placé à la racine du dépôt.

    Colonnes attendues (au minimum) :
    - COUNTRY_ISO3
    - COUNTRY_NAME_FR
    - COUNTRY_NAME_EN
    """
    candidates = [
        COUNTRY_XLSX,
        "COUNTRY_ISO3_with_EN.xlsx",
        os.path.join(".", "COUNTRY_ISO3_with_EN.xlsx"),
        os.path.join(".", "data", "COUNTRY_ISO3_with_EN.xlsx"),
        # Fallbacks (ancien nom éventuel)
        "COUNTRY_ISO3.xlsx",
        os.path.join(".", "COUNTRY_ISO3.xlsx"),
        os.path.join(".", "data", "COUNTRY_ISO3.xlsx"),
    ]
    for p in candidates:
        if os.path.exists(p):
            try:
                df = pd.read_excel(p, dtype=str).fillna("")
                # Normalisation des noms de colonnes
                df.columns = [str(c).strip() for c in df.columns]
                if "COUNTRY_ISO3" not in df.columns:
                    continue
                if "COUNTRY_NAME_FR" not in df.columns:
                    df["COUNTRY_NAME_FR"] = ""
                if "COUNTRY_NAME_EN" not in df.columns:
                    df["COUNTRY_NAME_EN"] = ""
                if "COUNTRY_NAME_PT" not in df.columns:
                    df["COUNTRY_NAME_PT"] = df["COUNTRY_NAME_EN"]
                if "COUNTRY_NAME_AR" not in df.columns:
                    df["COUNTRY_NAME_AR"] = df["COUNTRY_NAME_EN"]
                df["COUNTRY_ISO3"] = df["COUNTRY_ISO3"].astype(str).str.strip().str.upper()
                df["COUNTRY_NAME_FR"] = df["COUNTRY_NAME_FR"].astype(str).str.strip()
                df["COUNTRY_NAME_EN"] = df["COUNTRY_NAME_EN"].astype(str).str.strip()
                df["COUNTRY_NAME_PT"] = df["COUNTRY_NAME_PT"].astype(str).str.strip()
                df["COUNTRY_NAME_AR"] = df["COUNTRY_NAME_AR"].astype(str).str.strip()
                df = df[df["COUNTRY_ISO3"] != ""].copy()
                df.attrs["source_path"] = p
                return df
            except Exception:
                continue
    empty = pd.DataFrame(columns=["COUNTRY_ISO3", "COUNTRY_NAME_FR", "COUNTRY_NAME_EN", "COUNTRY_NAME_PT", "COUNTRY_NAME_AR"])
    empty.attrs["source_path"] = ""
    return empty


def _pick_country_name_col(lang: str) -> str:
    if lang == "fr":
        return "COUNTRY_NAME_FR"
    if lang == "pt":
        return "COUNTRY_NAME_PT"
    if lang == "ar":
        return "COUNTRY_NAME_AR"
    return "COUNTRY_NAME_EN"


def _pick_longlist_domain_col(lang: str) -> str:
    if lang == "fr":
        return "domain_label_fr"
    if lang == "pt":
        return "domain_label_pt"
    if lang == "ar":
        return "domain_label_ar"
    return "domain_label_en"


def _pick_longlist_stat_col(lang: str) -> str:
    if lang == "fr":
        return "stat_label_fr"
    if lang == "pt":
        return "stat_label_pt"
    if lang == "ar":
        return "stat_label_ar"
    return "stat_label_en"


def country_maps(df_c: pd.DataFrame) -> Tuple[List[str], Dict[str, str], Dict[str, str], Dict[str, str], Dict[str, str]]:
    """Retourne (iso3_list, iso3_to_fr, iso3_to_en, iso3_to_pt, iso3_to_ar)."""
    if df_c is None or df_c.empty:
        return [], {}, {}, {}, {}
    iso3_to_fr: Dict[str, str] = {}
    iso3_to_en: Dict[str, str] = {}
    iso3_to_pt: Dict[str, str] = {}
    iso3_to_ar: Dict[str, str] = {}
    for _, r in df_c.iterrows():
        iso3 = str(r.get("COUNTRY_ISO3", "")).strip().upper()
        if not iso3:
            continue
        iso3_to_fr[iso3] = str(r.get("COUNTRY_NAME_FR", "")).strip()
        iso3_to_en[iso3] = str(r.get("COUNTRY_NAME_EN", "")).strip()
        iso3_to_pt[iso3] = str(r.get("COUNTRY_NAME_PT", "")).strip()
        iso3_to_ar[iso3] = str(r.get("COUNTRY_NAME_AR", "")).strip()
    iso3_list = sorted(set(list(iso3_to_fr.keys()) + list(iso3_to_en.keys()) + list(iso3_to_pt.keys()) + list(iso3_to_ar.keys())))
    return iso3_list, iso3_to_fr, iso3_to_en, iso3_to_pt, iso3_to_ar


def country_label(
    iso3: str,
    lang: str,
    iso3_to_fr: Dict[str, str],
    iso3_to_en: Dict[str, str],
    iso3_to_pt: Optional[Dict[str, str]] = None,
    iso3_to_ar: Optional[Dict[str, str]] = None,
) -> str:
    if not iso3:
        return ""
    iso3_to_pt = iso3_to_pt or {}
    iso3_to_ar = iso3_to_ar or {}
    if lang == "fr":
        return (iso3_to_fr.get(iso3) or iso3_to_en.get(iso3) or iso3_to_pt.get(iso3) or iso3_to_ar.get(iso3) or iso3).strip()
    if lang == "pt":
        return (iso3_to_pt.get(iso3) or iso3_to_en.get(iso3) or iso3_to_fr.get(iso3) or iso3_to_ar.get(iso3) or iso3).strip()
    if lang == "ar":
        return (iso3_to_ar.get(iso3) or iso3_to_en.get(iso3) or iso3_to_fr.get(iso3) or iso3_to_pt.get(iso3) or iso3).strip()
    return (iso3_to_en.get(iso3) or iso3_to_fr.get(iso3) or iso3_to_pt.get(iso3) or iso3_to_ar.get(iso3) or iso3).strip()


def domains_from_longlist(df_long: pd.DataFrame, lang: str) -> List[Tuple[str, str]]:
    if df_long.empty:
        return []
    col = _pick_longlist_domain_col(lang)
    if col not in df_long.columns:
        col = "domain_label_en" if "domain_label_en" in df_long.columns else "domain_label_fr"
    tmp = df_long[["domain_code", col]].drop_duplicates().sort_values(["domain_code", col])
    return [(r["domain_code"], r[col]) for _, r in tmp.iterrows()]


def stats_for_domain(df_long: pd.DataFrame, domain_code: str, lang: str) -> List[Tuple[str, str]]:
    if df_long.empty or not domain_code:
        return []
    col = _pick_longlist_stat_col(lang)
    if col not in df_long.columns:
        col = "stat_label_en" if "stat_label_en" in df_long.columns else "stat_label_fr"
    tmp = df_long[df_long["domain_code"] == domain_code][["stat_code", col]].drop_duplicates()
    tmp = tmp.sort_values(["stat_code", col])
    return [(r["stat_code"], (str(r[col]).strip() if str(r[col]).strip() else r["stat_code"])) for _, r in tmp.iterrows()]


# =========================
# Stockage : SQLite (local)
# =========================


def domain_label_map(df_long: pd.DataFrame, lang: str) -> Dict[str, str]:
    """Map domain_code -> label in selected language."""
    if df_long is None or df_long.empty:
        return {}
    col = _pick_longlist_domain_col(lang)
    if col not in df_long.columns:
        col = "domain_label_en" if "domain_label_en" in df_long.columns else "domain_label_fr"
    m = {}
    for _, r in df_long.drop_duplicates("domain_code").iterrows():
        code = str(r["domain_code"])
        lbl = str(r.get(col, "")).strip()
        if not lbl:
            lbl = str(r.get("domain_label_fr", "")).strip() or code
        m[code] = lbl
    return m

def stat_label_map(df_long: pd.DataFrame, lang: str) -> Dict[str, str]:
    """Map stat_code -> label in selected language."""
    if df_long is None or df_long.empty:
        return {}
    col = _pick_longlist_stat_col(lang)
    if col not in df_long.columns:
        col = "stat_label_en" if "stat_label_en" in df_long.columns else "stat_label_fr"
    m = {}
    for _, r in df_long.drop_duplicates("stat_code").iterrows():
        code = str(r["stat_code"])
        lbl = str(r.get(col, "")).strip()
        if not lbl:
            lbl = str(r.get("stat_label_fr", "")).strip() or code
        m[code] = lbl
    return m

def build_publication_report_docx(lang: str, filtered_payloads: pd.DataFrame, by_domain: pd.DataFrame, by_stat: pd.DataFrame, scored_rows: pd.DataFrame) -> bytes:
    """
    Génère un rapport Word 'publication' enrichi :
    - profil des répondants
    - domaines TOP 5 (fréquences)
    - tableau agrégé des statistiques et scores moyens
    - graphiques (bar charts)
    - annexes
    """
    from docx import Document
    from docx.shared import Inches
    import matplotlib.pyplot as plt

    doc = Document()
    title = t(lang, "Rapport de synthèse – Consultation sur les statistiques prioritaires", "Summary report – Consultation on priority statistics")
    doc.add_heading(title, level=0)
    doc.add_paragraph(t(lang, f"Date : {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", f"Date: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"))
    doc.add_paragraph("")

    # Sample profile
    doc.add_heading(t(lang, "Profil des répondants", "Respondent profile"), level=1)
    n = len(filtered_payloads)
    doc.add_paragraph(t(lang, f"Nombre de réponses analysées : {n}", f"Number of responses analyzed: {n}"))

    # Countries
    if "pays" in filtered_payloads.columns:
        vc = filtered_payloads["pays"].fillna("").astype(str)
        vc = vc[vc.str.strip() != ""].value_counts().head(10)
        if len(vc) > 0:
            doc.add_paragraph(t(lang, "Top pays (10 premiers) :", "Top countries (top 10):"))
            for k, v in vc.items():
                doc.add_paragraph(f"- {k} : {v}", style=None)

    # Actor types
    if "type_acteur" in filtered_payloads.columns:
        vc = filtered_payloads["type_acteur"].fillna("").astype(str)
        vc = vc[vc.str.strip() != ""].value_counts()
        if len(vc) > 0:
            doc.add_paragraph(t(lang, "Répartition par type d’acteur :", "Distribution by stakeholder type:"))
            for k, v in vc.items():
                doc.add_paragraph(f"- {k} : {v}", style=None)

    # Domain aggregation
    doc.add_heading(t(lang, "Domaines prioritaires (scores moyens)", "Priority domains (mean scores)"), level=1)
    top_dom = by_domain.head(15).copy()
    # Table
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = t(lang, "Domaine", "Domain")
    hdr[1].text = t(lang, "Nb. soumissions", "Submissions")
    hdr[2].text = t(lang, "Nb. stats notées", "Scored indicators")
    hdr[3].text = t(lang, "Score moyen", "Mean score")
    for _, r in top_dom.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["domain_label"])
        row[1].text = str(int(r["n_submissions"]))
        row[2].text = str(int(r["n_stats"]))
        row[3].text = f"{float(r['mean_overall']):.2f}"

    # Chart domain
    try:
        fig = plt.figure()
        plt.bar(top_dom["domain_label"], top_dom["mean_overall"])
        plt.xticks(rotation=75, ha="right")
        plt.ylabel(t(lang, "Score moyen", "Mean score"))
        plt.tight_layout()
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format="png", dpi=150)
        plt.close(fig)
        img_stream.seek(0)
        doc.add_paragraph(t(lang, "Graphique : score moyen par domaine (top 15).", "Chart: mean score by domain (top 15)."))
        doc.add_picture(img_stream, width=Inches(6.5))
    except Exception:
        pass

    # Statistic aggregation
    doc.add_heading(t(lang, "Statistiques prioritaires (scores moyens)", "Priority indicators (mean scores)"), level=1)
    top_stat = by_stat.sort_values(["mean_overall", "n"], ascending=[False, False]).head(30).copy()
    table2 = doc.add_table(rows=1, cols=6)
    h = table2.rows[0].cells
    h[0].text = t(lang, "Domaine", "Domain")
    h[1].text = t(lang, "Statistique", "Indicator")
    h[2].text = t(lang, "N", "N")
    h[3].text = t(lang, "Demande", "Demand")
    h[4].text = t(lang, "Disponibilité", "Availability")
    h[5].text = t(lang, "Faisabilité", "Feasibility")
    for _, r in top_stat.iterrows():
        row = table2.add_row().cells
        row[0].text = str(r["domain_label"])
        row[1].text = str(r["stat_label"])
        row[2].text = str(int(r["n"]))
        row[3].text = f"{float(r['mean_demand']):.2f}"
        row[4].text = f"{float(r['mean_availability']):.2f}"
        row[5].text = f"{float(r['mean_feasibility']):.2f}"


    # Chart top overall indicators
    try:
        fig = plt.figure()
        plt.barh(top_stat["stat_label"].iloc[::-1], top_stat["mean_overall"].iloc[::-1])
        plt.xlabel(t(lang, "Score moyen", "Mean score"))
        plt.tight_layout()
        img2 = io.BytesIO()
        plt.savefig(img2, format="png", dpi=150)
        plt.close(fig)
        img2.seek(0)
        doc.add_paragraph(t(lang, "Graphique : score moyen par statistique (top 30).", "Chart: mean score by indicator (top 30)."))
        doc.add_picture(img2, width=Inches(6.5))
    except Exception:
        pass

    # Interpretation auto
    doc.add_heading(t(lang, "Interprétations automatiques", "Automatic interpretations"), level=1)
    # Simple rules
    best_dom = top_dom.iloc[0]
    doc.add_paragraph(
        t(
            lang,
            f"Le domaine le mieux noté est « {best_dom['domain_label']} » avec un score moyen de {best_dom['mean_overall']:.2f} (sur 3).",
            f"The highest-rated domain is “{best_dom['domain_label']}” with a mean score of {best_dom['mean_overall']:.2f} (out of 3)."
        )
    )
    best_stat = top_stat.iloc[0]
    doc.add_paragraph(
        t(
            lang,
            f"La statistique la mieux notée est « {best_stat['stat_label']} » (domaine : {best_stat['domain_label']}) avec un score moyen de {best_stat['mean_overall']:.2f}.",
            f"The highest-rated indicator is “{best_stat['stat_label']}” (domain: {best_stat['domain_label']}) with a mean score of {best_stat['mean_overall']:.2f}."
        )
    )

    # Annexes
    doc.add_heading(t(lang, "Annexes", "Annexes"), level=1)
    doc.add_paragraph(t(lang, "A1. Tableau détaillé (statistiques agrégées) – extrait", "A1. Detailed table (aggregated indicators) – excerpt"))
    annex = by_stat.head(50).copy()
    tab3 = doc.add_table(rows=1, cols=5)
    hh = tab3.rows[0].cells
    hh[0].text = t(lang, "Domaine", "Domain")
    hh[1].text = t(lang, "Statistique", "Indicator")
    hh[2].text = t(lang, "N", "N")
    hh[3].text = t(lang, "Score moyen", "Mean score")
    hh[4].text = t(lang, "Détail", "Detail")
    for _, r in annex.iterrows():
        rr = tab3.add_row().cells
        rr[0].text = str(r["domain_label"])
        rr[1].text = str(r["stat_label"])
        rr[2].text = str(int(r["n"]))
        rr[3].text = f"{float(r['mean_overall']):.2f}"
        if lang == "fr":
            rr[4].text = f"Demande={float(r['mean_demand']):.2f}, Disponibilité={float(r['mean_availability']):.2f}, Faisabilité={float(r['mean_feasibility']):.2f}"
        elif lang == "pt":
            rr[4].text = f"Procura={float(r['mean_demand']):.2f}, Disponibilidade={float(r['mean_availability']):.2f}, Viabilidade={float(r['mean_feasibility']):.2f}"
        elif lang == "ar":
            rr[4].text = f"الطلب={float(r['mean_demand']):.2f}, التوفر={float(r['mean_availability']):.2f}, الجدوى={float(r['mean_feasibility']):.2f}"
        else:
            rr[4].text = f"Demand={float(r['mean_demand']):.2f}, Availability={float(r['mean_availability']):.2f}, Feasibility={float(r['mean_feasibility']):.2f}"


    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()




def db_init() -> None:
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS submissions(
            submission_id TEXT PRIMARY KEY,
            submitted_at_utc TEXT,
            lang TEXT,
            email TEXT,
            payload_json TEXT
        )
    """)

    # Backward compatibility : add email column if existing DB was created with older schema
    try:
        cur.execute("PRAGMA table_info(submissions)")
        cols = [r[1] for r in cur.fetchall()]
        if "email" not in cols:
            cur.execute("ALTER TABLE submissions ADD COLUMN email TEXT")
    except Exception:
        pass

    # Helpful index (non-unique) for email lookups
    try:
        cur.execute("CREATE INDEX IF NOT EXISTS idx_submissions_email ON submissions(email)")
    except Exception:
        pass


    # Optionnel : empêcher les doublons email (best-effort)
    # Note : si des doublons existent déjà dans la base, la création de l’index UNIQUE échouera (sans bloquer l’app).
    try:
        cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS uq_submissions_email ON submissions(email)")
    except Exception:
        pass
    # Drafts (for mobile stability / resume)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS drafts(
            draft_id TEXT PRIMARY KEY,
            updated_at_utc TEXT,
            email TEXT,
            payload_json TEXT
        )
    """)
    try:
        cur.execute("CREATE INDEX IF NOT EXISTS idx_drafts_email ON drafts(email)")
    except Exception:
        pass

    # App config (e.g. hashed admin password override)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS app_config(
            k TEXT PRIMARY KEY,
            v TEXT,
            updated_at_utc TEXT
        )
    """)

    con.commit()
    con.close()



# =========================
# Admin auth helpers
# =========================

PBKDF2_ITERS = 200_000

def db_get_config(k: str) -> Optional[str]:
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("SELECT v FROM app_config WHERE k=? LIMIT 1", (k,))
    row = cur.fetchone()
    con.close()
    return row[0] if row and row[0] is not None else None


def db_set_config(k: str, v: str) -> None:
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute(
        "INSERT OR REPLACE INTO app_config(k, v, updated_at_utc) VALUES(?, ?, ?)",
        (k, v, now_utc_iso()),
    )
    con.commit()
    con.close()


def db_delete_config(k: str) -> None:
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("DELETE FROM app_config WHERE k=?", (k,))
    con.commit()
    con.close()


def _pbkdf2_sha256_hex(password: str, salt: bytes, iterations: int = PBKDF2_ITERS) -> str:
    dk = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, iterations)
    return dk.hex()


def _safe_eq(a: str, b: str) -> bool:
    try:
        return hmac.compare_digest(a or "", b or "")
    except Exception:
        return (a or "") == (b or "")


def _get_secret_or_env(name: str) -> Tuple[Optional[str], Optional[str]]:
    """Return (value, source) where source is 'secrets' | 'env' | None.

    Robust to different Streamlit versions and secrets layouts:
    - Direct key : ADMIN_PASSWORD = "..."
    - Nested table : [general] ADMIN_PASSWORD = "..."
    """
    val: Optional[str] = None
    src: Optional[str] = None

    # 1) Streamlit secrets
    try:
        secrets_obj = getattr(st, "secrets", None)
        if secrets_obj is not None:
            # Direct access (preferred)
            try:
                if hasattr(secrets_obj, "__contains__") and name in secrets_obj:
                    v = secrets_obj[name]
                    if v not in (None, ""):
                        val = str(v)
                        src = "secrets"
            except Exception:
                pass

            # Fallback : convert to dict then search (supports nested sections)
            if not val:
                try:
                    d = secrets_obj.to_dict() if hasattr(secrets_obj, "to_dict") else dict(secrets_obj)
                except Exception:
                    d = {}

                if isinstance(d, dict):
                    if name in d and d.get(name) not in (None, ""):
                        val = str(d.get(name))
                        src = "secrets"
                    else:
                        # Search nested dicts
                        for _k, _v in d.items():
                            if isinstance(_v, dict) and name in _v and _v.get(name) not in (None, ""):
                                val = str(_v.get(name))
                                src = "secrets"
                                break
    except Exception:
        # Leave val/src as None
        pass

    # 2) Environment variable (only if not found in secrets)
    if not val:
        env_val = os.environ.get(name, None)
        if env_val not in (None, ""):
            val = str(env_val)
            src = "env"

    return (val, src)


def get_admin_auth_source() -> Tuple[str, str]:
    """Human-readable indicator of current admin password source."""
    h = db_get_config("ADMIN_PASSWORD_HASH")
    s = db_get_config("ADMIN_PASSWORD_SALT")
    if h and s:
        return ("db", t(st.session_state.get("lang", "fr"), "base (haché)", "database (hashed)"))
    v, src = _get_secret_or_env("ADMIN_PASSWORD")
    if v and src:
        return (src, t(st.session_state.get("lang", "fr"), src, src))
    return ("none", t(st.session_state.get("lang", "fr"), "non configuré", "not configured"))


def verify_admin_password(pw: str) -> bool:
    pw = pw or ""
    h = db_get_config("ADMIN_PASSWORD_HASH")
    s = db_get_config("ADMIN_PASSWORD_SALT")
    it = db_get_config("ADMIN_PASSWORD_ITERS")
    if h and s:
        try:
            salt = bytes.fromhex(s)
            iterations = int(it) if it else PBKDF2_ITERS
            calc = _pbkdf2_sha256_hex(pw, salt, iterations)
            return _safe_eq(calc, h)
        except Exception:
            return False

    expected, _src = _get_secret_or_env("ADMIN_PASSWORD")
    if expected:
        return _safe_eq(pw, str(expected))
    return False


def verify_superadmin_password(pw: str) -> bool:
    pw = pw or ""
    expected, _src = _get_secret_or_env("SUPERADMIN_PASSWORD")
    if expected:
        return _safe_eq(pw, str(expected))
    return False


def set_admin_password(new_pw: str) -> None:
    """Set (hashed) admin password override in DB."""
    new_pw = (new_pw or "").strip()
    if not new_pw:
        raise ValueError("empty password")
    salt = secrets.token_bytes(16)
    h = _pbkdf2_sha256_hex(new_pw, salt, PBKDF2_ITERS)
    db_set_config("ADMIN_PASSWORD_HASH", h)
    db_set_config("ADMIN_PASSWORD_SALT", salt.hex())
    db_set_config("ADMIN_PASSWORD_ITERS", str(PBKDF2_ITERS))


def reset_admin_password_to_secrets_env() -> None:
    """Remove DB override so app falls back to secrets/env."""
    db_delete_config("ADMIN_PASSWORD_HASH")
    db_delete_config("ADMIN_PASSWORD_SALT")
    db_delete_config("ADMIN_PASSWORD_ITERS")

def db_email_exists(email: str) -> bool:
    email = (email or "").strip().lower()
    if not email or not os.path.exists(DB_PATH):
        return False
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("SELECT 1 FROM submissions WHERE lower(email)=? LIMIT 1", (email,))
    row = cur.fetchone()
    con.close()
    return row is not None


def db_save_submission(submission_id: str, lang: str, email: str, payload: Dict[str, Any]) -> None:
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
        INSERT OR REPLACE INTO submissions(submission_id, submitted_at_utc, lang, email, payload_json)
        VALUES(?, ?, ?, ?, ?)
    """, (submission_id, now_utc_iso(), lang, (email or "").strip().lower(), json.dumps(payload, ensure_ascii=False)))
    con.commit()
    con.close()



def db_save_draft(draft_id: str, email: str, payload: Dict[str, Any]) -> None:
    draft_id = (draft_id or "").strip()
    email = (email or "").strip().lower()
    if not draft_id or not email:
        return
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute(
        """
        INSERT OR REPLACE INTO drafts(draft_id, updated_at_utc, email, payload_json)
        VALUES(?, ?, ?, ?)
        """,
        (draft_id, now_utc_iso(), email, json.dumps(payload, ensure_ascii=False)),
    )
    con.commit()
    con.close()


def db_load_draft(draft_id: str) -> Optional[Dict[str, Any]]:
    draft_id = (draft_id or "").strip()
    if not draft_id or not os.path.exists(DB_PATH):
        return None
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("SELECT payload_json FROM drafts WHERE draft_id=? LIMIT 1", (draft_id,))
    row = cur.fetchone()
    con.close()
    if not row or not row[0]:
        return None
    try:
        return json.loads(row[0])
    except Exception:
        return None


def db_delete_draft(draft_id: str) -> None:
    draft_id = (draft_id or "").strip()
    if not draft_id or not os.path.exists(DB_PATH):
        return
    db_init()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("DELETE FROM drafts WHERE draft_id=?", (draft_id,))
    con.commit()
    con.close()

def db_read_submissions(limit: int = 2000) -> pd.DataFrame:
    if not os.path.exists(DB_PATH):
        return pd.DataFrame(columns=["submission_id", "submitted_at_utc", "lang", "email", "payload_json"])
    con = sqlite3.connect(DB_PATH)
    df = pd.read_sql_query(
        "SELECT submission_id, submitted_at_utc, lang, email, payload_json FROM submissions ORDER BY submitted_at_utc DESC LIMIT ?",
        con,
        params=(limit,),
    )
    con.close()
    return df


def db_dump_csv_bytes(limit: int = 2000000) -> bytes:
    """Export the SQLite submissions table to CSV bytes."""
    df = db_read_submissions(limit=limit)
    return df.to_csv(index=False).encode("utf-8-sig")



def flatten_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Create a 'flat' row for exports (comprehensive).
    - Keeps keys stable across FR/EN by mapping table items to canonical ids.
    - Serializes list/dict fields into '; ' / JSON strings as needed.
    """
    def _join_list(v: Any) -> str:
        if isinstance(v, list):
            return "; ".join([str(x) for x in v if x is not None and str(x).strip() != ""])
        return ""

    def _json(v: Any) -> str:
        try:
            return json.dumps(v, ensure_ascii=False)
        except Exception:
            return ""

    # Canonical mappings for table questions (FR/EN)
    GENDER_ITEM_MAP = {
        "Désagrégation par sexe": "sex",
        "Disaggregation by sex": "sex",
        "Sexe": "sex",
        "Sex": "sex",
        "Désagrégation par âge": "age",
        "Disaggregation by age": "age",
        "Âge": "age",
        "Age": "age",
        "Milieu urbain / rural": "urban_rural",
        "Urban / rural": "urban_rural",
        "Milieu urbain/rural": "urban_rural",
        "Urban/rural residence": "urban_rural",
        "Handicap": "disability",
        "Disability": "disability",
        "Quintile de richesse": "wealth_quintile",
        "Wealth quintile": "wealth_quintile",
        "Violences basées sur le genre (VBG)": "gbv",
        "Gender-based violence (GBV)": "gbv",
        "Temps domestique non rémunéré": "unpaid_domestic",
        "Unpaid domestic work": "unpaid_domestic",
    }
    CAPACITY_ITEM_MAP = {
        "Compétences statistiques disponibles": "skills_hr",
        "Available statistical skills": "skills_hr",
        "Accès aux données administratives": "access_admin_data",
        "Access to administrative data": "access_admin_data",
        "Financement disponible": "funding",
        "Available funding": "funding",
        "Financement": "funding",
        "Funding": "funding",
        "Outils numériques (collecte, traitement, diffusion)": "digital_tools",
        "Digital tools (collection, processing, dissemination)": "digital_tools",
        "Outils numériques": "digital_tools",
        "Digital tools": "digital_tools",
        "Cadre juridique pour le partage de données": "legal_framework",
        "Legal framework for data sharing": "legal_framework",
        "Cadre juridique": "legal_framework",
        "Legal framework": "legal_framework",
        "Coordination interinstitutionnelle": "institutional_coordination",
        "Inter-institutional coordination": "institutional_coordination",
        "Coordination institutionnelle": "institutional_coordination",
        "Institutional coordination": "institutional_coordination",
    }

    def _extract_table(table_obj: Any, mapping: Dict[str, str], prefix: str) -> Dict[str, Any]:
        out_tbl: Dict[str, Any] = {}
        # Ensure stable columns even when a respondent skips the section
        canons = sorted(set(mapping.values()))
        for canon in canons:
            out_tbl[f"{prefix}_{canon}"] = ""
            out_tbl[f"{prefix}_{canon}_spec"] = ""
        if not isinstance(table_obj, dict):
            return out_tbl
        for label, canon in mapping.items():
            cell = table_obj.get(label, None)
            if isinstance(cell, dict):
                out_tbl[f"{prefix}_{canon}"] = cell.get("code", "")
                out_tbl[f"{prefix}_{canon}_spec"] = cell.get("spec", "")
            elif isinstance(cell, str):
                out_tbl[f"{prefix}_{canon}"] = cell
        return out_tbl

    out: Dict[str, Any] = {}

    # Identification (Rubrique 2)
    out["organisation"] = payload.get("organisation", "")
    out["pays"] = payload.get("pays", "")
    out["pays_autre"] = payload.get("pays_autre", "")
    out["pays_name_fr"] = payload.get("pays_name_fr", "")
    out["pays_name_en"] = payload.get("pays_name_en", "")
    out["type_acteur"] = payload.get("type_acteur", "")
    out["fonction"] = payload.get("fonction", "")
    out["email"] = payload.get("email", "")
    out["lang"] = payload.get("lang", "")

    # Rubrique 3 : portée
    out["scope"] = payload.get("scope", "")
    out["scope_other"] = payload.get("scope_other", "")

    # Rubrique 4 : domaines
    pre = payload.get("preselected_domains", payload.get("preselection_domains", []))
    out["preselection_domains"] = _join_list(pre)
    out["nb_preselection_domains"] = len(pre) if isinstance(pre, list) else 0

    top5 = payload.get("top5_domains", [])
    for i in range(5):
        out[f"top_domain_{i+1}"] = top5[i] if i < len(top5) else ""

    # Rubrique 5 : stats et notation
    selected_stats = payload.get("selected_stats", [])
    out["nb_stats"] = len(selected_stats) if isinstance(selected_stats, list) else 0
    out["stats_list"] = _join_list(selected_stats)
    out["selected_by_domain_json"] = _json(payload.get("selected_by_domain", {}))
    out["scoring_json"] = _json(payload.get("scoring", {}))
    out["scoring_version"] = payload.get("scoring_version", "")

    # Rubrique 6 : perspective de genre (table)
    out.update(_extract_table(payload.get("gender_table", {}), GENDER_ITEM_MAP, "gender"))

    # Rubrique 8 : capacité & faisabilité (table)
    out.update(_extract_table(payload.get("capacity_table", {}), CAPACITY_ITEM_MAP, "capacity"))

    # Rubrique 9 : harmonisation & qualité
    out["quality_expectations"] = _join_list(payload.get("quality_expectations", []))
    out["quality_other"] = payload.get("quality_other", "")

    # Rubrique 10 : diffusion
    out["dissemination_channels"] = _join_list(payload.get("dissemination_channels", []))
    out["dissemination_other"] = payload.get("dissemination_other", "")

    # Rubrique 12 : questions ouvertes
    out["comment_1"] = payload.get("open_q1", "")
    out["missing_indicators"] = payload.get("open_q2", "")
    out["support_needs"] = payload.get("open_q3", "")

    # Dernière question : consultation d’autres personnes pour remplir le questionnaire
    out["consulted_colleagues"] = payload.get("consulted_colleagues", "")

    return out


# =========================
# Validation logic (quality controls)
# =========================

EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")


def validate_r2(lang: str) -> List[str]:
    errs: List[str] = []
    organisation = resp_get("organisation", "").strip()
    pays = resp_get("pays", "").strip()
    pays_autre = resp_get("pays_autre", "").strip()
    type_acteur = resp_get("type_acteur", "").strip()
    fonction = resp_get("fonction", "").strip()
    fonction_autre = resp_get("fonction_autre", "").strip()
    email = resp_get("email", "").strip()

    if not organisation:
        errs.append(t(lang, "Organisation : champ obligatoire.", "Organization: required field."))
    # Contrôle qualité : éviter les sigles seuls
    elif len(organisation) < 12:
        errs.append(t(lang, "Organisation : indiquez le libellé complet (au moins 12 caractères) et non le sigle.", "Organization: please provide the full name (at least 12 characters), not only an acronym."))
    if not pays:
        errs.append(t(lang, "Pays de résidence : champ obligatoire.", "Country of Residence: required field."))
    elif pays in {"OTHER", "__OTHER__"} and not pays_autre:
        errs.append(t(lang, "Autre pays (à préciser) : champ obligatoire.", "Other country (please specify): required field."))
    if not type_acteur:
        errs.append(t(lang, "Type d’acteur : champ obligatoire.", "Stakeholder type: required field."))
    if not fonction:
        errs.append(t(lang, "Fonction : champ obligatoire.", "Role/Function: required field."))
    if is_other_value(fonction):
        if not fonction_autre:
            errs.append(t(lang, "Fonction (Autre) : précisez.", "Role (Other): please specify."))
    if not email:
        errs.append(t(lang, "Email : champ obligatoire.", "Email: required field."))
    elif not EMAIL_RE.match(email):
        errs.append(t(lang, "Email : format invalide.", "Email: invalid format."))
    return errs



def validate_r3(lang: str) -> List[str]:
    errs: List[str] = []
    scope = (resp_get("scope", "") or "").strip()
    if not scope:
        errs.append(t(lang, "Rubrique 3 : veuillez sélectionner une portée.", "Section 3: please select a scope."))
        return errs
    snds = (resp_get("snds_status", "") or "").strip()
    if not snds:
        errs.append(
            t(
                lang,
                "Rubrique 3 : veuillez indiquer le statut de la SNDS / plan national statistique.",
                "Section 3: please indicate the status of the NSDS / national statistical plan.",
            )
        )

    if scope == "Other":
        other = (resp_get("scope_other", "") or "").strip()
        if not other:
            errs.append(t(lang, "Rubrique 3 : précisez l’option « Autre ».", "Section 3: please specify the \"Other\" option."))
    return errs

def validate_r4(lang: str) -> List[str]:
    errs: List[str] = []
    pre = resp_get("preselected_domains", [])
    top5 = resp_get("top5_domains", [])
    if not isinstance(pre, list):
        pre = []
    if not isinstance(top5, list):
        top5 = []
    if len(pre) < 5 or len(pre) > 10:
        errs.append(t(lang, "Rubrique 4 : pré-sélectionnez entre 5 et 10 domaines.", "Section 4: pre-select 5 to 10 domains."))
    if len(set(pre)) != len(pre):
        errs.append(t(lang, "Rubrique 4 : la pré-sélection contient des doublons.", "Section 4: duplicates found in pre-selection."))
    if len(top5) != 5:
        errs.append(t(lang, "Rubrique 4 : le TOP 5 doit contenir exactement 5 domaines.", "Section 4: TOP 5 must contain exactly 5 domains."))
    else:
        if len(set(top5)) != 5:
            errs.append(t(lang, "Rubrique 4 : le TOP 5 contient des doublons.", "Section 4: TOP 5 contains duplicates."))
        missing = [d for d in top5 if d not in pre]
        if missing:
            errs.append(t(lang, "Rubrique 4 : chaque domaine du TOP 5 doit provenir de la pré-sélection.", "Section 4: TOP 5 must be selected from pre-selection."))
    return errs


def validate_r5(lang: str) -> List[str]:
    errs: List[str] = []
    top5 = resp_get("top5_domains", [])
    selected_by_domain: Dict[str, List[str]] = resp_get("selected_by_domain", {})
    if not isinstance(selected_by_domain, dict):
        selected_by_domain = {}

    all_stats: List[str] = []
    for d in top5:
        stats = selected_by_domain.get(d, [])
        if not isinstance(stats, list):
            stats = []
        if len(stats) < 1:
            errs.append(t(lang, f"Rubrique 5 : choisissez au moins 1 statistique pour {d}.",
                          f"Section 5: select at least 1 indicator for {d}."))
        if len(stats) > 3:
            errs.append(t(lang, f"Rubrique 5 : maximum 3 statistiques pour {d}.",
                          f"Section 5: maximum 3 indicators for {d}."))
        all_stats.extend(stats)

    if len(all_stats) < 5 or len(all_stats) > 15:
        errs.append(t(lang, "Rubrique 5 : le total des statistiques doit être entre 5 et 15.",
                      "Section 5: total number of indicators must be between 5 and 15."))

    if len(set(all_stats)) != len(all_stats):
        errs.append(t(lang, "Rubrique 5 : une même statistique ne doit pas être sélectionnée plusieurs fois.",
                      "Section 5: the same indicator must not be selected more than once."))

    # scoring
    scoring: Dict[str, Dict[str, Any]] = resp_get("scoring", {})
    if not isinstance(scoring, dict):
        scoring = {}

    for s in all_stats:
        if s not in scoring:
            errs.append(t(lang, f"Rubrique 5 : vous devez noter la statistique {s}.",
                          f"Section 5: you must score indicator {s}."))
            continue
        for k in ["demand", "availability", "feasibility"]:
            sc_row = scoring.get(s, {}) or {}
            k_lbl = {
                "demand": t(lang, "demande", "demand"),
                "availability": t(lang, "disponibilité", "availability"),
                "feasibility": t(lang, "faisabilité", "feasibility"),
            }.get(k, k)

            # Backward compatibility: legacy key "gap" -> "availability"
            if k == "availability":
                v_raw = sc_row.get("availability", sc_row.get("gap", None))
            else:
                v_raw = sc_row.get(k, None)

            if v_raw is None or str(v_raw).strip() == "":
                errs.append(t(lang, f"Rubrique 5 : la note '{k_lbl}' manque pour {s}.",
                              f"Section 5: missing score '{k_lbl}' for {s}."))
            else:
                try:
                    v = int(v_raw)
                    if v < 0 or v > 3:
                        errs.append(t(lang, f"Rubrique 5 : note invalide pour {s} ({k_lbl}).",
                                      f"Section 5: invalid score for {s} ({k_lbl})."))
                except Exception:
                    errs.append(t(lang, f"Rubrique 5 : note invalide pour {s} ({k_lbl}).",
                                  f"Section 5: invalid score for {s} ({k_lbl})."))
    return errs


def validate_r6(lang: str) -> List[str]:
    errs: List[str] = []
    tbl = resp_get("gender_table", {})
    if not isinstance(tbl, dict) or not tbl:
        errs.append(t(lang, "Rubrique 6 : veuillez renseigner le tableau.", "Section 6: please complete the table."))
        return errs
    for k, v in tbl.items():
        if not v:
            errs.append(t(lang, f"Rubrique 6 : ligne non renseignée : {k}.", f"Section 6: missing answer for: {k}."))
    return errs


def validate_r8(lang: str) -> List[str]:
    errs: List[str] = []

    # Tableau capacité / faisabilité
    tbl = resp_get("capacity_table", {})
    if not isinstance(tbl, dict) or not tbl:
        errs.append(t(lang, "Rubrique 8 : veuillez renseigner le tableau.", "Section 8: please complete the table."))
        return errs
    for k, v in tbl.items():
        if not v:
            errs.append(t(lang, f"Rubrique 8 : ligne non renseignée : {k}.", f"Section 8: missing answer for: {k}."))

    return errs


def validate_r9(lang: str) -> List[str]:
    errs: List[str] = []
    sel = resp_get("quality_expectations", [])
    if not isinstance(sel, list) or len([x for x in sel if str(x).strip() != ""]) == 0:
        errs.append(t(lang, "Rubrique 9 : veuillez sélectionner au moins une option.", "Section 9: please select at least one option."))
        return errs
    if has_other_option(sel):
        other = (resp_get("quality_other", "") or "").strip()
        if not other:
            errs.append(t(lang, "Rubrique 9 : précisez l’option « Autre ».", "Section 9: please specify the \"Other\" option."))
    return errs


def validate_r10(lang: str) -> List[str]:
    errs: List[str] = []
    sel = resp_get("dissemination_channels", [])
    if not isinstance(sel, list) or len([x for x in sel if str(x).strip() != ""]) == 0:
        errs.append(t(lang, "Rubrique 10 : veuillez sélectionner au moins une option.", "Section 10: please select at least one option."))
        return errs
    if has_other_option(sel):
        other = (resp_get("dissemination_other", "") or "").strip()
        if not other:
            errs.append(t(lang, "Rubrique 10 : précisez l’option « Autre ».", "Section 10: please specify the \"Other\" option."))
    return errs


def validate_r11(lang: str) -> List[str]:
    errs: List[str] = []
    sel = resp_get("data_sources", [])
    if not isinstance(sel, list):
        sel = []

    sel_clean = [str(x).strip() for x in sel if str(x).strip()]
    if len(sel_clean) < 2:
        errs.append(t(lang, "Rubrique 11 : sélectionnez au moins 2 sources.", "Section 11: please select at least 2 sources."))
        return errs
    if len(sel_clean) > 4:
        errs.append(t(lang, "Rubrique 11 : sélectionnez au maximum 4 sources.", "Section 11: please select at most 4 sources."))

    if has_other_option(sel_clean):
        other = (resp_get("data_sources_other", "") or "").strip()
        if not other:
            errs.append(t(lang, "Rubrique 11 : précisez l’option « Autres ».", "Section 11: please specify the 'Other' option."))

    return errs


def validate_r12(lang: str) -> List[str]:
    errs: List[str] = []
    sub = int(st.session_state.get("r12_substep", 0) or 0)
    if sub < 3:
        errs.append(
            t(
                lang,
                "Rubrique 12 : veuillez traiter les questions ouvertes une à une (bouton « Question suivante ») jusqu’à la Confirmation.",
                "Section 12: please go through the open questions one by one (use the “Next question” button) until Confirmation.",
            )
        )

    cc = (resp_get("consulted_colleagues", "") or "").strip()
    if cc not in ("YES", "NO"):
        errs.append(
            t(
                lang,
                "Rubrique 12 : veuillez indiquer si vous avez consulté d’autres collègues (Oui/Non).",
                "Section 12: please indicate whether you consulted other colleagues (Yes/No).",
            )
        )
    return errs

def validate_all(lang: str) -> List[str]:
    errs = []
    errs.extend(validate_r2(lang))
    errs.extend(validate_r3(lang))
    errs.extend(validate_r4(lang))
    errs.extend(validate_r5(lang))
    errs.extend(validate_r6(lang))
    errs.extend(validate_r8(lang))
    errs.extend(validate_r9(lang))
    errs.extend(validate_r10(lang))
    errs.extend(validate_r11(lang))
    errs.extend(validate_r12(lang))
    # Open questions (text fields) remain optional; warnings are shown in Section 12 / Submit.
    return errs


# =========================
# Navigation
# =========================

def get_steps(lang: str) -> List[Tuple[str, str]]:
    # Rubric 7 added, plus final SEND tab
    return [
        ("R1", t(lang, "Rubrique 1 : Instructions", "Section 1: Instructions")),
        ("R2", t(lang, "Rubrique 2 : Identification du répondant", "Section 2: Respondent identification")),
        ("R3", t(lang, "Rubrique 3 : Portée de la réponse", "Section 3: Scope of response")),
        ("R4", t(lang, "Rubrique 4 : Domaines prioritaires", "Section 4: Priority domains")),
        ("R5", t(lang, "Rubrique 5 : Statistiques prioritaires et notation", "Section 5: Priority indicators and scoring")),
        ("R6", t(lang, "Rubrique 6 : Dimension genre", "Section 6: Gender dimension")),
        ("R7", t(lang, "Rubrique 7 : Priorités genre", "Section 7: Gender priorities")),
        ("R8", t(lang, "Rubrique 8 : Capacité et faisabilité (12–24 mois)", "Section 8: Capacity and feasibility (12–24 months)")),
        ("R9", t(lang, "Rubrique 9 : Harmonisation et qualité", "Section 9: Harmonization and quality")),
        ("R10", t(lang, "Rubrique 10 : Diffusion", "Section 10: Dissemination")),
        ("R11", t(lang, "Rubrique 11 : Sources de données pertinentes", "Section 11: Relevant data sources")),
        ("R12", t(lang, "Rubrique 12 : Questions ouvertes", "Section 12: Open questions")),
        ("SEND", t(lang, "ENVOYER", "SUBMIT")),
    ]


def render_sidebar(lang: str, steps: List[Tuple[str, str]]) -> None:
    st.sidebar.header(t(lang, "Navigation", "Navigation"))
    labels = [s[1] for s in steps]

    # Keep sidebar selection in sync with nav_idx
    st.session_state.nav_radio = int(st.session_state.nav_idx)
    chosen = st.sidebar.radio(
        t(lang, "Aller à", "Go to"),
        options=list(range(len(labels))),
        index=int(st.session_state.nav_idx),
        format_func=lambda i: labels[i],
        key="nav_radio"
    )

    # User clicked in sidebar
    if int(chosen) != int(st.session_state.nav_idx):
        st.session_state.nav_idx = int(chosen)

    st.sidebar.divider()
    st.sidebar.caption(
        t(
            lang,
            "Note : les contrôles qualité peuvent bloquer la progression si une contrainte n’est pas respectée.",
            "Note: quality checks may prevent moving forward when constraints are not met."
        )
    )

    st.sidebar.markdown("---")
    st.sidebar.caption(
        t(
            lang,
            "NSP : Ne sait pas (score 0). Utilisez NSP uniquement si l’information est indisponible.",
            "UK: Unknown (score 0). Use UK only when information is unavailable."
        )
    )
    st.sidebar.markdown("---")
    st.sidebar.subheader(t(lang, "Brouillon", "Draft"))
    if st.sidebar.button(t(lang, "Sauvegarder maintenant", "Save now")):
        ok, msg = autosave_draft(force=True)
        if ok:
            st.sidebar.success(t(lang, "Brouillon sauvegardé.", "Draft saved."))
        else:
            st.sidebar.error(t(lang, "Brouillon non sauvegardé.", "Draft not saved."))
    if st.session_state.get("draft_id"):
        st.sidebar.caption(
            t(
                lang,
                "Reprise : conservez l’URL de cette page (paramètre rid=...).",
                "Resume: keep this page URL (rid=... parameter)."
            )
        )



def nav_buttons(lang: str, steps: List[Tuple[str, str]], df_long: pd.DataFrame) -> None:
    """Bottom nav buttons, with blocking based on current step validations."""
    step_key = steps[st.session_state.nav_idx][0]
    errors: List[str] = []

    # Blocking rules per step
    if step_key == "R2":
        errors = validate_r2(lang)
    elif step_key == "R3":
        errors = validate_r3(lang)
    elif step_key == "R4":
        errors = validate_r4(lang)
    elif step_key == "R5":
        errors = validate_r5(lang)
    elif step_key == "R6":
        errors = validate_r6(lang)
    elif step_key == "R7":
        errors = validate_r7(lang)
    elif step_key == "R8":
        errors = validate_r8(lang)
    elif step_key == "R9":
        errors = validate_r9(lang)
    elif step_key == "R10":
        errors = validate_r10(lang)
    elif step_key == "R11":
        errors = validate_r11(lang)
    elif step_key == "R12":
        errors = validate_r12(lang)

    col1, col2, col3 = st.columns([1, 1, 3])
    with col1:
        prev_disabled = st.session_state.nav_idx <= 0
        if st.button(t(lang, "⬅ Précédent", "⬅ Previous"), disabled=prev_disabled):
            autosave_draft(force=True)
            st.session_state.nav_idx = max(0, st.session_state.nav_idx - 1)
            st.rerun()
    with col2:
        next_disabled = (st.session_state.nav_idx >= len(steps) - 1) or bool(errors)
        if st.button(t(lang, "Suivant ➡", "Next ➡"), disabled=next_disabled):
            autosave_draft(force=True)
            st.session_state.nav_idx = min(len(steps) - 1, st.session_state.nav_idx + 1)
            st.rerun()
    with col3:
        if errors:
            st.error("\n".join(errors))


# =========================
# UI : Rubrics
# =========================

def rubric_1(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 1 : Instructions", "Section 1: Instructions", "Secção 1: Instruções", "القسم 1: التعليمات"))
    st.markdown(
        t(
            lang,
            """
### Objectif
Ce questionnaire vise à recueillir votre avis sur **les statistiques socio-économiques prioritaires** à produire et diffuser au niveau continental.

### Comment répondre
1. **Identifiez** votre organisation (Rubrique 2).
2. **Pré-sélectionnez 5 à 10 domaines** et classez un **TOP 5** (Rubrique 4).
3. Pour chaque domaine du TOP 5 : choisissez **1 à 3 statistiques** et attribuez des **notes** (Rubrique 5).
4. Complétez les rubriques transversales : **genre**, **capacité/faisabilité**, **etc.**.
5. **N'hésitez pas à consulter les infobulles ⍰ pour plus de précisions.**

            """,
            """
### Purpose
This questionnaire collects your views on **priority socio-economic statistics** to be produced and disseminated at continental level.

### How to answer
1. **Identify** your organization (Section 2).
2. **Pre-select 5–10 domains** and rank a **TOP 5** (Section 4).
3. For each TOP 5 domain: select **1–3 indicators** and provide **scores** (Section 5).
4. Complete cross-cutting sections: **gender**, **capacity/feasibility**,  **etc.**.
5. **Feel free to consult the ⍰ tooltips for more details.**

            """
        )
    )


def rubric_2(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 2 : Identification du répondant", "Section 2: Respondent identification", "Secção 2: Identificação do respondente", "القسم 2: تعريف المجيب"))
    st.info(
        t(
            lang,
            "Merci de renseigner ces informations. Elles servent uniquement à l’analyse et ne seront pas publiées nominativement.",
            "Please provide these details. They are used for analysis and will not be published in a personally identifiable way."
        )
    )

    resp_set("lang", lang)

    st.text_input(t(lang, "Nom de l'organisation", "Organization Name"), key="org_input", value=resp_get("organisation", ""))
    resp_set("organisation", st.session_state.get("org_input", "").strip())
    st.caption(t(lang, "Merci d’indiquer le libellé complet (évitez le sigle seul).", "Please provide the full organization name (avoid acronym only)."))
    col1, col2 = st.columns(2)
    with col1:
        # Pays de résidence : liste déroulante (ISO3 + libellés FR/EN) + option Autre
        df_countries = load_countries()
        iso3_list, iso3_to_fr, iso3_to_en, iso3_to_pt, iso3_to_ar = country_maps(df_countries)

        prev_country = (resp_get("pays", "") or "").strip()
        prev_country_other = (resp_get("pays_autre", "") or "").strip()
        prev_iso3 = (prev_country.split("|", 1)[0].strip().upper() if "|" in prev_country else prev_country.strip().upper())

        # Compatibilité ascendante : si ancienne valeur libre non-ISO3, basculer sur "OTHER"
        if prev_iso3 and prev_iso3 not in iso3_list and prev_iso3 not in {"OTHER", "__OTHER__"}:
            if not prev_country_other:
                prev_country_other = prev_country
            prev_iso3 = "OTHER"
        elif prev_iso3 == "__OTHER__":
            prev_iso3 = "OTHER"

        if not iso3_list:
            # Fallback si le fichier pays est introuvable : champ libre
            st.text_input(t(lang, "Pays de résidence", "Country of residence"), key="country_input", value=resp_get("pays", ""))
            resp_set("pays", st.session_state.get("country_input", "").strip())
            resp_set("pays_autre", "")
            resp_set("pays_name_fr", "")
            resp_set("pays_name_en", "")
        else:
            options = [""] + sorted(iso3_list, key=lambda x: country_label(x, lang, iso3_to_fr, iso3_to_en, iso3_to_pt, iso3_to_ar).lower()) + ["OTHER"]

            chosen_iso3 = st.selectbox(
                t(lang, "Pays de résidence", "Country of residence"),
                options=options,
                index=options.index(prev_iso3) if prev_iso3 in options else 0,
                format_func=lambda x: (
                    t(lang, "— Sélectionner —", "— Select —") if x == ""
                    else t(lang, "Autre pays (à préciser)", "Other country (please specify)") if x == "OTHER"
                    else f"{country_label(x, lang, iso3_to_fr, iso3_to_en, iso3_to_pt, iso3_to_ar)} ({x})"
                ),
                help=t(lang, "Choisissez votre pays de résidence (liste ISO3) ou ‘Autre pays (à préciser)’.",
                       "Select your country of residence (ISO3 list) or ‘Other country (please specify)’."),
                key="country_iso3_select",
            )
            resp_set("pays", chosen_iso3)

            if chosen_iso3 == "OTHER":
                other_country = st.text_input(
                    t(lang, "Autre pays (à préciser)", "Other country (please specify)"),
                    key="country_other_input",
                    value=prev_country_other,
                )
                other_country = (other_country or "").strip()
                resp_set("pays_autre", other_country)
                # On remplit aussi les libellés pour l'export, sans forcer une ISO3
                resp_set("pays_name_fr", other_country)
                resp_set("pays_name_en", other_country)
            elif chosen_iso3:
                resp_set("pays_autre", "")
                # Libellés normalisés (utile pour les exports / analyses)
                resp_set("pays_name_fr", iso3_to_fr.get(chosen_iso3, "").strip())
                resp_set("pays_name_en", (iso3_to_en.get(chosen_iso3, "") or iso3_to_fr.get(chosen_iso3, "")).strip())
            else:
                resp_set("pays_autre", "")
                resp_set("pays_name_fr", "")
                resp_set("pays_name_en", "")

    with col2:
        st.text_input(
            t(lang, "Email", "Email"),
            key="email_input",
            value=resp_get("email", ""),
            help=t(
                lang,
                "Saisissez une adresse email valide (ex. nom@domaine.tld).",
                "Enter a valid email address (e.g., name@domain.tld).",
            ),
        )
        resp_set("email", st.session_state.get("email_input", "").strip())

    # Brouillon : crée un identifiant de reprise dès que l’email est renseigné
    ensure_draft_id()
    autosave_draft(force=False)

    # Afficher le message de reprise dès la première session après saisie de l’email
    _email_now = resp_get("email", "")
    if _email_now and ("@" in _email_now) and not st.session_state.get("draft_resume_notice_shown"):
        st.warning(
            t(
                lang,
                "La saisie est sauvegardée. En cas de suspension de moins de 48 heures, reprenez là où vous vous étiez arrêté en ré-ouvrant le lien contenant rid (à conserver / mettre en favori / retrouver dans l'historique).",
                "Your input is saved. If you pause for less than 48 hours, resume where you left off by reopening the link containing rid (keep it / bookmark it / find it in your browser history).",
            )
        )
        st.session_state["draft_resume_notice_shown"] = True
    type_options = [
        ("NSO", {"fr": "Institut national de statistique", "en": "National Statistical Office"}),
        ("Ministry", {"fr": "Ministère / Service statistique sectoriel", "en": "Ministry / Sector statistical unit"}),
        ("REC", {"fr": "Communauté économique régionale", "en": "Regional Economic Community"}),
        ("AU", {"fr": "Union Africaine (UA)", "en": "African Union (AU)"}),
        ("CivilSoc", {"fr": "Société civile", "en": "Civil society"}),
        ("DevPartner", {"fr": "Partenaire technique et financier", "en": "Development partner"}),
        ("Academia", {"fr": "Université / Recherche", "en": "Academia / Research"}),
        ("Other", {"fr": "Autre", "en": "Other"}),
    ]
    type_labels = [t(lang, x[1]["fr"], x[1]["en"]) for x in type_options]
    type_keys = [x[0] for x in type_options]

    # Type d’acteur : pas de pré-remplissage (placeholder)
    type_opts = [""] + type_keys
    prev_type = resp_get("type_acteur", "")
    idx = type_opts.index(prev_type) if prev_type in type_opts else 0

    chosen_type = st.selectbox(
        t(lang, "Type d’acteur", "Stakeholder type"),
        options=type_opts,
        index=idx,
        format_func=lambda k: (t(lang, "— Sélectionner —", "— Select —") if k == "" else type_labels[type_keys.index(k)]),
        help=t(lang, "Choisissez la catégorie correspondant le mieux à votre organisation.", 
               "Choose the category that best matches your organization.")
    )
    resp_set("type_acteur", chosen_type)
# Fonction dropdown : pas de pré-remplissage (placeholder)
    role_opts = get_role_options(lang)
    role_options = [""] + role_opts
    prev_role = resp_get("fonction", "")
    role_idx = role_options.index(prev_role) if prev_role in role_options else 0

    chosen_role = st.selectbox(
        t(lang, "Fonction", "Role/Function"),
        options=role_options,
        index=role_idx,
        format_func=lambda x: (t(lang, "— Sélectionner —", "— Select —") if x == "" else x),
        help=t(lang, "Indiquez votre fonction principale dans l’organisation.", "Indicate your main role in the organization."),
    )
    resp_set("fonction", chosen_role)
    if is_other_value(chosen_role):
        st.text_input(t(lang, "Préciser (fonction)", "Specify (role)"),
                      key="fonction_autre_input", value=resp_get("fonction_autre", ""))
        resp_set("fonction_autre", st.session_state.get("fonction_autre_input", "").strip())
    else:
        resp_set("fonction_autre", "")

    # Live errors
    errs = validate_r2(lang)
    if errs:
        st.warning(t(lang, "Veuillez corriger les éléments ci-dessous :", "Please fix the following:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def rubric_3(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 3 : Portée de la réponse", "Section 3: Scope of response", "Secção 3: Âmbito da resposta", "القسم 3: نطاق الإجابة"))
    st.markdown(
        t(
            lang,
            "Indiquez le périmètre principal de votre réponse. Cela aide à interpréter vos priorités.",
            "Indicate the main scope of your response. This helps interpret your priorities.",
            "Indique o âmbito principal da sua resposta. Isto ajuda a interpretar as suas prioridades.",
            "حددوا النطاق الرئيسي لإجابتكم. فهذا يساعد على تفسير أولوياتكم."
        )
    )

    scope_opts_raw = [
        ("National", {"fr": "National", "en": "National"}),
        ("Regional", {"fr": "Régional (CER)", "en": "Regional (REC)"}),
        ("Continental", {"fr": "Continental (UA)", "en": "Continental (AU)"}),
        ("Global", {"fr": "International", "en": "International"}),
        ("Other", {"fr": "Autre", "en": "Other"}),
    ]
    scope_labels = {k: t(lang, v["fr"], v["en"]) for k, v in scope_opts_raw}
    scope_keys = [k for k, _ in scope_opts_raw]
    scope_options = [""] + scope_keys

    prev_scope = resp_get("scope", "")
    scope_idx = scope_options.index(prev_scope) if prev_scope in scope_options else 0

    chosen_scope = st.selectbox(
        t(lang, "Portée", "Scope"),
        options=scope_options,
        index=scope_idx,
        format_func=lambda k: (t(lang, "— Sélectionner —", "— Select —") if k == "" else scope_labels.get(k, k)),
        help=t(
            lang,
            "Indiquez le périmètre principal de votre réponse : national, régional (CER), continental (UA) ou international.",
            "Indicate the main scope of your response: national, regional (REC), continental (AU), or international."
        )
    )
    resp_set("scope", chosen_scope)

    # SNDS / Plan statistique national (obligatoire)
    snds_opts = ["", "YES", "NO", "PREP", "IMPL_PREP", "NSP"]
    snds_labels = {
        "YES": t(lang, "Oui", "Yes"),
        "NO": t(lang, "Non", "No"),
        "PREP": t(lang, "En préparation", "In preparation"),
        "IMPL_PREP": t(lang, "En cours de mise en œuvre ET nouvelle en préparation", "Under implementation AND new one in preparation"),
        "NSP": t(lang, "NSP", "DK"),
    }
    prev_snds = (resp_get("snds_status", "") or "").strip()
    idx_snds = snds_opts.index(prev_snds) if prev_snds in snds_opts else 0
    chosen_snds = st.selectbox(
        t(
            lang,
            "Statut de la SNDS / plan national statistique en cours",
            "Status of the current NSDS / national statistical plan",
        ),
        options=snds_opts,
        index=idx_snds,
        format_func=lambda k: (
            t(lang, "— Sélectionner —", "— Select —") if k == "" else snds_labels.get(k, k)
        ),
        help=t(
            lang,
            "Indiquez si une stratégie / plan statistique national est en cours, non, en préparation, ou NSP.",
            "Indicate whether an NSDS / national statistical plan is current, not in place, under preparation, or DK.",
        ),
        key="snds_status_select",
    )
    resp_set("snds_status", chosen_snds)


    # Contrôle qualité (alerte) : cohérence acteur × portée
    _actor = (resp_get("type_acteur", "") or "").strip()
    _scope = (resp_get("scope", "") or "").strip()
    if _actor in ["NSO", "Ministry"] and _scope and _scope != "National":
        st.warning(
            t(
                lang,
                "Alerte : vous avez indiqué « Institut national de statistique » ou « Ministère », mais la portée n’est pas « National ». Merci de vérifier la cohérence.",
                "Warning: you selected “National Statistical Office” or “Ministry”, but the scope is not “National”. Please check consistency."
            )
        )

    if resp_get("scope") == "Other":
        st.text_input(t(lang, "Préciser", "Specify"), key="scope_other_input", value=resp_get("scope_other", ""))
        resp_set("scope_other", st.session_state.get("scope_other_input", "").strip())
    else:
        resp_set("scope_other", "")



def rubric_4(lang: str, df_long: pd.DataFrame) -> None:
    st.subheader(t(lang, "Rubrique 4 : Domaines prioritaires", "Section 4: Priority domains", "Secção 4: Domínios prioritários", "القسم 4: المجالات ذات الأولوية"))

    st.info(
        t(
            lang,
            "Veuillez d’abord choisir 5 à 10 domaines (pré-sélection). Ensuite, choisissez exactement 5 domaines dans ce sous-ensemble (TOP 5).\n\nConseil : choisissez les domaines où la demande politique est forte.",
            "First select 5 to 10 domains (pre-selection). Then choose exactly 5 domains within that subset (TOP 5).\n\nTip: choose domains where policy demand is strong."
        )
    )

    domains = domains_from_longlist(df_long, lang)
    if not domains:
        st.error(
            t(
                lang,
                "La liste des domaines n’est pas disponible (longlist introuvable ou vide).",
                "Domain list is not available (longlist not found or empty).",
            )
        )
        st.caption(
            t(
                lang,
                "Vérifiez que le dépôt contient : data/indicator_longlist.csv (prioritaire) ou data/longlist.xlsx (ou ces fichiers à la racine).",
                "Check that the repository contains: data/indicator_longlist.csv (preferred) or data/longlist.xlsx (or these files at repo root).",
            )
        )
        return

    code_to_label = {c: lbl for c, lbl in domains}

    # Build display labels without showing codes (codes are stored internally)
    labels = [code_to_label[c] for c, _ in domains]
    # Disambiguate duplicates if any (rare)
    seen = {}
    for i, (c, _) in enumerate(domains):
        lbl = code_to_label[c]
        seen[lbl] = seen.get(lbl, 0) + 1
    display_labels = []
    label_to_code = {}
    for c, _ in domains:
        lbl = code_to_label[c]
        disp = lbl if seen[lbl] == 1 else f"{lbl} ({c})"
        display_labels.append(disp)
        label_to_code[disp] = c

    st.markdown(
        t(
            lang,
            """
### Étape 1 : Pré-sélection
Sélectionnez **entre 5 et 10 domaines** (sans doublons).
            """,
            """
### Step 1: Pre-selection
Select **5 to 10 domains** (no duplicates).
            """
        )
    )

    pre_default_codes = resp_get("preselected_domains", [])
    pre_default_disp = []
    for c in pre_default_codes:
        lbl = code_to_label.get(c, "")
        if not lbl:
            continue
        disp = lbl if seen.get(lbl, 1) == 1 else f"{lbl} ({c})"
        if disp in label_to_code:
            pre_default_disp.append(disp)

    # Avoid "first click not kept" by initializing widget state once (no default on every rerun)
    if "r4_preselection_ms" not in st.session_state:
        st.session_state["r4_preselection_ms"] = pre_default_disp

    pre_disp = st.multiselect(
        t(lang, "Pré-sélection (5–10 domaines)", "Pre-selection (5–10 domains)"),
        options=display_labels,
        max_selections=10,
        key="r4_preselection_ms",
        help=t(lang, "Choisissez au maximum 10 domaines. Une fois 10 domaines sélectionnés, les nouveaux clics seront ignorés.", 
               "Select up to 10 domains. Once 10 domains are selected, additional clicks are ignored.")
    )

    pre_codes = [label_to_code[x] for x in pre_disp]
    resp_set("preselected_domains", pre_codes)

    st.divider()
    st.markdown(
        t(
            lang,
            """
### Étape 2 : Classement TOP 5
Classez exactement **5 domaines** parmi votre pré-sélection.
            """,
            """
### Step 2: Rank TOP 5
Rank exactly **5 domains** from your pre-selection.
            """
        )
    )

    if len(pre_codes) < 5:
        st.warning(t(lang, "Sélectionnez d’abord au moins 5 domaines dans la pré-sélection.",
                     "Please pre-select at least 5 domains first."))
        resp_set("top5_domains", [])
        return

    top5: List[str] = []

    # Ranking with 5 selectboxes (no prefill + no duplicates)
    chosen_prev: List[str] = []
    for i in range(5):
        key = f"top5_rank_{i+1}"

        # Options for this rank = preselection minus already chosen
        remaining = [c for c in pre_codes if c not in chosen_prev]
        options = [""] + remaining  # "" placeholder (no prefill)

        prev = resp_get(key, "")
        if prev and prev in remaining:
            idx = options.index(prev)
        else:
            idx = 0

        choice = st.selectbox(
            t(lang, f"Rang {i+1}", f"Rank {i+1}"),
            options=options,
            index=idx,
            format_func=lambda c: (t(lang, "— Sélectionner —", "— Select —") if c == "" else code_to_label.get(c, c)),
            help=t(
                lang,
                "Choisissez un domaine unique pour chaque rang. Les domaines déjà choisis ne sont plus proposés aux rangs suivants.",
                "Choose a unique domain for each rank. Already selected domains are removed from the next ranks.",
            ),
            key=key,
        )

        if choice != "":
            top5.append(choice)
            chosen_prev.append(choice)

    resp_set("top5_domains", top5)


    errs = validate_r4(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))



def rubric_5(lang: str, df_long: pd.DataFrame) -> None:
    st.subheader(t(lang, "Rubrique 5 : Statistiques prioritaires et notation", "Section 5: Priority indicators and scoring", "Secção 5: Estatísticas prioritárias e pontuação", "القسم 5: الإحصاءات ذات الأولوية والتقييم"))

    top5 = resp_get("top5_domains", [])
    if not top5 or len(top5) != 5:
        st.warning(t(lang, "Veuillez d’abord finaliser le TOP 5 des domaines (Rubrique 4).",
                     "Please complete TOP 5 domains first (Section 4)."))
        return

    # mapping for domain display
    dom_map = {c: lbl for c, lbl in domains_from_longlist(df_long, lang)}

    st.markdown(
        t(
            lang,
            "La rubrique se remplit en deux temps : (A) sélection des statistiques, puis (B) notation multicritères.",
            "This section is completed in two steps: (A) indicator selection, then (B) multi-criteria scoring.",
            "Esta secção é preenchida em duas etapas : (A) seleção dos indicadores, depois (B) pontuação multicritério.",
            "يُستكمل هذا القسم على مرحلتين : (أ) اختيار المؤشرات، ثم (ب) التقييم متعدد المعايير."
        )
    )

    selected_by_domain: Dict[str, List[str]] = resp_get("selected_by_domain", {})
    if not isinstance(selected_by_domain, dict):
        selected_by_domain = {}

    scoring: Dict[str, Dict[str, Any]] = resp_get("scoring", {})
    if not isinstance(scoring, dict):
        scoring = {}

    # Ensure dict keys exist
    for d in top5:
        if d not in selected_by_domain:
            selected_by_domain[d] = []

    st.markdown("### " + t(lang, "Étape A : Sélection des statistiques", "Step A: Select indicators"))
    st.caption(
        t(
            lang,
            "Pour chaque domaine du TOP 5, choisissez 1 à 3 statistiques (total attendu : 5 à 15). Une statistique ne doit pas apparaître dans deux domaines.",
            "For each TOP 5 domain, select 1 to 3 indicators (expected total: 5 to 15). The same indicator must not appear under two domains.",
        )
    )

    # UI selection per domain (codes hidden)
    for d in top5:
        st.markdown(f"#### {dom_map.get(d, d)}")

        stats_opts = stats_for_domain(df_long, d, lang)
        stat_code_to_label = {c: lbl for c, lbl in stats_opts}

        # build display labels without showing stat codes
        labels = [stat_code_to_label[c] for c, _ in stats_opts]
        seen = {}
        for c, _ in stats_opts:
            lbl = stat_code_to_label[c]
            seen[lbl] = seen.get(lbl, 0) + 1
        display_labels = []
        label_to_code = {}
        for c, _ in stats_opts:
            lbl = stat_code_to_label[c]
            disp = lbl if seen[lbl] == 1 else f"{lbl} ({c})"
            display_labels.append(disp)
            label_to_code[disp] = c

        default_codes = selected_by_domain.get(d, [])
        default_disp = []
        for c in default_codes:
            lbl = stat_code_to_label.get(c, "")
            if not lbl:
                continue
            disp = lbl if seen.get(lbl, 1) == 1 else f"{lbl} ({c})"
            if disp in label_to_code:
                default_disp.append(disp)

        key_ms = f"stats_ms_{d}"

        # Init widget state once (avoid "first click" issues)
        if key_ms not in st.session_state:
            st.session_state[key_ms] = default_disp

        picked_disp = st.multiselect(
            t(lang, "Choisir 1 à 3 statistiques", "Select 1 to 3 indicators"),
            options=display_labels,
            max_selections=3,
            key=key_ms,
            help=t(lang, "Sélectionnez au minimum 1 et au maximum 3 statistiques pour ce domaine.",
                   "Select at least 1 and at most 3 indicators for this domain.")
        )

        picked_codes = [label_to_code[x] for x in picked_disp]
        selected_by_domain[d] = picked_codes


    # Uniqueness check
    flattened = []
    for d in top5:
        flattened.extend(selected_by_domain.get(d, []))
    duplicates = [x for x in set(flattened) if flattened.count(x) > 1]
    if duplicates:
        st.error(
            t(
                lang,
                "Une ou plusieurs statistiques sont sélectionnées dans plusieurs domaines. Veuillez corriger.",
                "One or more indicators are selected under multiple domains. Please correct."
            )
        )

    resp_set("selected_by_domain", selected_by_domain)
    resp_set("selected_stats", flattened)

    # Map codes to labels for display in scoring
    global_map = {}
    for d in top5:
        for c, lbl in stats_for_domain(df_long, d, lang):
            global_map[c] = lbl

    st.divider()
    st.markdown("### " + t(lang, "Étape B : Notation multicritères", "Step B: Multi-criteria scoring"))
    st.caption(
        t(
            lang,
            "Pour chaque statistique sélectionnée, attribuez une note (0–3) sur : demande politique, disponibilité actuelle (bonne = score plus élevé) et faisabilité à 12–24 mois.",
            "For each selected indicator, assign a score (0–3) for: political demand, current availability (good = higher score), and feasibility in 12–24 months.",
        )
    )

    for s in flattened:
        if s not in scoring or not isinstance(scoring.get(s), dict):
            scoring[s] = {}

        # Backward compatibility: legacy key "gap" (écart) -> "availability"
        # On normalise sur l'échelle v3 (Bonne=3)
        if "availability" not in scoring[s] and "gap" in scoring[s]:
            scoring[s]["availability"] = normalize_availability(scoring[s].get("gap", 0), 0)

        # Ensure keys exist
        for k in ["demand", "availability", "feasibility"]:
            if k not in scoring[s]:
                scoring[s][k] = None

        st.markdown(f"**{global_map.get(s, s)}**")

        c1, c2, c3 = st.columns(3)
        opts = [None, 1, 2, 3, 0]  # None = placeholder (no prefill). 0 = NSP / DK

        with c1:
            prev = scoring[s].get("demand", None)
            idx = 0
            try:
                if prev is not None and prev != "":
                    idx = opts.index(int(prev))
            except Exception:
                idx = 0
            scoring[s]["demand"] = st.selectbox(
                t(lang, "Demande politique", "Political demand"),
                options=opts,
                index=idx,
                format_func=score_format(lang, "demand"),
                help=t(
                    lang,
                    "Définition : importance de l’indicateur pour le pilotage des politiques publiques, la redevabilité et les priorités.",
                    "Definition: importance for steering public policies, accountability and priorities.",
                ),
                key=f"sc_dem_{s}",
            )

        with c2:
            prev = scoring[s].get("availability", None)
            idx = 0
            try:
                if prev is not None and prev != "":
                    idx = opts.index(int(prev))
            except Exception:
                idx = 0
            scoring[s]["availability"] = st.selectbox(
                t(lang, "Disponibilité actuelle", "Current availability"),
                options=opts,
                index=idx,
                format_func=score_format(lang, "availability"),
                help=t(
                    lang,
                    "Définition : l’indicateur est-il déjà produit régulièrement, avec une couverture et une qualité suffisantes, et sous une forme utilisable ? (Bonne = score plus élevé).",
                    "Definition: is the indicator already produced regularly with sufficient coverage and quality, in a usable form? (Good = higher score).",
                ),
                key=f"sc_avail_{s}",
            )

        with c3:
            prev = scoring[s].get("feasibility", None)
            idx = 0
            try:
                if prev is not None and prev != "":
                    idx = opts.index(int(prev))
            except Exception:
                idx = 0
            scoring[s]["feasibility"] = st.selectbox(
                t(lang, "Faisabilité (12–24 mois)", "Feasibility (12–24 months)"),
                options=opts,
                index=idx,
                format_func=score_format(lang, "feasibility"),
                help=t(
                    lang,
                    "Définition : capacité réaliste à améliorer ou produire l’indicateur d’ici 12–24 mois, compte tenu des sources, capacités et prérequis.",
                    "Definition: realistic ability to improve or produce the indicator within 12–24 months, considering sources, capacities and prerequisites.",
                ),
                key=f"sc_fea_{s}",
            )

    resp_set("scoring", scoring)

    errs = validate_r5(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def rubric_6(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 6 : Dimension genre", "Section 6: Gender dimension", "Secção 6: Dimensão de género", "القسم 6: البعد الجنساني"))
    st.markdown(
        t(
            lang,
            "Indiquez si les statistiques prioritaires doivent intégrer ces dimensions (Oui/Non/Selon indicateur/NSP).",
            "Indicate whether priority indicators should integrate these dimensions (Yes/No/Indicator-specific/UK)."
        )
    )

    options = [
        (t(lang, "Oui", "Yes"), "YES"),
        (t(lang, "Non", "No"), "NO"),
        (t(lang, "Selon indicateur", "Indicator-specific"), "SPEC"),
        (t(lang, UK_FR, UK_EN, UK_PT, UK_AR), "UK"),
    ]
    labels = [x[0] for x in options]
    code_map = {x[0]: x[1] for x in options}

    items_fr = [
        "Désagrégation par sexe",
        "Désagrégation par âge",
        "Milieu urbain / rural",
        "Handicap",
        "Quintile de richesse",
        "Violences basées sur le genre (VBG)",
        "Temps domestique non rémunéré",
    ]
    items_en = [
        "Disaggregation by sex",
        "Disaggregation by age",
        "Urban / rural",
        "Disability",
        "Wealth quintile",
        "Gender-based violence (GBV)",
        "Unpaid domestic work",
    ]
    items = items_fr if lang == "fr" else items_en

    tbl = resp_get("gender_table", {})
    if not isinstance(tbl, dict):
        tbl = {}

    for it in items:
        rev_map = {v: k for k, v in code_map.items()}
        prev_code = tbl.get(it, None)
        idx = labels.index(rev_map[prev_code]) if prev_code in rev_map else None
        chosen = st.radio(tr(lang, it), options=labels, index=idx, horizontal=True, key=f"gender_{it}")
        tbl[it] = code_map.get(chosen, None)

    resp_set("gender_table", tbl)

    errs = validate_r6(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def validate_r7(lang: str) -> List[str]:
    errs: List[str] = []
    p1 = (resp_get("gender_priority_1", "") or "").strip()
    p2 = (resp_get("gender_priority_2", "") or "").strip()
    p3 = (resp_get("gender_priority_3", "") or "").strip()
    other = (resp_get("gender_priority_other", "") or "").strip()

    if not p1:
        errs.append(
            t(
                lang,
                "Rubrique 7 : veuillez sélectionner au moins une priorité genre (Priorité 1).",
                "Section 7: please select at least one gender priority (Priority 1).",
            )
        )

    # No rank 3 without rank 2
    if p3 and not p2:
        errs.append(
            t(
                lang,
                "Rubrique 7 : veuillez renseigner la Priorité 2 avant la Priorité 3.",
                "Section 7: please fill Priority 2 before Priority 3.",
            )
        )

    # Uniqueness
    chosen = [x for x in [p1, p2, p3] if x]
    if len(set(chosen)) != len(chosen):
        errs.append(
            t(
                lang,
                "Rubrique 7 : les priorités genre doivent être différentes (pas de doublons).",
                "Section 7: gender priorities must be distinct (no duplicates).",
            )
        )

    # Other text required if OTHER selected
    if "OTHER" in chosen and not other:
        errs.append(
            t(
                lang,
                "Rubrique 7 : précisez l’option « Autre ».",
                "Section 7: please specify the 'Other' option.",
            )
        )

    return errs


def rubric_7(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 7 : Priorités genre", "Section 7: Gender priorities", "Secção 7: Prioridades de género", "القسم 7: أولويات النوع الاجتماعي"))
    st.markdown(
        t(
            lang,
            "Sélectionnez de 1 à 3 priorités genre en commençant par la plus importante.",
            "Select 1 to 3 gender priorities, starting with the most important.",
        )
    )

    gp_opts = ["", "ECO", "SERV", "GBV", "PART_DEC", "CARE", "OTHER"]
    gp_labels = {
        "ECO": t(lang, "Autonomisation économique", "Economic empowerment"),
        "SERV": t(lang, "Accès aux services", "Access to services"),
        "GBV": t(lang, "Violences basées sur le genre (VBG)", "Gender based violence (GBV)"),
        "PART_DEC": t(lang, "Participation aux instances décisionnelles", "Participation in decision-making bodies"),
        "CARE": t(lang, "Temps domestique non rémunéré", "Unpaid domestic and care work"),
        "OTHER": t(lang, "Autre", "Other"),
        "": t(lang, "— Sélectionner —", "— Select —"),
    }

    # Rank 1 (required)
    prev1 = (resp_get("gender_priority_1", "") or "").strip()
    idx1 = gp_opts.index(prev1) if prev1 in gp_opts else 0
    p1 = st.selectbox(
        t(lang, "Vos trois (3) priorités genre – Priorité 1 (obligatoire)", "Your three (3) gender priorities – Priority 1 (required)"),
        options=gp_opts,
        index=idx1,
        format_func=lambda k: gp_labels.get(k, k),
        key="gender_priority_1_select",
    )
    resp_set("gender_priority_1", p1)
    # Backward compatibility (previous single-priority field)
    resp_set("gender_priority_main", p1)

    # Rank 2 (optional), exclude already chosen (except empty)
    opts2 = [""] + [k for k in gp_opts if k not in ["", p1]]
    prev2 = (resp_get("gender_priority_2", "") or "").strip()
    idx2 = opts2.index(prev2) if prev2 in opts2 else 0
    p2 = st.selectbox(
        t(lang, "Priorité 2 (optionnelle)", "Priority 2 (optional)"),
        options=opts2,
        index=idx2,
        format_func=lambda k: gp_labels.get(k, k),
        key="gender_priority_2_select",
    )
    resp_set("gender_priority_2", p2)

    # Rank 3 (optional), exclude already chosen (except empty)
    opts3 = [""] + [k for k in gp_opts if k not in ["", p1, p2]]
    prev3 = (resp_get("gender_priority_3", "") or "").strip()
    idx3 = opts3.index(prev3) if prev3 in opts3 else 0
    p3 = st.selectbox(
        t(lang, "Priorité 3 (optionnelle)", "Priority 3 (optional)"),
        options=opts3,
        index=idx3,
        format_func=lambda k: gp_labels.get(k, k),
        key="gender_priority_3_select",
    )
    resp_set("gender_priority_3", p3)

    chosen_any = [x for x in [p1, p2, p3] if x]
    if "OTHER" in chosen_any:
        other = st.text_input(
            t(lang, "Autre : préciser", "Other: please specify"),
            key="gender_priority_other_input",
            value=resp_get("gender_priority_other", ""),
        )
        resp_set("gender_priority_other", (other or "").strip())
    else:
        resp_set("gender_priority_other", "")

    errs = validate_r7(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))

def rubric_8(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 8 : Capacité et faisabilité (12–24 mois)", "Section 8: Capacity and feasibility (12–24 months)", "Secção 8: Capacidade e viabilidade (12–24 meses)", "القسم 8: القدرات وإمكانية التنفيذ (12–24 شهراً)"))
    st.markdown(
        t(
            lang,
            "Évaluez le niveau de disponibilité et d’adéquation des moyens pour produire les statistiques prioritaires dans les 12–24 mois à venir.",
            "Assess the availability and adequacy of resources to produce priority statistics in the coming 12–24 months."
        )
    )

    scale = [
        (t(lang, "Élevé", "High"), "HIGH"),
        (t(lang, "Moyen", "Medium"), "MED"),
        (t(lang, "Faible", "Low"), "LOW"),
        (t(lang, UK_FR, UK_EN, UK_PT, UK_AR), "UK"),
    ]
    labels = [x[0] for x in scale]
    code_map = {x[0]: x[1] for x in scale}

    st.caption(t(lang, "Échelle : Élevé = capacité suffisante et opérationnelle ; Moyen = partiellement disponible ; Faible = insuffisant ; NSP = ne sait pas.",
                   "Scale: High = sufficient and operational; Medium = partially available; Low = insufficient; DK = does not know."))

    items_fr = [
        "Compétences statistiques disponibles",
        "Accès aux données administratives",
        "Financement disponible",
        "Outils numériques (collecte, traitement, diffusion)",
        "Cadre juridique pour le partage de données",
        "Coordination interinstitutionnelle",
    ]
    items_en = [
        "Available statistical skills",
        "Access to administrative data",
        "Available funding",
        "Digital tools (collection, processing, dissemination)",
        "Legal framework for data sharing",
        "Inter-institutional coordination",
    ]
    items = items_fr if lang == "fr" else items_en

    helps_fr = [
        "Ressources humaines : disponibilité de statisticiens/analystes qualifiés et expérience pertinente.",
        "Accès aux données administratives : disponibilité, qualité, régularité et conditions d’accès pour usage statistique.",
        "Financement : budget disponible et soutenable pour la production, y compris opérations de collecte/traitement.",
        "Outils numériques : disponibilité et adéquation des outils pour collecte, traitement, stockage, diffusion, interopérabilité (logiciels, matériel, connectivité, sécurité).",
        "Cadre juridique : existence et applicabilité des textes/accords permettant le partage de données à des fins statistiques (lois, décrets, protocoles, MoU, clauses de confidentialité).",
        "Coordination : mécanismes de coordination interinstitutionnelle (comités, conventions, échanges réguliers, standards communs).",
    ]
    helps_en = [
        "Human resources: availability of qualified statisticians/analysts and relevant experience.",
        "Access to administrative data: availability, quality, timeliness and conditions of access for statistical use.",
        "Funding: available and sustainable budget for production, including collection/processing operations.",
        "Digital tools: availability and adequacy of tools for collection, processing, storage, dissemination, interoperability (software, hardware, connectivity, security).",
        "Legal framework: existence and enforceability of texts/agreements enabling data sharing for statistical purposes (laws, decrees, protocols, MoUs, confidentiality clauses).",
        "Coordination: inter-institutional coordination mechanisms (committees, agreements, regular exchanges, shared standards).",
    ]
    helps = helps_fr if lang == "fr" else helps_en

    tbl = resp_get("capacity_table", {})
    if not isinstance(tbl, dict):
        tbl = {}

    rev_map = {v: k for k, v in code_map.items()}
    for it, hp in zip(items, helps):
        prev_code = tbl.get(it, None)
        idx = labels.index(rev_map[prev_code]) if prev_code in rev_map else None
        chosen = st.radio(tr(lang, it), options=labels, index=idx, horizontal=True, key=f"cap_{it}", help=tr(lang, hp))
        tbl[it] = code_map.get(chosen, None)

    resp_set("capacity_table", tbl)

    errs = validate_r8(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))


def rubric_9(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 9 : Harmonisation et qualité", "Section 9: Harmonization and quality", "Secção 9: Harmonização e qualidade", "القسم 9: المواءمة والجودة"))
    st.markdown(
        t(
            lang,
            "Indiquez 1 à 3 exigences attendues en matière d’harmonisation et d’assurance qualité.",
            "Indicate 1 to 3 expectations regarding harmonization and quality assurance."
        )
    )

    opts_fr = [
        "Manuels de normes et méthodes communes (par domaine) disponibles",
        "Cadre d’assurance qualité fonctionnel",
        "Procédures de validation et certification des données",
        "Mécanismes de cohérence des données nationales entre secteurs",
        "Renforcement des capacités techniques du SSN",
        "Renforcement du leadership de l’INS au sein du SSN",
        "Groupes techniques spécialisés (GTS/UA) opérationnels",
        "Autre (préciser) ",
     ]
    opts_en = [
        "Manuals on common standards and methods (by domain) available",
        "Functional quality assurance framework (quality toolkit) ",
        "Data validation and certification procedures (certified quality) ",
        "Toolkit / mechanisms for cross-sector consistency of national data",
        "Strengthening NSS technical capacity",
        "Strengthening NSO leadership within the NSS",
        "Specialized Technical Groups (STGs/AU) operational",
        "Other (specify) ",
    ]
    opts = opts_fr if lang == "fr" else opts_en
    # Stabilité mobile : initialiser le widget une seule fois
    if "r9_multiselect" not in st.session_state:
        st.session_state["r9_multiselect"] = resp_get("quality_expectations", [])
    sel = st.multiselect(t(lang, "Sélectionnez", "Select"), options=opts, key="r9_multiselect", max_selections=3, format_func=lambda x: tr(lang, x))
    resp_set("quality_expectations", sel)
    if has_other_option(sel):
        st.text_input(t(lang, "Préciser (Autre)", "Specify (Other)"),
                      key="q9_other_input", value=resp_get("quality_other", ""))
        resp_set("quality_other", st.session_state.get("q9_other_input", "").strip())
    else:
        resp_set("quality_other", "")

    errs = validate_r9(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))

    # Auto-save draft (mobile)
    autosave_draft(force=False)



def rubric_10(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 10 : Diffusion", "Section 10: Dissemination", "Secção 10: Disseminação", "القسم 10: النشر"))
    st.markdown(
        t(
            lang,
            "Indiquez 1 à 3 canaux de diffusion jugés les plus utiles pour les statistiques prioritaires.",
            "Indicate 1 to 3 dissemination channels you find most useful for priority statistics."
        )
    )
    opts_fr = [
        "Portail web / tableaux de bord",
        "Communiqués / notes de conjoncture",
        "Microdonnées anonymisées (accès sécurisé)",
        "API / Open data",
        "Ateliers et webinaires",
        "Autre",
    ]
    opts_en = [
        "Web portal / dashboards",
        "Press releases / bulletins",
        "Anonymized microdata (secure access)",
        "API / Open data",
        "Workshops and webinars",
        "Other",
    ]
    opts = opts_fr if lang == "fr" else opts_en
    # Éviter les problèmes de clic (init du widget une seule fois)
    if "r10_multiselect" not in st.session_state:
        st.session_state["r10_multiselect"] = resp_get("dissemination_channels", [])
    sel = st.multiselect(
        t(lang, "Sélectionnez", "Select"),
        options=opts,
        max_selections=3,
        key="r10_multiselect",
        format_func=lambda x: tr(lang, x),
        help=t(lang, "Choisissez les canaux de diffusion les plus utiles.", "Select the most useful dissemination channels.")
    )
    resp_set("dissemination_channels", sel)
    if has_other_option(sel):
        st.text_input(t(lang, "Préciser (Autre)", "Specify (Other)"),
                      key="q10_other_input", value=resp_get("dissemination_other", ""))
        resp_set("dissemination_other", st.session_state.get("q10_other_input", "").strip())
    else:
        resp_set("dissemination_other", "")

    errs = validate_r10(lang)
    if errs:
        st.warning(t(lang, "Contrôles qualité :", "Quality checks:"))
        st.write("\n".join([f"- {e}" for e in errs]))

    # Auto-save draft (mobile)
    autosave_draft(force=False)



def rubric_11(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 11 : Sources de données pertinentes", "Section 11: Relevant data sources", "Secção 11: Fontes de dados pertinentes", "القسم 11: مصادر البيانات ذات الصلة"))
    st.markdown(
        t(
            lang,
            "Sélectionnez **2 à 4** sources de données les plus importantes pour produire les statistiques prioritaires.",
            "Select **2 to 4** of the most important data sources to produce the priority statistics.",
        )
    )

    opts_fr = [
        "Enquêtes ménages",
        "Enquêtes entreprises",
        "Recensements",
        "Données administratives",
        "Registres état-civil",
        "Données géospatiales",
        "Données privées",
        "Autres",
    ]
    opts_en = [
        "Household surveys",
        "Enterprise surveys",
        "Censuses",
        "Administrative data",
        "Civil registration and vital statistics (CRVS)",
        "Geospatial data",
        "Private data",
        "Other",
    ]

    options = opts_fr if lang == "fr" else opts_en

    prev = resp_get("data_sources", [])
    if not isinstance(prev, list):
        prev = []

    sel = st.multiselect(
        t(
            lang,
            "2 à 4 sources de données les plus pertinentes",
            "2 to 4 most relevant data sources",
        ),
        options=options,
        default=[x for x in prev if x in options],
        max_selections=4,
        format_func=lambda x: tr(lang, x),
        help=t(
            lang,
            "Choisissez entre 2 et 4 options. Si vous choisissez Autres, précisez.",
            "Choose between 2 and 4 options. If you select Other, please specify.",
        ),
        key="data_sources_multiselect",
    )
    resp_set("data_sources", sel)

    other_label = t(lang, "Autres", "Other", "Outras", "أخرى")
    if has_other_option(sel):
        other = st.text_input(
            t(lang, "Autres : préciser", "Other: please specify"),
            key="data_sources_other_input",
            value=resp_get("data_sources_other", ""),
        )
        resp_set("data_sources_other", (other or "").strip())
    else:
        resp_set("data_sources_other", "")

    errs = validate_r11(lang)
    if errs:
        for e in errs:
            st.error(e)

def rubric_12(lang: str) -> None:
    st.subheader(t(lang, "Rubrique 12 : Questions ouvertes", "Section 12: Open questions", "Secção 12: Questões abertas", "القسم 12: الأسئلة المفتوحة"))
    st.markdown(
        t(
            lang,
            "Ces questions sont **optionnelles**. Vous pouvez les laisser vides. Toutefois, elles sont présentées **une à une** pour faciliter la saisie de vos opinions ou sentiments sur le système statistique.",
            "These questions are **optional**. You may leave them blank. They are presented **one by one** to facilitate completion.",
            "Estas perguntas são **opcionais**. Pode deixá-las em branco. Contudo, são apresentadas **uma a uma** para facilitar a introdução das suas opiniões ou perceções sobre o sistema estatístico.",
            "هذه الأسئلة **اختيارية**. يمكنكم تركها فارغة. ومع ذلك، تُعرض **واحدة تلو الأخرى** لتسهيل إدخال آرائكم أو انطباعاتكم حول النظام الإحصائي."
        )
    )

    if "r12_substep" not in st.session_state:
        st.session_state["r12_substep"] = 0  # 0..2=open questions, 3=confirmation

    sub = int(st.session_state.get("r12_substep", 0) or 0)

    if sub == 0:
        st.markdown("#### " + t(lang, "Question 1 / 3", "Question 1 / 3"))
        q1 = st.text_area(
            t(lang, "1) Commentaires / recommandations clés", "1) Key comments / recommendations", "1) Comentários / recomendações principais", "1) التعليقات / التوصيات الأساسية"),
            value=resp_get("open_q1", ""),
            height=160,
            key="open_q1_input",
        )
        resp_set("open_q1", (q1 or "").strip())
        if not resp_get("open_q1", ""):
            st.warning(t(lang, "Alerte : la question 1 est vide (vous pouvez tout de même continuer).",
                         "Warning: question 1 is empty (you can still proceed)."))

    elif sub == 1:
        st.markdown("#### " + t(lang, "Question 2 / 3", "Question 2 / 3"))
        q2 = st.text_area(
            t(
                lang,
                "2) Un ou des indicateur(s) statistique(s) socio-économique(s) essentiel(s) manquant(s) et justification(s)",
                "2) One or more missing essential socio-economic statistical indicator(s) and justification(s)",
                "2) Um ou mais indicadores estatísticos socioeconómicos essenciais em falta e respetiva justificação",
                "2) مؤشر أو أكثر من المؤشرات الإحصائية الاجتماعية والاقتصادية الأساسية غير المدرجة مع المبررات",
            ),
            value=resp_get("open_q2", ""),
            height=160,
            key="open_q2_input",
        )
        resp_set("open_q2", (q2 or "").strip())
        if not resp_get("open_q2", ""):
            st.warning(t(lang, "Alerte : la question 2 est vide (vous pouvez tout de même continuer).",
                         "Warning: question 2 is empty (you can still proceed)."))

    elif sub == 2:
        st.markdown("#### " + t(lang, "Question 3 / 3", "Question 3 / 3"))
        q3 = st.text_area(
            t(lang, "3) Besoins de soutien (technique, financier, etc.)", "3) Support needs (technical, financial, etc.)", "3) Necessidades de apoio (técnico, financeiro, etc.)", "3) احتياجات الدعم (الفني، المالي، إلخ)"),
            value=resp_get("open_q3", ""),
            height=160,
            key="open_q3_input",
        )
        resp_set("open_q3", (q3 or "").strip())
        if not resp_get("open_q3", ""):
            st.warning(t(lang, "Alerte : la question 3 est vide (vous pouvez tout de même continuer).",
                         "Warning: question 3 is empty (you can still proceed)."))

    else:
        st.markdown("#### " + t(lang, "Confirmation", "Confirmation"))
        st.info(
            t(
                lang,
                "Dernière étape : merci d’indiquer si vous avez consulté d’autres collègues. Cette question est obligatoire.",
                "Final step: please indicate whether you consulted other colleagues. This question is mandatory.",
            )
        )

        cc_opts = ["", "YES", "NO"]
        cc_labels = {"YES": t(lang, "Oui", "Yes"), "NO": t(lang, "Non", "No")}
        prev_cc = (resp_get("consulted_colleagues", "") or "").strip()
        idx_cc = cc_opts.index(prev_cc) if prev_cc in cc_opts else 0
        chosen_cc = st.selectbox(
            t(
                lang,
                "Avez-vous consulté d’autres collègues pour remplir ce questionnaire ?",
                "Did you consult other colleagues to complete this questionnaire?",
            ),
            options=cc_opts,
            index=idx_cc,
            format_func=lambda k: (t(lang, "— Sélectionner —", "— Select —") if k == "" else cc_labels.get(k, k)),
            key="consulted_colleagues_select",
        )
        resp_set("consulted_colleagues", chosen_cc)

    st.markdown("---")
    col1, col2, col3 = st.columns([1, 1, 3])
    with col1:
        prev_disabled = (sub <= 0)
        if st.button(t(lang, "⬅ Question précédente", "⬅ Previous question"), disabled=prev_disabled, key="r12_prev_btn"):
            st.session_state["r12_substep"] = max(0, sub - 1)
            st.rerun()
    with col2:
        if sub < 2:
            next_label = t(lang, "Question suivante ➡", "Next question ➡")
        elif sub == 2:
            next_label = t(lang, "Aller à la confirmation ➡", "Go to confirmation ➡")
        else:
            next_label = t(lang, "OK (rubrique terminée)", "OK (section completed)")
        next_disabled = (sub >= 3)
        if st.button(next_label, disabled=next_disabled, key="r12_next_btn"):
            st.session_state["r12_substep"] = min(3, sub + 1)
            st.rerun()
    with col3:
        st.caption(t(lang, "Progression : 1/3 → 2/3 → 3/3 → Confirmation.",
                     "Progress: 1/3 → 2/3 → 3/3 → Confirmation."))

    errs = validate_r12(lang)
    if errs:
        for e in errs:
            st.error(e)

def rubric_send(lang: str, df_long: pd.DataFrame) -> None:
    st.subheader(t(lang, "ENVOYER le questionnaire", "SUBMIT questionnaire", "ENVIAR questionário", "إرسال الاستبيان"))

    errors = validate_all(lang)
    if errors:
        st.error(t(lang, "Le questionnaire contient des erreurs bloquantes :",
                   "There are blocking errors:"))
        st.write("\n".join([f"- {e}" for e in errors]))
        st.info(t(lang, "Retournez aux rubriques concernées via la navigation.",
                  "Go back to the relevant sections using navigation."))
        return

    # Optional warnings
    if not resp_get("open_q1", "") or not resp_get("open_q2", "") or not resp_get("open_q3", ""):
        st.warning(t(lang,
                     "Certaines questions ouvertes sont vides (optionnel).",
                     "Some open questions are empty (optional)."))

    st.info(t(
        lang,
        "Tout est prêt. Cliquez sur **ENVOYER** pour soumettre votre questionnaire.",
        "Everything is ready. Click **SUBMIT** to send your questionnaire."
    ))

    # Empêcher les envois multiples (par email + par session)
    email = (resp_get("email", "") or "").strip()
    already_in_session = bool(st.session_state.get("submitted_once", False))

    # Pré-contrôle email (local SQLite)
    already_in_db = False
    if email:
        try:
            already_in_db = db_email_exists(email)
        except Exception:
            # Ne pas bloquer si la base locale est momentanément inaccessible (rare)
            st.warning(t(
                lang,
                "Impossible de vérifier pour l’instant si cet email a déjà répondu (base locale). "
                "Vous pouvez continuer ; en cas de doublon, l’envoi pourra être refusé au moment de l’enregistrement.",
                "Cannot check right now whether this email already submitted (local database). "
                "You can continue; duplicates may be rejected during saving."
            ))
            already_in_db = False

    if already_in_db:
        st.error(t(
            lang,
            "Ce questionnaire a déjà été envoyé avec cet email. Un seul envoi est autorisé.",
            "This questionnaire has already been submitted with this email. Only one submission is allowed."
        ))

    if already_in_session and not already_in_db:
        st.info(t(
            lang,
            "Ce navigateur a déjà effectué un envoi. Pour un nouvel envoi, utilisez un autre email / session.",
            "This browser session already submitted once. For a new submission, use another email / session."
        ))

    disable_submit = already_in_db or already_in_session

    if st.button(t(lang, "✅ ENVOYER et enregistrer", "✅ SUBMIT and save"), disabled=disable_submit):
        submission_id = str(uuid.uuid4())
        payload = st.session_state.responses.copy()
        payload["submission_id"] = submission_id
        payload["submitted_at_utc"] = now_utc_iso()
        payload["scoring_version"] = SCORING_VERSION

        try:
            db_save_submission(submission_id, lang, email, payload)
        except Exception as e:
            st.error(t(lang, "Échec d'enregistrement (base locale) : ",
                       "Save failed (local database): ") + str(e))

            # Anti-blocage : fournir une sauvegarde JSON immédiate
            json_text = json.dumps(payload, ensure_ascii=False, indent=2)
            st.download_button(
                label=t(lang, "⬇️ Télécharger la sauvegarde JSON", "⬇️ Download JSON backup"),
                data=json_text.encode("utf-8"),
                file_name=f"submission_{submission_id}.json",
                mime="application/json",
                key=f"dl_json_{submission_id}",
            )
            st.text_area(
                t(lang, "Copie du JSON (à conserver)", "JSON copy (keep it)"),
                value=json_text,
                height=260,
                key=f"json_copy_{submission_id}",
            )
            st.info(t(lang,
                      "Conservez ce fichier. Vous pourrez le transmettre pour intégration ultérieure.",
                      "Keep this file. You can share it later for integration."))
            return

        st.success(t(
            lang,
            "Merci ! Votre questionnaire a été enregistré.",
            "Thank you! Your submission has been saved."
        ))

        # Fournir aussi une copie JSON (utile si l’hébergement redémarre)
        json_text = json.dumps(payload, ensure_ascii=False, indent=2)
        st.download_button(
            label=t(lang, "⬇️ Télécharger une copie JSON", "⬇️ Download a JSON copy"),
            data=json_text.encode("utf-8"),
            file_name=f"submission_{submission_id}.json",
            mime="application/json",
            key=f"dl_json_ok_{submission_id}",
        )

        # Bloquer les envois multiples pour cette session
        st.session_state.submitted_once = True
        st.caption(f"ID : {submission_id}")
        st.info(t(lang, "Envoi terminé. Vous pouvez fermer cette page.",
                  "Submission complete. You can close this page."))
        st.session_state.submission_id = submission_id


# =========================
# Admin dashboard
# =========================

def admin_login(lang: str) -> None:
    st.subheader(t(lang, "Administration", "Administration"))
    pw = st.text_input(t(lang, "Mot de passe admin", "Admin password"), type="password")

    # Si le répondant saisit le nom du secret au lieu du mot de passe
    if (pw or "").strip() in ("ADMIN_PASSWORD", "SUPERADMIN_PASSWORD"):
        st.info(
            t(
                lang,
                "Astuce : saisissez le mot de passe réel, pas le nom du secret (ex. ADMIN_PASSWORD).",
                "Tip: enter the actual password value, not the secret name (e.g., ADMIN_PASSWORD).",
            )
        )

    # Diagnostic discret : présence du mot de passe (sans afficher la valeur)
    src_code, src_lbl = get_admin_auth_source()
    if src_code == "none":
        st.caption(
            t(
                lang,
                "ADMIN_PASSWORD : non configuré (secrets, variable d’environnement ou mot de passe haché en base).",
                "ADMIN_PASSWORD : not configured (secrets, environment variable, or hashed password in DB).",
            )
        )
    else:
        st.caption(
            t(
                lang,
                f"ADMIN_PASSWORD : configuré via {src_lbl}.",
                f"ADMIN_PASSWORD : configured via {src_lbl}.",
            )
        )


    # Diagnostic avancé (invisible par défaut) : ajouter ?diag=1 à l’URL
    # Utile pour vérifier rapidement que SUPERADMIN_PASSWORD est bien injecté (sans afficher la valeur).
    qp = get_query_params()
    if qp.get("diag", ["0"])[0] == "1":
        v_sa, src_sa = _get_secret_or_env("SUPERADMIN_PASSWORD")
        if v_sa and src_sa:
            st.caption(
                t(
                    lang,
                    f"SUPERADMIN_PASSWORD : configuré via {src_sa}.",
                    f"SUPERADMIN_PASSWORD : configured via {src_sa}.",
                )
            )
        else:
            st.caption(
                t(
                    lang,
                    "SUPERADMIN_PASSWORD : non configuré (secrets ou variable d’environnement).",
                    "SUPERADMIN_PASSWORD : not configured (secrets or environment variable).",
                )
            )

    if st.button(t(lang, "Se connecter", "Login")):
        # Superadmin : même écran, mais rôle différent (invisible tant que non authentifié)
        if verify_superadmin_password(pw):
            st.session_state.admin_authed = True
            st.session_state.admin_role = "superadmin"
            st.success(t(lang, "Connexion réussie.", "Logged in."))
            st.rerun()
        elif verify_admin_password(pw):
            st.session_state.admin_authed = True
            st.session_state.admin_role = "admin"
            st.success(t(lang, "Connexion réussie.", "Logged in."))
            st.rerun()
        else:
            st.error(
                t(
                    lang,
                    "Mot de passe incorrect ou secret ADMIN_PASSWORD manquant.",
                    "Incorrect password or missing ADMIN_PASSWORD secret.",
                )
            )


def admin_dashboard(lang: str) -> None:
    st.subheader(t(lang, "Tableau de bord admin", "Admin dashboard"))

    # Load data from SQLite
    df = db_read_submissions(limit=20000)
    st.metric(t(lang, "Nombre de réponses", "Number of responses"), len(df))

    if df.empty:
        st.info(t(lang, "Aucune réponse pour le moment.", "No responses yet."))
        return

    # Parse payloads
    payloads = []
    for _, r in df.iterrows():
        try:
            payloads.append(json.loads(r["payload_json"]))
        except Exception:
            payloads.append({})

    # Flat view for quick export
    flat = pd.DataFrame([flatten_payload(p) for p in payloads])
    flat.insert(0, "submission_id", df["submission_id"].values)
    flat.insert(1, "submitted_at_utc", df["submitted_at_utc"].values)

    is_super = st.session_state.get("admin_role") == "superadmin"

    if is_super:
        tab_quick, tab_super, tab_sec = st.tabs([
            t(lang, "Vue rapide", "Quick view"),
            t(lang, "Analyse avancée", "Advanced analysis"),
            t(lang, "Sécurité", "Security"),
        ])
    else:
        tab_quick, = st.tabs([t(lang, "Vue rapide", "Quick view")])

    with tab_quick:
        st.dataframe(flat, use_container_width=True)

        # Export Excel (flat + raw)
        out = io.BytesIO()
        with pd.ExcelWriter(out, engine="openpyxl") as writer:
            flat.to_excel(writer, sheet_name="submissions", index=False)
            df.to_excel(writer, sheet_name="raw_json", index=False)
        out.seek(0)

        st.download_button(
            t(lang, "Exporter en Excel", "Export to Excel"),
            data=out.getvalue(),
            file_name="consultation_stat_niang_export.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        # Export CSV (UTF-8 avec BOM pour Excel)
        csv_flat = flat.to_csv(index=False).encode("utf-8-sig")
        csv_raw = df.to_csv(index=False).encode("utf-8-sig")

        st.download_button(
            t(lang, "Exporter en CSV (vue à plat)", "Export CSV (flat view)"),
            data=csv_flat,
            file_name="consultation_submissions_flat.csv",
            mime="text/csv",
        )

        st.download_button(
            t(lang, "Exporter en CSV (brut : JSON)", "Export CSV (raw : JSON)"),
            data=csv_raw,
            file_name="consultation_submissions_raw.csv",
            mime="text/csv",
        )

        # Export JSON Lines (JSONL) : 1 ligne = 1 soumission
        jsonl_bytes = ("\n".join([json.dumps(p, ensure_ascii=False) for p in payloads if isinstance(p, dict)]) + "\n").encode("utf-8")
        st.download_button(
            t(lang, "Exporter en JSONL (1 ligne = 1 soumission)", "Export JSONL (1 line per submission)"),
            data=jsonl_bytes,
            file_name="consultation_submissions.jsonl",
            mime="application/json",
        )

        # Télécharger la base SQLite
        st.caption(t(lang, f"Base locale : {DB_PATH}", f"Local database : {DB_PATH}"))
        try:
            with open(DB_PATH, "rb") as f:
                db_bytes = f.read()
            st.download_button(
                t(lang, "Télécharger la base locale (responses.db)", "Download local database (responses.db)"),
                data=db_bytes,
                file_name="responses.db",
                mime="application/x-sqlite3",
            )
        except Exception:
            st.warning(t(lang, "Base SQLite introuvable pour le moment.", "SQLite database not found yet."))

        # Export ZIP complet (DB + CSV + JSON par soumission)
        zip_buf = io.BytesIO()
        try:
            with zipfile.ZipFile(zip_buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
                # DB
                try:
                    with open(DB_PATH, "rb") as f:
                        z.writestr("responses.db", f.read())
                except Exception:
                    pass

                # CSV
                z.writestr("consultation_submissions_flat.csv", csv_flat)
                z.writestr("consultation_submissions_raw.csv", csv_raw)
                z.writestr("consultation_submissions.jsonl", jsonl_bytes)

                # JSON individuels
                for p in payloads:
                    sid = str(p.get("submission_id") or "")
                    if not sid:
                        continue
                    z.writestr(
                        f"json/submission_{sid}.json",
                        json.dumps(p, ensure_ascii=False, indent=2).encode("utf-8"),
                    )

            zip_buf.seek(0)
            st.download_button(
                t(lang, "Exporter en ZIP (DB + CSV + JSON)", "Export ZIP (DB + CSV + JSON)"),
                data=zip_buf.getvalue(),
                file_name="consultation_export.zip",
                mime="application/zip",
            )
        except Exception:
            st.warning(t(lang, "Export ZIP indisponible (erreur technique).", "ZIP export unavailable (technical error)."))

        st.info(
            t(
                lang,
                "Astuce : utilisez ?admin=1 dans l’URL pour afficher l’espace admin.",
                "Tip: use ?admin=1 in the URL to access the admin space."
            )
        )

    if not is_super:
        return

    with tab_super:
        # --- Build a richer dataset for analysis
        df_super = pd.DataFrame(payloads)
        df_super["submission_id"] = df["submission_id"].values
        df_super["submitted_at_utc"] = pd.to_datetime(df["submitted_at_utc"], errors="coerce", utc=True)

        # Filters
        st.markdown("### " + t(lang, "Filtres", "Filters"))

        colf1, colf2, colf3 = st.columns(3)
        with colf1:
            countries = sorted([c for c in df_super.get("pays", pd.Series(dtype=str)).dropna().unique().tolist() if str(c).strip() != ""])
            sel_countries = st.multiselect(t(lang, "Pays de résidence", "Country of residence"), options=countries, default=[], key="f_country")
        with colf2:
            actors = sorted([a for a in df_super.get("type_acteur", pd.Series(dtype=str)).dropna().unique().tolist() if str(a).strip() != ""])
            sel_actors = st.multiselect(t(lang, "Type d’acteur", "Stakeholder type"), options=actors, default=[], key="f_actor")
        with colf3:
            # Period filter
            min_dt = df_super["submitted_at_utc"].min()
            max_dt = df_super["submitted_at_utc"].max()
            if pd.isna(min_dt) or pd.isna(max_dt):
                min_dt = pd.Timestamp.utcnow() - pd.Timedelta(days=30)
                max_dt = pd.Timestamp.utcnow()
            date_range = st.date_input(
                t(lang, "Période", "Period"),
                value=(min_dt.date(), max_dt.date()),
                key="f_period"
            )

        filtered = df_super.copy()
        if sel_countries:
            filtered = filtered[filtered["pays"].isin(sel_countries)]
        if sel_actors:
            filtered = filtered[filtered["type_acteur"].isin(sel_actors)]
        if isinstance(date_range, (list, tuple)) and len(date_range) == 2:
            # Streamlit date_input returns datetime.date; our column may be tz-aware (UTC).
            col = pd.to_datetime(filtered["submitted_at_utc"], utc=True, errors="coerce")

            start_d = pd.Timestamp(date_range[0])
            end_d = pd.Timestamp(date_range[1]) + pd.Timedelta(days=1) - pd.Timedelta(seconds=1)

            if start_d.tz is None:
                start_d = start_d.tz_localize("UTC")
            else:
                start_d = start_d.tz_convert("UTC")

            if end_d.tz is None:
                end_d = end_d.tz_localize("UTC")
            else:
                end_d = end_d.tz_convert("UTC")

            filtered = filtered[(col >= start_d) & (col <= end_d)]

        st.caption(t(lang, f"Réponses filtrées : {len(filtered)}", f"Filtered responses: {len(filtered)}"))

        if filtered.empty:
            st.warning(t(lang, "Aucune réponse dans ce filtre.", "No responses match these filters."))
            return

        # Longlist for labels (domain/stat)
        df_long = load_longlist()
        dom_lbl = domain_label_map(df_long, lang)
        stat_lbl = stat_label_map(df_long, lang)

        # --- Build aggregated prioritization table
        rows = []
        for _, p in filtered.iterrows():
            top5 = p.get("top5_domains", []) or []
            sel_by_dom = p.get("selected_by_domain", {}) or {}
            scoring = p.get("scoring", {}) or {}
            sid = p.get("submission_id", "")
            for d, stats in (sel_by_dom.items() if isinstance(sel_by_dom, dict) else []):
                if not isinstance(stats, list):
                    continue
                for s in stats:
                    sc = scoring.get(s, {})
                    avail_raw = (sc.get("availability", sc.get("gap", 0)))
                    avail = normalize_availability(avail_raw, p.get("scoring_version", 0))
                    dem = int(sc.get("demand", 0) or 0)
                    fea = int(sc.get("feasibility", 0) or 0)
                    overall = (avail + dem + fea) / 3.0
                    rows.append({
                        "submission_id": sid,
                        "pays": p.get("pays", ""),
                        "type_acteur": p.get("type_acteur", ""),
                        "domain_code": d,
                        "domain_label": dom_lbl.get(d, d),
                        "stat_code": s,
                        "stat_label": stat_lbl.get(s, s),
                        "availability": avail, "demand": dem, "feasibility": fea, "overall": overall
                    })

        df_rows = pd.DataFrame(rows)
        if df_rows.empty:
            st.warning(t(lang, "Aucune statistique notée dans ces réponses.", "No scored indicators in these responses."))
            return

        # Aggregation
        by_stat = df_rows.groupby(["domain_code", "domain_label", "stat_code", "stat_label"], as_index=False).agg(
            n=("submission_id", "nunique"),
            mean_availability=("availability", "mean"),
            mean_demand=("demand", "mean"),
            mean_feasibility=("feasibility", "mean"),
            mean_overall=("overall", "mean"),
        ).sort_values(["domain_code", "mean_overall", "n"], ascending=[True, False, False])

        by_domain = df_rows.groupby(["domain_code", "domain_label"], as_index=False).agg(
            n_stats=("stat_code", "count"),
            n_submissions=("submission_id", "nunique"),
            mean_overall=("overall", "mean"),
        ).sort_values(["mean_overall", "n_submissions"], ascending=[False, False])

        st.markdown("### " + t(lang, "Tableau de priorisation agrégé", "Aggregated prioritization table"))
        st.dataframe(by_stat, use_container_width=True, height=420)

        # Export aggregated Excel
        out2 = io.BytesIO()
        with pd.ExcelWriter(out2, engine="openpyxl") as writer:
            by_domain.to_excel(writer, sheet_name="by_domain", index=False)
            by_stat.to_excel(writer, sheet_name="by_statistic", index=False)
            df_rows.to_excel(writer, sheet_name="scored_rows", index=False)
        out2.seek(0)

        st.download_button(
            t(lang, "Télécharger l’agrégé (Excel)", "Download aggregated (Excel)"),
            data=out2.getvalue(),
            file_name="prioritization_aggregated.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        # Rich Word report
        st.markdown("### " + t(lang, "Rapport Word (publication)", "Word report (publication)"))
        st.caption(t(lang, "Génère un rapport enrichi avec graphiques et annexes.", "Generates an enriched report with charts and annexes."))

        if st.button(t(lang, "Générer le rapport Word", "Generate Word report"), key="btn_word_publication"):
            try:
                doc_bytes = build_publication_report_docx(lang, filtered, by_domain, by_stat, df_rows)
                st.download_button(
                    t(lang, "Télécharger le rapport (.docx)", "Download report (.docx)"),
                    data=doc_bytes,
                    file_name="rapport_publication_priorisation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                st.success(t(lang, "Rapport généré.", "Report generated."))


            except Exception as e:


                st.error(f"Word : {e}")


    with tab_sec:
        st.subheader(t(lang, "Sécurité", "Security"))

        st.markdown("### " + t(lang, "Mot de passe admin", "Admin password"))
        src_code, src_lbl = get_admin_auth_source()
        if src_code == "db":
            st.info(t(lang, "Mot de passe admin : défini en base (haché).", "Admin password: set in database (hashed)."))
        elif src_code in ("secrets", "env"):
            st.info(t(lang, f"Mot de passe admin : défini via {src_lbl}.", f"Admin password: set via {src_lbl}."))
        else:
            st.warning(t(lang, "Mot de passe admin : non configuré.", "Admin password: not configured."))

        st.markdown(t(
            lang,
            "#### Changer le mot de passe admin (réservé au superadmin)",
            "#### Change admin password (superadmin only)"
        ))

        new1 = st.text_input(t(lang, "Nouveau mot de passe admin", "New admin password"), type="password", key="new_admin_pw1")
        new2 = st.text_input(t(lang, "Confirmer le nouveau mot de passe", "Confirm new password"), type="password", key="new_admin_pw2")

        if st.button(t(lang, "Mettre à jour le mot de passe admin", "Update admin password")):
            p1 = (new1 or "").strip()
            p2 = (new2 or "").strip()
            if len(p1) < 10:
                st.error(t(lang, "Mot de passe trop court (≥ 10 caractères).", "Password too short (≥ 10 characters)."))
            elif p1 != p2:
                st.error(t(lang, "Les deux champs ne correspondent pas.", "The two fields do not match."))
            else:
                try:
                    set_admin_password(p1)
                    st.success(t(lang, "Mot de passe admin mis à jour.", "Admin password updated."))
                    st.info(t(lang, "Le nouveau mot de passe est actif immédiatement (sans redeploy).",
                              "The new password is active immediately (no redeploy needed)."))
                except Exception as e:
                    st.error(f"{e}")

        st.divider()
        st.markdown(t(
            lang,
            "#### Réinitialiser vers Secrets / variable d’environnement",
            "#### Reset to Secrets / environment variable"
        ))
        st.caption(t(
            lang,
            "Cette action supprime le mot de passe haché stocké en base. L’app utilisera alors ADMIN_PASSWORD (secrets/env).",
            "This removes the hashed password stored in DB. The app will then use ADMIN_PASSWORD (secrets/env)."
        ))

        if st.button(t(lang, "Réinitialiser le mot de passe admin", "Reset admin password")):
            reset_admin_password_to_secrets_env()
            v, src = _get_secret_or_env("ADMIN_PASSWORD")
            if v and src:
                st.success(t(lang, f"Réinitialisation effectuée : retour à {src}.", f"Reset done: back to {src}."))
            else:
                st.warning(t(lang, "Réinitialisation effectuée, mais ADMIN_PASSWORD n’est pas défini dans secrets/env.",
                            "Reset done, but ADMIN_PASSWORD is not defined in secrets/env."))





def apply_language_direction(lang: str) -> None:
    if lang == "ar":
        st.markdown(
            """
            <style>
            .block-container, .stMarkdown, .stText, label, p, h1, h2, h3, h4, h5, h6 { direction: rtl; text-align: right; }
            div[data-testid="stHorizontalBlock"] { direction: rtl; }
            </style>
            """,
            unsafe_allow_html=True,
        )


def render_language_switch(lang: str) -> str:
    st.markdown(
        """
        <style>
        .lang-banner {
            border: 1px solid rgba(49, 51, 63, 0.2);
            padding: 0.9rem 1rem 0.35rem 1rem;
            border-radius: 0.9rem;
            margin-bottom: 1rem;
            background: rgba(240, 242, 246, 0.65);
        }
        .lang-banner-title {
            font-size: 1rem;
            font-weight: 700;
            margin-bottom: 0.35rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    st.markdown(f"<div class='lang-banner'><div class='lang-banner-title'>{t(lang, '🌐 Choisissez votre langue', '🌐 Choose your language', '🌐 Escolha o seu idioma', '🌐 اختر اللغة')}</div></div>", unsafe_allow_html=True)
    options = LANG_OPTIONS
    idx = options.index(lang) if lang in options else 0
    chosen = st.radio(
        label=t(lang, "Langue / Language", "Language", "Idioma", "اللغة"),
        options=options,
        index=idx,
        horizontal=True,
        format_func=lambda x: LANG_LABELS.get(x, x),
        key="visible_lang_switch",
        label_visibility="collapsed",
    )
    st.session_state.lang = chosen
    return chosen


# =========================
# Main
# =========================

def main() -> None:
    # set_page_config déjà défini en haut du fichier
    init_session()
    # Initialise la base locale (créée dans le même dossier que ce script)
    db_init()
    maybe_restore_draft()

    # Language toggle (prominent on main screen + compact reminder in sidebar)
    lang = render_language_switch(st.session_state.lang)
    st.sidebar.markdown("### 🌐 " + t(lang, "Langue", "Language", "Idioma", "اللغة"))
    st.sidebar.caption(" / ".join([LANG_LABELS[k] for k in LANG_OPTIONS]))
    apply_language_direction(lang)

    # Admin access
    qp = get_query_params()
    is_admin = "admin" in qp and qp["admin"] and qp["admin"][0] in ["1", "true", "yes"]

    df_long = load_longlist()

    # Header
    st.title(t(lang, APP_TITLE_FR, APP_TITLE_EN, APP_TITLE_PT, APP_TITLE_AR))
    st.caption(t(lang, "Application unifiée (FR/EN/PT/AR) – codes masqués – contrôles qualité intégrés.",
                 "Unified app (FR/EN/PT/AR) – hidden codes – built-in quality controls.",
                 "Aplicação multilingue (FR/EN/PT/AR) – códigos ocultos – controlos de qualidade integrados.",
                 "تطبيق متعدد اللغات (FR/EN/PT/AR) – رموز مخفية – ضوابط جودة مدمجة."))
    if st.session_state.get("draft_exists") and not st.session_state.get("draft_resume_notice_shown"):
        st.warning(
            t(
                lang,
                "La saisie est sauvegardée. En cas de suspension de moins de 48 heures, reprenez-la là où vous vous étiez arrêté en ré-ouvrant le lien contenant rid (à conserver / mettre en favori / retrouver dans l’historique).",
                "Your entry is saved. If you pause for less than 48 hours, resume where you left off by reopening the link containing rid (bookmark / save / find it in your browser history).",
            ),
            icon="💾",
        )
        st.session_state.draft_resume_notice_shown = True

    # Sidebar navigation
    steps = get_steps(lang)
    render_sidebar(lang, steps)

    # Admin view
    if is_admin:
        if not st.session_state.admin_authed:
            admin_login(lang)
            return
        admin_dashboard(lang)
        return

    # Normal flow
    step_key = steps[st.session_state.nav_idx][0]

    if step_key == "R1":
        rubric_1(lang)
    elif step_key == "R2":
        rubric_2(lang)
    elif step_key == "R3":
        rubric_3(lang)
    elif step_key == "R4":
        rubric_4(lang, df_long)
    elif step_key == "R5":
        rubric_5(lang, df_long)
    elif step_key == "R6":
        rubric_6(lang)
    elif step_key == "R7":
        rubric_7(lang)
    elif step_key == "R8":
        rubric_8(lang)
    elif step_key == "R9":
        rubric_9(lang)
    elif step_key == "R10":
        rubric_10(lang)
    elif step_key == "R11":
        rubric_11(lang)
    elif step_key == "R12":
        rubric_12(lang)
    elif step_key == "SEND":
        rubric_send(lang, df_long)

    st.divider()
    nav_buttons(lang, steps, df_long)


if __name__ == "__main__":
    main()
